"""
================================================================================
TOUR ITINERARY RECOMMENDATION ENGINE — PHASE 2
v20260521c
================================================================================
Paste an agent email and get a structured itinerary recommendation back,
drawn from 6,520 confirmed historical tours in master.db.

USAGE IN COLAB:
  1. Run Cell 1 (Setup) once per session
  2. Paste email into Cell 2 and run

DATA SOURCES:
  master.db               — 6,520 historical tours (hotels, activities, monuments)
  hotel_portfolio.csv     — 2,802 hotels · 354 cities (tier, heritage flag)
  activity_details.csv    — 3,105 activities · 172 cities (theme_score, DKC flag)
  fnb_details.csv         — 1,987 F&B entries · 255 cities (tci_score, hotel linkage)

DECISION HIERARCHY:
  Hotels      → market + city + tier → market + city → all markets + city
  Activities  → DB market + city → DB all markets → CSV portfolio (theme_score)
  Monuments   → market + city → all markets + city
  F&B         → CSV portfolio by city, ranked by tci_score (hotel-linked + standalone)
  Vehicles    → num_pax rule-based
  Guides      → flag if language escort requested

CHANGELOG:
v20260522f  N4 + N5 — Score cap + cross-month date range (26 May 2026).
            N4: Score cap was missing from historical_query.py (unversioned
            canonical). historical_query_v20260522b.py had the fix
            (min(round(score*100), 100)) but the unversioned file loaded
            by Colab did not — causing 278% display for Trendtours (B2 #G2).
            Fix: use historical_query_v20260522b.py in Colab (copy as
            historical_query.py). No engine code change needed.
            N5: Cross-month spelled-out date range not parsed. B2 #002
            "arriving in Mumbai on 19 Oct 2026 and departing from Delhi on
            04 November" = 16 nights — was not detected. P1-06 only handled
            same-month ranges ("January 6–20"). New block added after P1-06:
            matches two separate day+month mentions (day-first and month-first)
            up to 80 chars apart, with year on the second mention. EN/ES/PT/FR
            month names. Validated: "19 Oct ... 04 November 2026" → 16n ✅.
v20260522e  P1-06 Spelled-out month date-range parsing.
            "January 6–20, 2027" style dates were not detected — engine
            was falling back to city-count estimate (8n instead of 14n).
            Root cause: existing B26 parser only handled numeric formats
            (DD/MM/YYYY). New block added immediately after B26:
            - Month-first: "January 6–20, 2027" / "March 14 – 28, 2027"
            - Day-first: "6–20 January 2027" / "14 – 28 March 2027"
            - Em-dash (–), hyphen (-), spaced variants all matched
            - EN/ES/PT/FR month names all covered
            - Only fires when B26 numeric parser did not already match
            - Sets duration_nights, travel_start, travel_end,
              _duration_from_date_range (authoritative — suppresses
              free-text "N nights" override downstream)
            Validated: "January 6–20, 2027" → 14n ✅
v20260522d  Structured return + vehicle on intent dict.
            1. intent['_vehicle'] written after vehicle computation —
               vehicle was computed and displayed but never stored.
               Makes vehicle available to downstream consumers
               (Word doc, API, Outlook surface).
            2. recommend() now returns a clean structured dict instead
               of raw intent. FastAPI can return recommend(email) directly.
               Keys: parsed, itineraries, hotels, activities, monuments,
               fnb, vehicle, domestic_flights, proposal_path, llm_used.
               Breaking change for Colab display cells — use shim:
               intent = result['parsed']
               intent['historical_options'] = result['itineraries'] etc.
v20260522c  Verification regression sweep — 3 new findings from Batch 1/2/3
            re-run against v20260522b (22 May 2026) closed:
            1. N1 — False heritage upgrade on bare 'fort' substring. Heritage
               keywords split into STANDALONE (heritage, palace, haveli,
               boutique — fire on their own) and AMBIGUOUS (fort, historic,
               historical, traditional, authentic, character, royal — only
               fire when within ~40 chars of a stay-context word like 'hotel',
               'stay', 'property', 'palaces' etc.). Closes B3 #002 ('Amber Fort
               then drive' false-fire) and B3 #004 ('Fort Kochi' false-fire).
               Existing positives (B3 #002 heritage detection for legit
               palace/haveli requests) preserved.
            2. N2 — Rajasthan sub-region false-trigger on locator mentions
               ('Chand Baori, Abhneri, Rajasthan'). Sub-region keyword scan
               now requires destination/stay intent context (tour/trip/visit/
               nights/days/in/to + duration tokens) within ~30 chars of the
               keyword. Closes B3 #003 ('rajasthan' sub-region falsely set).
               Applied uniformly across all 17 sub-regions for consistency.
            3. N3 — Fatehpur Sikri duplicated in route (B3 #002 tail:
               'Agra 1n → Fatehpur Sikri 1n' alongside earlier
               'Fatehpur Sikri (day)'). fix_transit_stops() now coerces
               Fatehpur Sikri to 0n, dedupes consecutive Fatehpur Sikri
               entries, and transfers any freed nights to Agra (preferred)
               or Jaipur. Route total preserved.
            Also: Spanish/Portuguese budget tier keywords added —
            'presupuesto bajo', 'bajo presupuesto', 'presupuesto reducido',
            'presupuesto limitado', 'orçamento baixo', 'baixo orçamento'.
            Closes the B3 #003 budget-signal gap (Finding #16 in B3 review).
v20260522b  Quick Wins sweep — 5 open items closed:
            1. P2-04 third wildlife trigger site (~13205, _wl_requested in
               central India city-injection block) — bare 'tiger'/'jungle'
               replaced with tightened pattern matching display site (~12248)
               and route-validation site (~13315). All three P2-04 sites now
               consistent.
            2. P1-09 rajistan typo variants — 'rajistan', 'rajisthan', 'rajstan'
               added to all four north_india keyword locations: explicit body
               region check (~2214), REGION_KEYWORDS dict (~2393), _north_kws
               post-parse cleanup (~11804), cascade _north_india_explicit
               check (~12754).
            3. P2-08 score cap — historical_query print_options + print_wildlife_
               options now cap displayed score at 100% via min(round(score*100),
               100). BL2-γ city-set match can push _score above 1.0; display
               was showing 104% etc.
            4. P2-11 1-pax vehicle flag — recommend_vehicle(1) now returns a
               note flagging Innova as oversized for solo traveller, suggesting
               sedan (Swift Dzire / Etios) available on request.
            5. P2-12 duration overshoot note — two-tier system replaces the old
               single >5n warning: ℹ soft when |delta| > 3n, ⚠ hard when
               overshoot > 30% of requested duration (e.g. 14n primary for 10n
               request = 40% → hard warning).
v20260522a  Engine review session — no code changes. Third wildlife trigger
            site identified as unpatched (line ~13200).
v20260521c  P2-04 Wildlife false-trigger fix — Batch 3 review action item.
            Bare 'tiger' and 'jungle' in the wildlife trigger keyword list
            were false-positiving on tourist landmarks and metaphors:
              - "Tiger's Nest" (Paro Taktsang, Bhutan monastery) tripped
                'tiger' in B3#003 (MEX/fytviajes Spanish religious-circuit
                email) — engine showed all 5 wildlife zones for a yoga +
                ashram pilgrimage request to India + Bhutan.
              - Same pattern reproduced on B2#009 (FRA/Cercle Des Voyages).
              - Third reproduction was the original Batch 3 finding.
            Fix at both trigger sites (display @ ~12211 + route-validation
            @ ~13267):
              - 'tiger' → 'tiger reserve', 'tiger safari', 'tiger sighting',
                'tigress' (specific wildlife phrases)
              - 'jungle' → only matches when paired with wildlife noun
                (safari/lodge/camp/trek/drive/resort/stay/experience)
                via word-bounded regex.
              - Strong unambiguous keywords retained (wildlife, safari,
                leopard, rhino, game drive, national park).
            Validated against four scenarios:
              ✓ B3#003 Tiger's Nest + ashram → no wildlife fire
              ✓ "tiger safari at Ranthambore" → wildlife fires
              ✓ "Corbett Tiger Reserve" → wildlife fires
              ✓ "Sunrise at Tiger Hill Darjeeling" → no wildlife fire
v20260521b  Parser hotfix (2 items) — v20260521a follow-up after live testing
            against B2#006 (Mrs & Mr) and B3#001 (12 días):
                1. P1-02 Mrs & Mr ordering — _mr_mrs_pattern was positioned
                   as a fallback after the numeric pax_patterns loop. Noisy
                   phrases like "1 guest" / "1 traveller" / "1 person" in
                   the email body locked num_pax=1 first, causing the Mrs
                   & Mr block to be skipped entirely (output: 1 pax / 1
                   single instead of 2 pax / 1 double). Fix: lift the
                   _mr_mrs_pattern check ABOVE the numeric loop. Explicit
                   couple salutations beat noisy single-pax matches.
                2. P1-05/P1-11 Spanish días accent — regex \\bdias?\\b did
                   not match "días" because the accented í is not in \\w
                   so the word boundary placement failed. "aprox 15 días"
                   survived the substitution as-is and duration parsing
                   returned None. Fix: \\bdias?\\b → \\bd[íi]as?\\b. Tested
                   against both "12 días" (B3#001) and "aprox 15 días"
                   (B3#003).
            ENGINE_VERSION constant ALSO bumped — was stale at "20260519b"
            in v20260521a (docstring header was updated, constant was not).
            Now both read "20260521b".
v20260521a  Parser fixes (6 items) — Batch 1/2/3 review findings:
                1. Word-form duration normalisation (EN/FR/DE/ES/PT):
                   "two weeks"→14n, "zwei Wochen"→14n, "12 días"→12n,
                   "dos semanas"→14n, "quinze jours"→14n etc. Runs before
                   numeric pattern matching. Also adds DE/ES approximation
                   word stripping (ungefähr, aproximadamente, aprox).
                2. Couple keyword inference sentence-boundary agnostic:
                   "I have a couple." now fires. Added "ein Pärchen",
                   "tengo una pareja", "i have a couple" to pattern.
                3. Mrs & Mr pax inference: "Mrs & Mr Smith", "Mr & Mrs",
                   "Madame et Monsieur", "Señor y Señora", "Herr und Frau"
                   → 2 pax / 1 double.
                4. Rooms→pax crossvalidation: when pax=None but rooms
                   parsed (e.g. singles:1), infer pax from room count.
                5. French/Spanish room types: "chambre single/double",
                   "chambre twin", "habitación doble/individual" now parsed.
                   "1 chambre single" without digit → 1 single room, 1 pax.
                6. [None] market label fix: source=None in hotel/activity/
                   monument display now shows "No hotel data — all markets"
                   instead of "[None]".
v20260519b  Wildlife historical routes fix — SIMILAR HISTORICAL ROUTES now
                              runs a focused query for wildlife multi-region requests
                              using original regions/sub_region context (not the
                              injected 30+ city wildlife set). Fixes empty ranked
                              routes section for emails like "Rajasthan and wildlife". When email names a state/circuit
                              ('Rajasthan', 'Kerala' etc.), intent['sub_region']
                              is set and historical_query uses SUB_REGION_TO_CITIES
                              for overlap instead of the full parent region. Gateway
                              cities (Delhi, Mumbai etc.) excluded from overlap —
                              neutral, not penalised. Golden Triangle earns its rank
                              by Rajasthan city content, not by brand. 15 sub-regions
                              defined. GATEWAY_CITIES set in historical_query.
  v20260514c  Repeated-city merge in _apply_city_night_rules (Step 0).
                              Historical sequences with bookend patterns
                              (e.g. Delhi → Shimla → Delhi → ... → Delhi)
                              had each occurrence clamped independently,
                              leaving an unresolvable surplus that pushed
                              all cities to their minimum floor (Delhi 1n
                              instead of 2n). Fix: sum repeated city nights
                              before clamping, preserving first-occurrence
                              order. Delhi 3× raw 4n each → merged 12n →
                              clamped to 3n → rebalanced to 2n correctly.
  v20260514b  City Night Rules CSV — new city_rules.csv support file (6th
                              data source alongside hotel_portfolio, agency_crm,
                              Preconfigured Tours, activity_details, fnb_details).
                              load_city_rules(path) loads CSV into _CITY_RULES cache
                              {city_lower → {min, max, type, notes}}. Applied in
                              _apply_city_night_rules() which clamps each city's
                              nightly split to [min, max] then rebalances total to
                              match requested duration. Operates per-tier — same
                              clamp logic for agent / market / global options.
                              recommend() gains city_rules_path parameter.
                              Display-time ±4n filter from v20260514a replaced by
                              this data-level fix.
  v20260514a  Duration display filter — cascade alternatives now filtered to
                              ±4n of the requested duration before display.
                              Previously all 10 cascade options were shown
                              regardless of duration fit (e.g. 31n option showing
                              for a 10n request). Primary route flagged if it
                              also falls outside the window. Fixes the "Option 3
                              shows 19n for 10n request" issue.
  v20260513b  BL2-α hotfix — Banner truth.
                              "✦ Using <source> booking pattern" where <source>
                              is derived from the cascade primary option's
                              source_pass (agent / market / region / global).
                              Fixes the case where cascade fires on a UNKNOWN-
                              market / no-agent email and primary route comes
                              from global TCI data — previously mislabeled as
                              "agent's historical pattern" (no agent in scope).
  v20260513a  BL2-α  — Cascade as primary route source. Replaces REGIONAL_ROUTES
                       template lookup at the Type 1 / Type 2 integration site
                       with historical_query.assemble_options() output. New
                       cities_source='cascade'. REGIONAL_ROUTES retained as
                       deepest fallback (cascade-empty / cascade-exception).
                       Closes Push Criterion 2b. Requires historical_query v0.5.6_BL2β.
  v20260512a  B27-INT  — Wire historical_query.py v0.5.4b into recommend().
                         Replaces legacy get_similar_routes() with three-pass
                         scored Options (agent → market → global), MAX-merge,
                         confidence labels, nightly_split. Output exposed at
                         intent['historical_options'] for Word doc / frontend /
                         Outlook surfaces.
  v20260508a  B4 fix   — per-agent email index + two-pass cascade with city
                         disambiguation. Defensive portfolio auto-load.
================================================================================
"""

ENGINE_VERSION = "20260522f"

import sqlite3
import re
from collections import defaultdict

# ── CONFIG ────────────────────────────────────────────────────────────────────
import os as _os
from pathlib import Path as _Path
_here = _Path(__file__).parent
DB_PATH = str(_here / 'master_v2_anon.db') if (_here / 'master_v2_anon.db').exists() else str(_here / 'master_v2.db')
PORTFOLIO_PATH  = '/content/hotel_portfolio.csv'
ACTIVITY_PATH   = '/content/activity_details.csv'
FNB_PATH        = '/content/fnb_details.csv'
AGENCY_PATH     = '/content/agency_crm.xlsx'          # CRM agency portfolio (Agent_CRM Excel)
CITY_RULES_PATH = '/content/city_rules.csv'           # City night rules (min/max per city)

# Module-level portfolio cache
_PORTFOLIO_BY_CITY = {}   # city -> [{name, tier, official, code}]
_PORTFOLIO_NAME_TIER = {} # name.lower() -> {tier, official, code}  (last-write-wins for duplicate names)
_PORTFOLIO_NAME_CITY_TIER = {} # (name.lower(), city.lower()) -> {tier, official, code}

# Module-level city night rules cache
# Loaded from city_rules.csv — {city_lower → {min, max, type, notes}}
# Fallback when city not in CSV: {min: 1, max: 3, type: 'touring', notes: ''}
_CITY_RULES = {}

# Module-level agency cache
_AGENCY_BY_DOMAIN  = {}    # domain.lower() -> {account_code, agency_name, market, city, brand, ...}
                           #   one canonical record per domain (brand-priority winner)
_AGENCY_BY_DOMAIN_LIST = {} # domain.lower() -> [list of all records for that domain]
                           #   used for city-based disambiguation when a domain holds
                           #   multiple Account Numbers (e.g. audleytravel.com → ACC0131
                           #   Witney UK and ACC0132 Boston US)
_AGENCY_BY_EMAIL   = {}    # full_email.lower() -> full record  (per-agent CRM row, distinguishes
                           #   ACC0131 stephanie.buckland@audleytravel.com from
                           #   ACC0132 thalia.marini@audleytravel.com — same domain, different markets)
_AGENCY_CITY_MARKET = {}   # city.lower() -> market code  (from CRM city field, non-India/Nepal)
_AGENCY_BY_NAME    = {}    # agency_name.lower() -> CRM record  (for PDF parser name-based lookup)

# ── Gemini LLM rate limiter ────────────────────────────────────────────────────
# Free tier limits: 15 requests/min · 1,500 requests/day
# Engine hard stops at 12/min and 1,400/day — 100 buffer each
import time as _time_mod
_LLM_DAILY_LIMIT   = 1_400   # hard stop before free tier 1,500/day
_LLM_MINUTE_LIMIT  = 12      # hard stop before free tier 15/min
_LLM_DAILY_COUNT   = 0       # resets when _LLM_DAY changes
_LLM_MINUTE_COUNT  = 0       # resets every minute
_LLM_DAY           = None    # tracks current date
_LLM_MINUTE_START  = 0.0     # tracks start of current minute window
_LLM_CACHE         = {}      # hash(email[:500]) -> parsed result

def _llm_quota_check():
    """Check if a Gemini LLM call is allowed under free tier limits.
    Returns True if call is allowed, False if quota exceeded.
    Updates counters and resets as needed.
    """
    global _LLM_DAILY_COUNT, _LLM_MINUTE_COUNT, _LLM_DAY, _LLM_MINUTE_START
    import datetime as _dt_mod

    today = _dt_mod.date.today().isoformat()

    # Reset daily counter on new day
    if _LLM_DAY != today:
        _LLM_DAY          = today
        _LLM_DAILY_COUNT  = 0

    # Reset minute counter if window has elapsed
    now = _time_mod.time()
    if now - _LLM_MINUTE_START >= 60:
        _LLM_MINUTE_START = now
        _LLM_MINUTE_COUNT = 0

    # Check limits
    if _LLM_DAILY_COUNT >= _LLM_DAILY_LIMIT:
        print(f"  ℹ LLM daily quota reached ({_LLM_DAILY_COUNT}/{_LLM_DAILY_LIMIT}) — using regex parser")
        return False
    if _LLM_MINUTE_COUNT >= _LLM_MINUTE_LIMIT:
        print(f"  ℹ LLM minute quota reached ({_LLM_MINUTE_COUNT}/{_LLM_MINUTE_LIMIT}) — using regex parser")
        return False

    # Increment counters
    _LLM_DAILY_COUNT  += 1
    _LLM_MINUTE_COUNT += 1
    return True

def _llm_quota_status():
    """Return current LLM usage stats for display."""
    return f"{_LLM_DAILY_COUNT}/{_LLM_DAILY_LIMIT} daily · {_LLM_MINUTE_COUNT}/{_LLM_MINUTE_LIMIT} this minute"

# Module-level activity & F&B caches
_ACTIVITY_BY_CITY = {}    # city -> [{name, theme_score, tour_category, is_dkc}]
_FNB_BY_CITY      = {}    # city -> [{name, restaurant_type, part_of_hotel, hotel_code, tci_score, ambience_score, tour_category}]
_PORTFOLIO_CODE_NAME = {}  # business_code -> hotel name (for F&B hotel lookup)

def load_portfolio(portfolio_path=PORTFOLIO_PATH):
    """Load hotel portfolio CSV into module-level cache.

    B1 fix (v20260519a): when a business_code appears multiple times with
    conflicting official_classification, prefer the 'heritage' row over any
    other classification.  We do a two-pass load: first pass collects all rows
    keyed by business_code; second pass picks the best row per code.

    B6 fix (v20260519a): normalise 'Chettinadu' → 'Chettinad' on load so the
    portfolio city_name matches the engine's canonical spelling everywhere
    (CITY_CODE_MAP, regional maps, city keyword index).
    """
    global _PORTFOLIO_BY_CITY, _PORTFOLIO_NAME_TIER, _PORTFOLIO_NAME_CITY_TIER
    if _PORTFOLIO_BY_CITY:
        return  # already loaded
    try:
        import csv as _csv

        # B6: city name normalisations (portfolio CSV → engine canonical)
        _CITY_NAME_NORM = {
            'Chettinadu': 'Chettinad',
        }

        # B1: collect all rows per business_code, then pick best
        _rows_by_code = {}  # business_code -> list of row dicts
        with open(portfolio_path, newline='', encoding='utf-8') as f:
            for row in _csv.DictReader(f):
                city = _CITY_NAME_NORM.get(row['city_name'].strip(), row['city_name'].strip())
                name = row['name'].strip()
                tier = row['tci_classification'].strip()
                off  = row['official_classification'].strip()
                code = row['business_code'].strip()
                _rows_by_code.setdefault(code, []).append(
                    {'city': city, 'name': name, 'tier': tier, 'official': off, 'code': code}
                )

        # B1: for each business_code, prefer the heritage row if one exists
        def _best_row(rows):
            heritage_rows = [r for r in rows if r['official'] == 'heritage']
            return heritage_rows[0] if heritage_rows else rows[0]

        for code, rows in _rows_by_code.items():
            r    = _best_row(rows)
            city = r['city']
            name = r['name']
            tier = r['tier']
            off  = r['official']
            if city not in _PORTFOLIO_BY_CITY:
                _PORTFOLIO_BY_CITY[city] = []
            existing = {h['name'] for h in _PORTFOLIO_BY_CITY[city]}
            if name not in existing:
                _PORTFOLIO_BY_CITY[city].append(
                    {'name': name, 'tier': tier, 'official': off, 'code': code}
                )
            _PORTFOLIO_NAME_TIER[name.lower()] = {'tier': tier, 'official': off, 'code': code, 'supplierName': name}
            _PORTFOLIO_NAME_CITY_TIER[(name.lower(), city.lower())] = {'tier': tier, 'official': off, 'code': code, 'supplierName': name}

        print(f'  ✓ Portfolio loaded: {sum(len(v) for v in _PORTFOLIO_BY_CITY.values())} hotels across {len(_PORTFOLIO_BY_CITY)} cities')
    except FileNotFoundError:
        pass  # portfolio optional — engine works without it


def load_city_rules(city_rules_path=CITY_RULES_PATH):
    """Load city_rules.csv into module-level _CITY_RULES cache.

    CSV columns: city, region, min_nights, max_nights, type, notes
    Cache key: city name lowercased (matches against cascade city_sequence values).

    Called once from recommend(); subsequent calls are no-ops (cache check).
    Silent on FileNotFoundError — engine falls back to defaults per city.

    Fallback when city not in CSV: min=1, max=3, type='touring'.
    Transit cities (min=0, max=0) are respected — clamp will zero-cap them.
    """
    global _CITY_RULES
    if _CITY_RULES:
        return
    try:
        import csv as _csv
        loaded = 0
        with open(city_rules_path, newline='', encoding='utf-8') as f:
            for row in _csv.DictReader(f):
                city = row['city'].strip()
                if not city:
                    continue
                try:
                    min_n = int(row['min_nights'])
                    max_n = int(row['max_nights'])
                except (ValueError, KeyError):
                    min_n, max_n = 1, 3
                _CITY_RULES[city.lower()] = {
                    'min':   min_n,
                    'max':   max_n,
                    'type':  row.get('type', 'touring').strip(),
                    'notes': row.get('notes', '').strip(),
                    'city':  city,   # preserve original casing for display
                }
                loaded += 1
        print(f'  ✓ City rules loaded: {loaded} cities')
    except FileNotFoundError:
        print('  ℹ  city_rules.csv not found — using built-in defaults')


def _get_city_rule(city):
    """Return city night rule dict for a city name.

    Looks up _CITY_RULES by lowercased city name.
    Falls back to {min:1, max:3, type:'touring'} when not found.
    """
    return _CITY_RULES.get(city.lower(), {'min': 1, 'max': 3, 'type': 'touring', 'notes': ''})


def _apply_city_night_rules(nightly_split, requested_duration=None):
    """Clamp each city's nights to [rule_min, rule_max] then rebalance total
    to match requested_duration.

    Applied per cascade option (agent / market / global tier equally).

    Args:
        nightly_split:     list of (city, nights) tuples from _nightly_split
        requested_duration: int | None — target total nights from intent

    Returns:
        list of (city, nights) tuples with clamped + rebalanced nights

    Logic:
        0. Merge repeated city entries (e.g. Delhi bookend appearing 3× in
           a historical sequence) — sum their nights, preserve first-occurrence
           order. Without this, each repeated entry gets clamped independently
           and the surplus rebalancer can't converge to the requested duration.
        1. Clamp each city to its rule [min, max]
        2. If requested_duration known:
             a. surplus  (total > requested) → bleed from cities above their min,
                         prioritising stay-put cities last (they absorb cuts least)
             b. deficit  (total < requested) → add to cities below their max,
                         prioritising stay-put cities first (they absorb extras best)
        3. If requested_duration unknown: return clamped values as-is

    Transit cities (rule max=0) are always kept at 0 — never absorb surplus/deficit.
    """
    if not nightly_split:
        return nightly_split

    # Step 0 — merge repeated city entries (preserve first-occurrence order)
    # Historical routes sometimes have bookend patterns: Delhi → ... → Delhi → ... → Delhi
    # Each occurrence must be summed before clamping so the clamp sees the full
    # city allocation, and the rebalancer can converge to requested_duration.
    seen_order = []
    merged_nights = {}
    for city, nights in nightly_split:
        key = city.lower()
        if key not in merged_nights:
            seen_order.append(city)
            merged_nights[key] = 0
        merged_nights[key] += nights
    nightly_split = [(city, merged_nights[city.lower()]) for city in seen_order]

    # Step 1 — clamp to [rule_min, rule_max]
    clamped = []
    for city, nights in nightly_split:
        rule = _get_city_rule(city)
        clamped_n = max(rule['min'], min(nights, rule['max']))
        clamped.append([city, clamped_n, rule])   # keep rule for rebalance

    if requested_duration is None:
        return [(c, n) for c, n, _ in clamped]

    total = sum(n for _, n, _ in clamped)
    diff  = requested_duration - total   # positive = deficit, negative = surplus

    if diff == 0:
        return [(c, n) for c, n, _ in clamped]

    # Step 2a — absorb surplus (total > requested): reduce nights
    if diff < 0:
        surplus = -diff
        # Reduce from touring cities first (highest nights → lowest), stay-put last
        priority = sorted(
            range(len(clamped)),
            key=lambda i: (
                0 if clamped[i][2]['type'] != 'stay-put' else 1,  # touring first
                -(clamped[i][1] - clamped[i][2]['min'])            # most headroom first
            )
        )
        for i in priority:
            if surplus <= 0:
                break
            city, nights, rule = clamped[i]
            if rule['max'] == 0:   # transit — never touch
                continue
            headroom = nights - rule['min']
            cut = min(headroom, surplus)
            clamped[i][1] -= cut
            surplus -= cut

    # Step 2b — absorb deficit (total < requested): add nights
    if diff > 0:
        deficit = diff
        # Add to stay-put cities first, then touring, never transit
        priority = sorted(
            range(len(clamped)),
            key=lambda i: (
                0 if clamped[i][2]['type'] == 'stay-put' else
                1 if clamped[i][2]['type'] == 'touring'  else 2,
                -(clamped[i][2]['max'] - clamped[i][1])   # most headroom first
            )
        )
        for i in priority:
            if deficit <= 0:
                break
            city, nights, rule = clamped[i]
            if rule['max'] == 0:   # transit — never touch
                continue
            headroom = rule['max'] - nights
            add = min(headroom, deficit)
            clamped[i][1] += add
            deficit -= add

    return [(c, n) for c, n, _ in clamped]


# ── TIER KEYWORDS ─────────────────────────────────────────────────────────────
TIER_KEYWORDS = {
    'lux':          ['luxury', 'luxurious', '5 star', '5*', '5star', '5-star',
                     'five star', 'ultra', 'premium', 'higher end', 'high end',
                     'top end', 'top-end', 'highest end', 'no budget', 'budget flexible',
                     'luxury stay', 'luxury hotel', 'luxury property', 'luxury accommodation',
                     # Common typos
                     'luxry', 'luxery', 'luxurey', 'luxurious stay', 'luxurious hotel',
                 # French
                 'luxe', '5 etoiles', '5 étoiles', 'cinq etoiles', 'cinq étoiles',
                 'haut de gamme', 'charme', 'prestige', 'grand luxe', 'palace hotel',
                 # German
                 'luxus', '5 sterne', 'fuenf sterne'],
    'first-class':  ['first class', 'first-class', '4 star', '4*', '4star', '4-star',
                     'four star', 'superior', 'deluxe'],
    'moderate':     ['moderate', '3 star', '3*', '3star', '3-star', 'three star', 'good hotel'],
    'budget':       ['budget', 'économique', 'economique', 'bon marché', 'pas cher',
                     'low cost', 'low-cost', 'cheap', 'affordable', 'economy',
                     'basic hotel', '2 star', '2*', '2star', 'two star',
                     # German
                     'günstig', 'preiswert', 'einfach',
                     # Italian
                     'economico', 'conveniente',
                     # Spanish (v20260522c: presupuesto bajo / presupuesto reducido /
                     # presupuesto limitado / bajo presupuesto added — closes the
                     # gap in B3 #003 'presupuesto bajo' detection)
                     'barato', 'económico', 'presupuesto bajo', 'bajo presupuesto',
                     'presupuesto reducido', 'presupuesto limitado',
                     # Portuguese (v20260522c: BRA market parity)
                     'orçamento baixo', 'baixo orçamento', 'barato'],
}

HERITAGE_KEYWORDS = [
    'heritage', 'palace', 'haveli', 'fort', 'historic', 'historical',
    'traditional', 'authentic', 'character', 'boutique', 'royal'
]

HERITAGE_NAME_SIGNALS = [
    'palace', 'haveli', 'fort', 'mahal', 'garh', 'bagh', 'niwas',
    'vilas', 'bhawan', 'kothi', 'villa', 'heritage', 'historic'
]

# Hotels suppressed from multi-night touring recommendations
# These are airport hotels, transit properties, or out-of-city locations
# not suitable for clients staying 2+ nights in a city for touring
SUPPRESS_FOR_TOURING = {
    # Delhi — airport / out-of-city
    'jw marriott aerocity',
    'ibis delhi aerocity',
    'pullman aerocity',
    'andaz delhi aerocity',
    'lemon tree premier aerocity',
    'holiday inn aerocity',
    'aloft aerocity',
    'novotel aerocity',
    'hyatt regency delhi aerocity',
    # Gurgaon properties often listed under Delhi
    'leela ambience',
    'trident gurgaon',
    'oberoi gurgaon',
    'hyatt gurgaon',
    'westin gurgaon',
    'the westin gurgaon',
    # Mumbai — airport
    'hyatt regency mumbai',
    'renaissance mumbai',
    'sahara star',
    'novotel mumbai',
    'grand hyatt santacruz',
    # Bangalore — airport
    'sheraton grand bangalore',
    # Chennai — airport
    'hilton chennai',
}

LANGUAGE_GUIDE_PATTERNS = [
    r'italian[\s-]speaking',
    r'french[\s-]speaking',
    r'german[\s-]speaking',
    r'japanese[\s-]speaking',
    r'spanish[\s-]speaking',
    r'francophone',
    r'guide francophone',
    r'deutschsprachig',
    r'escort guide',
    r'speaking guide',
    r'language.*guide',
    r'guide.*language',
    r'guide.*francoph',
    # Catch-all: "French guide", "guide français", "French or English guide"
    r'french.*guide',
    r'guide.*french',
    r'guide.*français',
    r'français.*guide',
    # Italian / German loose patterns
    r'italian.*guide',
    r'guide.*italian',
    r'german.*guide',
    r'guide.*german',
    r'guida italiana',
    r'guida madrelingua',
    r'reiseleiter',
    # Accompanying / tour escort guide (non-language-specific)
    r'accompanying guide',
    r'accompany.*guide',
    r'guide.*accompan',
    r'tour.*escort',
    r'escort.*tour',
    r'tour manager',
    r'tour leader',
    r'full.?time.*guide',
    r'guide.*full.?time',
]

# ── CSV CITY ALIAS MAP ─────────────────────────────────────────────────────────
# Maps engine canonical city names → CSV city names where they differ
_CSV_CITY_ALIASES = {
    'Kochi':              'Kochi (Cochin)',
    'Kochi (Cochin)':     'Kochi (Cochin)',
    'Mysuru':             'Mysuru (Mysore)',
    'Puducherry':         'Puducherry (Pondicherry)',
}

def _csv_city(city):
    """Return the CSV file's city name for a given engine canonical city."""
    return _CSV_CITY_ALIASES.get(city, city)


_ACTIVITY_SUPPRESS_KEYWORDS = {
    'conveyance charges', 'transferring to old city', 'transfer to old city',
    'tuk tuk conveyance', 'tuk-tuk conveyance', 'conveyance for transfer',
    'charges extra', 'supplement over', 'per vehicle charges',
    'outsourced', 'rate extra', 'tipping', 'porterage', 'luggage transfer',
    'life jacket', 'chair for', 'guide charges', 'guide fee', 'hd guide charges',
    'entrance fee', 'entry charges', 'toll charges', 'parking charges', 'fuel surcharge',
    'mineral water', 'drinking water', 'vip meet', 'meet & greet',
    'transfer from/for airport', 'transfer from airport', 'airport transfer',
    'naturalist charges', 'naturalist fee', 'jeep charges', 'gypsy charges',
    'tuk-tuk for visit', 'tuk tuk for visit',
    # Chandigarh toy train / Kalka railway logistics
    'cushion handling', 'cushion for toy train', 'water jar',
    'catering charges at kalka', 'assistance at kalka', 'kalka railway station',
    # Ambala / food delivery logistics
    'assistance at ambala', 'ambala railway station',
    'delivery of food', 'delivery of meal',
    'packed food', 'picnic hamper', 'meal within city',
}

def _activity_suppressed(name):
    """Return True if activity name looks like a logistics/billing item."""
    nl = name.lower()
    if nl in SUPPRESS_ACTIVITIES:
        return True
    return any(kw in nl for kw in _ACTIVITY_SUPPRESS_KEYWORDS)


def load_agency_portfolio(agency_path=AGENCY_PATH):
    """Load CRM agency portfolio (Agent_CRM Excel) into _AGENCY_BY_DOMAIN cache.

    Expected columns (CRM export format):
        Account Number  — CRM account code  e.g. ACC0221
        Account Name    — agency display name
        City            — agency city
        Country         — full country name  e.g. France
        Brand           — SITA / DFT / TCI / Nepal / TGV / VILASITA / ARANYA
        Email           — primary contact email

    Duplicate domain handling: when multiple accounts share a domain,
    the record with the best brand priority is used (SITA > DFT > TCI > others).
    Ambiguous matches are flagged in engine output.

    Falls back gracefully if file not found.
    """
    global _AGENCY_BY_DOMAIN, _AGENCY_BY_DOMAIN_LIST, _AGENCY_BY_EMAIL, _AGENCY_CITY_MARKET
    if _AGENCY_BY_DOMAIN:
        return

    # ── Complete country → market code map (all 70 CRM countries) ────────────
    COUNTRY_TO_MARKET = {
        # Core booking markets
        'france': 'FRA', 'germany': 'DEU', 'united kingdom': 'GBR',
        'italy': 'ITA', 'spain': 'ESP', 'united states of america': 'USA',
        'japan': 'JPN', 'australia': 'AUS', 'switzerland': 'CHE',
        'netherlands': 'NLD', 'belgium': 'BEL', 'poland': 'POL',
        'romania': 'ROU', 'russia': 'RUS', 'canada': 'CAN', 'india': 'IND',
        # Extended markets
        'argentina': 'ARG', 'austria': 'AUT', 'azerbaijan': 'AZE',
        'bahrain': 'BHR', 'brazil': 'BRA', 'bulgaria': 'BGR',
        'chile': 'CHL', 'colombia': 'COL', 'cyprus': 'CYP',
        'czech republic': 'CZE', 'denmark': 'DNK', 'ecuador': 'ECU',
        'finland': 'FIN', 'georgia': 'GEO', 'greece': 'GRC',
        'hungary': 'HUN', 'indonesia': 'IDN', 'ireland': 'IRL',
        'israel': 'ISR', 'kazakhstan': 'KAZ', 'kenya': 'KEN',
        'kuwait': 'KWT', 'kyrgyzstan': 'KGZ', 'lithuania': 'LTU',
        'luxembourg': 'LUX', 'malaysia': 'MYS', 'mauritius': 'MUS',
        'mexico': 'MEX', 'moldova': 'MDA', 'monaco': 'MCO',
        'nepal': 'NPL', 'new zealand': 'NZL', 'norway': 'NOR',
        'oman': 'OMN', 'peru': 'PER', 'philippines': 'PHL',
        'portugal': 'PRT', 'qatar': 'QAT', 'reunion': 'REU',
        'saudi arabia': 'SAU', 'serbia': 'SRB', 'singapore': 'SGP',
        'slovakia': 'SVK', 'slovenia': 'SVN', 'south africa': 'ZAF',
        'sweden': 'SWE', 'tanzania': 'TZA', 'thailand': 'THA',
        'tunisia': 'TUN', 'turkey': 'TUR', 'ukraine': 'UKR',
        'united arab emirates': 'ARE', 'uruguay': 'URY', 'vietnam': 'VNM',
        # Aliases / alternate spellings
        'united states': 'USA', 'usa': 'USA', 'uk': 'GBR',
        'great britain': 'GBR', 'england': 'GBR', 'russian federation': 'RUS',
        'holland': 'NLD', 'deutschland': 'DEU', 'italia': 'ITA',
        'españa': 'ESP', 'uae': 'ARE',
    }

    SKIP_DOMAINS = {
        'gmail.com', 'yahoo.com', 'yahoo.co.in', 'yahoo.co.uk', 'yahoo.fr',
        'yahoo.it', 'yahoo.de', 'yahoo.es', 'yahoo.com.au', 'yahoo.ca',
        'yahoo.com.br', 'yahoo.com.ar', 'yahoo.co.jp', 'yahoo.com.mx',
        'hotmail.com', 'hotmail.co.uk', 'hotmail.fr', 'hotmail.it',
        'outlook.com', 'live.com', 'msn.com', 'icloud.com', 'me.com',
        'mail.ru', 'yandex.ru', 'aol.com',
        # SITA internal
        'sita.in', 'distantfrontiers.in', 'tci.co.in', 'govacation-india.com',
        'go-vacation.com', 'vilasita.com', 'tgvindia.com',
    }

    # Brand priority for duplicate domain resolution — lower = preferred
    # Nepal brand excluded entirely — these are Nepal-product specialists,
    # not relevant for India tour market detection
    BRAND_PRIORITY = {'SITA': 0, 'DFT': 1, 'TCI': 2,
                      'TGV': 3, 'VILASITA': 4, 'ARANYA': 5}
    EXCLUDED_BRANDS = {'Nepal'}

    try:
        import pandas as _pd
        df = _pd.read_excel(agency_path, sheet_name=0)

        # Normalise column names
        df.columns = [str(c).strip() for c in df.columns]

        # Flexible column lookup
        def _find_col(candidates):
            for c in candidates:
                if c in df.columns:
                    return c
            return None

        col_acct  = _find_col(['Account Number', 'account_number', 'AccountNumber'])
        col_name  = _find_col(['Account Name', 'account_name', 'Agency', 'Name'])
        col_city  = _find_col(['City', 'city', 'Agency City'])
        col_ctry  = _find_col(['Country', 'country', 'Country / Market', 'Market'])
        col_brand = _find_col(['Brand', 'brand'])
        col_email = _find_col(['Email', 'email', 'Email Address', 'Contact Email'])

        if col_email is None:
            print('  ⚠ Agency portfolio: no email column found — skipping')
            return

        # Collect all records into a staging dict keyed by domain
        # Keeps the best brand-priority record for each domain
        staging = {}   # domain → {account_code, agency_name, market, city, brand, country_raw, all_accounts}
        skipped_generic = skipped_no_market = 0

        for _, row in df.iterrows():
            email_val = str(row[col_email]).strip().lower() if _pd.notna(row.get(col_email)) else ''
            if not email_val or '@' not in email_val:
                continue
            domain = email_val.split('@')[-1].strip()
            if not domain or '.' not in domain or domain in SKIP_DOMAINS:
                skipped_generic += 1
                continue

            brand = str(row[col_brand]).strip() if col_brand and _pd.notna(row.get(col_brand)) else ''
            if brand in EXCLUDED_BRANDS:
                continue   # skip Nepal brand entirely

            raw_country = str(row[col_ctry]).strip() if col_ctry and _pd.notna(row.get(col_ctry)) else ''
            market_code = COUNTRY_TO_MARKET.get(raw_country.lower(), '')
            if not market_code:
                skipped_no_market += 1
                continue

            acct_code = str(row[col_acct]).strip() if col_acct and _pd.notna(row.get(col_acct)) else ''
            name      = str(row[col_name]).strip() if col_name and _pd.notna(row.get(col_name)) else ''
            city      = str(row[col_city]).strip() if col_city and _pd.notna(row.get(col_city)) else ''
            brand     = str(row[col_brand]).strip() if col_brand and _pd.notna(row.get(col_brand)) else ''
            brand_rank = BRAND_PRIORITY.get(brand, 99)

            if domain not in staging:
                staging[domain] = {
                    'account_code': acct_code,
                    'agency_name':  name,
                    'market':       market_code,
                    'city':         city,
                    'brand':        brand,
                    'brand_rank':   brand_rank,
                    'country_raw':  raw_country,
                    'email':        email_val,
                    'all_accounts': [(acct_code, name, brand)],
                }
            else:
                # Keep better brand-priority record
                staging[domain]['all_accounts'].append((acct_code, name, brand))
                if brand_rank < staging[domain]['brand_rank']:
                    staging[domain].update({
                        'account_code': acct_code,
                        'agency_name':  name,
                        'market':       market_code,
                        'city':         city,
                        'brand':        brand,
                        'brand_rank':   brand_rank,
                        'country_raw':  raw_country,
                        'email':        email_val,
                    })

            # ── B4 fix (v20260508): per-row email index ──────────────────────
            # Distinguishes accounts that share a domain. e.g. Audley Travel
            # has ACC0131 (UK, stephanie.buckland@audleytravel.com) and ACC0132
            # (US, thalia.marini@audleytravel.com) — same domain, different
            # markets. _AGENCY_BY_DOMAIN keeps a single canonical record per
            # domain (brand-priority winner); _AGENCY_BY_EMAIL keeps every row
            # for exact-email lookup. parse_email tries email first, then
            # falls back to domain.
            full_record = {
                'account_code': acct_code,
                'agency_name':  name,
                'market':       market_code,
                'city':         city,
                'brand':        brand,
                'country_raw':  raw_country,
                'email':        email_val,
                'ambiguous':    False,
                'all_accounts': [],
            }
            _AGENCY_BY_EMAIL[email_val] = full_record
            # _AGENCY_BY_DOMAIN_LIST: every record per domain — used by
            # parse_email to disambiguate multi-account domains via signature
            # city when exact email match misses.
            _AGENCY_BY_DOMAIN_LIST.setdefault(domain, []).append(full_record)

        # Promote to final cache, flagging ambiguous domains
        ambiguous_count = 0
        for domain, rec in staging.items():
            is_ambiguous = len(rec['all_accounts']) > 1
            if is_ambiguous:
                ambiguous_count += 1
            _AGENCY_BY_DOMAIN[domain] = {
                'account_code': rec['account_code'],
                'agency_name':  rec['agency_name'],
                'market':       rec['market'],
                'city':         rec['city'],
                'brand':        rec['brand'],
                'country_raw':  rec['country_raw'],
                'email':        rec['email'],
                'ambiguous':    is_ambiguous,
                'all_accounts': rec['all_accounts'] if is_ambiguous else [],
            }

        markets_covered = len(set(v['market'] for v in _AGENCY_BY_DOMAIN.values()))
        print(f'  ✓ CRM agency portfolio loaded: {len(_AGENCY_BY_DOMAIN)} agencies · '
              f'{len(_AGENCY_BY_EMAIL)} agents · '
              f'{markets_covered} markets · {ambiguous_count} ambiguous domains '
              f'({skipped_generic} generic skipped · {skipped_no_market} unknown country skipped)')

        # ── Hardcoded domain overrides (agencies not in CRM with salessupport@ etc) ──
        _HARDCODED_DOMAINS = {
            'tourlane.com': {'account_code': 'ACC1816', 'agency_name': 'Sensation Travel GMBH',
                             'market': 'DEU', 'city': 'Berlin', 'brand': 'SITA',
                             'country_raw': 'Germany', 'ambiguous': False, 'all_accounts': []},
        }
        for _hd, _hv in _HARDCODED_DOMAINS.items():
            if _hd not in _AGENCY_BY_DOMAIN:
                _AGENCY_BY_DOMAIN[_hd] = _hv

        # ── Build name → record index for PDF parser lookup ────────────────
        # Normalises agency names: lowercase, strip punctuation/whitespace
        import re as _re_name
        def _norm_name(s):
            return _re_name.sub(r'[\s\-_&,\.\']+', ' ', s.lower()).strip()
        for rec in _AGENCY_BY_DOMAIN.values():
            n = rec.get('agency_name', '')
            if n:
                _AGENCY_BY_NAME[_norm_name(n)] = rec

        # ── Build city → market lookup from CRM city field ─────────────────
        # Scans email address/signature for known international city names
        # as a secondary market confirmation layer
        SKIP_CITY_COUNTRIES = {'India', 'Nepal'}
        city_tally = {}   # city.lower() → {market: count}
        for _, row in df.iterrows():
            brand = str(row[col_brand]).strip() if col_brand and _pd.notna(row.get(col_brand)) else ''
            if brand in EXCLUDED_BRANDS:
                continue
            country = str(row[col_ctry]).strip() if col_ctry and _pd.notna(row.get(col_ctry)) else ''
            if country in SKIP_CITY_COUNTRIES:
                continue
            market = COUNTRY_TO_MARKET.get(country.lower(), '')
            if not market:
                continue
            city_raw = str(row[col_city]).strip() if col_city and _pd.notna(row.get(col_city)) else ''
            if not city_raw or city_raw.lower() in ('none', 'nan', ''):
                continue
            city_key = city_raw.strip().lower()
            if city_key not in city_tally:
                city_tally[city_key] = {}
            city_tally[city_key][market] = city_tally[city_key].get(market, 0) + 1

        for city_key, market_counts in city_tally.items():
            total = sum(market_counts.values())
            top_market = max(market_counts, key=market_counts.get)
            # Only store if dominant market is ≥80% — avoids ambiguous cities
            if market_counts[top_market] / total >= 0.8:
                _AGENCY_CITY_MARKET[city_key] = top_market

        print(f'  ✓ City→market lookup built: {len(_AGENCY_CITY_MARKET)} cities')

    except FileNotFoundError:
        pass
    except Exception as e:
        print(f'  ⚠ Agency portfolio load error: {e}')


def load_activity_portfolio(activity_path=ACTIVITY_PATH):
    """Load activity CSV into _ACTIVITY_BY_CITY cache.
    Deduplicates: for any name that exists both as plain and [A DKC Experience],
    keeps only the DKC version (highest score among duplicates).
    Falls back gracefully if file not found.
    """
    global _ACTIVITY_BY_CITY
    if _ACTIVITY_BY_CITY:
        return
    try:
        import csv as _csv
        raw = {}  # (city, canonical_name) -> best record
        with open(activity_path, newline='', encoding='utf-8') as f:
            for row in _csv.DictReader(f):
                city     = row['start_city_name'].strip()
                name     = row['name'].strip()
                score    = float(row['theme_score']) if row['theme_score'] else None
                category = row['tour_category'].strip()
                # Skip rows with no score or no category (incomplete records)
                if score is None or not category:
                    continue
                is_dkc   = '[A DKC Experience]' in name
                # Skip logistics/billing items
                if _activity_suppressed(name):
                    continue
                # Canonical name = strip DKC suffix for dedup key
                canon    = name.replace(' [A DKC Experience]', '').strip().lower()
                key      = (city, canon)
                existing = raw.get(key)
                # Prefer DKC version; among same type prefer higher score
                if existing is None:
                    raw[key] = {'name': name, 'theme_score': score,
                                'tour_category': category, 'is_dkc': is_dkc}
                else:
                    # Upgrade to DKC if this row is DKC and existing is not
                    if is_dkc and not existing['is_dkc']:
                        raw[key] = {'name': name, 'theme_score': score,
                                    'tour_category': category, 'is_dkc': is_dkc}
                    # Same DKC status → keep higher score
                    elif is_dkc == existing['is_dkc'] and score > existing['theme_score']:
                        existing['theme_score'] = score
        # Build city dict
        for (city, _), rec in raw.items():
            if city not in _ACTIVITY_BY_CITY:
                _ACTIVITY_BY_CITY[city] = []
            _ACTIVITY_BY_CITY[city].append(rec)
        # Sort each city by theme_score desc
        for city in _ACTIVITY_BY_CITY:
            _ACTIVITY_BY_CITY[city].sort(key=lambda x: x['theme_score'], reverse=True)
        total = sum(len(v) for v in _ACTIVITY_BY_CITY.values())
        print(f'  ✓ Activity portfolio loaded: {total} activities across {len(_ACTIVITY_BY_CITY)} cities')
    except FileNotFoundError:
        pass
    except Exception as e:
        print(f"  ⚠ Activity portfolio load error: {e}")


def load_fnb_portfolio(fnb_path=FNB_PATH):
    """Load F&B CSV into _FNB_BY_CITY and _PORTFOLIO_CODE_NAME caches.
    Falls back gracefully if file not found.
    """
    global _FNB_BY_CITY, _PORTFOLIO_CODE_NAME
    if _FNB_BY_CITY:
        return
    try:
        import csv as _csv
        # Also build hotel code→name map from portfolio if available
        if _PORTFOLIO_BY_CITY and not _PORTFOLIO_CODE_NAME:
            for city_hotels in _PORTFOLIO_BY_CITY.values():
                for h in city_hotels:
                    if h.get('code'):
                        _PORTFOLIO_CODE_NAME[h['code']] = h['name']
        with open(fnb_path, newline='', encoding='utf-8') as f:
            for row in _csv.DictReader(f):
                city         = row['city_name'].strip()
                name         = row['name'].strip()
                if not name or name.lower() == 'restaurant':
                    continue  # skip unnamed / generic entries
                rtype        = row['restaurant_type'].strip()
                hotel_linked = row['part_of_hotel'].strip().lower() == 'true'
                hotel_code   = row['hotel'].strip()
                tci          = float(row['tci_score'] or 0)
                ambience     = float(row['ambience_score'] or 0)
                category     = row['tour_category'].strip()
                if city not in _FNB_BY_CITY:
                    _FNB_BY_CITY[city] = []
                _FNB_BY_CITY[city].append({
                    'name':         name,
                    'restaurant_type': rtype,
                    'hotel_linked': hotel_linked,
                    'hotel_code':   hotel_code,
                    'tci_score':    tci,
                    'ambience_score': ambience,
                    'tour_category': category,
                })
        # Sort each city by tci_score desc
        for city in _FNB_BY_CITY:
            _FNB_BY_CITY[city].sort(key=lambda x: x['tci_score'], reverse=True)
        total = sum(len(v) for v in _FNB_BY_CITY.values())
        print(f'  ✓ F&B portfolio loaded: {total} restaurants across {len(_FNB_BY_CITY)} cities')
    except FileNotFoundError:
        pass
    except Exception as e:
        print(f"  ⚠ F&B portfolio load error: {e}")


def get_activities_csv(city, category_filter=None, top_n=5):
    """Return top activities from CSV for a city.
    Falls back through CSV city alias if needed.
    Returns (list of dicts, source_label) or ([], None).
    """
    csv_city = _csv_city(city)
    rows = _ACTIVITY_BY_CITY.get(csv_city) or _ACTIVITY_BY_CITY.get(city) or []
    if category_filter:
        filtered = [r for r in rows if r['tour_category'] == category_filter]
        if filtered:
            rows = filtered
    return rows[:top_n], ('CSV portfolio' if rows else None)


def get_fnb(city, top_n=4):
    """Return top F&B entries for a city sorted by tci_score.
    Shows both hotel-linked and standalone, with hotel name when available.
    Returns list of dicts.
    """
    csv_city = _csv_city(city)
    rows = _FNB_BY_CITY.get(csv_city) or _FNB_BY_CITY.get(city) or []
    return rows[:top_n]


# ── VEHICLE RULES ─────────────────────────────────────────────────────────────
def recommend_vehicle(num_pax):
    if num_pax == 1:
        return 'Toyota Innova Crysta', '1 pax — Innova oversized; sedan (Swift Dzire / Etios) available on request'
    elif num_pax <= 3:
        return 'Toyota Innova Crysta', '1-3 pax standard vehicle'
    elif num_pax <= 6:
        return 'Force Urbania 7', '4-6 pax standard vehicle'
    elif num_pax <= 9:
        return 'Force Tempo Traveller 9 / A/C Coach 18', '7-9 pax — coach recommended'
    elif num_pax <= 15:
        return 'Regular A/C Coach 18 or 45', '10-15 pax — size by exact count'
    else:
        return 'Regular A/C Coach 45', '16+ pax — full coach'

# ── EMAIL PARSER ──────────────────────────────────────────────────────────────
def detect_and_translate(email_text):
    """Translation disabled — returns original text unchanged."""
    return email_text, False

def parse_email(email_text, agency_path=None):
    """
    Extracts structured intent from a raw agent email.
    Returns a dict with all parsed fields.

    Args:
        email_text   : raw email text
        agency_path  : optional path to agency_crm.xlsx — auto-loads the
                       agency portfolio if not yet loaded in this process.
                       Without it, the function tries the AGENCY_PATH
                       default and prints a warning if the dict is still
                       empty after that.
    """
    # ── B4 fix (v20260508): defensive portfolio auto-load ───────────────
    # parse_email was previously silent when _AGENCY_BY_DOMAIN was empty,
    # falling through to keyword scoring with no warning. Any caller that
    # invoked parse_email() without first calling load_agency_portfolio()
    # — test harnesses, B27 prototype, direct user code — would have all
    # agency lookups silently miss. This block makes the failure visible.
    if not _AGENCY_BY_DOMAIN:
        try:
            if agency_path:
                load_agency_portfolio(agency_path)
            else:
                load_agency_portfolio()  # uses AGENCY_PATH default
        except Exception:
            pass
    if not _AGENCY_BY_DOMAIN:
        print("  ⚠ Agency portfolio empty — agency lookup will be skipped, "
              "market detection falling back to keyword scoring")

    text = email_text.lower()
    result = {}
    result['_raw_text'] = email_text.lower()

    # --- Source market detection ---
    # Step 1: Agency portfolio lookup (authoritative — wins outright if matched)
    #
    # B4 fix (v20260508): build a list of email candidates from two signals
    # and run a TWO-PASS cascade with city disambiguation in Pass 2.
    #
    # Why two passes:
    #   Some agencies have multiple Account Numbers under the same domain —
    #   e.g. Audley Travel has ACC0131 (UK, stephanie.buckland@audleytravel.com)
    #   and ACC0132 (US, thalia.marini@audleytravel.com), same domain, different
    #   markets. Domain alone can't tell them apart. Three signals help:
    #     (1) exact email of sender or signature → per-agent precision
    #     (2) city in signature/body → disambiguates multi-account domains
    #     (3) domain alone → falls back to canonical record (ambiguous flagged)
    #
    # Pass 1: exact email match against _AGENCY_BY_EMAIL (per-agent index).
    # Pass 2: domain match. If domain has multiple records, scan body for
    #         any of those agencies' cities; first city hit wins. Otherwise
    #         use the canonical brand-priority record for that domain.
    #
    # Signals for candidate emails (in priority order):
    #   (a) From: header line — anchored. Old regex scanned whole text and
    #       could pick up a wrong <email> from a quoted reply or CC field.
    #   (b) Signature / body emails — fallback when From: is a generic shared
    #       inbox (Calestani: From=asia@asia365.ch generic, agent's specific
    #       email lives in signature) or when From: is a SITA forwarder.
    #
    # SITA internal domains are always skipped from candidates.
    _SITA_INTERNAL = {'sita.in', 'distantfrontiers.in', 'tci.co.in',
                      'govacation-india.com', 'go-vacation.com',
                      'vilasita.com', 'tgvindia.com'}

    # Each candidate is (source_label, full_email, domain)
    _candidates = []

    def _add_candidate(source, full_email):
        full_email = full_email.strip().lower()
        if '@' not in full_email:
            return
        dom = full_email.split('@', 1)[1]
        if dom in _SITA_INTERNAL:
            return
        if any(c[1] == full_email for c in _candidates):
            return
        _candidates.append((source, full_email, dom))

    # (a) From: header line — anchored
    _from_line_match = re.search(r'^\s*from\s*:[^\n]*', text,
                                 re.MULTILINE | re.IGNORECASE)
    if _from_line_match:
        _from_line = _from_line_match.group(0)
        for _m in re.finditer(
                r'[a-z0-9._%+\-]+@[a-z0-9.\-]+\.[a-z]{2,}',
                _from_line, re.IGNORECASE):
            _add_candidate('from_header', _m.group(0))

    # (b) Signature / body emails — collect in order of appearance
    for _m in re.finditer(
            r'[a-z0-9._%+\-]+@[a-z0-9.\-]+\.[a-z]{2,}',
            text, re.IGNORECASE):
        _add_candidate('signature', _m.group(0))

    # ── Two-pass cascade ────────────────────────────────────────────────
    _agency_match = None
    _matched_via = ''
    _sender_domain = ''

    # Pass 1: exact email match (per-agent CRM index)
    if _AGENCY_BY_EMAIL:
        for _src, _full, _dom in _candidates:
            _candidate = _AGENCY_BY_EMAIL.get(_full)
            if _candidate and _candidate.get('market'):
                _agency_match = _candidate
                _matched_via = f'{_src}_email'
                _sender_domain = _dom
                break

    # Pass 2: domain fallback with multi-account city disambiguation
    #
    # When _AGENCY_BY_DOMAIN_LIST has multiple records for a domain
    # (e.g. ACC0131 Witney UK + ACC0132 Boston US under audleytravel.com),
    # we look in the email body for any of the agency cities. First city
    # match wins. If no city resolves, fall back to the canonical
    # brand-priority record (with the existing agency_ambiguous flag set
    # for downstream awareness).
    if not _agency_match and _AGENCY_BY_DOMAIN:
        _seen_doms = set()
        for _src, _full, _dom in _candidates:
            if _dom in _seen_doms:
                continue
            _seen_doms.add(_dom)
            _records = _AGENCY_BY_DOMAIN_LIST.get(_dom, [])
            if not _records:
                continue

            # Single record per domain — unambiguous, use directly
            if len(_records) == 1:
                _agency_match = _records[0]
                _matched_via = f'{_src}_domain'
                _sender_domain = _dom
                break

            # Multiple records — try city disambiguation.
            # Find which agency-city appears EARLIEST in the body. More
            # intuitive than record-iteration order: a forwarded email
            # mentioning "Boston office forwarding to Witney" picks Boston
            # (the originator, mentioned first), not Witney.
            _city_match = None
            _earliest_pos = len(text) + 1
            for _rec in _records:
                _city = (_rec.get('city') or '').strip().lower()
                if len(_city) >= 4:
                    _m = re.search(
                        r'\b' + re.escape(_city) + r'\b', text)
                    if _m and _m.start() < _earliest_pos:
                        _earliest_pos = _m.start()
                        _city_match = _rec

            if _city_match:
                _agency_match = _city_match
                _matched_via = f'{_src}_domain_city'
                _sender_domain = _dom
                break

            # Multi-record, no city signal — fall back to canonical
            _canonical = _AGENCY_BY_DOMAIN.get(_dom)
            if _canonical and _canonical.get('market'):
                _agency_match = _canonical
                _matched_via = f'{_src}_domain'
                _sender_domain = _dom
                break

    if _agency_match and _agency_match.get('market'):
        result['source_market']        = _agency_match['market']
        result['agency_name']          = _agency_match['agency_name']
        result['agency_city']          = _agency_match['city']
        result['agency_domain']        = _sender_domain
        result['agency_account_code']  = _agency_match.get('account_code', '')
        result['agency_brand']         = _agency_match.get('brand', '')
        result['agency_ambiguous']     = _agency_match.get('ambiguous', False)
        result['agency_all_accounts']  = _agency_match.get('all_accounts', [])
        result['_market_source']       = 'agency_portfolio'
        result['_agency_match_via']    = _matched_via
        # 'from_header_email'       — exact email match, agent in From: line
        # 'signature_email'         — exact email match, agent in body
        # 'from_header_domain_city' — multi-record domain, city in body
        #                             matched a specific account
        # 'signature_domain_city'   — multi-record domain (signature source),
        #                             city in body matched
        # 'from_header_domain'      — domain match, single record OR
        #                             multi-record canonical fallback
        # 'signature_domain'        — domain match (signature source)
    else:
        # Step 2: Keyword scoring fallback
        result['_market_source'] = 'keyword_scoring'
        market_signals = {
            'GBR': ['united kingdom', 'england', 'english', 'british', 'gbr', '.co.uk',
                    ' uk ', '(uk)', 'uk\n',
                    'london', 'manchester', 'birmingham', 'edinburgh', 'glasgow', 'scotland', 'wales',
                    'fly from london', 'fly to london', 'from london', 'to london',
                    'heathrow', 'gatwick', 'stansted', 'luton'],
            'ITA': ['italy', 'italian', 'italia', '.it', 'cartorange', 'milan',
                    'rome', 'florence', 'venice', 'naples'],
            'DEU': ['germany', 'german', 'deutschland', '.de', 'munich', 'berlin',
                    'frankfurt', 'hamburg', 'lufthansa', 'gmbh'],
            'FRA': ['france', 'french', 'français', '.fr', 'paris', 'lyon'],
            'USA': ['usa', 'united states', 'american', '.us'],
            'JPN': ['japan', 'japanese', '.jp', 'tokyo', 'osaka'],
            'AUS': ['australia', 'australian', '.au', 'sydney', 'melbourne'],
            'CHE': ['switzerland', 'swiss', '.ch', 'zurich', 'geneva'],
            'NLD': ['netherlands', 'dutch', '.nl', 'amsterdam'],
            'ESP': ['spain', 'spanish', '.es', 'madrid', 'barcelona'],
            'BEL': ['belgium', 'belgian', '.be', 'brussels'],
            'POL': ['poland', 'polish', 'pol', '.pl', 'warsaw'],
            'ROU': ['romania', 'romanian', 'rou', '.ro', 'bucharest'],
            'RUS': ['russia', 'russian', 'rus', '.ru', 'moscow'],
            'CAN': ['canada', 'canadian', 'can', '.ca', 'toronto', 'vancouver'],
            'IND': ['india', 'indian', 'ind', '.in', 'sita.in'],
            'PRT': ['portugal', 'portuguese', 'prt', '.pt', 'lisbon', 'porto', 'faro'],
            'AUT': ['austria', 'austrian', 'aut', '.at', 'vienna', 'wien', 'salzburg', 'innsbruck'],
        }
        SITA_INTERNAL_DOMAINS = {
            'sita.in', 'distantfrontiers.in', 'tci.co.in', 'govacation-india.com'
        }
        HIGH_VALUE_DOMAINS = {
            'FRA': {
                'visiteurs.com', 'amplitudes.com', 'altiplano-voyage.com',
                'voyages-du-monde.com', 'comptoir.fr', 'terre-entiere.com',
                'marco-vasco.com', 'kuoni.fr', 'nouvelles-frontieres.fr',
            },
            'GBR': {
                'abercrombiekent.co.uk', 'steppestravel.com', 'kuoni.co.uk',
                'transindus.co.uk', 'audleytravel.com', 'hayes-jarvis.com',
            },
            'DEU': {'studiosus.com', 'dertour.de', 'neckermann.de'},
            'ITA': {'cartorange.com', 'italiacharme.it'},
        }
        for _mkt, _domains in HIGH_VALUE_DOMAINS.items():
            for _d in _domains:
                if _d not in market_signals[_mkt]:
                    market_signals[_mkt].append(_d)
        scoring_lines = [
            line for line in text.split('\n')
            if not any(d in line for d in SITA_INTERNAL_DOMAINS)
        ]
        scoring_text = '\n'.join(scoring_lines)
        market_scores = {}
        for market, signals in market_signals.items():
            score = 0
            for s in signals:
                if s not in scoring_text:
                    continue
                hv = HIGH_VALUE_DOMAINS.get(market, set())
                weight = 3 if s in hv else 1
                # Short abbreviations (≤4 chars, no dot) need word-boundary match
                # to prevent 'aut' matching 'automne', 'rou' matching inside words etc.
                if len(s) <= 4 and not s.startswith('.'):
                    if re.search(r'\b' + re.escape(s) + r'\b', scoring_text):
                        score += weight
                elif s in scoring_text:
                    score += weight
            if score > 0:
                market_scores[market] = score
        if market_scores:
            result['source_market'] = max(market_scores, key=market_scores.get)
        else:
            result['source_market'] = 'UNKNOWN'

    # ── Step 3: City-based market confirmation / fallback ─────────────────────
    # Scans the email text for known international city names from the CRM
    # city→market lookup.  Used in two ways:
    #   a) Confirmation: if city market matches portfolio/keyword result → no change
    #   b) Rescue: if market is still UNKNOWN, city match promotes a suggestion
    # Never overrides a portfolio-matched market (Step 1 is authoritative).
    if _AGENCY_CITY_MARKET:
        _city_market_found = None
        _city_found = None
        for city_key, city_mkt in _AGENCY_CITY_MARKET.items():
            # Word-boundary match — avoid matching "nice" in "service" etc.
            if re.search(r'\b' + re.escape(city_key) + r'\b', text):
                _city_market_found = city_mkt
                _city_found = city_key
                break   # first match wins (longest city names should come first ideally)

        if _city_market_found:
            if result.get('_market_source') == 'agency_portfolio':
                # Portfolio match is authoritative — just record city as confirmation
                result['_city_confirmed_market'] = _city_market_found
                result['_city_found'] = _city_found
            elif result.get('source_market') == 'UNKNOWN':
                # Rescue: promote city-based suggestion
                result['source_market']   = _city_market_found
                result['_market_source']  = 'city_lookup'
                result['_city_found']     = _city_found
            elif result['source_market'] != _city_market_found:
                # Mismatch between keyword scoring and city — flag for awareness
                result['_city_market_conflict'] = (
                    f"keyword→{result['source_market']} vs city({_city_found})→{_city_market_found}"
                )

    # --- Duration / nights ---
    # Pre-process approximation words (EN/FR/DE/ES)
    text = re.sub(r'environ\s+', '', text)        # French "approximately"
    text = re.sub(r'approximately\s+', '', text)
    text = re.sub(r'around\s+', '', text)
    text = re.sub(r'ungef[äa]hr\s+', '', text)    # German "approximately"
    text = re.sub(r'aproximadamente\s+', '', text) # Spanish "approximately"
    text = re.sub(r'aprox\.?\s+', '', text)        # Spanish "approx"
    # Word-form duration normalisation (EN/FR/DE/ES/PT)
    # Must run BEFORE numeric pattern matching so "two weeks" → "14 nights" etc.
    _word_dur_map = [
        (r'\bfour\s+weeks?\b',     '28 nights'),
        (r'\bthree\s+weeks?\b',    '21 nights'),
        (r'\btwo\s+weeks?\b',      '14 nights'),
        (r'\bone\s+week\b',        '7 nights'),
        (r'\ba\s+fortnight\b',     '14 nights'),
        (r'\bdeux\s+semaines?\b',  '14 nuits'),
        (r'\bune\s+semaine\b',     '7 nuits'),
        (r'\bquinze\s+jours?\b',   '14 nuits'),
        (r'\bzwei\s+wochen\b',     '14 nächte'),
        (r'\beine\s+woche\b',      '7 nächte'),
        (r'\bdos\s+semanas?\b',    '14 noches'),
        (r'\buna\s+semana\b',      '7 noches'),
        (r'\bd[íi]as?\b',          'nights'),    # Spanish días/dia → nights (handles accented í)
        (r'\bnoches?\b',           'nights'),    # Spanish noches → nights
    ]
    for _pat, _repl in _word_dur_map:
        text = re.sub(_pat, _repl, text, flags=re.IGNORECASE)

    duration_patterns = [
        r'(\d+)\s*[-–to à]+\s*(\d+)\s*nights?',
        r'(\d+)\s*nights?',
        r'(\d+)\s*[-–to à]+\s*(\d+)\s*(?:nuits?|noche)',   # French/Spanish
        r'(\d+)\s*(?:nuits?|noche)',                           # French/Spanish single
        r'(\d+)\s*[-–to à]+\s*(\d+)\s*(?:nächte?|nacht)',   # German
        r'(\d+)\s*(?:nächte?|nacht)',                          # German single
        r'(\d+)\s*(?:notti|notte)',                            # Italian
        r'(\d+)\s*[-–to]+\s*(\d+)\s*days?',
        r'(\d+)\s*days?',
        r'(\d+)\s*semaines?',                                  # French weeks
        r'(\d+)\s*weeks?',
    ]
    result['duration_nights'] = None

    # B26: Date-range parsing — most authoritative duration signal when present.
    # Handles "from DD/MM/YYYY until DD/MM/YYYY" in EN/FR/DE/ES/IT/PT.
    # If matched, sets duration_nights AND travel_start/travel_end for festival
    # overlap matching downstream.
    _date_range_pattern = re.compile(
        r'(?:from|du|del|vom|between)\s+(?:[a-z]+\s+)?'
        r'(\d{1,2})[/\.\s\-]+(\d{1,2})[/\.\s\-]+(\d{4})\s+'
        r'(?:until|to|au|al|bis|and|y|et|i)\s+(?:[a-z]+\s+)?'
        r'(\d{1,2})[/\.\s\-]+(\d{1,2})[/\.\s\-]+(\d{4})',
        re.IGNORECASE
    )
    _dr_match = _date_range_pattern.search(text)
    if _dr_match:
        try:
            from datetime import date as _date
            d1, m1, y1, d2, m2, y2 = _dr_match.groups()
            _start = _date(int(y1), int(m1), int(d1))
            _end   = _date(int(y2), int(m2), int(d2))
            _nights_from_range = (_end - _start).days
            if 1 <= _nights_from_range <= 60:
                result['duration_nights'] = (_nights_from_range, _nights_from_range)
                result['travel_start'] = _start.isoformat()
                result['travel_end']   = _end.isoformat()
                result['_duration_from_date_range'] = True
        except (ValueError, TypeError):
            pass  # invalid date — fall through to other patterns

    # P1-06: Spelled-out month date-range parsing — "January 6–20, 2027" style.
    # Handles same-month ranges with spelled month name in EN/ES/PT/FR.
    # Covers month-first ("January 6–20, 2027") and day-first ("6–20 January 2027").
    # Em-dash (–), hyphen (-), and spaced variants all matched.
    # Only fires if numeric date-range parser above did not already set duration.
    if not result.get('_duration_from_date_range'):
        _MONTH_NAMES = {
            'january':1,'february':2,'march':3,'april':4,'may':5,'june':6,
            'july':7,'august':8,'september':9,'october':10,'november':11,'december':12,
            'jan':1,'feb':2,'mar':3,'apr':4,'jun':6,'jul':7,'aug':8,
            'sep':9,'oct':10,'nov':11,'dec':12,
            # Spanish
            'enero':1,'febrero':2,'marzo':3,'abril':4,'mayo':5,'junio':6,
            'julio':7,'agosto':8,'septiembre':9,'octubre':10,'noviembre':11,'diciembre':12,
            # Portuguese
            'janeiro':1,'fevereiro':2,u'mar\xe7o':3,'abril':4,'maio':5,'junho':6,
            'julho':7,'agosto':8,'setembro':9,'outubro':10,'novembro':11,'dezembro':12,
            # French
            'janvier':1,u'f\xe9vrier':2,'fevrier':2,'mars':3,'avril':4,'mai':5,'juin':6,
            'juillet':7,u'ao\xfbt':8,'aout':8,'septembre':9,'octobre':10,'novembre':11,
            u'd\xe9cembre':12,'decembre':12,
        }
        _MON_RE = (
            r'january|february|march|april|may|june|july|august|september|october|november|december'
            r'|jan|feb|mar|apr|jun|jul|aug|sep|oct|nov|dec'
            r'|enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre'
            r'|janeiro|fevereiro|mar\xe7o|abril|maio|junho|julho|setembro|outubro|novembro|dezembro'
            r'|janvier|f\xe9vrier|fevrier|mars|avril|mai|juin|juillet|ao\xfbt|aout|novembre|d\xe9cembre|decembre'
        )
        # Month-first: "January 6–20, 2027"
        _p_mf = re.compile(
            rf'({_MON_RE})\s+(\d{{1,2}})\s*[\u2013\-]\s*(\d{{1,2}})[,\s]+(\d{{4}})',
            re.IGNORECASE
        )
        # Day-first: "6–20 January 2027"
        _p_df = re.compile(
            rf'(\d{{1,2}})\s*[\u2013\-]\s*(\d{{1,2}})\s+({_MON_RE})[,\s]+(\d{{4}})',
            re.IGNORECASE
        )
        _smdr_match = _p_mf.search(text) or _p_df.search(text)
        if _smdr_match:
            try:
                from datetime import date as _date2
                _g = _smdr_match.groups()
                # month-first: (mon_str, d1, d2, yr) | day-first: (d1, d2, mon_str, yr)
                if _g[0].isdigit():
                    _d1_s, _d2_s, _mon_s, _yr_s = _g[0], _g[1], _g[2], _g[3]
                else:
                    _mon_s, _d1_s, _d2_s, _yr_s = _g[0], _g[1], _g[2], _g[3]
                _mon_n = _MONTH_NAMES.get(_mon_s.lower())
                if _mon_n:
                    _start2 = _date2(int(_yr_s), _mon_n, int(_d1_s))
                    _end2   = _date2(int(_yr_s), _mon_n, int(_d2_s))
                    _nights2 = (_end2 - _start2).days
                    if 1 <= _nights2 <= 60:
                        result['duration_nights'] = (_nights2, _nights2)
                        result['travel_start'] = _start2.isoformat()
                        result['travel_end']   = _end2.isoformat()
                        result['_duration_from_date_range'] = True
            except (ValueError, TypeError):
                pass  # invalid date — fall through

    # N5: Cross-month spelled-out date range — "arriving 19 Oct ... departing 04 November 2026".
    # Handles two separate spelled month+day mentions in the same sentence/paragraph
    # where the second mention carries the year. EN/ES/PT/FR month names.
    # Only fires when no date-range has been detected yet.
    if not result.get('_duration_from_date_range'):
        _MONTH_NAMES_X = {
            'january':1,'february':2,'march':3,'april':4,'may':5,'june':6,
            'july':7,'august':8,'september':9,'october':10,'november':11,'december':12,
            'jan':1,'feb':2,'mar':3,'apr':4,'jun':6,'jul':7,'aug':8,
            'sep':9,'oct':10,'nov':11,'dec':12,
            'enero':1,'febrero':2,'marzo':3,'abril':4,'mayo':5,'junio':6,
            'julio':7,'agosto':8,'septiembre':9,'octubre':10,'noviembre':11,'diciembre':12,
            'janeiro':1,'fevereiro':2,u'mar\xe7o':3,'abril':4,'maio':5,'junho':6,
            'julho':7,'agosto':8,'setembro':9,'outubro':10,'novembro':11,'dezembro':12,
            'janvier':1,u'f\xe9vrier':2,'fevrier':2,'mars':3,'avril':4,'mai':5,'juin':6,
            'juillet':7,u'ao\xfbt':8,'aout':8,'septembre':9,'octobre':10,'novembre':11,
            u'd\xe9cembre':12,'decembre':12,
            # Short Spanish/Portuguese
            'oct':10,'nov':11,'dic':12,'ene':1,'feb':2,'mar':3,'abr':4,'jun':6,
            'jul':7,'ago':8,'sep':9,
        }
        _MON_RE_X = (
            r'january|february|march|april|may|june|july|august|september|october|november|december'
            r'|jan|feb|mar|apr|jun|jul|aug|sep|oct|nov|dec'
            r'|enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre'
            r'|janeiro|fevereiro|mar\xe7o|abril|maio|junho|julho|setembro|outubro|novembro|dezembro'
            r'|janvier|f\xe9vrier|fevrier|mars|avril|mai|juin|juillet|ao\xfbt|aout|novembre|d\xe9cembre|decembre'
        )
        # Matches: "19 Oct ... 04 November 2026" or "19 Oct 2026 ... 04 November 2026"
        # Two passes: numeric-day first, then month-first
        _p_cross_dn = re.compile(
            rf'(\d{{1,2}})\s+({_MON_RE_X})(?:\s+(\d{{4}}))?.{{1,120}}?(\d{{1,2}})\s+({_MON_RE_X})\s+(\d{{4}})',
            re.IGNORECASE
        )
        _p_cross_mn = re.compile(
            rf'({_MON_RE_X})\s+(\d{{1,2}})(?:[,\s]+(\d{{4}}))?.{{1,120}}?({_MON_RE_X})\s+(\d{{1,2}})[,\s]+(\d{{4}})',
            re.IGNORECASE
        )
        for _pat_x, _mode_x in [(_p_cross_dn, 'dn'), (_p_cross_mn, 'mn')]:
            _mx = _pat_x.search(text)
            if _mx:
                try:
                    from datetime import date as _date3
                    _gx = _mx.groups()
                    if _mode_x == 'dn':
                        # (d1, mon1, yr1_opt, d2, mon2, yr2)
                        _d1x, _m1x, _yr1x, _d2x, _m2x, _yr2x = _gx
                        _yr_x = int(_yr2x)
                        _yr1_x = int(_yr1x) if _yr1x else _yr_x
                    else:
                        # (mon1, d1, yr1_opt, mon2, d2, yr2)
                        _m1x, _d1x, _yr1x, _m2x, _d2x, _yr2x = _gx
                        _yr_x = int(_yr2x)
                        _yr1_x = int(_yr1x) if _yr1x else _yr_x
                    _mon1_n = _MONTH_NAMES_X.get(_m1x.lower())
                    _mon2_n = _MONTH_NAMES_X.get(_m2x.lower())
                    if _mon1_n and _mon2_n:
                        _start3 = _date3(_yr1_x, _mon1_n, int(_d1x))
                        _end3   = _date3(_yr_x,  _mon2_n, int(_d2x))
                        _nights3 = (_end3 - _start3).days
                        if 1 <= _nights3 <= 60:
                            result['duration_nights'] = (_nights3, _nights3)
                            result['travel_start'] = _start3.isoformat()
                            result['travel_end']   = _end3.isoformat()
                            result['_duration_from_date_range'] = True
                            break
                except (ValueError, TypeError):
                    pass

    # First try: find a single explicit total duration statement
    # (e.g. "13 nights", "2 weeks", "10-14 nights").
    # B26: skip first-pass if date-range already established the duration —
    # date ranges are more authoritative than free-text "N nights" mentions.
    _total_duration_found = bool(result.get('_duration_from_date_range'))
    if not _total_duration_found:
        for pattern in duration_patterns:
            match = re.search(pattern, text)
            if match:
                if match.lastindex == 2:
                    n1, n2 = int(match.group(1)), int(match.group(2))
                    if 'semaine' in pattern or 'week' in pattern:
                        n1, n2 = n1 * 7, n2 * 7
                    # B-A2: reject implausible values (typically year-shaped, e.g. "2026 day")
                    if not (1 <= n1 <= 90 and 1 <= n2 <= 90):
                        continue
                    result['duration_nights'] = (n1, n2)
                else:
                    n = int(match.group(1))
                    if 'semaine' in pattern or 'week' in pattern:
                        n = n * 7
                    # B-A2: reject implausible values (typically year-shaped, e.g. "2026 day")
                    if not (1 <= n <= 90):
                        continue
                    result['duration_nights'] = (n, n)
                _total_duration_found = True
                break

    # Second pass: decide between trip TOTAL vs SUM of per-city allocations.
    # B13 logic:
    # - If scope keyword (in india / in south india / total / throughout) follows
    #   a value within 50 chars -> that value is the trip TOTAL (Chorten case)
    # - Else if multiple distinct values exist -> SUM as per-city allocations
    #   (Melodie case: 2 nuits Delhi + 3 nuits Bandhavgarh + 3 nuits Kanha = 8)
    # - Else first-pass result stands
    _city_night_pattern = re.compile(
        r'(\d+)\s*(?:nights?|nuits?|nächte?|nacht|notti|notte|noche)',
        re.IGNORECASE
    )
    _all_with_pos = [(int(m.group(1)), m.start())
                     for m in _city_night_pattern.finditer(text)
                     if 1 <= int(m.group(1)) <= 30]

    _scope_keywords = (
        r'\bin\s+(?:south|north|central|east|west)?\s*india\b|'
        r'\bin\s+(?:kerala|rajasthan|tamil\s+nadu|karnataka)\b|'
        r'\bin\s+total\b|\btotal\b|\bin\s+country\b|\bthroughout\b|'
        r'\ben\s+inde\b'
    )
    _scope_total = None
    for n_value, pos in _all_with_pos:
        _window = text[pos:pos + 50]
        if re.search(_scope_keywords, _window, re.IGNORECASE):
            _scope_total = n_value
            break

    # Subject+body duplicate detection: if exactly TWO matches with the same
    # value AND they are far apart (>120 chars), that's likely the duration
    # appearing in subject and body. Else keep all matches as per-city
    # allocations (Mélodie case: 2 nuits Delhi + 3 nuits Bandhavgarh + 3 nuits Kanha).
    _deduped_pos = list(_all_with_pos)
    if len(_all_with_pos) == 2:
        (n1, p1), (n2, p2) = _all_with_pos
        if n1 == n2 and abs(p2 - p1) > 120:
            # Subject/body duplicate of same value
            _deduped_pos = [(n1, p1)]
    _deduped_counts = [n for n, _ in _deduped_pos]

    # Range protection: if first-pass found a range (e.g. "10-14 nights"),
    # don't let scope-total or per-city sum override it. Range stays as tuple.
    # B26 protection: same for date-range derived duration — it's authoritative.
    _existing_dur = result.get('duration_nights')
    _is_existing_range = (_existing_dur and isinstance(_existing_dur, tuple)
                          and _existing_dur[0] != _existing_dur[1])
    _is_from_date_range = bool(result.get('_duration_from_date_range'))

    if _scope_total is not None and not _is_existing_range and not _is_from_date_range:
        # Scope keyword present -> trip TOTAL (overrides first-pass single value)
        result['duration_nights'] = (_scope_total, _scope_total)
        _total_duration_found = True
    elif len(_deduped_pos) >= 2 and _scope_total is None and not _is_from_date_range:
        # Multiple per-city allocations -> sum.
        # Preserve first-pass range tuple (10-14 nights) without overriding.
        _existing = result.get('duration_nights')
        _is_range = (_existing and isinstance(_existing, tuple)
                     and _existing[0] != _existing[1])
        if not _is_range:
            _summed = sum(_deduped_counts)
            result['duration_nights'] = (_summed, _summed)

    # --- Pax count ---
    # Step 1: Detect "NAME + N PAX" pattern — means N+1 total (e.g. "Calestani + 1 PAX" = 2)
    _plus_pax = re.search(r'[A-Za-z]+\s*\+\s*(\d+)\s*pax', text, re.IGNORECASE)
    if _plus_pax:
        result['num_pax'] = int(_plus_pax.group(1)) + 1
    else:
        # P1-02 (lifted from fallback to priority position 21 May): explicit couple
        # salutation patterns are a strong signal — two named people. Must run BEFORE
        # the numeric pax loop, otherwise noisy phrases like "1 guest" / "1 traveller"
        # / "1 person" steal priority and lock num_pax=1, leaving Mrs & Mr skipped.
        _mr_mrs_pattern = re.compile(
            r'\b(?:mrs?\s*[&/]\s*mrs?|mr\s*[&/]\s*mrs?|mrs?\s+and\s+mrs?'
            r'|madame\s+et\s+monsieur|monsieur\s+et\s+madame'
            r'|se[ñn]or\s+y\s+se[ñn]ora|se[ñn]ora\s+y\s+se[ñn]or'
            r'|herr\s+und\s+frau|frau\s+und\s+herr)\b',
            re.IGNORECASE
        )
        if _mr_mrs_pattern.search(text):
            result['num_pax'] = 2
            result['rooms']   = {'doubles': 1, 'singles': 0, 'twins': 0}
            print("  ℹ Pax inferred from 'Mrs & Mr' pattern: 2 (1 double)")
        else:
            pax_patterns = [
                r'(\d+)\s*(?:pax|passenger|person|people|adult|travell|voyageur|voyageurs?|personne|passagers?|passager|guest)',
                r'number\s+of\s+(?:guest|pax|passenger|person|travell)s?\s*[:\-]?\s*(\d+)',
                r'(?:guest|pax|passenger|person|travell)s?\s*[:\-]\s*(\d+)',   # "Guests: 2"
                r'(?:nombre\s+de\s+(?:personnes?|passagers?|voyageurs?))\s*[:\-]?\s*(\d+)',  # French: nombre de personnes: 2
                r'(\d+)\s*(?:pers\.?|persones?)',  # pers. abbreviation
                r'(\d+)\s*x\s*\d+',
                r'for\s+(\d+)(?!\s*(?:nights?|nuits?|nächte?|notte|days?|weeks?|semaines?))',
                r'group\s+of\s+(\d+)',
                r'\bx\s*(\d+)\b',   # B-A1: word-boundary anchored — prevents flight numbers (LX146) leaking
            ]
            result['num_pax'] = None
            for pattern in pax_patterns:
                match = re.search(pattern, text)
                if match:
                    result['num_pax'] = int(match.group(1))
                    break
        # Fallback: count salutations (Mr/Mrs/Miss/Ms/Dr) as pax
        if not result['num_pax']:
            _salutations = re.findall(
                r'\b(?:mr|mrs|miss|ms|dr|m\.|mme|mlle)\.?\s+[a-z]',
                text, re.IGNORECASE
            )
            if len(_salutations) >= 2:
                result['num_pax'] = len(_salutations)
                print(f"  ℹ Pax derived from salutations: {result['num_pax']}")

        # Fallback: contextual defaults when pax genuinely unknown
        # "guest" / "guests" → 2 pax · "family" / "famille" / "familia" → 4 pax
        # Also handles multilingual: "ein Paar", "pareja", "couple", "un couple"
        if not result['num_pax']:
            _family_pattern = re.compile(
                r'\b(?:family|famille|familia|familj|familie|family\s+of'
                r'|famil[íi]a|eine\s+familie)\b',
                re.IGNORECASE
            )
            _couple_pattern = re.compile(
                r'\b(?:a\s+couple|un\s+couple|une\s+couple|ein\s+paar|ein\s+p[äa]rchen|p[äa]rchen'
                r'|pareja|coppia|a\s+pair|one\s+couple|1\s+couple'
                r'|have\s+a\s+couple|i\s+have\s+a\s+couple'
                r'|tengo\s+una\s+pareja|tengo\s+un\s+matrimonio)\b',
                re.IGNORECASE
            )
            _guest_pattern = re.compile(
                r'\b(?:a\s+guest|the\s+guests?|our\s+guests?|my\s+guest'
                r'|clients?\s+(?:are|is)\s+a\s+(?:couple|pair)|guests?)\b',
                re.IGNORECASE
            )
            if _family_pattern.search(text):
                result['num_pax'] = 4
                result['rooms']   = {'doubles': 2, 'singles': 0, 'twins': 0}
                print(f"  ℹ Pax inferred from 'family' keyword: 4 (2 doubles)")
            elif _couple_pattern.search(text):
                result['num_pax'] = 2
                result['rooms']   = {'doubles': 1, 'singles': 0, 'twins': 0}
                print(f"  ℹ Pax inferred from 'couple/pair' keyword: 2 (1 double)")
            elif _guest_pattern.search(text):
                result['num_pax'] = 2
                result['rooms']   = {'doubles': 1, 'singles': 0, 'twins': 0}
                print(f"  ℹ Pax inferred from 'guest' keyword: 2 (1 double)")

    # --- Room config ---
    doubles = re.search(r'(\d+)\s*(?:double|chambre double|dbl)', text)
    couples = re.search(r'(\d+)\s*couple', text)
    if not doubles and couples:
        doubles = couples  # 2 couples = 2 double rooms
    singles = re.search(r'(\d+)\s*(?:single|sgl|chambre\s+single|chambre\s+simple|habitaci[oó]n\s+individual)', text, re.IGNORECASE)
    twins   = re.search(r'(\d+)\s*(?:twin|twn|chambre\s+twin|habitaci[oó]n\s+twin)', text, re.IGNORECASE)

    # Only overwrite rooms if explicit keywords found — preserve values set by
    # family/couple/guest inference above
    if doubles or singles or twins:
        result['rooms'] = {
            'doubles': int(doubles.group(1)) if doubles else (int(couples.group(1)) if couples else 0),
            'singles': int(singles.group(1)) if singles else 0,
            'twins':   int(twins.group(1)) if twins else 0,
        }
    elif 'rooms' not in result:
        result['rooms'] = {'doubles': 0, 'singles': 0, 'twins': 0}

    # French "1 chambre single" without explicit digit — treat as 1 single
    if not singles and re.search(r'\bchambre\s+single\b', text, re.IGNORECASE):
        result['rooms']['singles'] = max(result['rooms'].get('singles', 0), 1)
        if not result['num_pax']:
            result['num_pax'] = 1
            print("  ℹ Pax inferred from 'chambre single': 1 (1 single)")
    # Handle "X couples" — each couple = 2 pax, 1 double room
    couples_match = re.search(r'(\d+)\s*couples?', text)
    if couples_match:
        n_couples = int(couples_match.group(1))
        if result['rooms']['doubles'] == 0:
            result['rooms']['doubles'] = n_couples
        # Always derive pax from couples — overrides single-digit match
        result['num_pax'] = n_couples * 2

    # Default: if pax=2 and no rooms detected, assume 1 double
    total_rooms = sum(result['rooms'].values())
    if total_rooms == 0 and result['num_pax'] == 2:
        result['rooms']['doubles'] = 1
    # Default: if pax=1 and no rooms detected, assume 1 single
    elif total_rooms == 0 and result['num_pax'] == 1:
        result['rooms']['singles'] = 1
    # Fix P1-03: if pax still None but rooms were parsed, infer pax from rooms
    if not result.get('num_pax'):
        _r = result.get('rooms', {})
        _inferred_pax = _r.get('doubles', 0) * 2 + _r.get('singles', 0) + _r.get('twins', 0) * 2
        if _inferred_pax > 0:
            result['num_pax'] = _inferred_pax
            print(f"  ℹ Pax inferred from room count: {_inferred_pax}")

    # B25: detect children/kids separately from adults.
    # Adult count → num_pax (already computed). Child count → num_children.
    # For vehicle and room capacity, what matters is total bodies, so we also
    # set num_pax_total = adults + children. Pricing/booking logic downstream
    # may treat children differently (child fares, age-based rates).
    _kids_pattern = re.compile(
        r'(\d+)\s*(?:kids?|child(?:ren)?|enfants?|niñ[oa]s?|kinder|bambini|crianças?)\b',
        re.IGNORECASE
    )
    _kids_match = _kids_pattern.search(text)
    if _kids_match:
        _n_kids = int(_kids_match.group(1))
        if 1 <= _n_kids <= 30:  # sanity bound
            result['num_children'] = _n_kids
            result['num_adults']   = result.get('num_pax') or 0
            # Total bodies = adults + children (used for vehicle + room sizing)
            result['num_pax_total'] = result['num_adults'] + _n_kids
        else:
            result['num_children'] = 0
    else:
        result['num_children'] = 0

    # --- Hotel tier ---
    # Normalise line breaks for multi-word tier keywords (e.g. 'higher\nend' → 'higher end')
    _tier_text = re.sub(r'[\r\n]+', ' ', text)
    # B14: detect star RANGES first ("4-5 star", "4 to 5 star") — these should
    # map to the LOWER tier in the range (more conservative; agent indicates
    # flexibility down to N-star, so we offer N-star + can upgrade)
    _star_range = re.search(
        r'(\d)\s*[-–to]+\s*(\d)\s*(?:[eé]toiles?|stars?|sterne|stelle|estrellas?|\*)',
        _tier_text, re.IGNORECASE
    )
    # Then check single numeric star ratings (4 étoiles, 4 stars, 4*, 5 sterne etc.)
    _star_match = re.search(
        r'(\d)\s*(?:[eé]toiles?|stars?|sterne|stelle|estrellas?|\*)',
        _tier_text, re.IGNORECASE
    )
    result['tier'] = None
    result['tier_keywords_found'] = []
    if _star_range:
        # Range like "4-5 star" — use lower bound for tier mapping
        _lower = min(int(_star_range.group(1)), int(_star_range.group(2)))
        if _lower >= 5:
            result['tier'] = 'lux'
        elif _lower == 4:
            result['tier'] = 'first-class'
        elif _lower == 3:
            result['tier'] = 'moderate'
        result['tier_keywords_found'] = [_star_range.group(0)]
    elif _star_match:
        _stars = int(_star_match.group(1))
        if _stars >= 5:
            result['tier'] = 'lux'
        elif _stars == 4:
            result['tier'] = 'first-class'
        elif _stars == 3:
            result['tier'] = 'moderate'
        result['tier_keywords_found'] = [_star_match.group(0)]
    if not result['tier']:
        for tier, keywords in TIER_KEYWORDS.items():
            found = [k for k in keywords if k in _tier_text]
            if found:
                result['tier'] = tier
                result['tier_keywords_found'] = found
                break

    # --- Heritage flag ---
    # v20260522c (N1): two-tier detection.
    # STANDALONE keywords are unambiguous and trigger heritage on their own:
    #   heritage, palace, haveli, boutique
    # AMBIGUOUS keywords commonly appear as monument / place names and
    # only trigger when they sit within ~40 chars of a stay-context word:
    #   fort, historic, historical, traditional, authentic, character, royal
    # This stops false-fires on 'Amber Fort then drive' (B3 #002 monument),
    # 'Fort Kochi' (B3 #004 place name), 'Royal Bengal Tiger', etc. while
    # still firing on 'fort hotel', 'historic property', 'royal palace stay'.
    _HERITAGE_STANDALONE = ['heritage', 'palace', 'haveli', 'boutique']
    _HERITAGE_AMBIGUOUS = ['fort', 'historic', 'historical', 'traditional',
                           'authentic', 'character', 'royal']
    _HERITAGE_STAY_CONTEXT = ['hotel', 'hotels', 'stay', 'stays', 'property',
                              'properties', 'accommodation', 'accommodations',
                              'lodging', 'resort', 'resorts', 'palace', 'palaces',
                              'haveli', 'havelis', 'heritage']
    _heritage_standalone_pat = re.compile(
        r'\b(?:' + '|'.join(re.escape(k) for k in _HERITAGE_STANDALONE) + r')\b',
        re.IGNORECASE
    )
    _amb_alt = '|'.join(re.escape(k) for k in _HERITAGE_AMBIGUOUS)
    _ctx_alt = '|'.join(re.escape(k) for k in _HERITAGE_STAY_CONTEXT)
    _heritage_ambig_pat = re.compile(
        rf'\b(?:{_ctx_alt})\b[\w\s,/\-]{{0,40}}?\b({_amb_alt})\b'
        rf'|\b({_amb_alt})\b[\w\s,/\-]{{0,40}}?\b(?:{_ctx_alt})\b',
        re.IGNORECASE
    )
    heritage_found = list(_heritage_standalone_pat.findall(text))
    for _m in _heritage_ambig_pat.finditer(text):
        _kw = _m.group(1) or _m.group(2)
        if _kw:
            heritage_found.append(_kw)
    # Deduplicate while preserving order
    _seen = set()
    heritage_found = [k for k in heritage_found if not (_seen.add(k.lower()) or k.lower() in _seen - {k.lower()})]
    result['heritage_requested'] = len(heritage_found) > 0
    result['heritage_keywords'] = heritage_found

    # B24: Heritage upgrade rule.
    # When heritage hotels are explicitly requested AND the parser couldn't
    # detect a tier (or detected only 'moderate'), upgrade tier to 'first-class'.
    # Rationale: heritage properties (palaces, havelis, forts) are predominantly
    # first-class or above in TCI catalogue. Showing budget/moderate hotels for
    # a heritage request mismatches agent intent.
    # Trigger keywords are already in HERITAGE_KEYWORDS — checking flag is enough.
    if result.get('heritage_requested') and result.get('tier') in (None, 'moderate'):
        _prior_tier = result.get('tier')
        result['tier'] = 'first-class'
        result['_tier_upgraded_for_heritage'] = True
        result['_tier_before_upgrade'] = _prior_tier

    # --- Language guide ---
    result['language_guide'] = None
    for pattern in LANGUAGE_GUIDE_PATTERNS:
        match = re.search(pattern, text)
        if match:
            result['language_guide'] = match.group(0)
            break

    # --- Dates ---
    date_patterns = [
        r'(\d{1,2})\s*(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s*(?:20\d\d)?',
        r'(20\d\d[-/]\d{2}[-/]\d{2})',
    ]
    dates_found = []
    for pattern in date_patterns:
        dates_found.extend(re.findall(pattern, text))
    result['dates_raw'] = dates_found[:4]

    # --- Budget ---
    # Budget: strip date/header lines that contain year numbers before scanning
    # Prevents "mars 2026" matching "rs 2026" as currency
    _budget_text = re.sub(
        r'^.*(envoy[eé]|sent|date|janvier|f[eé]vrier|mars|avril|mai|juin|'
        r'juillet|ao[uû]t|septembre|octobre|novembre|d[eé]cembre|'
        r'january|february|march|april|may|june|july|august|'
        r'september|october|november|december).*$',
        '', text, flags=re.IGNORECASE | re.MULTILINE
    )
    # Budget: search near budget keywords first to avoid false matches
    # Use _budget_text (date lines stripped) to prevent "mars 2026" → "rs 2026"
    _budget_ctx = re.search(
        r'budget[^\n]{0,60}?(\d[\d\.,]{2,})\s*(?:€|euros?|usd|\$|eur|inr|£|rs\.?|rupees?)'
        r'|(?:€|euros?|usd|\$|eur|inr|£|rs\.?|rupees?)\s*(\d[\d\.,]{2,})(?:[^\n]{0,30}budget)?'
        r'|(\d[\d\.,]{2,})\s*(?:€|euros?|usd|\$|eur|inr|£|rs\.?|rupees?)\s*(?:per\s*pers|par\s*pers|pp|p\.p)',
        _budget_text, re.IGNORECASE
    )
    if not _budget_ctx:
        # Fallback: any currency amount >= 100 — word boundary on 'rs' avoids 'mars','hours' etc.
        _budget_ctx = re.search(
            r'([1-9]\d{2,}[\d\.,]*)\s*(?:€|euros?|usd|\$|eur|inr|£|(?<!\w)rs\.?|rupees?)'
            r'|(?:€|euros?|usd|\$|eur|inr|£|(?<!\w)rs\.?|rupees?)\s*([1-9]\d{2,}[\d\.,]*)',
            _budget_text, re.IGNORECASE
        )
    result['budget_raw'] = _budget_ctx.group(0).strip() if _budget_ctx else None

    # --- City normalisation map ---
    # Maps any agent/client term → exact DB city name
    # Ordered: more specific signals first to avoid false matches
    CITY_NORM = {
        # Delhi
        'New Delhi':                        ['new delhi', 'delhi', 'ndls', 'igidel',
                                             'indira gandhi', 'connaught', 'imperial hotel delhi',
                                             'oberoi new delhi', 'the oberoi delhi'],
        # Agra — careful not to match "Taj Mahal Palace Mumbai"
        'Agra':                             ['agra', 'amar vilas', 'amarvilas',
                                             'oberoi amarvilas', 'taj agra',
                                             'taj sunrise', 'fatehabad road',
                                             'agra fort', 'sikandra'],
        # Jaipur
        'Jaipur':                           ['jaipur', 'pink city', 'rambagh',
                                             'raj vilas', 'rajvilas', 'jai mahal',
                                             'amber fort', 'amer fort',
                                             'hawa mahal', 'jantar mantar jaipur'],
        # Udaipur
        'Udaipur (RJ)':                     ['udaipur', 'lake city', 'lake pichola',
                                             'pichola', 'fateh garh', 'fateh vilas',
                                             'fateh niwas', 'taj lake palace',
                                             'oberoi udaivilas', 'udaivilas'],
        # Jodhpur
        'Jodhpur':                          ['jodhpur', 'blue city', 'mehrangarh',
                                             'umaid bhawan', 'raas jodhpur',
                                             'clock tower jodhpur'],
        # Jaisalmer
        'Jaisalmer':                        ['jaisalmer', 'golden city', 'desert fort',
                                             'sam sand dunes', 'thar desert jaisalmer'],
        # Varanasi
        'Varanasi':                         ['varanasi', 'banaras', 'benares', 'kashi',
                                             'ganga aarti', 'dashashwamedh', 'sarnath',
                                             'brijrama', 'nadesar'],
        # Khajuraho
        'Khajuraho':                        ['khajuraho', 'chandela', 'western temples'],
        # Mumbai
        'Mumbai':                           ['mumbai', 'bombay', 'taj mahal palace',
                                             'marine drive', 'gateway of india',
                                             'bandra', 'juhu', 'nariman point',
                                             'itc maratha', 'taj lands end',
                                             'four seasons mumbai', 'st regis mumbai'],
        # Goa
        'Goa':                              ['goa', 'panaji', 'panjim', 'calangute',
                                             'baga', 'anjuna', 'candolim', 'vasco'],
        # Kochi
        'Kochi (Cochin)':                   ['kochi', 'cochin', 'ernakulam', 'fort kochi',
                                             'forte kochi', 'chinese fishing nets'],
        # Kerala backwaters
        'Alappuzha (Alleppey)':             ['alleppey', 'alappuzha', 'allepey',
                                             'backwater', 'houseboat alleppey'],
        # Mararikulam
        'Mararikulam':                      ['marari', 'mararikulam', 'carnoustie',
                                             'marari beach'],
        # Kumarakom
        'Kumarakom':                        ['kumarakom', 'vembanad', 'coconut lagoon'],
        # Thekkady
        'Thekkady (Periyar/Kumily)':        ['thekkady', 'periyar', 'kumily',
                                             'spice garden thekkady'],
        # Munnar
        'Munnar':                           ['munnar', 'tea garden', 'eravikulam'],
        # Kovalam
        'Kovalam':                          ['kovalam', 'lighthouse beach', 'trivandrum beach'],
        # Thiruvananthapuram
        'Thiruvananthapuram (Trivandrum)':  ['trivandrum', 'thiruvananthapuram',
                                             'thiruvanath', 'tvm'],
        # Chennai
        'Chennai':                          ['chennai', 'madras', 'egmore', 'marina beach'],
        # Mamallapuram
        'Mamallapuram (Mahabalipuram)':     ['mamallapuram', 'mahabalipuram', 'mahabs',
                                             'shore temple', 'mamalla'],
        # Puducherry
        'Puducherry (Pondicherry)':         ['pondicherry', 'puducherry', 'pondy', 'auroville'],
        # Madurai
        'Madurai':                          ['madurai', 'meenakshi', 'madurai temple'],
        # Thanjavur
        'Thanjavur (Tanjore)':              ['thanjavur', 'tanjore', 'tanjore temple',
                                             'brihadeeswarar'],
        # Tamil Nadu temple trail — Chola circuit enroute stops
        'Chidambaram':                      ['chidambaram', 'nataraja temple', 'nataraja',
                                             'thillai'],
        'Gangaikondacholapuram':            ['gangaikondacholapuram', 'gangai konda',
                                             'gangaikonda', 'brihadisvara gangaikonda'],
        'Darasuram':                        ['darasuram', 'airavatesvara', 'airavateswarar',
                                             'airavatesvara temple'],
        # Chettinad — Tamil Nadu mansion heritage circuit
        'Chettinad':                        ['chettinad', 'chettinadu', 'karaikudi',
                                             'chettiar mansion', 'chettinad mansion'],
        # Kumbakonam — Chola temple circuit
        'Kumbakonam':                       ['kumbakonam'],
        # Kolkata
        'Kolkata':                          ['kolkata', 'calcutta', 'howrah', 'victoria memorial'],
        # Bengaluru
        'Bengaluru':                        ['bengaluru', 'bangalore', 'blr'],
        # Hyderabad
        'Hyderabad':                        ['hyderabad', 'golconda', 'charminar', 'hyd'],
        # Prayagraj
        'Prayagraj':                        ['prayagraj', 'allahabad', 'sangam'],
        # Lucknow
        'Lucknow':                          ['lucknow', 'nawabs', 'lko'],
        # Amritsar
        'Amritsar':                         ['amritsar', 'golden temple', 'wagah border',
                                             'punjab amritsar'],
        # Shimla
        'Shimla':                           ['shimla', 'simla', 'oberoi cecil',
                                             'wildflower hall'],
        # Dharamshala
        'Dharamshala':                      ['dharamshala', 'dharamsala', 'mcleod ganj',
                                             'dalai lama', 'kangra'],
        # Chandigarh
        'Chandigarh':                       ['chandigarh', 'rock garden chandigarh'],
        # Leh
        'Leh':                              ['leh', 'ladakh', 'nubra', 'pangong'],
        # Srinagar
        'Srinagar':                         ['srinagar', 'dal lake', 'kashmir', 'gulmarg'],
        # Rishikesh
        'Rishikesh':                        ['rishikesh', 'rishikesh yoga', 'haridwar rishikesh'],
        # Haridwar
        'Haridwar':                         ['haridwar', 'har ki pauri'],
        # Ranthambore
        'Ranthambore':                      ['ranthambore', 'ranthambhore', 'tiger reserve ranthambore', 'ranthambore national park', 'ranthambore np'],
        # Corbett
        'Corbett':                          ['corbett', 'jim corbett', 'dhikala'],
        # Mandawa
        'Mandawa':                          ['mandawa', 'shekhawati'],
        # Bikaner
        'Bikaner':                          ['bikaner', 'junagarh fort'],
        # Pushkar
        'Pushkar':                          ['pushkar', 'brahma temple pushkar'],
        # Bundi
        'Bundi':                            ['bundi', 'taragarh fort'],
        # Narlai — heritage stopover between Jodhpur and Udaipur (Rawla Narlai)
        'Narlai':                           ['narlai', 'rawla narlai', 'narlai rawla'],
        # Rohet — heritage stopover near Jodhpur (Rohet Garh)
        'Rohet':                            ['rohet', 'rohet garh', 'rohetgarh'],
        # Jawai — leopard country between Jodhpur and Udaipur
        'Jawai':                             ['jawai', 'jawai bandh', 'leopard jawai',
                                              'jawai leopard', 'sujan jawai'],
        # Deogarh — heritage stopover between Jodhpur and Udaipur
        'Deogarh':                           ['deogarh', 'deogarh mahal', 'dev shree',
                                              'castle deogarh'],
        # Shahpura — heritage stopover between Jodhpur and Udaipur
        'Shahpura':                          ['shahpura', 'shahpura bagh', 'bagh shahpura'],
        # Samode — palace village near Jaipur, standalone stop on longer tours
        'Samode':                            ['samode', 'samode palace', 'samode haveli', 'samode bagh'],
        # Kumbhalgarh
        'Kumbhalgarh':                      ['kumbhalgarh', 'kumbalgarh', 'great wall rajasthan'],
        # Gwalior
        'Gwalior':                          ['gwalior', 'gwalior fort', 'jai vilas palace'],
        # Panna — nearest tiger park to Khajuraho
        'Panna':                            ['panna', 'panna tiger', 'panna national park',
                                             'park to see tigers', 'tiger park khajuraho'],
        # Bandhavgarh
        'Bandhavgarh':                      ['bandhavgarh', 'bandhavgarh tiger', 'bandhavgarh national park'],
        # Kanha
        'Kanha':                            ['kanha', 'kanha tiger', 'kanha national park'],
        # Orchha
        'Orchha':                           ['orchha', 'orchcha'],
        # Khajuraho already above
        # Aurangabad
        'Aurangabad (MH)':                  ['aurangabad', 'ellora', 'ajanta', 'ajanta caves',
                                             'ellora caves', 'bibi ka maqbara'],
        # Ahmedabad
        'Ahmedabad':                        ['ahmedabad', 'ahmedabad gujarat', 'sabarmati'],
        # Gurgaon
        'Gurgaon':                          ['gurgaon', 'gurugram'],
        # Fatehpur Sikri — transit only, no hotel
        'Fatehpur Sikri':                   ['fatehpur sikri', 'fatehpursikri', 'fatehpur sikri'],
        # Nepal
        'Kathmandu':                        ['kathmandu', 'nepal', 'katmandu', 'pokhara'],
        # Bhutan
        'Thimpu':                           ['thimpu', 'thimphu', 'bhutan', 'paro'],
        # Sri Lanka
        'Sri Jayawardenepura Kotte':        ['colombo', 'sri lanka', 'ceylon', 'kandy',
                                             'sigiriya', 'galle'],
        # Darjeeling
        'Darjeeling':                       ['darjeeling', 'darjiling', 'tea estate darjeeling'],
        # Hampi
        'Hampi (Hosapete)':                 ['hampi', 'hosapete', 'vijayanagara'],
        # Coorg
        'Coorg (Kodagu)':                   ['coorg', 'kodagu', 'madikeri'],
        # Mysore
        'Mysuru (Mysore)':                  ['mysore', 'mysuru', 'mysore palace'],
    }

    # Transit cities — do not query hotels for these
    TRANSIT_CITIES = {'Fatehpur Sikri', 'Orchha'}  # day visits unless client explicitly requests overnight

    result['cities_detected'] = []
    seen = set()

    has_mumbai  = any(s in text for s in ['mumbai', 'bombay', 'taj mahal palace',
                                           'marine drive', 'gateway of india',
                                           'itc maratha', 'taj lands end'])
    has_jaipur  = any(s in text for s in ['jaipur', 'pink city', 'rambagh', 'raj vilas',
                                           'rajvilas', 'jai mahal'])

    # Short signals that need word-boundary matching to avoid false positives
    # e.g. 'goa' in 'goal', 'puri' in 'impurity', 'oman' in 'woman'
    BOUNDARY_SIGNALS = {
        'goa', 'puri', 'oman', 'leh', 'una', 'gaya', 'pala', 'diu',
        'mandi', 'alwar', 'alwar', 'arni', 'bari', 'kota', 'pali',
    }

    for db_city, signals in CITY_NORM.items():
        if db_city in seen:
            continue
        for signal in signals:
            # Use word-boundary match for short/common signals
            if signal in BOUNDARY_SIGNALS:
                if not re.search(r'\b' + re.escape(signal) + r'\b', text, re.IGNORECASE):
                    continue
            elif signal not in text:
                continue
            # Agra / Mumbai collision — "taj mahal" alone only triggers Agra if Mumbai not present
            if db_city == 'Agra' and signal == 'taj mahal' and has_mumbai:
                continue
            # "amber fort" only triggers Jaipur if jaipur context present
            if db_city == 'Jaipur' and signal == 'amber fort' and not has_jaipur:
                continue
            result['cities_detected'].append(db_city)
            seen.add(db_city)
            break

    # --- European date format DD.MM.YYYY duration detection ─────────────────
    import re as _re
    if not result.get('duration_nights'):
        _all_eu = _re.findall(r'(\d{2})\.(\d{2})\.(\d{4})', text)
        if len(_all_eu) >= 2:
            try:
                from datetime import date as _date2
                # B-A2: use chronological min/max, not textual first/last —
                # email may list dates out of order (flight lines after trip range)
                _dates = [_date2(int(y), int(m), int(d)) for d, m, y in _all_eu]
                _d1 = min(_dates)
                _d2 = max(_dates)
                if _d2 > _d1 and (_d2 - _d1).days <= 90:
                    _dur = (_d2 - _d1).days
                    result['duration_nights'] = (_dur, _dur)
                    result['travel_start'] = str(_d1)
                    result['travel_end']   = str(_d2)
                    result['_duration_from_pnr'] = True
            except Exception:
                pass

    # --- Flight PNR parser ---
    # Three formats supported:
    #   Spaced:       AF 226V 13MAR CDG DEL
    #   Slash:        AF 284H 17OCT CDG/DEL HK4
    #   Concatenated: 1AF284X03APRFCDGDEL  or  2AF/6E3777K17APRFVNSDEL
    

    INDIAN_AIRPORTS = {
        'DEL', 'BOM', 'MAA', 'BLR', 'CCU', 'HYD', 'AMD', 'JAI', 'VNS', 'LKO',
        'PAT', 'IXB', 'GAU', 'SXR', 'IXC', 'ATQ', 'AGR', 'JLR', 'PNQ', 'GOI',
        'TRV', 'UDR', 'JDH', 'JSA', 'BKB', 'KJB', 'IMF', 'IXZ',
    }

    MONTH_MAP = {
        'JAN': 1, 'FEB': 2, 'MAR': 3, 'APR': 4, 'MAY': 5, 'JUN': 6,
        'JUL': 7, 'AUG': 8, 'SEP': 9, 'OCT': 10, 'NOV': 11, 'DEC': 12,
    }

    def _parse_pnr_date(date_str):
        """Parse '03APR' or '17OCT' → date object. Future-year logic applied."""
        import datetime as _dt
        m = _re.match(r'(\d{1,2})([A-Z]{3})', date_str.upper())
        if not m:
            return None
        day, mon = int(m.group(1)), MONTH_MAP.get(m.group(2))
        if not mon:
            return None
        year = _dt.date.today().year
        try:
            d = _dt.date(year, mon, day)
            if d < _dt.date.today():
                d = _dt.date(year + 1, mon, day)
            return d
        except ValueError:
            return None

    # Pattern 1: spaced + slash formats — AF 226V 13MAR CDG DEL  /  AF 284H 17OCT CDG/DEL
    pnr_spaced = _re.compile(
        r'\b([A-Z]{2})\s*\d+[A-Z]?\s+(\d{1,2}[A-Z]{3})\s+([A-Z]{3})\s*[→>/\-]?\s*([A-Z]{3})',
        _re.IGNORECASE
    )
    # Pattern 2: concatenated AF style — 1AF284X03APRFCDGDEL  or  2AF/6E3777K17APRFVNSDEL
    pnr_concat = _re.compile(
        r'(?:^|(?<=\s))\d[A-Z]{2}(?:/[A-Z0-9]{2,3})?\d+[A-Z](\d{1,2}[A-Z]{3})[A-Z]([A-Z]{3})([A-Z]{3})',
        _re.IGNORECASE | _re.MULTILINE
    )
    # Pattern 3: trailing-date variant — EK 510 DXB DEL 14FEB
    # (date comes AFTER city pair — common in Emirates / older GDS displays)
    pnr_trailing = _re.compile(
        r'\b([A-Z]{2})\s*\d+[A-Z]?\s+([A-Z]{3})\s*[→>/\-]?\s*([A-Z]{3})\s+(\d{1,2}[A-Z]{3})\b',
        _re.IGNORECASE
    )
    # Pattern 4 (B-A3): Natural-language flight prose — extracts gateway IATAs only.
    # "Arrival in DEL with LX146", "Departure from JDH with AI815".
    # Used when agent writes flights as sentences rather than GDS strings.
    # Adds to pnr_cities (gateway promotion) but NOT pnr_segments (no origin pair).
    pnr_prose_arr = _re.compile(
        r'\b(?:arrival|arriving|arr\.?)\s+(?:in|at|to|into)\s+([A-Z]{3})\b',
        _re.IGNORECASE
    )
    pnr_prose_dep = _re.compile(
        r'\b(?:departure|departing|dep\.?|departs?)\s+(?:from|ex|out\s+of)\s+([A-Z]{3})\b',
        _re.IGNORECASE
    )

    pnr_cities   = []
    pnr_dates    = []
    pnr_segments = []  # (date_str, orig, dest) for duration derivation

    for m in pnr_spaced.finditer(text):
        date_str = m.group(2).upper()
        orig, dest = m.group(3).upper(), m.group(4).upper()
        pnr_cities += [orig, dest]
        pnr_dates.append(date_str)
        pnr_segments.append((date_str, orig, dest))

    for m in pnr_concat.finditer(text):
        date_str = m.group(1).upper()
        orig, dest = m.group(2).upper(), m.group(3).upper()
        pnr_cities += [orig, dest]
        pnr_dates.append(date_str)
        pnr_segments.append((date_str, orig, dest))

    # Trailing-date variant — only run if no other format matched
    # (avoids double-counting segments)
    if not pnr_segments:
        for m in pnr_trailing.finditer(text):
            orig, dest = m.group(2).upper(), m.group(3).upper()
            date_str = m.group(4).upper()
            pnr_cities += [orig, dest]
            pnr_dates.append(date_str)
            pnr_segments.append((date_str, orig, dest))

    # B-A3: Natural-language prose patterns — always run, populate pnr_cities only.
    # Validates each capture against INDIAN_AIRPORTS so non-airport tokens
    # (e.g. "Departure from BAR" referencing a place) are rejected.
    # Also tracks arrival/departure direction on result for downstream
    # route-builder use (entry/exit gateway placement).
    for m in pnr_prose_arr.finditer(text):
        code = m.group(1).upper()
        if code in INDIAN_AIRPORTS:
            if code not in pnr_cities:
                pnr_cities.append(code)
            result.setdefault('_pnr_arrival_codes', []).append(code)
    for m in pnr_prose_dep.finditer(text):
        code = m.group(1).upper()
        if code in INDIAN_AIRPORTS:
            if code not in pnr_cities:
                pnr_cities.append(code)
            result.setdefault('_pnr_departure_codes', []).append(code)

    # Derive travel_start / travel_end from first Indian arrival + last Indian departure
    if pnr_segments:
        india_arrivals    = []
        india_departures  = []
        for date_str, orig, dest in pnr_segments:
            d = _parse_pnr_date(date_str)
            if d:
                if dest in INDIAN_AIRPORTS:
                    india_arrivals.append(d)
                # Only count as India departure if flying OUT of India
                # (orig Indian, dest non-Indian) — avoids domestic legs inflating travel_end
                if orig in INDIAN_AIRPORTS and dest not in INDIAN_AIRPORTS:
                    india_departures.append(d)
        if india_arrivals and india_departures:
            travel_start   = min(india_arrivals)
            intl_departure = max(india_departures)  # date of DEL→CDG etc.

            # Overnight connection detection:
            # If there is a domestic segment on (intl_departure - 1 day), the client's
            # last actual day in India is that domestic departure date, not the connection date.
            # e.g. VNS→DEL 17APR night + DEL→CDG 18APR 00:50 → travel_end = 17APR
            import datetime as _dt
            day_before    = intl_departure - _dt.timedelta(days=1)
            domestic_dates = set()
            for _ds, _orig, _dest in pnr_segments:
                if _orig in INDIAN_AIRPORTS and _dest in INDIAN_AIRPORTS:
                    _d = _parse_pnr_date(_ds)
                    if _d:
                        domestic_dates.add(_d)

            travel_end = day_before if day_before in domestic_dates else intl_departure

            nights = (travel_end - travel_start).days
            if nights > 0:
                result['travel_start'] = travel_start.isoformat()
                result['travel_end']   = travel_end.isoformat()
                # B13: only set duration from PNR if body parser didn't already
                # find an explicit "N nights" / "N nuits" / "N nächte" statement.
                # Body-stated duration always wins — PNR is a calendar gap that
                # may include arrival/buffer days.
                if not result.get('duration_nights'):
                    result['duration_nights'] = (nights, nights)
                    result['_duration_from_pnr'] = True
                else:
                    result['_duration_from_body'] = True
                    result['_pnr_calendar_gap']   = nights

    if pnr_cities:
        result['pnr_cities'] = pnr_cities
        result['pnr_dates']  = pnr_dates
        IATA_MAP = {
            'DEL': 'New Delhi',   'BOM': 'Mumbai',       'COK': 'Kochi (Cochin)',
            'BLR': 'Bengaluru',   'MAA': 'Chennai',      'HYD': 'Hyderabad',
            'CCU': 'Kolkata',     'AMD': 'Ahmedabad',    'JAI': 'Jaipur',
            'VNS': 'Varanasi',    'LKO': 'Lucknow',      'PAT': 'Patna',
            'IXB': 'Darjeeling',  'GAU': 'Guwahati',     'IMF': 'Imphal',
            'SXR': 'Srinagar',    'IXC': 'Chandigarh',   'ATQ': 'Amritsar',
            'KNU': 'Kanpur',      'AGR': 'Agra',         'JLR': 'Jabalpur',
            'PNQ': 'Pune',        'GOI': 'Goa',          'TRV': 'Thiruvananthapuram (Trivandrum)',
            'IXZ': 'Andaman',     'UDR': 'Udaipur (RJ)', 'JDH': 'Jodhpur',
            'JSA': 'Jaisalmer',   'BKB': 'Bikaner',      'KJB': 'Khajuraho',
        }
        # B16: PNR cities should only be promoted to cities_detected when
        # they're consistent with the body's stated region. Otherwise PNR
        # cities are gateways, not touring cities.
        # Example: PNR DEL→DXB but body says "south India" — Delhi is just
        # the international gateway, not a stop in the itinerary.
        # NOTE: this block runs BEFORE region detection, so we scan the email
        # text directly for explicit region keywords rather than reading
        # regions_detected (which would be polluted by PNR cities themselves).
        _PNR_CITY_REGION = {
            'New Delhi': 'north_india', 'Mumbai': 'west_india',
            'Jaipur': 'north_india', 'Varanasi': 'north_india',
            'Udaipur (RJ)': 'north_india', 'Jodhpur': 'north_india',
            'Jaisalmer': 'north_india', 'Bikaner': 'north_india',
            'Khajuraho': 'central_india', 'Lucknow': 'north_india',
            'Amritsar': 'north_india', 'Chandigarh': 'north_india',
            'Agra': 'north_india', 'Jabalpur': 'central_india',
            'Kanpur': 'north_india', 'Patna': 'north_india',
            'Kochi (Cochin)': 'south_india', 'Bengaluru': 'south_india',
            'Chennai': 'south_india', 'Hyderabad': 'south_india',
            'Thiruvananthapuram (Trivandrum)': 'south_india',
            'Goa': 'west_india', 'Pune': 'west_india',
            'Ahmedabad': 'west_india',
            'Kolkata': 'east_india', 'Darjeeling': 'east_india',
            'Guwahati': 'northeast_india', 'Imphal': 'northeast_india',
            'Srinagar': 'himalayan', 'Andaman': 'andaman',
        }
        # Detect EXPLICIT body region statements directly from email text
        _text_lower = text.lower()
        _explicit_body_regions = set()
        if re.search(r'\bsouth\s+(?:india|indian)\b', _text_lower) or any(
                kw in _text_lower for kw in ['kerala', 'tamil nadu', 'tamilnadu',
                                              'tamil naadu', 'karnataka', 'karnatka',
                                              'south indian', 'south of india']):
            _explicit_body_regions.add('south_india')
        if re.search(r'\bnorth\s+(?:india|indian)\b', _text_lower) or any(
                kw in _text_lower for kw in ['rajasthan', 'rajistan', 'rajisthan', 'rajstan',
                                              'golden triangle', 'north indian',
                                              'north of india']):
            _explicit_body_regions.add('north_india')
        if any(kw in _text_lower for kw in ['east india', 'eastern india', 'west bengal', 'odisha']):
            _explicit_body_regions.add('east_india')
        if any(kw in _text_lower for kw in ['west india', 'western india', 'maharashtra', 'gujarat']):
            _explicit_body_regions.add('west_india')
        if any(kw in _text_lower for kw in ['northeast india', 'north east india', 'assam', 'meghalaya']):
            _explicit_body_regions.add('northeast_india')

        _has_explicit_body_region = bool(_explicit_body_regions)
        _pnr_gateway_skipped = []  # for transparency banner

        for code in pnr_cities:
            mapped = IATA_MAP.get(code)
            if not mapped or mapped in result.get('cities_detected', []):
                continue
            # Check region consistency
            _city_region = _PNR_CITY_REGION.get(mapped)
            if (_has_explicit_body_region
                    and _city_region
                    and _city_region not in _explicit_body_regions):
                # Region mismatch — PNR city is gateway only, NOT a touring city
                _pnr_gateway_skipped.append((mapped, _city_region))
                continue
            result.setdefault('cities_detected', []).append(mapped)

        if _pnr_gateway_skipped:
            # Dedup — same city can appear multiple times if PNR has it
            # as both arrival and departure
            _seen_gw = set()
            _deduped_gw = []
            for c, r in _pnr_gateway_skipped:
                if c not in _seen_gw:
                    _seen_gw.add(c)
                    _deduped_gw.append((c, r))
            result['_pnr_gateway_only'] = _deduped_gw

        # Extract domestic flight sectors from PNR
        # A domestic segment = both origin and destination are Indian airports
        domestic_sectors = []
        seen_sectors = set()
        for date_str, orig, dest in pnr_segments:
            if orig in INDIAN_AIRPORTS and dest in INDIAN_AIRPORTS:
                key = (orig, dest)
                if key not in seen_sectors:
                    domestic_sectors.append({'orig': orig, 'dest': dest, 'date': date_str})
                    seen_sectors.add(key)
        if domestic_sectors:
            result['domestic_sectors'] = domestic_sectors
    
    # --- URL detection and fetch ---
    # Detect URLs in email (including __url__ markdown format) and fetch page
    # to extract city/itinerary information
    _url_pattern = re.compile(
        r'(?:__|\[)?https?://([^\s\]|>]+)(?:__|\])?',
        re.IGNORECASE
    )
    result['urls_detected'] = []
    for m in _url_pattern.finditer(email_text):
        url = m.group(0).strip('_[]')
        if not url.startswith('http'):
            url = 'https://' + m.group(1)
        # Skip email addresses only — fetch all other URLs including itinerary pages
        if '@' in url:
            continue
        # Skip bare homepages but allow itinerary/product pages
        bare_domains = {'https://kuoni.co.uk', 'https://www.kuoni.co.uk',
                        'https://amplitudes.com', 'https://visiteurs.com'}
        if url.rstrip('/') in bare_domains:
            continue
        result['urls_detected'].append(url)

    # Fetch each URL and extract cities from page text
    if result['urls_detected']:
        try:
            import urllib.request as _urllib
            for _url in result['urls_detected'][:2]:  # max 2 URLs
                try:
                    req = _urllib.Request(_url, headers={'User-Agent': 'Mozilla/5.0'})
                    with _urllib.urlopen(req, timeout=8) as resp:
                        page_text = resp.read().decode('utf-8', errors='ignore').lower()
                    # Extract cities from page text using CITY_NORM signals
                    _found_from_url = []
                    _seen_url = set(result.get('cities_detected', []))
                    for db_city, signals in CITY_NORM.items():
                        if db_city in _seen_url:
                            continue
                        for sig in signals:
                            if len(sig) > 3 and sig in page_text:
                                _found_from_url.append(db_city)
                                _seen_url.add(db_city)
                                break
                    if _found_from_url:
                        result.setdefault('cities_detected', []).extend(_found_from_url)
                        print(f"  ✓ URL fetch: {_url[:60]} — found {len(_found_from_url)} cities")
                except Exception as _ue:
                    print(f"  ⚠ URL fetch failed: {str(_ue)[:60]}")
        except Exception:
            pass

    # --- Transport mentions ---
    result['transport_notes'] = []
    if any(w in text for w in ['domestic flight', 'fly', 'flight']):
        result['transport_notes'].append('domestic flights mentioned')
    if any(w in text for w in ['train', 'rail', 'express', 'rajdhani', 'shatabdi',
                                'gatimaan', 'gaatiman', 'duronto', 'uday express']):
        result['transport_notes'].append('train mentioned')
    # Named train detection
    named_trains = {
        'Gatimaan Express (Delhi→Agra)': ['gatimaan', 'gaatiman'],
        'Shatabdi Express':              ['shatabdi'],
        'Rajdhani Express':              ['rajdhani'],
        'Uday Express (Agra→Khajuraho)': ['uday express'],
        'Palace on Wheels':              ['palace on wheels'],
    }
    result['named_trains'] = []
    for train, signals in named_trains.items():
        if any(s in text for s in signals):
            result['named_trains'].append(train)
    if 'by road' in text or 'road transfer' in text or 'drive to' in text:
        result['transport_notes'].append('road transfer mentioned')

    # --- Agent-specified hotels ---
    # Known hotel name fragments to scan for in email text
    KNOWN_HOTELS = [
        'oberoi new delhi', 'the oberoi', 'oberoi amarvilas', 'amar vilas',
        'haveli dharampura', 'dharampura haveli',
        'the imperial', 'imperial hotel',
        'taj mahal palace', 'taj palace', 'taj hotel',
        'taj lake palace', 'taj lands end',
        'oberoi khajuraho', 'oberoi rajgarh', 'rajgarh palace',
        'rambagh palace', 'raj vilas', 'oberoi rajvilas', 'samode haveli', 'samode palace',
        'jai mahal', 'umaid bhawan', 'fateh garh', 'fateh vilas',
        'raas', 'carnoustie', 'marari beach', 'marari',
        'oberoi udaivilas', 'udaivilas',
        'leela palace', 'leela ambience',
        'aman', 'amanbagh', 'amankora',
        'suryaa', 'claridges', 'le meridien',
        'radisson', 'hilton', 'marriott', 'hyatt', 'sheraton',
        'itc maurya', 'itc maratha', 'itc windsor',
        'trident', 'vivanta', 'ginger',
        'itc mughal', 'ananda in the himalayas', 'ananda spa',
        'aahana resort', 'aahana', 'dhyaana farms', 'dhyaana',
        'mayfair darjeeling', 'mayfair', 'windamere hotel', 'windamere',
        'raas jodhpur', 'raas',
        'brijrama', 'nadesar palace',
        'grand hyatt', 'four seasons', 'st regis',
    ]
    result['agent_specified_hotels'] = []
    for hotel in KNOWN_HOTELS:
        if hotel in text:
            result['agent_specified_hotels'].append(hotel.title())

    # --- Specific activity requests ---
    activity_signals = {
        'Cycle Rickshaw Ride (Delhi)':          ['rickshaw', 'cycle rickshaw'],
        'Ganga Aarti Ceremony':                  ['aarti', 'ganga aarti'],
        'Ganges Sunrise Boat Cruise':            ['sunrise cruise', 'boat.*gange', 'gange.*boat'],
        'Exclusive Boat Ride on Lake Pichola':   ['lake pichola', 'pichola', 'sunset cruise', 'boat.*udaipur'],
        'Jeep Ride at Amber Fort':               ['amber fort', 'jeep.*amber', 'amber.*jeep'],
        # Taj Mahal Sunset is a monument visit preference, not a separate activity — handled in monuments
        'Cooking Class':                          ['cooking class', 'cookery', 'cook.*class'],
        'Sarnath Visit':                          ['sarnath'],
        'Sound & Light Show':                     ['sound.*light', 'light.*show'],
        'Safari / Jeep Safari':                   ['safari', 'jeep safari', 'jungle'],
    }
    result['activities_requested'] = []
    for activity, patterns in activity_signals.items():
        if any(re.search(p, text) for p in patterns):
            result['activities_requested'].append(activity)

    # --- Regional keywords ---
    # ── REGION DETECTION ─────────────────────────────────────────────────
    # Primary: derive from detected cities using CITY_TO_REGION
    # Fallback: keyword signals for explicit region mentions
    REGION_KEYWORDS = {
        'north_india':   ['north india', 'northern india', 'golden triangle',
                          'rajasthan', 'rajistan', 'rajisthan', 'rajstan',
                          'delhi agra jaipur'],
        'south_india':   ['south india', 'southern india', 'kerala', 'tamil nadu',
                          'tamilnadu', 'tamil naadu',
                          'karnataka', 'karnatka', 'down south'],
        'east_india':    ['east india', 'eastern india', 'west bengal', 'odisha'],
        'central_india': ['central india', 'madhya pradesh', 'madya pradesh',
                          'madhyapradesh', 'tiger reserve',
                          'central india wildlife'],
        'west_india':    ['west india', 'western india', 'gujarat', 'maharashtra'],
        'northeast_india': ['north east india', 'northeast india', 'assam',
                            'meghalaya', 'sikkim', 'arunachal'],
        'himalaya':      ['himalaya', 'himalayas', 'leh ladakh', 'uttarakhand'],
        'nepal':         ['nepal', 'kathmandu', 'pokhara', 'everest'],
        'bhutan':        ['bhutan', 'paro', 'thimphu'],
        'sri_lanka':     ['sri lanka', 'ceylon', 'colombo', 'kandy'],
    }
    regions_from_cities = set()
    for city in result.get('cities_detected', []):
        r = CITY_TO_REGION.get(city)
        if r:
            regions_from_cities.add(r)
    regions_from_keywords = set()
    for region, signals in REGION_KEYWORDS.items():
        if any(s in text for s in signals):
            regions_from_keywords.add(region)

    # Regex-based region detection — catches patterns like
    # "north & south india", "north and south india", "north+south india"
    if re.search(r'\bnorth\b.{0,25}\bindia\b', text, re.IGNORECASE):
        regions_from_keywords.add('north_india')
    if re.search(r'\bsouth\b.{0,25}\bindia\b', text, re.IGNORECASE):
        regions_from_keywords.add('south_india')

    result['regions_detected'] = list(regions_from_cities | regions_from_keywords)

    # --- E2-T4: Already visited city detection ---
    # Detects past visits mentioned in email and excludes those cities from recommendations.
    # Three layers: individual cities, regions, named circuits.

    VISITED_TRIGGERS = [
        r'(?:already|previously|before|last\s+(?:year|time|trip|visit))\s+(?:\w+\s+){0,5}(?:visited|seen|done|been\s+to|traveled\s+to|travelled\s+to)',
        r'(?:visited|been\s+to|seen|done|traveled\s+to|travelled\s+to)\s+(?:\w+\s+){0,6}(?:before|already|previously|last\s+(?:year|time|trip|visit))',
        r'(?:visited|been\s+to|seen|done|traveled\s+to|travelled\s+to)\s+(?:\w+\s+){0,5}in\s+(?:20\d\d|19\d\d)',
        r'(?:\w+\s+){0,4}(?:in|back\s+in)\s+(?:20\d\d|19\d\d)',
        r'previous\s+(?:trip|visit|tour|holiday|vacation)\s+(?:included|covered|was|to)',
        r'(?:they|client|he|she|we)\s+(?:have|had|\'ve|\'d)\s+(?:already\s+)?(?:visited|seen|done|been)',
        r'have\s+(?:already\s+)?(?:done|visited|seen|been\s+to)',
        r'\b(?:done|visited|seen)\s+(?:it|this|that|there)\b',
    ]

    # Phrases that signal future intent — split line here, only scan text BEFORE this marker
    FUTURE_MARKERS = re.compile(
        r'\b(?:looking\s+for|want|wants|would\s+like|interested\s+in|seeking|'
        r'now\s+want|this\s+time|now\s+(?:they\s+)?(?:want|wish)|'
        r'instead|rather|prefer|new\s+(?:destination|experience|route))\b',
        re.IGNORECASE
    )

    NAMED_CIRCUITS = {
        'golden triangle': ['Agra', 'Jaipur'],  # Delhi kept as transit
        'taj mahal':       ['Agra'],
        'rajasthan':       ['Jaipur', 'Jodhpur', 'Udaipur (RJ)', 'Jaisalmer',
                            'Bikaner', 'Pushkar', 'Mandawa', 'Ranthambore'],
        'kerala':          ['Kochi (Cochin)', 'Alappuzha (Alleppey)', 'Kumarakom',
                            'Thekkady (Periyar/Kumily)', 'Munnar', 'Kovalam'],
        'north india':     ['Agra', 'Jaipur', 'Varanasi'],
        'south india':     ['Kochi (Cochin)', 'Chennai', 'Madurai', 'Mysuru (Mysore)',
                            'Bengaluru', 'Thekkady (Periyar/Kumily)', 'Munnar'],
        'goa':             ['Goa'],
        'himalayas':       ['Shimla', 'Manali', 'Dharamshala', 'Leh'],
        'ladakh':          ['Leh'],
        'kashmir':         ['Srinagar'],
    }

    REGION_CITIES = {
        'rajasthan':   NAMED_CIRCUITS['rajasthan'],
        'kerala':      NAMED_CIRCUITS['kerala'],
        'goa':         ['Goa'],
        'ladakh':      ['Leh'],
        'kashmir':     ['Srinagar'],
        'himachal':    ['Shimla', 'Manali', 'Dharamshala'],
        'uttarakhand': ['Rishikesh', 'Haridwar', 'Mussoorie', 'Nainital'],
    }

    excluded_cities = set()
    exclusion_reasons = []

    for line in email_text.split('\n'):
        line_lower = line.lower().strip()
        if not line_lower:
            continue

        # Split at future-intent markers — only scan the "past" clause
        future_match = FUTURE_MARKERS.search(line_lower)
        scan_text = line_lower[:future_match.start()] if future_match else line_lower

        has_trigger = any(re.search(t, scan_text) for t in VISITED_TRIGGERS)
        has_negation = bool(re.search(
            r'\b(?:no|not|skip|avoid|exclude|don\'t\s+want|do\s+not\s+want)\b', scan_text
        ))

        if not (has_trigger or has_negation):
            continue

        for circuit, cities in NAMED_CIRCUITS.items():
            if circuit in scan_text:
                for c in cities:
                    if c not in excluded_cities:
                        excluded_cities.add(c)
                        exclusion_reasons.append(f'{c} (circuit: {circuit})')

        for region_kw, cities in REGION_CITIES.items():
            if region_kw in scan_text:
                for c in cities:
                    if c not in excluded_cities:
                        excluded_cities.add(c)
                        exclusion_reasons.append(f'{c} (region: {region_kw})')

        for db_city, signals in CITY_NORM.items():
            if db_city in excluded_cities:
                continue
            if any(s in scan_text for s in signals):
                excluded_cities.add(db_city)
                exclusion_reasons.append(f'{db_city} (mentioned in visited context)')

    # Delhi protection — only hard-exclude if explicitly stated
    delhi_hard_exclude = any(
        re.search(r'\b(?:no|skip|avoid|not)\b.{0,20}\bdelhi\b', line.lower())
        for line in email_text.split('\n')
    )
    if 'New Delhi' in excluded_cities and not delhi_hard_exclude:
        excluded_cities.discard('New Delhi')
        exclusion_reasons = [r for r in exclusion_reasons if not r.startswith('New Delhi')]

    if excluded_cities:
        result['cities_excluded'] = sorted(excluded_cities)
        result['_exclusion_reasons'] = exclusion_reasons
        result['cities_detected'] = [
            c for c in result['cities_detected'] if c not in excluded_cities
        ]
        if result.get('inline_city_order'):
            result['inline_city_order'] = [
                c for c in result['inline_city_order'] if c not in excluded_cities
            ]
    else:
        result['cities_excluded'] = []

    # --- General India revisit detection (regex — does not require LLM) ---
    # Catches phrases like "visit India again", "second visit", "been to India before"
    # without needing specific city mentions. Sets _previously_visited_india flag
    # so South India alternatives are offered alongside North India primary route.
    _REVISIT_PATTERNS = [
        r'\bvisit(?:ing|ed)?\s+india\s+again\b',
        r'\bagain\s+to\s+india\b',
        r'\bsecond\s+(?:time|visit|trip|tour)\s+(?:to\s+)?india\b',
        r'\bindia\s+(?:for\s+the\s+)?second\s+time\b',
        r'\b(?:revisit|re-visit)\s+india\b',
        r'\bback\s+to\s+india\b',
        r'\bbeen\s+to\s+india\s+before\b',
        r'\bprevious(?:ly)?\s+(?:visited|been\s+to)\s+india\b',
        r'\bvisited\s+india\s+(?:before|previously|already|last\s+(?:year|time))\b',
        r'\bthey\s+(?:have|had)\s+(?:already\s+)?(?:visited|seen|done)\s+india\b',
        r'\b(?:we|they|client)\s+(?:know|knew)\s+india\b',
        r'\bindia\s+again\b',
        r'\bagain\s+india\b',
        r'\bwant.*india.*again\b',
        r'\bvisiting\s+india\s+again\b',
    ]
    _raw_for_revisit = email_text.lower()
    if any(re.search(p, _raw_for_revisit, re.IGNORECASE) for p in _REVISIT_PATTERNS):
        result['_previously_visited_india'] = True

    # --- Structured section itinerary parser ---
    # Detects agent emails where each city is listed on its own line with a night count:
    #   DELHI - 2 nights          AGRA - 2 nights ?
    #   DELHI – 2 NUITS           JAIPUR - 3 NIGHTS
    # Builds a nightly_split directly, preserving agent's stated order and night counts.
    # Fires when ≥3 city+night pairs detected across consecutive lines.
    _section_itin = []
    # Pattern A: "CITY - N nights"  (standard)
    _section_pattern = re.compile(
        r'^([A-Za-z\u00c0-\u00ff\s\(\)\-]+?)'
        r'\s*[-\u2013\u2014]\s*'
        r'(\d+)\s*(?:or\s*\d+)?\s*'
        r'(?:nights?|nuits?|notti?|nts?)',
        re.IGNORECASE
    )
    # Pattern B: "N Nts CITY" reversed — Kuoni/UK style
    _section_pattern_rev = re.compile(
        r'^(\d+)\s*(?:or\s*\d+)?\s*'
        r'(?:nights?|nts?|nt)\s+'
        r'([A-Za-z\u00c0-\u00ff\s\(\)]+)',
        re.IGNORECASE
    )
    # Pattern C: "* N nuits à CITY" — French bullet style
    _section_pattern_fr = re.compile(
        r'^[\*\-•]\s*(\d+)\s*(?:nuits?|noche|notti?)\s+(?:à|a|en|in|at)?\s*'
        r'([A-Za-z\u00c0-\u00ff\s\(\)]+)',
        re.IGNORECASE
    )
    for line in email_text.split('\n'):
        line_stripped = line.strip()
        if not line_stripped:
            continue
        m = _section_pattern.match(line_stripped)
        if m:
            city_raw    = m.group(1).strip().lower()
            night_count = int(m.group(2))
        else:
            m = _section_pattern_rev.match(line_stripped)
            if m:
                night_count = int(m.group(1))
                city_raw    = m.group(2).strip().lower()
            else:
                m = _section_pattern_fr.match(line_stripped)
                if m:
                    night_count = int(m.group(1))
                    city_raw    = m.group(2).strip().lower()
                else:
                    continue
        # Match city_raw against CITY_NORM
        matched_city = None
        for db_city, signals in CITY_NORM.items():
            if any(s in city_raw or city_raw in s for s in signals):
                matched_city = db_city
                break
            # Also try direct match on db_city name
            if city_raw in db_city.lower() or db_city.lower() in city_raw:
                matched_city = db_city
                break
        # B22: consecutive-only dedup. Allow non-adjacent repeats (e.g. Delhi
        # bookend: trip starts and ends in Delhi). Only block immediate adjacent
        # duplicates which would be parser noise (same line matched twice).
        if matched_city:
            _last_city = _section_itin[-1][0] if _section_itin else None
            if matched_city != _last_city:
                _section_itin.append((matched_city, night_count))

    if len(_section_itin) >= 3:
        result['nightly_split_agent'] = _section_itin
        # Merge cities into cities_detected preserving order
        existing = set(result.get('cities_detected', []))
        for city, _ in _section_itin:
            if city not in existing:
                result.setdefault('cities_detected', []).append(city)
                existing.add(city)
        # Override duration with sum of agent-specified nights
        total_agent_nights = sum(n for _, n in _section_itin)
        result['duration_nights'] = (total_agent_nights, total_agent_nights)
        print(f"  ✓ Structured section itinerary detected ({len(_section_itin)} cities, "
              f"{total_agent_nights} nights) — using agent-specified route")

    # --- Inline ordered city sequence (em-dash / dash / hyphen separated, no night counts) ---
    # Detects agent lines like:
    #   "Delhi – Amritsar – Dharamsala – Varanasi"  (em-dash)
    #   "Delhi - Amritsar - Dharamsala - Varanasi"  (space-hyphen-space)
    # Preserves agent's stated travel order rather than applying geographic sort.
    inline_order = []
    for line in email_text.split('\n'):
        # Split on em-dash, horizontal bar, OR space-hyphen-space
        parts = re.split(r'\s*[–—]\s*|\s+-\s+', line)
        if len(parts) < 3:
            continue
        matched = []
        for part in parts:
            # Strip suffixes: "/Delhi", "(day use)" etc.
            part_clean = re.sub(r'[/\(].*$', '', part).strip().lower()
            if not part_clean:
                continue
            for db_city, signals in CITY_NORM.items():
                if any(s in part_clean for s in signals):
                    if db_city not in matched:
                        matched.append(db_city)
                    break
        if len(matched) >= 3 and len(matched) > len(inline_order):
            inline_order = matched
    if inline_order:
        result['inline_city_order'] = inline_order

    return result


# ── DATABASE QUERIES ──────────────────────────────────────────────────────────
def touring_filter(rows):
    """Remove airport/transit hotels not suitable for multi-night touring stays."""
    filtered = [r for r in rows if r[0].lower() not in SUPPRESS_FOR_TOURING]
    filtered = [r for r in filtered if not any(
        kw in r[0].lower() for kw in ['aerocity', 'airport', 'terminal']
    )]
    return filtered if filtered else rows

def get_hotels(con, city, market, tier=None, heritage=False, top_n=3, agency_account=None, agency_name=None):
    """
    Returns hotel recommendations for a city using the decision hierarchy:
    0. Agency + City: what THIS agency books here — when ≥30 agency tours exist (sizable data)
    1. Market + City + Tier
    2. Market + City (all tiers)
    3. All markets + City (fallback)

    Agency takes precedence over market when sizable data is available (≥30 tours).
    Below threshold, market data is used — agency data is too sparse to be reliable.
    """
    AGENCY_MIN_TOURS = 30   # minimum tours before agency overrides market

    cur = con.cursor()
    results = []
    source = None

    def heritage_filter(rows):
        if not heritage:
            return rows
        heritage_rows = [r for r in rows if any(
            s in r[0].lower() for s in HERITAGE_NAME_SIGNALS
        )]
        return heritage_rows if len(heritage_rows) >= top_n else rows

    # ── Priority 0: Agency + City ─────────────────────────────────────────
    # Use agency_name (stored in DB as tours.agent_name) when sizable data exists.
    # Falls back to market if agency has fewer than AGENCY_MIN_TOURS total bookings.
    _agency_lookup = agency_name
    if _agency_lookup and not results:
        # Check total tour count for this agency — sizable data threshold
        cur.execute('''SELECT COUNT(*) FROM tours WHERE agent_name = ?''', (_agency_lookup,))
        _total = cur.fetchone()
        _agency_tours = _total[0] if _total else 0

        if _agency_tours >= AGENCY_MIN_TOURS:
            # Sizable data — agency takes priority over market
            cur.execute('''
                SELECT s.service_name, s.tci_classification, COUNT(*) as bookings
                FROM services s
                JOIN tours t ON s.file_code = t.file_code
                WHERE t.agent_name = ? AND s.city_name = ?
                  AND s.record_type = 'Accommodation'
                  AND s.service_name != ''
                GROUP BY s.service_name
                ORDER BY bookings DESC
                LIMIT ?
            ''', (_agency_lookup, city, top_n * 3))
            rows = cur.fetchall()
            if rows:
                rows = [r for r in rows if r[0].lower() not in SUPPRESS_FOR_TOURING
                        and not any(kw in r[0].lower() for kw in ['aerocity','airport','terminal'])]
                rows = heritage_filter(rows)[:top_n]
                if len(rows) >= 2:
                    results = rows
                    source = f'★ {_agency_lookup} ({_agency_tours} tours) · {city}'
        # else: agency data sparse — fall through to market

    # Priority 1: market + city + tier
    if tier:
        cur.execute('''
            SELECT service_name, tci_classification, COUNT(*) as bookings
            FROM services
            WHERE source_market = ? AND city_name = ?
              AND record_type = 'Accommodation'
              AND tci_classification = ?
              AND service_name != ''
            GROUP BY service_name
            ORDER BY bookings DESC
            LIMIT ?
        ''', (market, city, tier, top_n * 3))
        rows = cur.fetchall()
        if rows:
            rows = [r for r in rows if r[0].lower() not in SUPPRESS_FOR_TOURING and not any(kw in r[0].lower() for kw in ["aerocity", "airport", "terminal"])]
            rows = heritage_filter(rows)[:top_n]
            results = rows
            source = f'{market} market · {city} · {tier} tier'

    # Priority 2: market + city (all tiers)
    if not results:
        cur.execute('''
            SELECT service_name, tci_classification, COUNT(*) as bookings
            FROM services
            WHERE source_market = ? AND city_name = ?
              AND record_type = 'Accommodation'
              AND service_name != ''
            GROUP BY service_name
            ORDER BY bookings DESC
            LIMIT ?
        ''', (market, city, top_n * 3))
        rows = cur.fetchall()
        if rows:
            rows = [r for r in rows if r[0].lower() not in SUPPRESS_FOR_TOURING and not any(kw in r[0].lower() for kw in ["aerocity", "airport", "terminal"])]
            rows = heritage_filter(rows)[:top_n]
            results = rows
            source = f'{market} market · {city} · all tiers'

    # Priority 3: all markets + city (fallback)
    if not results:
        cur.execute('''
            SELECT service_name, tci_classification, COUNT(*) as bookings
            FROM services
            WHERE city_name = ?
              AND record_type = 'Accommodation'
              AND service_name != ''
            GROUP BY service_name
            ORDER BY bookings DESC
            LIMIT ?
        ''', (city, top_n * 3))
        rows = cur.fetchall()
        if rows:
            rows = [r for r in rows if r[0].lower() not in SUPPRESS_FOR_TOURING and not any(kw in r[0].lower() for kw in ["aerocity", "airport", "terminal"])]
            rows = heritage_filter(rows)[:top_n]
            results = rows
            source = f'⚠ No {market} data — all markets · {city}'

    # E3-T6: If market results thin (< 3), supplement with GBR fallback
    if len(results) < 3 and market != 'GBR':
        cur.execute('''
            SELECT service_name, tci_classification, COUNT(*) as bookings
            FROM services
            WHERE source_market = 'GBR' AND city_name = ?
              AND record_type = 'Accommodation'
              AND service_name != ''
            GROUP BY service_name
            ORDER BY bookings DESC
            LIMIT ?
        ''', (city, top_n * 3))
        gbr_rows = cur.fetchall()
        gbr_rows = [r for r in gbr_rows if r[0].lower() not in SUPPRESS_FOR_TOURING
                    and not any(kw in r[0].lower() for kw in ['aerocity', 'airport', 'terminal'])]
        gbr_rows = heritage_filter(gbr_rows)
        # Add GBR hotels not already in results
        existing = {r[0].lower() for r in results}
        for r in gbr_rows:
            if r[0].lower() not in existing and len(results) < top_n:
                results.append(r)
                existing.add(r[0].lower())
        if len(results) > 0 and source and 'all markets' in source:
            source = source  # keep warning
        elif len(results) >= 3 and source and 'all markets' in source:
            source = f'⚠ Sparse {market} data — supplemented with GBR · {city}'

    # ── PORTFOLIO SUPPLEMENT ─────────────────────────────────────────────
    # E1-T5: Add portfolio hotels not already in results
    # E1-T6: Correct tier from portfolio TCI classification
    # E1-T7: Heritage flag from official_classification
    HERITAGE_OFFICIAL = {'heritage', 'luxury-heritage-boutique-hotel', '5str-lux'}

    if _PORTFOLIO_BY_CITY:
        # Correct tiers for existing results using portfolio data
        corrected = []
        for (name, db_tier, bookings) in results:
            # City-aware lookup first — prevents 'The Imperial' (Kushinagar budget)
            # overwriting 'The Imperial' (New Delhi lux)
            p = _PORTFOLIO_NAME_CITY_TIER.get((name.lower(), city.lower())) \
                or _PORTFOLIO_NAME_TIER.get(name.lower())
            corrected_tier = p['tier'] if p else db_tier
            corrected.append((name, corrected_tier, bookings))
        results = corrected

        # Supplement with portfolio hotels not in results
        portfolio_hotels = _PORTFOLIO_BY_CITY.get(city, [])
        existing_names = {r[0].lower() for r in results}
        supplements = []
        for h in portfolio_hotels:
            if h['name'].lower() in existing_names:
                continue
            # Filter by tier if specified
            if tier and h['tier'] != tier:
                continue
            supplements.append((h['name'], h['tier'], 0))  # 0 bookings = never booked
        # Add up to (top_n - len(results)) supplements, sorted by tier priority
        TIER_ORDER = {'lux': 0, 'first-class': 1, 'moderate': 2, 'budget': 3}
        supplements.sort(key=lambda x: TIER_ORDER.get(x[1], 9))
        slots = max(0, top_n - len(results))
        results = results + supplements[:slots]

    return results, source


# Activities to suppress — operational/billing items, not genuine experiences
SUPPRESS_ACTIVITIES = {
    'life jacket charges', 'life jacket', 'lifejacket',
    'chair for evening aarti', 'chair for guide during evening aarti',
    'chair for guide', 'chair for aarti',
    'tuk tuk conveyance charges', 'conveyance charges',
    'tuk tuk conveyance', 'transfer charges',
    'vip meet & greet', 'vip meet and greet',
    'meet & greet', 'meet and greet',
    'porter charges', 'porterage',
    'supplement', 'surcharge', 'fuel surcharge',
    'toll charges', 'parking charges',
    'guide charges', 'guide fee', 'hd guide charges',
    'entrance fee', 'entry charges',
    'tips', 'gratuity',
    'mineral water', 'drinking water',
    'luggage transfer', 'baggage transfer',
}

def suppress_activities(rows):
    """Filter out utility/billing items from activity results.
    Uses the shared _ACTIVITY_SUPPRESS_KEYWORDS set so additions there
    apply to both DB rows and CSV portfolio rows.
    """
    return [
        r for r in rows
        if not _activity_suppressed(r[0])
    ]

def get_activities(con, city, market, top_n=5):
    """Top activities for market + city.
    Fallback hierarchy:
      1. DB: market + city
      2. DB: all markets + city
      3. CSV activity portfolio (theme_score ranked, DKC preferred)
    """
    cur = con.cursor()

    cur.execute('''
        SELECT service_name, COUNT(*) as bookings
        FROM services
        WHERE source_market = ? AND city_name = ?
          AND record_type = 'Activity'
          AND service_name != ''
        GROUP BY service_name
        ORDER BY bookings DESC
        LIMIT ?
    ''', (market, city, top_n))
    rows = cur.fetchall()
    rows = suppress_activities(rows)

    if rows:
        return rows[:top_n], f'{market} market'

    # Fallback 1: all markets in DB
    cur.execute('''
        SELECT service_name, COUNT(*) as bookings
        FROM services
        WHERE city_name = ?
          AND record_type = 'Activity'
          AND service_name != ''
        GROUP BY service_name
        ORDER BY bookings DESC
        LIMIT ?
    ''', (city, top_n * 3))
    rows = cur.fetchall()
    rows = [r for r in rows if r[0].lower() not in SUPPRESS_FOR_TOURING
            and not any(kw in r[0].lower() for kw in ['aerocity', 'airport', 'terminal'])]
    rows = suppress_activities(rows)
    if rows:
        return rows[:top_n], f'⚠ No {market} data — all markets'

    # Fallback 2: CSV activity portfolio
    csv_rows, _ = get_activities_csv(city, top_n=top_n)
    if csv_rows:
        converted = [(r['name'], r['theme_score']) for r in csv_rows]
        return converted, '📋 CSV portfolio'

    return [], None


def get_monuments(con, city, market, top_n=5):
    """Top monuments for market + city, with fallback to all markets."""
    cur = con.cursor()

    cur.execute('''
        SELECT service_name, COUNT(*) as bookings
        FROM services
        WHERE source_market = ? AND city_name = ?
          AND record_type = 'Monument'
          AND service_name != ''
        GROUP BY service_name
        ORDER BY bookings DESC
        LIMIT ?
    ''', (market, city, top_n))
    rows = cur.fetchall()

    if not rows:
        cur.execute('''
            SELECT service_name, COUNT(*) as bookings
            FROM services
            WHERE city_name = ?
              AND record_type = 'Monument'
              AND service_name != ''
            GROUP BY service_name
            ORDER BY bookings DESC
            LIMIT ?
        ''', (city, top_n))
        rows = cur.fetchall()
        rows = [r for r in rows if r[0].lower() not in SUPPRESS_FOR_TOURING and not any(kw in r[0].lower() for kw in ["aerocity", "airport", "terminal"])]
        return rows[:top_n], f'⚠ No {market} data — all markets'

    return rows, f'{market} market'


def get_similar_routes(con, cities, market, top_n=3):
    """Find historically similar routes for this market."""
    cur = con.cursor()
    # Build a partial sequence match
    city_seq = ' → '.join(cities)
    cur.execute('''
        SELECT city_sequence, COUNT(*) as bookings
        FROM tours
        WHERE source_market = ?
          AND city_sequence != ''
        GROUP BY city_sequence
        ORDER BY bookings DESC
        LIMIT 20
    ''', (market,))
    all_routes = cur.fetchall()

    # Score by how many requested cities appear in each historical route
    scored = []
    for route, count in all_routes:
        matches = sum(1 for c in cities if c in route)
        if matches > 0:
            scored.append((route, count, matches))

    scored.sort(key=lambda x: (x[2], x[1]), reverse=True)
    return scored[:top_n]


def get_bookend_routes(con, entry_city, exit_city, market, min_nights, max_nights, top_n=3):
    """
    Find historical routes that start at entry_city and end at exit_city
    within the requested duration range, with full nightly splits.
    """
    from datetime import date as _date
    cur = con.cursor()

    cur.execute("""
        SELECT t.file_code, t.city_sequence,
               s.city_name, s.check_in, s.check_out
        FROM tours t
        JOIN services s ON t.file_code = s.file_code
        WHERE t.source_market = ?
        AND t.city_sequence LIKE ?
        AND t.city_sequence LIKE ?
        AND s.record_type = 'Accommodation'
        AND s.check_in != '' AND s.check_out != ''
    """, (market, f'%{entry_city}%', f'%{exit_city}%'))
    rows = cur.fetchall()

    fallback = False
    if not rows:
        cur.execute("""
            SELECT t.file_code, t.city_sequence,
                   s.city_name, s.check_in, s.check_out
            FROM tours t
            JOIN services s ON t.file_code = s.file_code
            WHERE t.city_sequence LIKE ?
            AND t.city_sequence LIKE ?
            AND s.record_type = 'Accommodation'
            AND s.check_in != '' AND s.check_out != ''
        """, (f'%{entry_city}%', f'%{exit_city}%'))
        rows = cur.fetchall()
        fallback = True

    tour_data = defaultdict(lambda: {'sequence': '', 'cities': defaultdict(int)})
    for file_code, seq, city, check_in, check_out in rows:
        try:
            ci = _date.fromisoformat(check_in)
            co = _date.fromisoformat(check_out)
            nights = (co - ci).days
            if nights > 0:
                tour_data[file_code]['sequence'] = seq
                tour_data[file_code]['cities'][city] += nights
        except:
            pass

    route_signatures = defaultdict(int)
    for fc, data in tour_data.items():
        if not data['cities']:
            continue
        seq_cities = [c.strip() for c in data['sequence'].split('->')]
        ordered = [(c, data['cities'][c]) for c in seq_cities if c in data['cities']]
        if not ordered:
            continue
        if entry_city not in [c for c, n in ordered]:
            continue
        if exit_city not in [c for c, n in ordered]:
            continue
        total_nights = sum(n for _, n in ordered)
        if min_nights <= total_nights <= max_nights:
            sig = tuple(ordered)
            route_signatures[sig] += 1

    sorted_routes = sorted(route_signatures.items(), key=lambda x: x[1], reverse=True)
    results = [(list(sig), count) for sig, count in sorted_routes[:top_n]]
    return results, fallback


def fix_transit_stops(route, market=None):
    """
    Apply transit rules:
    - Delhi after Gurgaon = 0 nights (wash & change only)
    - Orchha = 0 nights (day visit en route to Khajuraho) — night transferred to Gwalior if FRA market
    - Fatehpur Sikri = 0 nights (day visit en route Agra↔Jaipur) — never overnight
      (v20260522c, N3): nights freed by Fatehpur Sikri coercion are redistributed
      to Agra (the natural overnight before/after) so the route total stays correct.
    - Duplicate transit-city entries collapsed to a single 0n entry to stop
      'Agra 1n → Fatehpur Sikri 1n' tail-duplication when the email mentions
      Fatehpur Sikri as a touring stop (B3 #002).
    """
    cities = [c for c, n in route]
    orchha_night_freed = 0
    fatehpur_night_freed = 0
    result = []
    for i, (city, nights) in enumerate(route):
        if city == 'New Delhi' and i > 0 and route[i-1][0] == 'Gurgaon':
            result.append((city, 0))  # transit only
        elif city == 'Orchha' and nights > 0:
            orchha_night_freed = nights  # save freed night to give to Gwalior
            result.append((city, 0))  # day visit en route
        elif city == 'Fatehpur Sikri' and nights > 0:
            fatehpur_night_freed += nights
            result.append((city, 0))
        else:
            result.append((city, nights))
    # v20260522c (N3): dedupe consecutive Fatehpur Sikri entries.
    # If the email body and the route builder both inject Fatehpur Sikri,
    # we can end up with two FS entries — collapse to one 0n entry.
    _deduped = []
    _seen_fs = False
    for c, n in result:
        if c == 'Fatehpur Sikri':
            if _seen_fs:
                # Already added; if this entry had a positive night, transfer to Agra
                fatehpur_night_freed += n
                continue
            _seen_fs = True
            _deduped.append((c, 0))  # always 0 — transit-only
        else:
            _deduped.append((c, n))
    result = _deduped
    # Transfer Fatehpur Sikri's freed nights to the adjacent Agra (preferred)
    # or to Jaipur if Agra not in route. If neither, drop the nights silently
    # (better to keep route total slightly under than to over-stay a transit).
    if fatehpur_night_freed > 0:
        _target = 'Agra' if any(c == 'Agra' for c, _ in result) else (
            'Jaipur' if any(c == 'Jaipur' for c, _ in result) else None
        )
        if _target:
            new_result = []
            done = False
            for c, n in result:
                if c == _target and not done:
                    new_result.append((c, n + fatehpur_night_freed))
                    done = True
                else:
                    new_result.append((c, n))
            result = new_result
    # Transfer Orchha's freed night to Gwalior — FRA market only
    # If client mentioned Gwalior, add night to existing entry
    # If not mentioned, auto-inject Gwalior after Orchha and before Khajuraho
    # Only apply Gwalior compensation for FRA market with long itineraries (>20n)
    route_total = sum(n for _, n in result)
    if orchha_night_freed > 0 and market == 'FRA' and (route_total + orchha_night_freed) > 20:
        if 'Gwalior' in cities:
            result = [(c, n + orchha_night_freed if c == 'Gwalior' else n) for c, n in result]
        else:
            new_result = []
            inserted = False
            for c, n in result:
                new_result.append((c, n))
                if c == 'Orchha' and not inserted:
                    new_result.append(('Gwalior', orchha_night_freed))
                    inserted = True
            if not inserted:
                # Orchha was 0n already — insert before Khajuraho
                final = []
                for c, n in new_result:
                    if c == 'Khajuraho' and not inserted:
                        final.append(('Gwalior', orchha_night_freed))
                        inserted = True
                    final.append((c, n))
                new_result = final
            result = new_result
    return result

def insert_delhi_after_corbett(route):
    """Insert New Delhi (0n transit) after Corbett — only way out is back to Delhi to fly south."""
    cities = [c for c, n in route]
    if 'Corbett' not in cities:
        return route
    # Only insert if Delhi not already immediately after Corbett
    corbett_idx = cities.index('Corbett')
    if corbett_idx + 1 < len(cities) and cities[corbett_idx + 1] == 'New Delhi':
        return route  # already there
    result = []
    for i, (city, nights) in enumerate(route):
        result.append((city, nights))
        if city == 'Corbett':
            result.append(('New Delhi', 0))  # transit back to Delhi to fly south
    return result


def insert_fatehpur_sikri(route):
    """Insert Fatehpur Sikri (0 nights) between Agra and Jaipur if both present consecutively."""
    result = []
    cities = [c for c, n in route]
    for i, (city, nights) in enumerate(route):
        result.append((city, nights))
        # Insert after Agra if next city is Jaipur
        if city == 'Agra' and i + 1 < len(route) and route[i+1][0] == 'Jaipur':
            result.append(('Fatehpur Sikri', 0))
        # Insert after Jaipur if next city is Agra (reverse direction)
        if city == 'Jaipur' and i + 1 < len(route) and route[i+1][0] == 'Agra':
            result.append(('Fatehpur Sikri', 0))
    return result


# ── CITY → REGION MAPPING ─────────────────────────────────────────────────────
# Derived from State_Region_Data_Production.csv
CITY_TO_REGION = {
    # North India — Delhi, UP, Rajasthan, Punjab, Haryana, HP, Uttarakhand, J&K, Ladakh, Bihar
    'New Delhi': 'north_india', 'Gurgaon': 'north_india', 'Noida': 'north_india',
    'Agra': 'north_india', 'Varanasi': 'north_india', 'Lucknow': 'north_india',
    'Jaipur': 'north_india', 'Udaipur (RJ)': 'north_india', 'Jodhpur': 'north_india',
    'Jaisalmer': 'north_india', 'Bikaner': 'north_india', 'Pushkar': 'north_india',
    'Mandawa': 'north_india', 'Alsisar': 'north_india', 'Ranthambore': 'north_india',
    'Ajmer': 'north_india', 'Chittorgarh': 'north_india', 'Kumbhalgarh': 'north_india',
    'Ranakpur': 'north_india', 'Rohet': 'north_india', 'Narlai': 'north_india',
    'Jawai': 'north_india', 'Deogarh': 'north_india', 'Barli': 'north_india',
    'Fatehpur Sikri': 'north_india', 'Mathura': 'north_india', 'Vrindavan': 'north_india',
    'Amritsar': 'north_india', 'Chandigarh': 'north_india', 'Shimla': 'north_india',
    'Manali': 'north_india', 'Dharamshala': 'north_india', 'Dalhousie': 'north_india',
    'Mussoorie': 'north_india', 'Rishikesh': 'north_india', 'Haridwar': 'north_india',
    'Nainital': 'north_india', 'Jim Corbett': 'north_india', 'Dehradun': 'north_india',
    'Leh': 'north_india', 'Ladakh': 'north_india', 'Srinagar': 'north_india',
    'Patna': 'north_india', 'Bodh Gaya': 'north_india', 'Nalanda': 'north_india',
    # Central India — MP, Chhattisgarh, Jharkhand
    'Gwalior': 'central_india', 'Orchha': 'central_india', 'Khajuraho': 'central_india',
    'Panna': 'central_india', 'Bandhavgarh': 'central_india', 'Kanha': 'central_india',
    'Pench': 'central_india', 'Satpura': 'central_india', 'Tadoba': 'central_india',
    'Bhopal': 'central_india', 'Indore': 'central_india', 'Ujjain': 'central_india',
    'Pachmarhi': 'central_india', 'Jabalpur': 'central_india',
    'Raipur': 'central_india', 'Jagdalpur': 'central_india', 'Ranchi': 'central_india',
    # North India wildlife (in addition to Ranthambore/Corbett already listed above)
    'Corbett': 'north_india',  # alias for Jim Corbett
    # South India — Kerala, TN, Karnataka, AP, Telangana
    'Kochi (Cochin)': 'south_india', 'Thiruvananthapuram': 'south_india',
    'Alappuzha (Alleppey)': 'south_india', 'Kumarakom': 'south_india',
    'Thekkady (Periyar/Kumily)': 'south_india', 'Munnar': 'south_india',
    'Kovalam': 'south_india', 'Varkala': 'south_india', 'Wayanad': 'south_india',
    'Kabini': 'south_india', 'Nagarhole': 'south_india',
    'Kozhikode': 'south_india', 'Thrissur': 'south_india',
    'Chennai': 'south_india', 'Mamallapuram (Mahabalipuram)': 'south_india',
    'Thanjavur (Tanjore)': 'south_india', 'Madurai': 'south_india',
    'Kumbakonam': 'south_india', 'Chettinad': 'south_india',
    'Chidambaram': 'south_india', 'Gangaikondacholapuram': 'south_india',
    'Darasuram': 'south_india',
    'Ooty (Udhagamandalam)': 'south_india', 'Coimbatore': 'south_india',
    'Puducherry (Pondicherry)': 'south_india',
    'Bengaluru': 'south_india', 'Mysuru (Mysore)': 'south_india',
    'Hampi': 'south_india', 'Badami': 'south_india', 'Coorg': 'south_india',
    'Hyderabad': 'south_india', 'Tirupati': 'south_india', 'Andaman': 'south_india',
    # West India — Maharashtra, Goa, Gujarat
    'Mumbai': 'west_india', 'Pune': 'west_india', 'Aurangabad (MH)': 'west_india',
    'Nashik': 'west_india', 'Lonavala': 'west_india',
    'Goa': 'west_india', 'Panjim': 'west_india',
    'Ahmedabad': 'west_india', 'Vadodara': 'west_india', 'Kutch': 'west_india',
    'Rann of Kutch': 'west_india', 'Diu': 'west_india', 'Bhavnagar': 'west_india',
    # East India — West Bengal, Odisha
    'Kolkata': 'east_india', 'Darjeeling': 'east_india', 'Kalimpong': 'east_india',
    'Sundarbans': 'east_india', 'Bishnupur': 'east_india',
    'Bhubaneswar': 'east_india', 'Puri': 'east_india', 'Konark': 'east_india',
    # North East India
    'Guwahati': 'northeast_india', 'Kaziranga': 'northeast_india',
    'Shillong': 'northeast_india', 'Cherrapunji': 'northeast_india',
    'Gangtok': 'northeast_india', 'Pelling': 'northeast_india',
    'Tawang': 'northeast_india', 'Dimapur': 'northeast_india',
    # International
    'Kathmandu': 'nepal', 'Pokhara': 'nepal', 'Chitwan': 'nepal',
    'Paro': 'bhutan', 'Thimphu': 'bhutan', 'Punakha': 'bhutan',
    'Colombo': 'sri_lanka', 'Kandy': 'sri_lanka', 'Sigiriya': 'sri_lanka',
    'Dubai': 'uae', 'Singapore': 'singapore',
    'Bangkok': 'thailand', 'Phuket': 'thailand', 'Tokyo': 'japan',
    'Nairobi': 'kenya',
}

# ── MARKET → CURRENCY MAP ─────────────────────────────────────────────────────
MARKET_CURRENCY = {
    'GBR': 'GBP', 'USA': 'USD', 'AUS': 'AUD', 'NZL': 'NZD',
    'CHE': 'EUR', 'JPN': 'JPY', 'CAN': 'CAD', 'IND': 'INR',
    'RUS': 'RUB', 'BRA': 'BRL', 'MXN': 'MXN', 'ZAF': 'ZAR',
    'NOR': 'NOK', 'SWE': 'SEK', 'DNK': 'DKK', 'POL': 'PLN',
    'HUN': 'HUF', 'CZE': 'CZK', 'ROU': 'RON', 'UKR': 'UAH',
    'AZE': 'AZN', 'KAZ': 'KZT', 'KGZ': 'KGS', 'GEO': 'GEL',
    'SGP': 'SGD', 'MYS': 'MYR', 'THA': 'THB', 'IDN': 'IDR',
    'PHL': 'PHP', 'VNM': 'VND', 'NPL': 'NPR',
    # All others → EUR
}


# ── MONUMENT COORDINATES (for proximity sorting) ──────────────────────────────
import math as _math

_MONUMENT_COORDS = {
    'Humayuns Tomb': (28.5933,77.2507), "Humayun's Tomb": (28.5933,77.2507),
    'Qutub Minar': (28.5244,77.1855), 'Qutb Minar': (28.5244,77.1855),
    'India Gate (Drive Past)': (28.6129,77.2295), 'India Gate': (28.6129,77.2295),
    'Red Fort [Drive Past]': (28.6562,77.2410), 'Red Fort, Delhi': (28.6562,77.2410), 'Red Fort': (28.6562,77.2410),
    'Jama Masjid': (28.6507,77.2334), 'Lotus Temple': (28.5535,77.2588),
    'Lodi Garden': (28.5931,77.2197), 'Raj Ghat': (28.6400,77.2497),
    'Rashtrapati Bhavan': (28.6143,77.1993), 'Parliament House': (28.6175,77.2090),
    'Taj Mahal (Including Mausoleum Visit)': (27.1751,78.0421), 'Taj Mahal': (27.1751,78.0421),
    'Agra Fort': (27.1800,78.0218),
    "Itmad-ud-Daula's Tomb": (27.1934,78.0370), "Itmad-ud-Daula's Tomb": (27.1934,78.0370),
    'Mehtab Bagh': (27.1706,78.0384), 'Ram Bagh': (27.2011,78.0306), 'Sikandra': (27.2036,77.9706),
    'Fatehpur Sikri': (27.0945,77.6604),
    'Amber Fort': (26.9855,75.8513), 'Amer Fort': (26.9855,75.8513),
    'Jantar Mantar (Observatory)': (26.9246,75.8242), 'Jantar Mantar': (26.9246,75.8242),
    'City Palace Museum': (26.9257,75.8237), 'City Palace': (26.9257,75.8237),
    'Palace of Wind (Hawa Mahal) - Drive Past': (26.9239,75.8267),
    'Palace of Wind  (Hawa Mahal)': (26.9239,75.8267), 'Hawa Mahal': (26.9239,75.8267),
    'Nahargarh Fort': (26.9426,75.8150), 'Albert Hall Museum': (26.9124,75.8185),
    'Sarnath': (25.3783,83.0235), 'Kashi Vishwanath Temple': (25.3109,83.0107),
    'Ramnagar Fort': (25.2833,83.0397),
    'Golden Temple': (31.6200,74.8765), 'Jallianwala Bagh': (31.6210,74.8753),
    'Wagah Border': (31.6042,74.5713),
    'Mehrangarh Fort': (26.2979,73.0188), 'Jaswant Thada': (26.2996,73.0181),
    'City Palace Udaipur': (24.5764,73.6836), 'Jagdish Temple': (24.5783,73.6822),
    'Gateway of India': (18.9220,72.8347),
}
_HOTEL_COORDS = {
    'The Imperial': (28.6242,77.2194), 'Taj Palace': (28.5985,77.1724),
    'The Claridges': (28.6004,77.2157), 'ITC Maurya': (28.5981,77.1708),
    'The Leela Ambience (Shahdara)': (28.6674,77.2838),
    'Taj Mahal Hotel': (28.6011,77.2198), 'The Oberoi': (28.6011,77.2198),
    'The Oberoi Amarvilas': (27.1739,78.0426),
    'Taj Hotel & Convention Centre': (27.1853,78.0064), 'Trident': (27.1770,78.0425),
    'Rambagh Palace': (26.9005,75.8128), 'Taj Amer': (26.9169,75.7989),
    'Samode Haveli': (26.9265,75.8290), 'Marriott': (26.9001,75.7860),
    'Dera Mandawa': (26.9280,75.8295),
}

def _mon_hav(lat1,lon1,lat2,lon2):
    R=6371; dlat,dlon=_math.radians(lat2-lat1),_math.radians(lon2-lon1)
    a=_math.sin(dlat/2)**2+_math.cos(_math.radians(lat1))*_math.cos(_math.radians(lat2))*_math.sin(dlon/2)**2
    return R*2*_math.asin(_math.sqrt(a))

def _mon_get_coords(name, table):
    c=table.get(name)
    if c: return c
    nl=name.lower()
    for k,v in table.items():
        if k.lower() in nl or nl in k.lower(): return v
    return None

def engine_sort_monuments_by_proximity(monuments, hotel_name=None, hotel_coords=None, max_per_day=4):
    anchor=hotel_coords or _mon_get_coords(hotel_name or '',_HOTEL_COORDS)
    with_c,without_c=[],[]
    for name,bookings in monuments:
        c=_mon_get_coords(name,_MONUMENT_COORDS)
        if c and anchor:
            with_c.append((name,bookings,round(_mon_hav(anchor[0],anchor[1],c[0],c[1]),1)))
        else:
            without_c.append((name,bookings,None))
    with_c.sort(key=lambda x:x[2])
    return (with_c+without_c)[:max_per_day]

# ── HOTEL SUPPLIER CODE MAP ─────────────────────────────────────────────────
# productCode -> supplierRef — extracted from 722 production queries
_HOTEL_SUPPLIER_MAP = {
    'HO10000': {'supplierCode': 'DR10014', 'supplierName': 'Andaz', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 2},
    'HO10001': {'supplierCode': 'DR10055', 'supplierName': 'Jai Mahal Palace', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10002': {'supplierCode': 'DR10027', 'supplierName': 'Atulyaa Taj', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10003': {'supplierCode': 'DR10053', 'supplierName': '47 Jobner Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10004': {'supplierCode': 'DR10106', 'supplierName': 'Four Points By Sheraton', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 1},
    'HO10005': {'supplierCode': 'DR10043', 'supplierName': 'Taj', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 5},
    'HO10006': {'supplierCode': 'DR10133', 'supplierName': 'Amritara Suryauday Haveli', 'parentGroup': 'Amritara Hotels and Resorts', 'parentGroupId': '1feba35d-2056-4407-8960-2ae840ea9393', 'supplierVersion': 2},
    'HO10007': {'supplierCode': 'DR10099', 'supplierName': 'Ekaa Villa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10008': {'supplierCode': 'DR10037', 'supplierName': 'Jaypee Palace', 'parentGroup': 'Jaypee Hotels & Resorts', 'parentGroupId': '7a196fbe-dae5-4cd4-994e-19705b331b63', 'supplierVersion': 1},
    'HO10009': {'supplierCode': 'DR10078', 'supplierName': 'Samode Haveli', 'parentGroup': 'Samode Hotels', 'parentGroupId': 'a7099f41-1dca-4fd0-90a5-6acce48cbc42', 'supplierVersion': 2},
    'HO10010': {'supplierCode': 'DR10081', 'supplierName': 'Rambagh Palace', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10011': {'supplierCode': 'DR10047', 'supplierName': 'Tajview - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10012': {'supplierCode': 'DR10054', 'supplierName': 'The Claridges', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10013': {'supplierCode': 'DR10056', 'supplierName': 'The Imperial', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 4},
    'HO10014': {'supplierCode': 'DR10071', 'supplierName': 'The Suryaa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10016': {'supplierCode': 'DR10083', 'supplierName': 'The Amayaa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10017': {'supplierCode': 'DR10126', 'supplierName': 'Radisson Jass', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 2},
    'HO10018': {'supplierCode': 'DR10052', 'supplierName': 'Lemon Tree Premier (New Delhi)', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 3},
    'HO10019': {'supplierCode': 'DR10032', 'supplierName': 'Syna Heritage', 'parentGroup': 'Syna Hotels & Resorts', 'parentGroupId': '9c5c6a44-f546-459b-8214-f6feb0bfb065', 'supplierVersion': 2},
    'HO10020': {'supplierCode': 'DR10004', 'supplierName': 'Ganges View', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10022': {'supplierCode': 'DR10167', 'supplierName': 'DoubleTree by Hilton', 'parentGroup': 'Hilton Hotels & Resorts', 'parentGroupId': '064b6ff4-e74c-41c8-ab61-660f04ce70d8', 'supplierVersion': 1},
    'HO10023': {'supplierCode': 'DR10199', 'supplierName': 'The Coral Court Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10026': {'supplierCode': 'DR10008', 'supplierName': 'Raffles', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 4},
    'HO10027': {'supplierCode': 'DR10143', 'supplierName': 'Brunton Boatyard', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 3},
    'HO10028': {'supplierCode': 'DR10166', 'supplierName': 'Poppy\'s Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10030': {'supplierCode': 'DR10230', 'supplierName': 'Indus Biznotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10032': {'supplierCode': 'DR10098', 'supplierName': 'Abad Whispering Palms', 'parentGroup': 'Abad Hotels & Resorts', 'parentGroupId': '5fb27ab2-5c97-48be-a6a6-05251369e711', 'supplierVersion': 2},
    'HO10033': {'supplierCode': 'DR10002', 'supplierName': 'Rivatas By Ideal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10034': {'supplierCode': 'DR10107', 'supplierName': 'Rhythm', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10035': {'supplierCode': 'DR10009', 'supplierName': 'Starlit Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10036': {'supplierCode': 'DR10135', 'supplierName': 'Pushp Villa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10038': {'supplierCode': 'DR10215', 'supplierName': 'Tatsarasa Resorts And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10040': {'supplierCode': 'DR10202', 'supplierName': 'Brijrama Palace', 'parentGroup': 'Brij Hotels', 'parentGroupId': '4957b920-a191-46d8-9122-ede09995fb10', 'supplierVersion': 2},
    'HO10043': {'supplierCode': 'DR10061', 'supplierName': 'Amet Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10047': {'supplierCode': 'DR10088', 'supplierName': 'Chunda Palace', 'parentGroup': 'Chunda Hotels', 'parentGroupId': '030530e0-04e5-4e48-9374-9251ec7451ff', 'supplierVersion': 3},
    'HO10048': {'supplierCode': 'DR10122', 'supplierName': 'Jagat Niwas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10054': {'supplierCode': 'DR10650', 'supplierName': 'JP Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10061': {'supplierCode': 'DR10172', 'supplierName': 'Ideal Beach Resort', 'parentGroup': 'Ideal Resorts', 'parentGroupId': 'fd141810-9bab-4dd6-8217-8b34f8192ed8', 'supplierVersion': 2},
    'HO10062': {'supplierCode': 'DR10203', 'supplierName': 'Shiv Niwas Palace', 'parentGroup': 'HRH Group of Hotels', 'parentGroupId': '005fb938-3619-46ea-96e6-729415793ad7', 'supplierVersion': 1},
    'HO10063': {'supplierCode': 'DR10110', 'supplierName': 'Kumarakom Lake Resort', 'parentGroup': 'Paul John Resorts and Hotels', 'parentGroupId': 'f1306273-db36-40e2-8707-8762678bc493', 'supplierVersion': 2},
    'HO10064': {'supplierCode': 'DR10134', 'supplierName': 'Hotel Mahendra Prakash', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10065': {'supplierCode': 'DR10003', 'supplierName': 'Ambady Estate', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10067': {'supplierCode': 'DR10121', 'supplierName': 'The Raintree Hotel (St Marry)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10069': {'supplierCode': 'DR10028', 'supplierName': 'Ranjit\'s Svaasa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10077': {'supplierCode': 'DR10210', 'supplierName': 'Taj Fateh Prakash Palace', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10078': {'supplierCode': 'DR10209', 'supplierName': 'Courtyard By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10079': {'supplierCode': 'DR10089', 'supplierName': 'Clarks Shiraz', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10081': {'supplierCode': 'DR10279', 'supplierName': 'Holiday Inn Cochin', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 1},
    'HO10086': {'supplierCode': 'DR10304', 'supplierName': 'Mandir Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10088': {'supplierCode': 'DR10324', 'supplierName': 'Shiv Vilas Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10089': {'supplierCode': 'DR10301', 'supplierName': 'Rajmahal Palace RAAS', 'parentGroup': 'Raas Hotels', 'parentGroupId': '0d88ba09-39d1-49aa-a183-130375ead42f', 'supplierVersion': 2},
    'HO10092': {'supplierCode': 'DR10252', 'supplierName': 'Dera Rawatsar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10094': {'supplierCode': 'DR10306', 'supplierName': 'Marriott Resort & Spa', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10097': {'supplierCode': 'DR10239', 'supplierName': 'Alsisar Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10101': {'supplierCode': 'DR10384', 'supplierName': 'Radisson Blu Plaza Delhi Airport', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO10103': {'supplierCode': 'DR10247', 'supplierName': 'InterContinental', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 2},
    'HO10104': {'supplierCode': 'DR10313', 'supplierName': 'Ramada (Khajuraho)', 'parentGroup': 'Wyndham Hotels & Resorts', 'parentGroupId': '1861249a-90a7-4244-a0dd-9a1fa184ed25', 'supplierVersion': 1},
    'HO10105': {'supplierCode': 'DR10265', 'supplierName': 'Blanket Hotel And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10106': {'supplierCode': 'DR10413', 'supplierName': 'Evolve Back', 'parentGroup': 'Evolve Back', 'parentGroupId': 'a90a1b01-9d73-4c45-b1af-9c567ee14b68', 'supplierVersion': 3},
    'HO10107': {'supplierCode': 'DR10277', 'supplierName': 'Haze And Kites Resort', 'parentGroup': 'Kondody Hotels', 'parentGroupId': '32efa702-3d63-4160-b825-4839e180fe6d', 'supplierVersion': 3},
    'HO10108': {'supplierCode': 'DR10269', 'supplierName': 'Ibis', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 2},
    'HO10109': {'supplierCode': 'DR10208', 'supplierName': 'Taj Aravali Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10111': {'supplierCode': 'DR10296', 'supplierName': 'Radisson Blu Jaipur', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 2},
    'HO10113': {'supplierCode': 'DR10300', 'supplierName': 'Radisson Jaipur City Center', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 2},
    'HO10114': {'supplierCode': 'DR10322', 'supplierName': 'Sarovar Portico', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 2},
    'HO10115': {'supplierCode': 'DR10211', 'supplierName': 'Trident', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 5},
    'HO10117': {'supplierCode': 'DR10268', 'supplierName': 'Holiday Inn Jaipur City Centre', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 1},
    'HO10120': {'supplierCode': 'DR10289', 'supplierName': 'Windermere Estate', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10121': {'supplierCode': 'DR10494', 'supplierName': 'ITC Windsor', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 4},
    'HO10124': {'supplierCode': 'DR10467', 'supplierName': 'Vivanta', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10125': {'supplierCode': 'DR10218', 'supplierName': 'The Fern Residency', 'parentGroup': 'Fern Hotels & Resorts', 'parentGroupId': '6be198c4-bfec-4170-bd7e-066cc7a5a9ef', 'supplierVersion': 1},
    'HO10126': {'supplierCode': 'DR10229', 'supplierName': 'Trident', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 3},
    'HO10127': {'supplierCode': 'DR10480', 'supplierName': 'Mayfair', 'parentGroup': 'Mayfair Hotels & Resorts', 'parentGroupId': '50b199a5-0468-4963-9fe6-7c2dc1a73d9d', 'supplierVersion': 2},
    'HO10130': {'supplierCode': 'DR10168', 'supplierName': 'Surya Kaiser Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10132': {'supplierCode': 'DR10493', 'supplierName': 'Fortune Park JP Celestial', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 1},
    'HO10134': {'supplierCode': 'DR10476', 'supplierName': 'Jims Jungle Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 4},
    'HO10135': {'supplierCode': 'DR10498', 'supplierName': 'Alila Diwa', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 2},
    'HO10136': {'supplierCode': 'DR10417', 'supplierName': 'Four Seasons', 'parentGroup': 'Four Seasons Hotels', 'parentGroupId': 'dbee834c-2903-4c1e-9218-5fb340ada0eb', 'supplierVersion': 4},
    'HO10139': {'supplierCode': 'DR10466', 'supplierName': 'Suba Palace', 'parentGroup': 'Choice Hotels', 'parentGroupId': '9dd46398-05e7-4eff-b89f-3edb669a4863', 'supplierVersion': 2},
    'HO10140': {'supplierCode': 'DR10468', 'supplierName': 'Sun N Sand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10142': {'supplierCode': 'DR10503', 'supplierName': 'Bloom Suites Goa', 'parentGroup': 'Bloom Hotel Group', 'parentGroupId': 'f23993ba-6c9d-4eb2-b36d-fd6abcd41ec7', 'supplierVersion': 1},
    'HO10143': {'supplierCode': 'DR10456', 'supplierName': 'The Sahil', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10144': {'supplierCode': 'DR10482', 'supplierName': 'Fariyas Hotels private limited', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10145': {'supplierCode': 'DR10086', 'supplierName': 'Taj Ganges', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10149': {'supplierCode': 'DR10653', 'supplierName': 'The Francis Residence', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10151': {'supplierCode': 'DR10011', 'supplierName': 'Fragrant Nature', 'parentGroup': 'Fragrant Nature Hotels and Resorts', 'parentGroupId': 'aa97c008-7a33-4c18-83dc-6dd821b009da', 'supplierVersion': 1},
    'HO10154': {'supplierCode': 'DR10654', 'supplierName': 'The Elgin (Darjeeling)', 'parentGroup': 'Elgin Hotels & Resorts', 'parentGroupId': '6b9a39f8-a5c5-4321-90f0-354cc2856e06', 'supplierVersion': 2},
    'HO10155': {'supplierCode': 'DR10128', 'supplierName': 'The Leela Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10158': {'supplierCode': 'DR10656', 'supplierName': 'Taj Cidade De Goa Heritage', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10161': {'supplierCode': 'DR10657', 'supplierName': 'Taj Fort Aguada Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10163': {'supplierCode': 'DR10658', 'supplierName': 'Taj Holiday Village Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10167': {'supplierCode': 'DR10342', 'supplierName': 'Taj Hari Mahal', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10168': {'supplierCode': 'DR10270', 'supplierName': 'Fragrant Nature', 'parentGroup': 'Fragrant Nature Hotels and Resorts', 'parentGroupId': 'aa97c008-7a33-4c18-83dc-6dd821b009da', 'supplierVersion': 2},
    'HO10169': {'supplierCode': 'DR10551', 'supplierName': 'Evolve Back', 'parentGroup': 'Evolve Back', 'parentGroupId': 'a90a1b01-9d73-4c45-b1af-9c567ee14b68', 'supplierVersion': 3},
    'HO10172': {'supplierCode': 'DR10668', 'supplierName': 'Spice Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10178': {'supplierCode': 'DR10404', 'supplierName': 'Samode Bagh', 'parentGroup': 'Samode Hotels', 'parentGroupId': 'a7099f41-1dca-4fd0-90a5-6acce48cbc42', 'supplierVersion': 2},
    'HO10179': {'supplierCode': 'DR10672', 'supplierName': 'Samode Palace', 'parentGroup': 'Samode Hotels', 'parentGroupId': 'a7099f41-1dca-4fd0-90a5-6acce48cbc42', 'supplierVersion': 2},
    'HO10181': {'supplierCode': 'DR10421', 'supplierName': 'Indeco (Swamimalai)', 'parentGroup': 'Indeco Hotels', 'parentGroupId': '7087fd2e-2e68-40a2-80e5-35a8a3f512a8', 'supplierVersion': 2},
    'HO10182': {'supplierCode': 'DR10667', 'supplierName': 'Xandari Pearl', 'parentGroup': 'Xandari Resorts', 'parentGroupId': '31e6528a-acf7-41e8-a5fe-48ec1150960e', 'supplierVersion': 2},
    'HO10184': {'supplierCode': 'DR10454', 'supplierName': 'Tilar Siro', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10185': {'supplierCode': 'DR10674', 'supplierName': 'Cardamom County', 'parentGroup': 'Xandari Resorts', 'parentGroupId': '31e6528a-acf7-41e8-a5fe-48ec1150960e', 'supplierVersion': 1},
    'HO10186': {'supplierCode': 'DR10420', 'supplierName': 'ITC Grand Central', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 3},
    'HO10190': {'supplierCode': 'DR10220', 'supplierName': 'The Leela Palace (Udaipur)', 'parentGroup': 'HLV', 'parentGroupId': 'c8bf0041-cad6-401c-8c7a-ad9dae79978b', 'supplierVersion': 1},
    'HO10191': {'supplierCode': 'DR10393', 'supplierName': 'Radisson Varanasi', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO10192': {'supplierCode': 'DR10646', 'supplierName': 'JW Marriott (Juhu)', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10193': {'supplierCode': 'DR10376', 'supplierName': 'Taj Palace', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10194': {'supplierCode': 'DR10378', 'supplierName': 'Eighth Bastion', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 2},
    'HO10195': {'supplierCode': 'DR10128', 'supplierName': 'The Leela Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10196': {'supplierCode': 'DR10581', 'supplierName': 'Fort Barli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10197': {'supplierCode': 'DR10251', 'supplierName': 'Dera Mandawa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10199': {'supplierCode': 'DR10271', 'supplierName': 'Ikaki Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10204': {'supplierCode': 'DR10690', 'supplierName': 'Ameya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10207': {'supplierCode': 'DR10699', 'supplierName': 'Beleza By The Beach', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10208': {'supplierCode': 'DR10711', 'supplierName': 'Heritage Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10209': {'supplierCode': 'DR10714', 'supplierName': 'Hoysala Village Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10210': {'supplierCode': 'DR10715', 'supplierName': 'Denissons', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10212': {'supplierCode': 'DR10429', 'supplierName': 'Travancore Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10213': {'supplierCode': 'DR10307', 'supplierName': 'Hotel Rawal Kot', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10214': {'supplierCode': 'DR10725', 'supplierName': 'Lakesong Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10215': {'supplierCode': 'DR10398', 'supplierName': 'Abode, Mumbai', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10216': {'supplierCode': 'DR10285', 'supplierName': 'Marigold', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10222': {'supplierCode': 'DR10538', 'supplierName': 'The St. Regis', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10223': {'supplierCode': 'DR10669', 'supplierName': 'Tall Trees', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10226': {'supplierCode': 'DR10013', 'supplierName': 'Grand Hotel D\' Europe', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10228': {'supplierCode': 'DR10079', 'supplierName': 'Regenta Central RS', 'parentGroup': 'Royal Orchid Hotels', 'parentGroupId': '025ed948-0a29-4839-9a2a-daaa298dbfad', 'supplierVersion': 3},
    'HO10229': {'supplierCode': 'DR10557', 'supplierName': 'Royal Orchid Metropole', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10231': {'supplierCode': 'DR10539', 'supplierName': 'Taj West End', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10232': {'supplierCode': 'DR10029', 'supplierName': 'Sandesh The Prince', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10234': {'supplierCode': 'DR10464', 'supplierName': 'The Residency Towers', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10235': {'supplierCode': 'DR10713', 'supplierName': 'Gateway Coonoor - IHCL Seleqtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10236': {'supplierCode': 'DR10697', 'supplierName': 'Taj Mahal', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10237': {'supplierCode': 'DR10564', 'supplierName': 'Banjaar Tola', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10238': {'supplierCode': 'DR10722', 'supplierName': 'Taj Green Cove Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10239': {'supplierCode': 'DR10683', 'supplierName': 'The Oberoi', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 3},
    'HO10240': {'supplierCode': 'DR10500', 'supplierName': 'The Oberoi', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 4},
    'HO10243': {'supplierCode': 'DR10693', 'supplierName': 'Evolve Back', 'parentGroup': 'Evolve Back', 'parentGroupId': 'a90a1b01-9d73-4c45-b1af-9c567ee14b68', 'supplierVersion': 5},
    'HO10244': {'supplierCode': 'DR10532', 'supplierName': 'Samode Safari Lodge', 'parentGroup': 'Samode Hotels', 'parentGroupId': 'a7099f41-1dca-4fd0-90a5-6acce48cbc42', 'supplierVersion': 2},
    'HO10246': {'supplierCode': 'DR10682', 'supplierName': 'Radisson Bengaluru City Center', 'parentGroup': 'GRT Hotels and Resorts', 'parentGroupId': '5dab556b-52dc-4962-8a00-6db6da461b7b', 'supplierVersion': 1},
    'HO10248': {'supplierCode': 'DR10179', 'supplierName': 'Radisson Resort Temple Bay', 'parentGroup': 'GRT Hotels and Resorts', 'parentGroupId': '5dab556b-52dc-4962-8a00-6db6da461b7b', 'supplierVersion': 1},
    'HO10253': {'supplierCode': 'DR10506', 'supplierName': 'JW Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 5},
    'HO10254': {'supplierCode': 'DR10041', 'supplierName': 'ITC Mughal', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 4},
    'HO10257': {'supplierCode': 'DR10354', 'supplierName': 'The Lodhi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10258': {'supplierCode': 'DR10737', 'supplierName': 'The Judge\'s Court', 'parentGroup': 'Welcomheritage Hotels', 'parentGroupId': '1ae3034f-d63b-46c8-82bd-1daefd69e317', 'supplierVersion': 3},
    'HO10259': {'supplierCode': 'DR10705', 'supplierName': 'The Zuri White Sands Goa Resort And Casino', 'parentGroup': 'Zuri Hotels & Resorts', 'parentGroupId': '6fba3d65-e81e-42fa-966c-a8e0da939bf5', 'supplierVersion': 1},
    'HO10262': {'supplierCode': 'DR10260', 'supplierName': 'Golden Tulip, Jaipur', 'parentGroup': 'Golden Tulip', 'parentGroupId': '63717936-1dba-4af1-b205-3314ad2b9be4', 'supplierVersion': 1},
    'HO10264': {'supplierCode': 'DR10338', 'supplierName': 'RAAS', 'parentGroup': 'Raas Hotels', 'parentGroupId': '0d88ba09-39d1-49aa-a183-130375ead42f', 'supplierVersion': 3},
    'HO10265': {'supplierCode': 'DR10676', 'supplierName': 'Amritara Shalimar Spice Garden', 'parentGroup': 'Amritara Hotels and Resorts', 'parentGroupId': '1feba35d-2056-4407-8960-2ae840ea9393', 'supplierVersion': 2},
    'HO10266': {'supplierCode': 'DR10529', 'supplierName': 'Welcomhotel By ITC Hotels, Rama International', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 4},
    'HO10269': {'supplierCode': 'DR10224', 'supplierName': 'Bel Morris', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10270': {'supplierCode': 'DR10518', 'supplierName': 'Amanbagh', 'parentGroup': 'Aman Hotels & Resorts', 'parentGroupId': 'f797cee6-8d10-4618-9d44-47e52a60d388', 'supplierVersion': 2},
    'HO10273': {'supplierCode': 'DR10287', 'supplierName': 'Narain Niwas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10275': {'supplierCode': 'DR10395', 'supplierName': 'The Malabar House', 'parentGroup': 'Malabar Escapes', 'parentGroupId': '91605486-ed4f-4620-af00-e3a00cd9b730', 'supplierVersion': 3},
    'HO10276': {'supplierCode': 'DR10000', 'supplierName': 'Libra Lords Inn', 'parentGroup': 'Lords Inn Hotels & Resorts', 'parentGroupId': '305c5d02-6e37-4992-afe5-399c4e0a43fe', 'supplierVersion': 3},
    'HO10277': {'supplierCode': 'DR10660', 'supplierName': 'Mandawa Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10278': {'supplierCode': 'DR10348', 'supplierName': 'The Oberoi', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 4},
    'HO10279': {'supplierCode': 'DR10544', 'supplierName': 'Dev Shree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 4},
    'HO10282': {'supplierCode': 'DR10827', 'supplierName': 'Raj Haveli Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10283': {'supplierCode': 'DR10678', 'supplierName': 'Hyatt Regency (Amritsar)', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 1},
    'HO10284': {'supplierCode': 'DR10648', 'supplierName': 'Radisson (Agra)', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO10285': {'supplierCode': 'DR10825', 'supplierName': 'Chirag', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10286': {'supplierCode': 'DR10372', 'supplierName': 'Pride Plaza', 'parentGroup': 'Pride Group of Hotels', 'parentGroupId': '61b141a4-32e5-4dc3-afcb-e3ed29489e73', 'supplierVersion': 2},
    'HO10288': {'supplierCode': 'DR10317', 'supplierName': 'Amar Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10289': {'supplierCode': 'DR10691', 'supplierName': 'Abad Plaza', 'parentGroup': 'Abad Hotels & Resorts', 'parentGroupId': '5fb27ab2-5c97-48be-a6a6-05251369e711', 'supplierVersion': 2},
    'HO10290': {'supplierCode': 'DR10346', 'supplierName': 'Vivanta', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10292': {'supplierCode': 'DR10800', 'supplierName': 'The Leela Ambience Gurugram Hotel & Residences', 'parentGroup': 'HLV', 'parentGroupId': 'c8bf0041-cad6-401c-8c7a-ad9dae79978b', 'supplierVersion': 2},
    'HO10293': {'supplierCode': 'DR10021', 'supplierName': 'Park Ocean', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10295': {'supplierCode': 'DR10389', 'supplierName': 'Roseate House', 'parentGroup': 'Roseate Hotels and Resorts', 'parentGroupId': '971316d9-03dd-4a2f-a02a-3c6b9ed36b5f', 'supplierVersion': 2},
    'HO10298': {'supplierCode': 'DR10826', 'supplierName': 'Lallgarh Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10300': {'supplierCode': 'DR10340', 'supplierName': 'Ratan Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10302': {'supplierCode': 'DR10355', 'supplierName': 'The Leela Palace (New Delhi)', 'parentGroup': 'HLV', 'parentGroupId': 'c8bf0041-cad6-401c-8c7a-ad9dae79978b', 'supplierVersion': 2},
    'HO10303': {'supplierCode': 'DR10811', 'supplierName': 'Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10304': {'supplierCode': 'DR10370', 'supplierName': 'Novotel', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 3},
    'HO10305': {'supplierCode': 'DR10780', 'supplierName': 'Welcomhotel By ITC Hotels (Dwarka)', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 4},
    'HO10306': {'supplierCode': 'DR10380', 'supplierName': 'Pullman', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 3},
    'HO10308': {'supplierCode': 'DR10357', 'supplierName': 'The Leela Ambience Convention Hotel', 'parentGroup': 'HLV', 'parentGroupId': 'c8bf0041-cad6-401c-8c7a-ad9dae79978b', 'supplierVersion': 4},
    'HO10309': {'supplierCode': 'DR10024', 'supplierName': 'Nachana Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10311': {'supplierCode': 'DR10227', 'supplierName': 'Ibis', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 2},
    'HO10312': {'supplierCode': 'DR10341', 'supplierName': 'The Rohet House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10313': {'supplierCode': 'DR10096', 'supplierName': 'The Oberoi Rajvilas', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 3},
    'HO10336': {'supplierCode': 'DR10610', 'supplierName': 'The Oberoi Cecil', 'parentGroup': 'Oberoi Hotels & Resorts', 'parentGroupId': None, 'supplierVersion': 1},
    'HO10355': {'supplierCode': 'DR10786', 'supplierName': 'Ramada', 'parentGroup': 'Wyndham Hotels & Resorts', 'parentGroupId': '1861249a-90a7-4244-a0dd-9a1fa184ed25', 'supplierVersion': 1},
    'HO10315': {'supplierCode': 'DR10747', 'supplierName': 'Classic Sarovar Portico', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 2},
    'HO10316': {'supplierCode': 'DR10364', 'supplierName': 'ITC Maurya', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 4},
    'HO10321': {'supplierCode': 'DR10084', 'supplierName': 'Ambassador - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10323': {'supplierCode': 'DR10142', 'supplierName': 'Trident', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 3},
    'HO10324': {'supplierCode': 'DR10801', 'supplierName': 'The Oberoi', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 4},
    'HO10325': {'supplierCode': 'DR10326', 'supplierName': 'The Leela Palace (Jaipur)', 'parentGroup': 'HLV', 'parentGroupId': 'c8bf0041-cad6-401c-8c7a-ad9dae79978b', 'supplierVersion': 1},
    'HO10326': {'supplierCode': 'DR10806', 'supplierName': 'Trident', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 4},
    'HO10327': {'supplierCode': 'DR10375', 'supplierName': 'Taj Malabar Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 5},
    'HO10328': {'supplierCode': 'DR10381', 'supplierName': 'Radisson Blu (Dwarka)', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 3},
    'HO10329': {'supplierCode': 'DR10241', 'supplierName': 'Lakes And Lagoon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10333': {'supplierCode': 'DR10205', 'supplierName': 'Keys By Lemon Tree Hotels, Kochi', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO10334': {'supplierCode': 'DR10533', 'supplierName': 'Deventure Sarovar Portico (Kapashera)', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 4},
    'HO10343': {'supplierCode': 'DR10639', 'supplierName': 'Marari Beach Resort', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 4},
    'HO10346': {'supplierCode': 'DR10710', 'supplierName': 'Taj City Centre', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10347': {'supplierCode': 'DR10891', 'supplierName': 'Residency Towers (Chennai)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10352': {'supplierCode': 'DR10770', 'supplierName': 'Purity', 'parentGroup': 'Malabar Escapes', 'parentGroupId': '91605486-ed4f-4620-af00-e3a00cd9b730', 'supplierVersion': 2},
    'HO10354': {'supplierCode': 'DR10748', 'supplierName': 'Khem Villas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10357': {'supplierCode': 'DR10892', 'supplierName': 'Taj Club House', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10359': {'supplierCode': 'DR10880', 'supplierName': 'Lemon Tree Premier (Mumbai)', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO10363': {'supplierCode': 'DR10894', 'supplierName': 'Taj Connemara', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10365': {'supplierCode': 'DR10721', 'supplierName': 'The Bangala', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10366': {'supplierCode': 'DR10444', 'supplierName': 'President - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10369': {'supplierCode': 'DR10258', 'supplierName': 'Diggi Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10370': {'supplierCode': 'DR10396', 'supplierName': 'The Oberoi Vanyavilas', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 5},
    'HO10371': {'supplierCode': 'DR10621', 'supplierName': 'Clarks (Khajuraho)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10372': {'supplierCode': 'DR10896', 'supplierName': 'Taj Coromandel', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10376': {'supplierCode': 'DR10755', 'supplierName': 'Taj Lands End', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10378': {'supplierCode': 'DR10495', 'supplierName': 'Lemon Tree Premier (Bengaluru)', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO10379': {'supplierCode': 'DR10469', 'supplierName': 'Taj (Santacruz)', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10383': {'supplierCode': 'DR10724', 'supplierName': 'Coconut Lagoon', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 1},
    'HO10387': {'supplierCode': 'DR10491', 'supplierName': 'The Leela (Mumbai)', 'parentGroup': 'HLV', 'parentGroupId': 'c8bf0041-cad6-401c-8c7a-ad9dae79978b', 'supplierVersion': 2},
    'HO10390': {'supplierCode': 'DR10323', 'supplierName': 'Shahpura House', 'parentGroup': 'Shahpura Hotels', 'parentGroupId': '0211fc75-9d49-4693-b223-e7fc6beafaa6', 'supplierVersion': 5},
    'HO10391': {'supplierCode': 'DR10158', 'supplierName': 'Courtyard By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10393': {'supplierCode': 'DR10112', 'supplierName': 'The Zuri Kumarakom Kerala Resort And Spa', 'parentGroup': 'Zuri Hotels & Resorts', 'parentGroupId': '6fba3d65-e81e-42fa-966c-a8e0da939bf5', 'supplierVersion': 2},
    'HO10397': {'supplierCode': 'DR10292', 'supplierName': 'Spice Village', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 3},
    'HO10399': {'supplierCode': 'DR10315', 'supplierName': 'Jaisalkot', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10400': {'supplierCode': 'DR10017', 'supplierName': 'Heritage Madurai', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10401': {'supplierCode': 'DR10898', 'supplierName': 'Abad Atrium', 'parentGroup': 'Abad Hotels & Resorts', 'parentGroupId': '5fb27ab2-5c97-48be-a6a6-05251369e711', 'supplierVersion': 2},
    'HO10402': {'supplierCode': 'DR10834', 'supplierName': 'Suryagarh', 'parentGroup': 'Suryagarh Collection', 'parentGroupId': 'ecbe107c-145b-4c83-b2f5-8fd2ae84c1e2', 'supplierVersion': 4},
    'HO10403': {'supplierCode': 'DR10505', 'supplierName': 'Taj', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10404': {'supplierCode': 'DR10768', 'supplierName': 'Efkay\'s Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10405': {'supplierCode': 'DR10369', 'supplierName': 'Le Meridien', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 2},
    'HO10406': {'supplierCode': 'DR10255', 'supplierName': 'Forte Kochi', 'parentGroup': 'Paul John Resorts and Hotels', 'parentGroupId': 'f1306273-db36-40e2-8707-8762678bc493', 'supplierVersion': 1},
    'HO10407': {'supplierCode': 'DR10514', 'supplierName': 'Taj (Yeshwantpur)', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10409': {'supplierCode': 'DR10727', 'supplierName': 'Gateway Madurai', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10411': {'supplierCode': 'DR10556', 'supplierName': 'Radisson Blu Plaza (Mysore)', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO10412': {'supplierCode': 'DR10854', 'supplierName': 'Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10414': {'supplierCode': 'DR10129', 'supplierName': 'Hotel Lakend', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10416': {'supplierCode': 'DR10875', 'supplierName': 'Taj Fisherman\'s Cove Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10417': {'supplierCode': 'DR10308', 'supplierName': 'Old Harbour', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10422': {'supplierCode': 'DR10733', 'supplierName': 'Savoy - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10425': {'supplierCode': 'DR10350', 'supplierName': 'Castle Mandawa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10426': {'supplierCode': 'DR10377', 'supplierName': 'The Dunes (Cochin)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10427': {'supplierCode': 'DR10888', 'supplierName': 'Anandha Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10431': {'supplierCode': 'DR10882', 'supplierName': 'Vividus', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10432': {'supplierCode': 'DR10181', 'supplierName': 'Atithi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10436': {'supplierCode': 'DR10272', 'supplierName': 'Narendra Bhawan', 'parentGroup': 'Suryagarh Collection', 'parentGroupId': 'ecbe107c-145b-4c83-b2f5-8fd2ae84c1e2', 'supplierVersion': 3},
    'HO10437': {'supplierCode': 'DR10213', 'supplierName': 'Taj Lake Palace', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10442': {'supplierCode': 'DR10775', 'supplierName': 'Sarovar Premiere', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 2},
    'HO10445': {'supplierCode': 'DR10390', 'supplierName': 'Sheraton', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 3},
    'HO10446': {'supplierCode': 'DR10782', 'supplierName': 'Maison Perumal', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 1},
    'HO10447': {'supplierCode': 'DR10893', 'supplierName': 'Colonel\'s Retreat (Defence Colony)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10448': {'supplierCode': 'DR10223', 'supplierName': 'The Oberoi Udaivilas', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 5},
    'HO10449': {'supplierCode': 'DR10178', 'supplierName': 'Lemon Tree Hotel- Sohna Road', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO10450': {'supplierCode': 'DR10144', 'supplierName': 'Lemon Tree Premier (Jaipur)', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO10456': {'supplierCode': 'DR10591', 'supplierName': 'Radisson Blu Resort Dharamshala', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO10457': {'supplierCode': 'DR10918', 'supplierName': 'Vivanta (Residency Road)', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10459': {'supplierCode': 'DR10116', 'supplierName': 'Crowne Plaza (New Delhi, Okhla)', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 1},
    'HO10460': {'supplierCode': 'DR10829', 'supplierName': 'Umaid Bhawan Palace', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10461': {'supplierCode': 'DR10945', 'supplierName': 'Mamalla Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10463': {'supplierCode': 'DR10942', 'supplierName': 'Villa Shanti', 'parentGroup': 'Rare India - Marketing Office (not to be used)', 'parentGroupId': 'd6a4f6e9-a08b-4de1-a32f-b5605811dec2', 'supplierVersion': 1},
    'HO10464': {'supplierCode': 'DR10545', 'supplierName': 'Trident (Nariman Point)', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 5},
    'HO10468': {'supplierCode': 'DR10177', 'supplierName': 'The Oberoi Amarvilas', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 4},
    'HO10469': {'supplierCode': 'DR10850', 'supplierName': 'Nahargarh', 'parentGroup': 'Alsisar Group of Hotels', 'parentGroupId': 'df3f1747-de95-4637-89d9-941da96af384', 'supplierVersion': 1},
    'HO10474': {'supplierCode': 'DR10913', 'supplierName': 'Abad Copper Castle', 'parentGroup': 'Abad Hotels & Resorts', 'parentGroupId': '5fb27ab2-5c97-48be-a6a6-05251369e711', 'supplierVersion': 2},
    'HO10476': {'supplierCode': 'DR10938', 'supplierName': 'The Residency Towers (Puducherry)', 'parentGroup': 'The Residency Group of Hotels', 'parentGroupId': 'a50d3204-182f-4c20-9a02-e1166b93f4d0', 'supplierVersion': 1},
    'HO10477': {'supplierCode': 'DR10964', 'supplierName': 'The Connaught - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10480': {'supplierCode': 'DR10939', 'supplierName': 'The Westin Pushkar Resort & Spa', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10481': {'supplierCode': 'DR10640', 'supplierName': 'The Oberoi Wildflower Hall', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 6},
    'HO10485': {'supplierCode': 'DR10195', 'supplierName': 'Svatma', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 4},
    'HO10486': {'supplierCode': 'DR10781', 'supplierName': 'Deogarh Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10487': {'supplierCode': 'DR10952', 'supplierName': 'Heritage Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10489': {'supplierCode': 'DR10329', 'supplierName': 'Ajit Bhawan', 'parentGroup': 'Ajit Group Of Hotels', 'parentGroupId': '056224b8-4b97-4170-ac08-cf200b4c17b8', 'supplierVersion': 3},
    'HO10490': {'supplierCode': 'DR10403', 'supplierName': 'Rohet Garh', 'parentGroup': 'House Of Rohet', 'parentGroupId': '79e67a64-7196-4974-893d-f3fe76c4b9b1', 'supplierVersion': 2},
    'HO10492': {'supplierCode': 'DR10400', 'supplierName': 'The Tigress Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10493': {'supplierCode': 'DR10561', 'supplierName': 'Rawla Narlai', 'parentGroup': 'Ajit Group Of Hotels', 'parentGroupId': '056224b8-4b97-4170-ac08-cf200b4c17b8', 'supplierVersion': 4},
    'HO10495': {'supplierCode': 'DR10950', 'supplierName': 'Dev Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10497': {'supplierCode': 'DR10787', 'supplierName': 'Jas Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10499': {'supplierCode': 'DR10851', 'supplierName': 'Ranthambore Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10500': {'supplierCode': 'DR10934', 'supplierName': 'Eastend', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10504': {'supplierCode': 'DR10955', 'supplierName': 'Dera Amer', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10507': {'supplierCode': 'DR10267', 'supplierName': 'Hilton (Jaipur)', 'parentGroup': 'Hilton Hotels & Resorts', 'parentGroupId': '064b6ff4-e74c-41c8-ab61-660f04ce70d8', 'supplierVersion': 1},
    'HO10508': {'supplierCode': 'DR10540', 'supplierName': 'Taj Krishna', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10511': {'supplierCode': 'DR10718', 'supplierName': 'Zone By The Park (Jaipur)', 'parentGroup': 'Apeejay Surrendra Park Hotels', 'parentGroupId': '1f4a03c7-d445-4cd2-816b-5d060c85220c', 'supplierVersion': 1},
    'HO10512': {'supplierCode': 'DR10739', 'supplierName': 'Fateh Bagh', 'parentGroup': 'HRH Group of Hotels', 'parentGroupId': '005fb938-3619-46ea-96e6-729415793ad7', 'supplierVersion': 1},
    'HO10513': {'supplierCode': 'DR10537', 'supplierName': 'Royal Orchid Central Kireeti', 'parentGroup': 'Royal Orchid Hotels', 'parentGroupId': '025ed948-0a29-4839-9a2a-daaa298dbfad', 'supplierVersion': 2},
    'HO10514': {'supplierCode': 'DR10527', 'supplierName': 'Lemon Tree, Aurangabad', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO10521': {'supplierCode': 'DR10965', 'supplierName': 'Taj Kumarakom Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10529': {'supplierCode': 'DR10944', 'supplierName': 'Brahma Horizon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10531': {'supplierCode': 'DR10542', 'supplierName': 'Taj Mahal Palace & Tower', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 5},
    'HO10537': {'supplierCode': 'DR10609', 'supplierName': 'Aloha On The Ganges Rishikesh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10543': {'supplierCode': 'DR10299', 'supplierName': 'Deo Bagh', 'parentGroup': 'Neemrana Hotels', 'parentGroupId': '20cca852-4d76-47ee-a2e7-ce543bf12867', 'supplierVersion': 1},
    'HO10544': {'supplierCode': 'DR10612', 'supplierName': 'Tadoba Jungle Camp', 'parentGroup': 'Jungle Camps India', 'parentGroupId': 'fcb5e0cb-47ca-42ba-a4d9-9a5ad75c4e50', 'supplierVersion': 3},
    'HO10547': {'supplierCode': 'DR11025', 'supplierName': 'Maidens', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 5},
    'HO10548': {'supplierCode': 'DR11035', 'supplierName': 'Lake Canopy', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10549': {'supplierCode': 'DR11013', 'supplierName': 'Chanoud Garh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10551': {'supplierCode': 'DR10361', 'supplierName': 'ITC Grand Chola', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 3},
    'HO10554': {'supplierCode': 'DR10237', 'supplierName': 'Alila Fort Bishangarh', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 2},
    'HO10555': {'supplierCode': 'DR10184', 'supplierName': 'ITC Rajputana', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 4},
    'HO10558': {'supplierCode': 'DR11005', 'supplierName': 'Uday Backwater Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10560': {'supplierCode': 'DR11012', 'supplierName': 'Courtyard By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10561': {'supplierCode': 'DR11016', 'supplierName': 'Royal Heritage Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10565': {'supplierCode': 'DR10992', 'supplierName': 'Taj (MG Road)', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10569': {'supplierCode': 'DR11019', 'supplierName': 'Six Senses Fort Barwara', 'parentGroup': 'Six Senses Hotels Resorts Spas', 'parentGroupId': '5e56de43-2c67-4cb9-98ac-ef7a4fdf8dcb', 'supplierVersion': 3},
    'HO10574': {'supplierCode': 'DR11008', 'supplierName': 'Peppervine', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10575': {'supplierCode': 'DR10719', 'supplierName': 'Fort Rajwada', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10578': {'supplierCode': 'DR10966', 'supplierName': 'Sujan The Serai', 'parentGroup': 'Sujan Luxury Hotels', 'parentGroupId': '4d2b4d25-57a4-412c-91f8-b94c10d07b82', 'supplierVersion': 1},
    'HO10579': {'supplierCode': 'DR10843', 'supplierName': 'The Mirador', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10583': {'supplierCode': 'DR11021', 'supplierName': 'Vythiri Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10585': {'supplierCode': 'DR10794', 'supplierName': 'Quality Inn (Gurgaon)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10586': {'supplierCode': 'DR11004', 'supplierName': 'Welcomhotel By ITC Hotels', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 4},
    'HO10591': {'supplierCode': 'DR10946', 'supplierName': 'Casino', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 1},
    'HO10597': {'supplierCode': 'DR10756', 'supplierName': 'Blue Jelly Cruises', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10600': {'supplierCode': 'DR10006', 'supplierName': 'White Water', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10622': {'supplierCode': 'DR10408', 'supplierName': 'Taj Bengal', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10623': {'supplierCode': 'DR11052', 'supplierName': 'Taj Exotica Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10625': {'supplierCode': 'DR10457', 'supplierName': 'Pratap Mahal - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10629': {'supplierCode': 'DR10528', 'supplierName': 'Gateway', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10630': {'supplierCode': 'DR11032', 'supplierName': 'Pal Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10631': {'supplierCode': 'DR11003', 'supplierName': 'Great Trails River view Resort', 'parentGroup': 'GRT Hotels and Resorts', 'parentGroupId': '5dab556b-52dc-4962-8a00-6db6da461b7b', 'supplierVersion': 1},
    'HO10635': {'supplierCode': 'DR10085', 'supplierName': 'Amar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10637': {'supplierCode': 'DR10425', 'supplierName': 'Dewalokam', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10638': {'supplierCode': 'DR10522', 'supplierName': 'Radisson Blu (Amritsar)', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO10640': {'supplierCode': 'DR10511', 'supplierName': 'ITC Grand Resort & Spa', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 3},
    'HO10644': {'supplierCode': 'DR10535', 'supplierName': 'Taj Falaknuma Palace', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10646': {'supplierCode': 'DR11048', 'supplierName': 'Kings Lodge', 'parentGroup': 'Pugdundee Safaris', 'parentGroupId': '3587777b-fb9e-455d-8e0a-4cf718e2d9c5', 'supplierVersion': 2},
    'HO10648': {'supplierCode': 'DR10993', 'supplierName': 'Ramgarh Lodge-IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10650': {'supplierCode': 'DR10250', 'supplierName': 'Crowne Plaza Kochi', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 1},
    'HO10653': {'supplierCode': 'DR10333', 'supplierName': 'Fairfield By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10655': {'supplierCode': 'DR11046', 'supplierName': 'Sawai Man Mahal', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10656': {'supplierCode': 'DR11034', 'supplierName': 'Brij Jawai', 'parentGroup': 'Brij Hotels', 'parentGroupId': '4957b920-a191-46d8-9122-ede09995fb10', 'supplierVersion': 3},
    'HO10663': {'supplierCode': 'DR11050', 'supplierName': 'Jehan Numa Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10667': {'supplierCode': 'DR10497', 'supplierName': 'Haveli Dharampura', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10668': {'supplierCode': 'DR10582', 'supplierName': 'Taj Bekal Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10675': {'supplierCode': 'DR11081', 'supplierName': 'Taj', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10679': {'supplierCode': 'DR11075', 'supplierName': 'Amritara Ambatty Greens', 'parentGroup': 'Amritara Hotels and Resorts', 'parentGroupId': '1feba35d-2056-4407-8960-2ae840ea9393', 'supplierVersion': 3},
    'HO10682': {'supplierCode': 'DR11087', 'supplierName': 'Vivanta (Panaji)', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10684': {'supplierCode': 'DR11083', 'supplierName': 'Taj Madikeri Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO10687': {'supplierCode': 'DR11092', 'supplierName': 'Taj Rishikesh Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10689': {'supplierCode': 'DR10865', 'supplierName': 'Fortune JP Palace', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 1},
    'HO10698': {'supplierCode': 'DR11089', 'supplierName': 'Baghvan', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10700': {'supplierCode': 'DR11095', 'supplierName': 'Vivanta', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10701': {'supplierCode': 'DR10629', 'supplierName': 'Taj Mahal', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10702': {'supplierCode': 'DR11060', 'supplierName': 'Fortune Sullivan Court', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 1},
    'HO10707': {'supplierCode': 'DR10327', 'supplierName': 'The Raj Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10713': {'supplierCode': 'DR11134', 'supplierName': 'Taj Skyline', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10716': {'supplierCode': 'DR11133', 'supplierName': 'Taj Cidade De Goa Horizon', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10717': {'supplierCode': 'DR11109', 'supplierName': 'Raajkutir - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10727': {'supplierCode': 'DR10749', 'supplierName': 'Grand Mercure', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 2},
    'HO10731': {'supplierCode': 'DR11103', 'supplierName': 'Trident', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 3},
    'HO10733': {'supplierCode': 'DR11136', 'supplierName': 'Sandal Suites', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO10734': {'supplierCode': 'DR11130', 'supplierName': 'The Elephant Court', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10755': {'supplierCode': 'DR11131', 'supplierName': 'Devi Ratn - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10757': {'supplierCode': 'DR10835', 'supplierName': 'Abad Turtle Beach', 'parentGroup': 'Abad Hotels & Resorts', 'parentGroupId': '5fb27ab2-5c97-48be-a6a6-05251369e711', 'supplierVersion': 2},
    'HO10762': {'supplierCode': 'DR11110', 'supplierName': 'Taj Amer', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10763': {'supplierCode': 'DR11151', 'supplierName': 'Amber Dale', 'parentGroup': 'Kondody Hotels', 'parentGroupId': '32efa702-3d63-4160-b825-4839e180fe6d', 'supplierVersion': 2},
    'HO10771': {'supplierCode': 'DR10418', 'supplierName': 'Grand Hyatt', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 2},
    'HO10773': {'supplierCode': 'DR11121', 'supplierName': 'Brij Laxman Sagar', 'parentGroup': 'Brij Hotels', 'parentGroupId': '4957b920-a191-46d8-9122-ede09995fb10', 'supplierVersion': 4},
    'HO10781': {'supplierCode': 'DR10426', 'supplierName': 'Hyatt Regency Thrissur', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 2},
    'HO10795': {'supplierCode': 'DR11119', 'supplierName': 'Bera Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 4},
    'HO10814': {'supplierCode': 'DR11171', 'supplierName': 'Daspan House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10818': {'supplierCode': 'DR11156', 'supplierName': 'Bhanwar Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10821': {'supplierCode': 'DR10762', 'supplierName': 'The Emerald Hotel & Service Apartment', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10823': {'supplierCode': 'DR10634', 'supplierName': 'Shivadya Resort And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10835': {'supplierCode': 'DR10967', 'supplierName': 'Hyatt Centric', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 3},
    'HO10838': {'supplierCode': 'DR11143', 'supplierName': 'Tranquil Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10842': {'supplierCode': 'DR10157', 'supplierName': 'Courtyard By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10844': {'supplierCode': 'DR10706', 'supplierName': 'Country Inn & Suites (Sector 12)', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 2},
    'HO10850': {'supplierCode': 'DR10374', 'supplierName': 'Radisson Blu', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 3},
    'HO10857': {'supplierCode': 'DR10694', 'supplierName': 'Glenburn Tea Estate', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10858': {'supplierCode': 'DR10510', 'supplierName': 'Shangri-La Eros', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO10878': {'supplierCode': 'DR10515', 'supplierName': 'The Metropolitan Hotel & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10881': {'supplierCode': 'DR10492', 'supplierName': 'Crowne Plaza (Rohini)', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 2},
    'HO10886': {'supplierCode': 'DR10906', 'supplierName': 'The Shalimar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10887': {'supplierCode': 'DR10462', 'supplierName': 'Sea Princess', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10889': {'supplierCode': 'DR10904', 'supplierName': 'Stepwell House', 'parentGroup': 'Raas Hotels', 'parentGroupId': '0d88ba09-39d1-49aa-a183-130375ead42f', 'supplierVersion': 2},
    'HO10890': {'supplierCode': 'DR10753', 'supplierName': 'Niranta Airport Transit Hotel And Lounge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10930': {'supplierCode': 'DR11206', 'supplierName': 'Kanha Earth Lodge', 'parentGroup': 'Pugdundee Safaris', 'parentGroupId': '3587777b-fb9e-455d-8e0a-4cf718e2d9c5', 'supplierVersion': 2},
    'HO10935': {'supplierCode': 'DR10311', 'supplierName': 'Rang Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10938': {'supplierCode': 'DR11211', 'supplierName': 'Glenburn Penthouse', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10944': {'supplierCode': 'DR10330', 'supplierName': 'Balsamand Lake Palace', 'parentGroup': 'Welcomheritage Hotels', 'parentGroupId': '1ae3034f-d63b-46c8-82bd-1daefd69e317', 'supplierVersion': 2},
    'HO10946': {'supplierCode': 'DR10523', 'supplierName': 'Ranjit Vilas', 'parentGroup': 'Welcomheritage Hotels', 'parentGroupId': '1ae3034f-d63b-46c8-82bd-1daefd69e317', 'supplierVersion': 1},
    'HO10947': {'supplierCode': 'DR11174', 'supplierName': 'Dileep Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10950': {'supplierCode': 'DR10310', 'supplierName': 'Port Muziris, A Tribute Portfolio', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10957': {'supplierCode': 'DR11041', 'supplierName': 'The Leela', 'parentGroup': 'HLV', 'parentGroupId': 'c8bf0041-cad6-401c-8c7a-ad9dae79978b', 'supplierVersion': 2},
    'HO10959': {'supplierCode': 'DR10283', 'supplierName': 'Le Meridien', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO10960': {'supplierCode': 'DR11199', 'supplierName': 'Jamtara Wilderness Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO10962': {'supplierCode': 'DR10344', 'supplierName': 'The Ummed (Jodhpur)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10967': {'supplierCode': 'DR11551', 'supplierName': 'Holiday Inn (Sector 90)', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 3},
    'HO10977': {'supplierCode': 'DR11320', 'supplierName': 'Vivanta', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO10978': {'supplierCode': 'DR11295', 'supplierName': 'Vivanta', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO10981': {'supplierCode': 'DR11222', 'supplierName': 'Elysium Resort & Spa', 'parentGroup': 'Welcomheritage Hotels', 'parentGroupId': '1ae3034f-d63b-46c8-82bd-1daefd69e317', 'supplierVersion': 4},
    'HO10985': {'supplierCode': 'DR10665', 'supplierName': 'Radisson', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO10997': {'supplierCode': 'DR11234', 'supplierName': 'Deventure Sarovar Portico (Patel Nagar)', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 2},
    'HO11000': {'supplierCode': 'DR11209', 'supplierName': 'Taj Surajkund Resort & Spa', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO11004': {'supplierCode': 'DR11298', 'supplierName': 'Travancore Court', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11005': {'supplierCode': 'DR11309', 'supplierName': 'Amritara Sadka', 'parentGroup': 'Amritara Hotels and Resorts', 'parentGroupId': '1feba35d-2056-4407-8960-2ae840ea9393', 'supplierVersion': 2},
    'HO11007': {'supplierCode': 'DR11316', 'supplierName': 'Holiday Inn (Racecourse)', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 1},
    'HO11017': {'supplierCode': 'DR10745', 'supplierName': 'Vanaashrya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11019': {'supplierCode': 'DR11178', 'supplierName': 'St. Regis', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 2},
    'HO11020': {'supplierCode': 'DR11256', 'supplierName': 'Aangan Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11029': {'supplierCode': 'DR10758', 'supplierName': 'Taj Swarna', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO11037': {'supplierCode': 'DR11354', 'supplierName': 'Pride Plaza', 'parentGroup': 'Pride Group of Hotels', 'parentGroupId': '61b141a4-32e5-4dc3-afcb-e3ed29489e73', 'supplierVersion': 1},
    'HO11038': {'supplierCode': 'DR11291', 'supplierName': 'Fortune Park BBD', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 1},
    'HO11040': {'supplierCode': 'DR11212', 'supplierName': 'Fortune Park JPS Grand', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 1},
    'HO11063': {'supplierCode': 'DR11389', 'supplierName': 'Clarks Inn', 'parentGroup': 'The Clarks Hotel & Resorts (Clarks Inn Group Of Hotels)', 'parentGroupId': 'deccdf58-cd0d-40db-842a-ab29b073a9b8', 'supplierVersion': 3},
    'HO11064': {'supplierCode': 'DR11404', 'supplierName': 'The Pride', 'parentGroup': 'Pride Group of Hotels', 'parentGroupId': '61b141a4-32e5-4dc3-afcb-e3ed29489e73', 'supplierVersion': 2},
    'HO11070': {'supplierCode': 'DR11369', 'supplierName': 'WelcomHeritage Badi Kothi', 'parentGroup': 'Welcomheritage Hotels', 'parentGroupId': '1ae3034f-d63b-46c8-82bd-1daefd69e317', 'supplierVersion': 1},
    'HO11075': {'supplierCode': 'DR11365', 'supplierName': 'Holiday Inn', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 2},
    'HO11079': {'supplierCode': 'DR11338', 'supplierName': 'Vivanta', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 4},
    'HO11082': {'supplierCode': 'DR11448', 'supplierName': 'The Oberoi Sukhvilas Spa Resort', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 4},
    'HO11083': {'supplierCode': 'DR11419', 'supplierName': 'Fortune Select JP Cosmos', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 1},
    'HO11084': {'supplierCode': 'DR10580', 'supplierName': 'Welcomhotel By ITC Hotels, Fort & Dunes', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 6},
    'HO11089': {'supplierCode': 'DR11453', 'supplierName': 'Grand Hyatt', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 3},
    'HO11091': {'supplierCode': 'DR11344', 'supplierName': 'Welcomhotel By ITC Hotels', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 5},
    'HO11093': {'supplierCode': 'DR11371', 'supplierName': 'Fortune Park', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 2},
    'HO11094': {'supplierCode': 'DR11459', 'supplierName': 'ITC Narmada', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 5},
    'HO11095': {'supplierCode': 'DR10617', 'supplierName': 'Trident', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 5},
    'HO11096': {'supplierCode': 'DR11461', 'supplierName': 'The Leela Bhartiya City', 'parentGroup': 'HLV', 'parentGroupId': 'c8bf0041-cad6-401c-8c7a-ad9dae79978b', 'supplierVersion': 2},
    'HO11097': {'supplierCode': 'DR10864', 'supplierName': 'Trident (Bandra Kurla)', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 5},
    'HO11101': {'supplierCode': 'DR11455', 'supplierName': 'The Postcard Cuelim', 'parentGroup': 'Postcard Hotels', 'parentGroupId': 'dd274aa6-45f2-4cbc-9f21-3b372f4b9cc3', 'supplierVersion': 2},
    'HO11110': {'supplierCode': 'DR10470', 'supplierName': 'Welcomhotel By ITC Hotels', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 4},
    'HO11111': {'supplierCode': 'DR11443', 'supplierName': 'Sujan Jawai', 'parentGroup': 'Sujan Luxury Hotels', 'parentGroupId': '4d2b4d25-57a4-412c-91f8-b94c10d07b82', 'supplierVersion': 2},
    'HO11120': {'supplierCode': 'DR11486', 'supplierName': 'Classic Diplomat', 'parentGroup': 'New Age Hotels & Resorts', 'parentGroupId': '08de6637-23fb-4214-bde0-0a6e67724180', 'supplierVersion': 2},
    'HO11129': {'supplierCode': 'DR11432', 'supplierName': 'RAAS Devigarh', 'parentGroup': 'Raas Hotels', 'parentGroupId': '0d88ba09-39d1-49aa-a183-130375ead42f', 'supplierVersion': 3},
    'HO11130': {'supplierCode': 'DR11521', 'supplierName': 'Bagh Villas Jungle Camp & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO11135': {'supplierCode': 'DR11545', 'supplierName': 'The Amar Mahal', 'parentGroup': 'Trulyy India', 'parentGroupId': '5d72e698-0565-44a7-8105-e11d07f1dd43', 'supplierVersion': 2},
    'HO11137': {'supplierCode': 'DR11319', 'supplierName': 'Aamaghati Wildlife Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11139': {'supplierCode': 'DR11538', 'supplierName': 'Utsav Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11144': {'supplierCode': 'DR11515', 'supplierName': 'The Rajbari Bawali', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO11149': {'supplierCode': 'DR10353', 'supplierName': 'Vivaana Culture', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11159': {'supplierCode': 'DR11360', 'supplierName': 'Shreyas Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO11162': {'supplierCode': 'DR11555', 'supplierName': 'Ananda In The Himalayas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11163': {'supplierCode': 'DR11345', 'supplierName': 'Fort Chanwa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11165': {'supplierCode': 'DR11510', 'supplierName': 'Haveli Hauz Khas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11176': {'supplierCode': 'DR10275', 'supplierName': 'K K Royal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11182': {'supplierCode': 'DR11600', 'supplierName': 'Ambassador Pallava', 'parentGroup': 'The Ambassador Group of Hotels', 'parentGroupId': '937345c7-9f63-4dd7-b4e5-cce4b3fa794d', 'supplierVersion': 2},
    'HO11183': {'supplierCode': 'DR11398', 'supplierName': 'Rakkh Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO11190': {'supplierCode': 'DR11421', 'supplierName': 'Kaav Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11193': {'supplierCode': 'DR11331', 'supplierName': 'Six Senses Vana', 'parentGroup': 'Six Senses Hotels Resorts Spas', 'parentGroupId': '5e56de43-2c67-4cb9-98ac-ef7a4fdf8dcb', 'supplierVersion': 3},
    'HO11194': {'supplierCode': 'DR11862', 'supplierName': 'Clarks Inn', 'parentGroup': 'The Clarks Hotel & Resorts (Clarks Inn Group Of Hotels)', 'parentGroupId': 'deccdf58-cd0d-40db-842a-ab29b073a9b8', 'supplierVersion': 3},
    'HO11197': {'supplierCode': 'DR10422', 'supplierName': 'ITC Maratha', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 6},
    'HO11201': {'supplierCode': 'DR11287', 'supplierName': 'Coorg Wilderness Resort & Spa', 'parentGroup': 'Paul John Resorts and Hotels', 'parentGroupId': 'f1306273-db36-40e2-8707-8762678bc493', 'supplierVersion': 2},
    'HO11203': {'supplierCode': 'DR11044', 'supplierName': 'Raj Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11211': {'supplierCode': 'DR11117', 'supplierName': 'Deshadan Cliff And Beach Resort', 'parentGroup': 'Deshadan Hotels', 'parentGroupId': 'db58a030-dbb0-4ade-8cf9-c76b05a376ac', 'supplierVersion': 1},
    'HO11218': {'supplierCode': 'DR11407', 'supplierName': 'The Serai', 'parentGroup': 'Coffee Day Hotels & Resorts', 'parentGroupId': 'c1ef0eb5-adbf-472f-a821-b1fc2c99a52d', 'supplierVersion': 2},
    'HO11219': {'supplierCode': 'DR11410', 'supplierName': 'The Serai', 'parentGroup': 'Coffee Day Hotels & Resorts', 'parentGroupId': 'c1ef0eb5-adbf-472f-a821-b1fc2c99a52d', 'supplierVersion': 2},
    'HO11221': {'supplierCode': 'DR11350', 'supplierName': 'Fort House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11222': {'supplierCode': 'DR11444', 'supplierName': 'Jalakara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11224': {'supplierCode': 'DR11418', 'supplierName': 'Gokulam Grand Turtle On The Beach', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11225': {'supplierCode': 'DR11194', 'supplierName': 'Planters Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11226': {'supplierCode': 'DR10472', 'supplierName': 'Tenerife Hill', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11232': {'supplierCode': 'DR11362', 'supplierName': 'Neeleshwar Hermitage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO11235': {'supplierCode': 'DR11426', 'supplierName': 'Lotus Houseboat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11237': {'supplierCode': 'DR10723', 'supplierName': 'Uday Samudra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11242': {'supplierCode': 'DR11280', 'supplierName': 'Chittoor Kottaram', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 2},
    'HO11244': {'supplierCode': 'DR10186', 'supplierName': 'Shenbaga Hotel And Convention Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11246': {'supplierCode': 'DR10700', 'supplierName': 'Chalston Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11248': {'supplierCode': 'DR11305', 'supplierName': 'Colonia Santa Maria', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11256': {'supplierCode': 'DR11604', 'supplierName': 'Hyatt', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 3},
    'HO11268': {'supplierCode': 'DR10349', 'supplierName': 'The Lalit Temple View', 'parentGroup': 'Lalit Suri Hospitality Group', 'parentGroupId': '7602f0eb-5c5e-4db0-a9b7-93eb0cd88996', 'supplierVersion': 4},
    'HO11270': {'supplierCode': 'DR11485', 'supplierName': 'Arches', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11273': {'supplierCode': 'DR10594', 'supplierName': 'Radisson Blu', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 2},
    'HO11276': {'supplierCode': 'DR10512', 'supplierName': 'The Lalit', 'parentGroup': 'Lalit Suri Hospitality Group', 'parentGroupId': '7602f0eb-5c5e-4db0-a9b7-93eb0cd88996', 'supplierVersion': 6},
    'HO11281': {'supplierCode': 'DR10836', 'supplierName': 'Holiday Inn Mumbai International Airport', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 2},
    'HO11302': {'supplierCode': 'DR10430', 'supplierName': 'The Lalit Great Eastern', 'parentGroup': 'Lalit Suri Hospitality Group', 'parentGroupId': '7602f0eb-5c5e-4db0-a9b7-93eb0cd88996', 'supplierVersion': 4},
    'HO11304': {'supplierCode': 'DR11877', 'supplierName': 'JW Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 4},
    'HO11306': {'supplierCode': 'DR11466', 'supplierName': 'The Lalit Golf & Spa Resort', 'parentGroup': 'Lalit Suri Hospitality Group', 'parentGroupId': '7602f0eb-5c5e-4db0-a9b7-93eb0cd88996', 'supplierVersion': 6},
    'HO11307': {'supplierCode': 'DR11299', 'supplierName': 'Soma Palmshore', 'parentGroup': 'Somatheeram Ayurveda Group', 'parentGroupId': 'c035bf80-541b-49b7-b5bf-e4825d602f56', 'supplierVersion': 1},
    'HO11322': {'supplierCode': 'DR11795', 'supplierName': 'Courtyard By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 4},
    'HO11325': {'supplierCode': 'DR10484', 'supplierName': 'The Lalit', 'parentGroup': 'Lalit Suri Hospitality Group', 'parentGroupId': '7602f0eb-5c5e-4db0-a9b7-93eb0cd88996', 'supplierVersion': 5},
    'HO11332': {'supplierCode': 'DR11755', 'supplierName': 'The Fog Resorts & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11340': {'supplierCode': 'DR11496', 'supplierName': 'The Nattika Beach Ayurveda Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 4},
    'HO11347': {'supplierCode': 'DR10990', 'supplierName': 'Hyatt Place', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 2},
    'HO11348': {'supplierCode': 'DR11797', 'supplierName': 'Hyatt Regency', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 3},
    'HO11349': {'supplierCode': 'DR10588', 'supplierName': 'The Lalit', 'parentGroup': 'Lalit Suri Hospitality Group', 'parentGroupId': '7602f0eb-5c5e-4db0-a9b7-93eb0cd88996', 'supplierVersion': 6},
    'HO11351': {'supplierCode': 'DR10219', 'supplierName': 'The Lalit Laxmi Vilas Palace', 'parentGroup': 'Lalit Suri Hospitality Group', 'parentGroupId': '7602f0eb-5c5e-4db0-a9b7-93eb0cd88996', 'supplierVersion': 6},
    'HO11355': {'supplierCode': 'DR10618', 'supplierName': 'The Lalit', 'parentGroup': 'Lalit Suri Hospitality Group', 'parentGroupId': '7602f0eb-5c5e-4db0-a9b7-93eb0cd88996', 'supplierVersion': 6},
    'HO11358': {'supplierCode': 'DR12047', 'supplierName': 'Fairfield By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 5},
    'HO11363': {'supplierCode': 'DR10499', 'supplierName': 'Holiday Inn Express (International Airport T3)', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 3},
    'HO11367': {'supplierCode': 'DR10141', 'supplierName': 'Holiday Inn (International Airport)', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 3},
    'HO11399': {'supplierCode': 'DR11472', 'supplierName': 'Windflower Resorts & Spa', 'parentGroup': 'Windflower Resorts & Spa', 'parentGroupId': '10d69eaf-730a-4df8-aa93-8693924eec64', 'supplierVersion': 3},
    'HO11426': {'supplierCode': 'DR11399', 'supplierName': 'Dev Garh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11438': {'supplierCode': 'DR12040', 'supplierName': 'Delhi Heritage Dharampura', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11440': {'supplierCode': 'DR12041', 'supplierName': 'Annai Resorts & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11444': {'supplierCode': 'DR12043', 'supplierName': 'Grand Mercure', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 2},
    'HO11445': {'supplierCode': 'DR11597', 'supplierName': 'Umed Bhawan Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11459': {'supplierCode': 'DR11317', 'supplierName': 'Courtyard By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO11463': {'supplierCode': 'DR12193', 'supplierName': 'WGH Poetree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11464': {'supplierCode': 'DR12209', 'supplierName': 'Hyatt Place', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 2},
    'HO11466': {'supplierCode': 'DR12060', 'supplierName': 'Fairfield By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO11468': {'supplierCode': 'DR10367', 'supplierName': 'The Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 4},
    'HO11469': {'supplierCode': 'DR12244', 'supplierName': 'Novotel', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 3},
    'HO11473': {'supplierCode': 'DR12144', 'supplierName': 'Taj Lakefront', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO11476': {'supplierCode': 'DR12215', 'supplierName': 'Norbu The Montanna - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO11481': {'supplierCode': 'DR11062', 'supplierName': 'Eros', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11491': {'supplierCode': 'DR12227', 'supplierName': 'Bundela Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11493': {'supplierCode': 'DR12229', 'supplierName': 'The Gordon House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11494': {'supplierCode': 'DR12196', 'supplierName': 'Taurus Sarovar Portico', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 3},
    'HO11496': {'supplierCode': 'DR12247', 'supplierName': 'Daiwik', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11504': {'supplierCode': 'DR10257', 'supplierName': 'Spice Routes', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11508': {'supplierCode': 'DR10110', 'supplierName': 'Kumarakom Lake Resort', 'parentGroup': 'Paul John Resorts and Hotels', 'parentGroupId': 'f1306273-db36-40e2-8707-8762678bc493', 'supplierVersion': 2},
    'HO11513': {'supplierCode': 'DR10442', 'supplierName': 'Somatheeram Research Institute & Ayurveda Hospital', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11515': {'supplierCode': 'DR12057', 'supplierName': 'Four Seasons', 'parentGroup': 'Four Seasons Hotels', 'parentGroupId': 'dbee834c-2903-4c1e-9218-5fb340ada0eb', 'supplierVersion': 3},
    'HO11516': {'supplierCode': 'DR10744', 'supplierName': 'Taj Sawai', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 5},
    'HO11517': {'supplierCode': 'DR10637', 'supplierName': 'Manvar Resort And Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO11526': {'supplierCode': 'DR12320', 'supplierName': 'Renaissance', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 1},
    'HO11527': {'supplierCode': 'DR12307', 'supplierName': 'Spicetree Rajakumari', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11528': {'supplierCode': 'DR12288', 'supplierName': 'The World Backwaters', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11530': {'supplierCode': 'DR12292', 'supplierName': 'Gorbandh Palace - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO11534': {'supplierCode': 'DR12291', 'supplierName': 'Gateway Varkala - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 1},
    'HO11536': {'supplierCode': 'DR12338', 'supplierName': 'Coral Country Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11537': {'supplierCode': 'DR12240', 'supplierName': 'Arcadia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO11539': {'supplierCode': 'DR12282', 'supplierName': 'Fortune Select Global', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 1},
    'HO11541': {'supplierCode': 'DR12222', 'supplierName': 'The Residency Tower', 'parentGroup': 'The Residency Group of Hotels', 'parentGroupId': 'a50d3204-182f-4c20-9a02-e1166b93f4d0', 'supplierVersion': 1},
    'HO11548': {'supplierCode': 'DR12207', 'supplierName': 'Chandra Raj Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11554': {'supplierCode': 'DR12232', 'supplierName': 'Modi Yoga Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO11558': {'supplierCode': 'DR12346', 'supplierName': 'Tree Of Life Birdsong Chalets', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO11559': {'supplierCode': 'DR12344', 'supplierName': 'Tree Of Life Bhadrajun House', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO11566': {'supplierCode': 'DR12361', 'supplierName': 'Radisson Blu GRT', 'parentGroup': 'GRT Hotels and Resorts', 'parentGroupId': '5dab556b-52dc-4962-8a00-6db6da461b7b', 'supplierVersion': 2},
    'HO11571': {'supplierCode': 'DR12364', 'supplierName': 'Heritage Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11573': {'supplierCode': 'DR12369', 'supplierName': 'Taj Deccan', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 1},
    'HO11575': {'supplierCode': 'DR12370', 'supplierName': 'Bari Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11577': {'supplierCode': 'DR12373', 'supplierName': 'Scenic Munnar - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO11578': {'supplierCode': 'DR12375', 'supplierName': 'Taj City Centre', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO11582': {'supplierCode': 'DR10546', 'supplierName': 'The Leaf', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11584': {'supplierCode': 'DR12409', 'supplierName': 'Blanket Days Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11585': {'supplierCode': 'DR12410', 'supplierName': 'Indian Summer House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11592': {'supplierCode': 'DR12396', 'supplierName': 'Park Inn By Radisson', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 3},
    'HO11595': {'supplierCode': 'DR12398', 'supplierName': 'WOW Crest - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 1},
    'HO11615': {'supplierCode': 'DR12345', 'supplierName': 'Alivaa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11617': {'supplierCode': 'DR12195', 'supplierName': 'Sunday Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11625': {'supplierCode': 'DR12618', 'supplierName': 'The PL Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11638': {'supplierCode': 'DR10371', 'supplierName': 'Palace Heights', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO11643': {'supplierCode': 'DR10382', 'supplierName': 'Radisson Blu Marina', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 4},
    'HO11663': {'supplierCode': 'DR13348', 'supplierName': 'Southern Panorama Cruises', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11683': {'supplierCode': 'DR10253', 'supplierName': 'Flora Airport Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11691': {'supplierCode': 'DR12942', 'supplierName': 'Ginger House Museum', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11692': {'supplierCode': 'DR12473', 'supplierName': 'Welcomhotel By ITC Hotels', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 5},
    'HO11699': {'supplierCode': 'DR12991', 'supplierName': 'Agonda Villas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11710': {'supplierCode': 'DR13404', 'supplierName': 'Killians', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11736': {'supplierCode': 'DR13395', 'supplierName': 'Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 1},
    'HO11756': {'supplierCode': 'DR13444', 'supplierName': 'Fortune Park', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 1},
    'HO11788': {'supplierCode': 'DR12548', 'supplierName': 'Heritage Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11789': {'supplierCode': 'DR12408', 'supplierName': 'Vibe Munnar Resorts & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11805': {'supplierCode': 'DR12486', 'supplierName': 'Vivanta (Miramar)', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO11811': {'supplierCode': 'DR12517', 'supplierName': 'Whispering Palms Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11812': {'supplierCode': 'DR12727', 'supplierName': 'The Chancery', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11866': {'supplierCode': 'DR12449', 'supplierName': 'Radisson Blu', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 2},
    'HO11867': {'supplierCode': 'DR12571', 'supplierName': 'AyurSoma Ayurveda Royal Retreat', 'parentGroup': 'Somatheeram Ayurveda Group', 'parentGroupId': 'c035bf80-541b-49b7-b5bf-e4825d602f56', 'supplierVersion': 1},
    'HO11873': {'supplierCode': 'DR12529', 'supplierName': 'Rockholm', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO11876': {'supplierCode': 'DR12885', 'supplierName': 'The Fern Habitat', 'parentGroup': 'Fern Hotels & Resorts', 'parentGroupId': '6be198c4-bfec-4170-bd7e-066cc7a5a9ef', 'supplierVersion': 2},
    'HO11877': {'supplierCode': 'DR13499', 'supplierName': 'The Kin', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11879': {'supplierCode': 'DR12504', 'supplierName': 'Gokulam Grand Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11892': {'supplierCode': 'DR13203', 'supplierName': 'Aalankrita Resort & Convention', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11902': {'supplierCode': 'DR12646', 'supplierName': 'Karwaan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11903': {'supplierCode': 'DR10475', 'supplierName': 'The Imperial', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11911': {'supplierCode': 'DR12425', 'supplierName': 'Novotel Hyderabad Airport', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 2},
    'HO11950': {'supplierCode': 'DR12477', 'supplierName': 'Novotel Mumbai International Airport', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 2},
    'HO11954': {'supplierCode': 'DR13127', 'supplierName': 'Royal Orchid Brindavan Garden Palace & Spa', 'parentGroup': 'Royal Orchid Hotels', 'parentGroupId': '025ed948-0a29-4839-9a2a-daaa298dbfad', 'supplierVersion': 1},
    'HO11967': {'supplierCode': 'DR12585', 'supplierName': 'Flamingo', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11969': {'supplierCode': 'DR12439', 'supplierName': 'Radisson Blu', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 3},
    'HO11972': {'supplierCode': 'DR13075', 'supplierName': 'The Baradari Palace', 'parentGroup': 'Neemrana Hotels', 'parentGroupId': '20cca852-4d76-47ee-a2e7-ce543bf12867', 'supplierVersion': 1},
    'HO11976': {'supplierCode': 'DR13008', 'supplierName': 'Suba International', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11992': {'supplierCode': 'DR13456', 'supplierName': 'Aralia International Airport', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11994': {'supplierCode': 'DR12418', 'supplierName': 'The Westin Mumbai Powai Lake', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 2},
    'HO12008': {'supplierCode': 'DR12554', 'supplierName': 'Chandys Windy Woods', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12009': {'supplierCode': 'DR12422', 'supplierName': 'Welcomhotel By ITC Hotels, Bay Island', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 3},
    'HO12010': {'supplierCode': 'DR13143', 'supplierName': 'Diplomat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12013': {'supplierCode': 'DR12658', 'supplierName': 'Gratitude Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12023': {'supplierCode': 'DR10424', 'supplierName': 'JW Marriott (Sahar)', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 3},
    'HO12025': {'supplierCode': 'DR12500', 'supplierName': 'Radisson Resort Pondicherry Bay', 'parentGroup': 'GRT Hotels and Resorts', 'parentGroupId': '5dab556b-52dc-4962-8a00-6db6da461b7b', 'supplierVersion': 1},
    'HO12050': {'supplierCode': 'DR12469', 'supplierName': 'Courtyard By Marriott', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 1},
    'HO12051': {'supplierCode': 'DR13513', 'supplierName': 'Sun N Sand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12053': {'supplierCode': 'DR11223', 'supplierName': 'Mayfair Tea Resort', 'parentGroup': 'Mayfair Hotels & Resorts', 'parentGroupId': '50b199a5-0468-4963-9fe6-7c2dc1a73d9d', 'supplierVersion': 1},
    'HO12060': {'supplierCode': 'DR13369', 'supplierName': 'Lemon Tree Suites (Sector 82 A)', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO12066': {'supplierCode': 'DR12433', 'supplierName': 'Hyatt Regency', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 2},
    'HO12077': {'supplierCode': 'DR13362', 'supplierName': 'Sreechithra Ayur Home', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12078': {'supplierCode': 'DR13400', 'supplierName': 'Radisson Blu (Hinjawadi)', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO12093': {'supplierCode': 'DR12496', 'supplierName': 'Abad Harmonia', 'parentGroup': 'Abad Hotels & Resorts', 'parentGroupId': '5fb27ab2-5c97-48be-a6a6-05251369e711', 'supplierVersion': 1},
    'HO12105': {'supplierCode': 'DR12578', 'supplierName': 'O By Tamara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12107': {'supplierCode': 'DR13083', 'supplierName': 'Antalya By Divine Ganga Cottage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12108': {'supplierCode': 'DR12424', 'supplierName': 'Vivanta', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 1},
    'HO12117': {'supplierCode': 'DR13407', 'supplierName': 'Mountain Shadows Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12127': {'supplierCode': 'DR13342', 'supplierName': 'Radisson Blu Hotel & Spa', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO12135': {'supplierCode': 'DR12557', 'supplierName': 'Hills & Hues', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12139': {'supplierCode': 'DR13287', 'supplierName': 'Mountain Courtyard', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12155': {'supplierCode': 'DR12693', 'supplierName': 'Best Western (Maharani Bagh)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12163': {'supplierCode': 'DR12527', 'supplierName': 'DoubleTree By Hilton', 'parentGroup': 'Hilton Hotels & Resorts', 'parentGroupId': '064b6ff4-e74c-41c8-ab61-660f04ce70d8', 'supplierVersion': 2},
    'HO12176': {'supplierCode': 'DR12308', 'supplierName': 'Ranthambhore Tiger Inn Comfort Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12187': {'supplierCode': 'DR12415', 'supplierName': 'Laalee', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12188': {'supplierCode': 'DR12407', 'supplierName': 'Manuscript Jhilwara Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12190': {'supplierCode': 'DR12501', 'supplierName': 'Accord Chrome', 'parentGroup': 'Accord Hotels and Resorts', 'parentGroupId': 'b9574f54-f933-40ba-99af-2b212340c07f', 'supplierVersion': 2},
    'HO12200': {'supplierCode': 'DR12414', 'supplierName': 'Raffles', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 1},
    'HO12201': {'supplierCode': 'DR12401', 'supplierName': 'Rhythm', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12205': {'supplierCode': 'DR12221', 'supplierName': 'Pushkara Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12210': {'supplierCode': 'DR12634', 'supplierName': 'Philipkutty\'s Farm', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12222': {'supplierCode': 'DR12664', 'supplierName': 'Cardinal Express Oxmo', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12235': {'supplierCode': 'DR12252', 'supplierName': 'Hyatt Regency', 'parentGroup': 'Hyatt Hotels & Resorts', 'parentGroupId': 'b837767b-767e-4e6f-aea9-5ce99ded691b', 'supplierVersion': 2},
    'HO12273': {'supplierCode': 'DR12696', 'supplierName': 'Niravi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12276': {'supplierCode': 'DR12702', 'supplierName': 'The Valle', 'parentGroup': 'Kondody Hotels', 'parentGroupId': '32efa702-3d63-4160-b825-4839e180fe6d', 'supplierVersion': 2},
    'HO12279': {'supplierCode': 'DR12714', 'supplierName': 'Cedar Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12283': {'supplierCode': 'DR12753', 'supplierName': 'Novotel Jaipur Convention Centre', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 2},
    'HO12284': {'supplierCode': 'DR12659', 'supplierName': 'Golden Tulip', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 1},
    'HO12285': {'supplierCode': 'DR12737', 'supplierName': 'Plutos', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12286': {'supplierCode': 'DR12751', 'supplierName': 'Sandesh Water Edge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12292': {'supplierCode': 'DR12754', 'supplierName': 'Devi Bhawan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12298': {'supplierCode': 'DR12750', 'supplierName': 'Africa Avenue (GK-1)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12302': {'supplierCode': 'DR12733', 'supplierName': 'Sakura Gold', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12317': {'supplierCode': 'DR12783', 'supplierName': 'Taj Cochin International Airport', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 3},
    'HO12319': {'supplierCode': 'DR12471', 'supplierName': 'Marli Hill', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 1},
    'HO12329': {'supplierCode': 'DR12726', 'supplierName': 'Grand Madurai', 'parentGroup': 'GRT Hotels and Resorts', 'parentGroupId': '5dab556b-52dc-4962-8a00-6db6da461b7b', 'supplierVersion': 1},
    'HO12334': {'supplierCode': 'DR10566', 'supplierName': 'Visalam', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO12335': {'supplierCode': 'DR10016', 'supplierName': 'Saratha Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12336': {'supplierCode': 'DR10764', 'supplierName': 'Chettinadu Court', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12340': {'supplierCode': 'DR12724', 'supplierName': 'Sterling', 'parentGroup': 'Sterling Holiday Resorts', 'parentGroupId': '979205b7-0edb-46f5-8de2-f32a34f7888b', 'supplierVersion': 1},
    'HO12348': {'supplierCode': 'DR11469', 'supplierName': 'Holiday Inn Resort (Cavelossim)', 'parentGroup': 'Intercontinental Hotels Group', 'parentGroupId': 'f446eae6-65b5-4707-ada6-470eb57f0223', 'supplierVersion': 3},
    'HO12350': {'supplierCode': 'DR12793', 'supplierName': 'The Fern', 'parentGroup': 'Fern Hotels & Resorts', 'parentGroupId': '6be198c4-bfec-4170-bd7e-066cc7a5a9ef', 'supplierVersion': 2},
    'HO12353': {'supplierCode': 'DR12720', 'supplierName': 'Umaid Mahal Palace', 'parentGroup': 'Jagat Singh Hotels', 'parentGroupId': '2ffcc35d-3f12-43d6-a323-3795a812b070', 'supplierVersion': 1},
    'HO12365': {'supplierCode': 'DR12809', 'supplierName': 'Brij Eternity', 'parentGroup': 'Leisure Hotels', 'parentGroupId': '0035025a-93fa-4882-858f-2d85f8225031', 'supplierVersion': 1},
    'HO12366': {'supplierCode': 'DR12800', 'supplierName': 'Cheetahgarh Resort & Spa', 'parentGroup': 'Welcomheritage Hotels', 'parentGroupId': '1ae3034f-d63b-46c8-82bd-1daefd69e317', 'supplierVersion': 1},
    'HO12375': {'supplierCode': 'DR12799', 'supplierName': 'The Lotus Palace', 'parentGroup': 'Apeejay Surrendra Park Hotels', 'parentGroupId': '1f4a03c7-d445-4cd2-816b-5d060c85220c', 'supplierVersion': 1},
    'HO12379': {'supplierCode': 'DR12871', 'supplierName': 'V Sarovar Portico', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 1},
    'HO12382': {'supplierCode': 'DR12837', 'supplierName': 'Niraamaya Retreats Surya Samudra', 'parentGroup': 'Niraamaya Life', 'parentGroupId': '10cbd2d0-c8ab-4d15-b40e-be4dfba57c0e', 'supplierVersion': 1},
    'HO12390': {'supplierCode': 'DR12834', 'supplierName': 'Niraamaya Retreats Backwaters & Beyond', 'parentGroup': 'Niraamaya Life', 'parentGroupId': '10cbd2d0-c8ab-4d15-b40e-be4dfba57c0e', 'supplierVersion': 1},
    'HO12392': {'supplierCode': 'DR12655', 'supplierName': 'TCI-Kathmandu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12395': {'supplierCode': 'DR12902', 'supplierName': 'Ramada By Wyndham', 'parentGroup': 'Wyndham Hotels & Resorts', 'parentGroupId': '1861249a-90a7-4244-a0dd-9a1fa184ed25', 'supplierVersion': 1},
    'HO12396': {'supplierCode': 'DR12831', 'supplierName': 'Fairmont', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 1},
    'HO12405': {'supplierCode': 'DR12908', 'supplierName': 'Saga Metho', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12406': {'supplierCode': 'DR12909', 'supplierName': 'Wisdom Himalayan Voyages Pvt Ltd', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12408': {'supplierCode': 'DR12926', 'supplierName': 'Clubside Tours & Travels Pvt Ltd', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12416': {'supplierCode': 'DR12929', 'supplierName': 'Mastiff Select Mandavya', 'parentGroup': 'Orange Tiger Hospitality Pvt. Ltd.', 'parentGroupId': '7415a63a-ad3a-440b-8aa4-c472fc389374', 'supplierVersion': 1},
    'HO12427': {'supplierCode': 'DR12958', 'supplierName': 'Clubside Tours & Travels Pvt Ltd', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12429': {'supplierCode': 'DR12956', 'supplierName': 'International ╨íentral Asia Travel LLC', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12431': {'supplierCode': 'DR12961', 'supplierName': 'Travel Link', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12432': {'supplierCode': 'DR12962', 'supplierName': 'Harsh Travels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12434': {'supplierCode': 'DR12965', 'supplierName': 'Barefoot Holidays', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12435': {'supplierCode': 'DR12968', 'supplierName': 'Tropiculture (Private) Limited', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12438': {'supplierCode': 'DR12967', 'supplierName': 'Adventure Escape Asia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12441': {'supplierCode': 'DR12972', 'supplierName': 'IV Sanctum', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12453': {'supplierCode': 'DR12987', 'supplierName': 'The Belvedere Himalayan Retreat', 'parentGroup': 'Leisure Hotels', 'parentGroupId': '0035025a-93fa-4882-858f-2d85f8225031', 'supplierVersion': 1},
    'HO12454': {'supplierCode': 'DR12989', 'supplierName': 'The Postcard In The Durrung Tea Estate', 'parentGroup': 'Postcard Hotels', 'parentGroupId': 'dd274aa6-45f2-4cbc-9f21-3b372f4b9cc3', 'supplierVersion': 2},
    'HO12467': {'supplierCode': 'DR12990', 'supplierName': 'Sawantwadi Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12469': {'supplierCode': 'DR12963', 'supplierName': 'Cavalry Villa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12474': {'supplierCode': 'DR12862', 'supplierName': 'Gulab Haveli', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 1},
    'HO12480': {'supplierCode': 'DR13007', 'supplierName': 'Gateway', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 1},
    'HO12486': {'supplierCode': 'DR12946', 'supplierName': 'Coco Shambhala', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12488': {'supplierCode': 'DR13051', 'supplierName': 'Bangaram Island Resort - IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 1},
    'HO12489': {'supplierCode': 'DR13016', 'supplierName': 'Urban Suites By BluSalzz', 'parentGroup': 'BluSalzz Hospitality', 'parentGroupId': '423a2b31-854f-48a6-875d-7652a45158e6', 'supplierVersion': 2},
    'HO12496': {'supplierCode': 'DR13014', 'supplierName': 'Purvi Discovery', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12502': {'supplierCode': 'DR13040', 'supplierName': 'The Travel Concierge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12504': {'supplierCode': 'DR13041', 'supplierName': 'Tour Planner Ltd.', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12505': {'supplierCode': 'DR12359', 'supplierName': 'Ganga Kinare', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12508': {'supplierCode': 'DR13025', 'supplierName': 'Tattwa Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12512': {'supplierCode': 'DR13026', 'supplierName': 'The Leela Palace', 'parentGroup': 'HLV', 'parentGroupId': 'c8bf0041-cad6-401c-8c7a-ad9dae79978b', 'supplierVersion': 1},
    'HO12513': {'supplierCode': 'DR12789', 'supplierName': 'Ranthambhore Bagh Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12515': {'supplierCode': 'DR12638', 'supplierName': 'Surpura Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12521': {'supplierCode': 'DR13034', 'supplierName': 'Lemon Tree Premier', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO12529': {'supplierCode': 'DR13052', 'supplierName': 'Amara Ayurveda Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12534': {'supplierCode': 'DR13072', 'supplierName': 'Park Inn By Radisson (Lagpat Nagar)', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO12537': {'supplierCode': 'DR12839', 'supplierName': 'Ramada Encore By Wyndham', 'parentGroup': 'Wyndham Hotels & Resorts', 'parentGroupId': '1861249a-90a7-4244-a0dd-9a1fa184ed25', 'supplierVersion': 1},
    'HO12546': {'supplierCode': 'DR12880', 'supplierName': 'Daulatgarh Resort And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12551': {'supplierCode': 'DR13097', 'supplierName': 'The Oberoi Rajgarh Palace', 'parentGroup': 'EIH', 'parentGroupId': '6232963a-4b06-488b-9922-186e72a0d5ef', 'supplierVersion': 2},
    'HO12560': {'supplierCode': 'DR13062', 'supplierName': 'Sunday Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12565': {'supplierCode': 'DR12046', 'supplierName': 'Carnoustie Ayurveda & Wellness Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12567': {'supplierCode': 'DR13116', 'supplierName': 'Broadway', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12572': {'supplierCode': 'DR11072', 'supplierName': 'Swaswara', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 3},
    'HO12580': {'supplierCode': 'DR13134', 'supplierName': 'Welcomhotel By ITC Hotels', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 1},
    'HO12583': {'supplierCode': 'DR13135', 'supplierName': 'The Nattika Beach Ayurveda & Yoga Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12585': {'supplierCode': 'DR13133', 'supplierName': 'Novotel (Nagar Road)', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 1},
    'HO12589': {'supplierCode': 'DR13140', 'supplierName': 'Ether Marari Beachfront', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12591': {'supplierCode': 'DR13141', 'supplierName': 'Blue Diamond Enterprises', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12597': {'supplierCode': 'DR13147', 'supplierName': 'Pullman', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 1},
    'HO12598': {'supplierCode': 'DR13149', 'supplierName': 'Regenta Central Javaji', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12599': {'supplierCode': 'DR13154', 'supplierName': 'Avantika By The Ganges ΓÇô IHCL SeleQtions', 'parentGroup': 'IHCL', 'parentGroupId': 'd43c515d-01c3-4f65-bd59-9a1b3d87444c', 'supplierVersion': 2},
    'HO12601': {'supplierCode': 'DR13156', 'supplierName': 'The Fern Samali Resort', 'parentGroup': 'Fern Hotels & Resorts', 'parentGroupId': '6be198c4-bfec-4170-bd7e-066cc7a5a9ef', 'supplierVersion': 1},
    'HO12607': {'supplierCode': 'DR13163', 'supplierName': 'Harbour View', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12615': {'supplierCode': 'DR13177', 'supplierName': 'Shangloo Travels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12618': {'supplierCode': 'DR13185', 'supplierName': 'The Evren', 'parentGroup': 'BluSalzz Hospitality', 'parentGroupId': '423a2b31-854f-48a6-875d-7652a45158e6', 'supplierVersion': 1},
    'HO12623': {'supplierCode': 'DR13194', 'supplierName': 'Fazlani Nature\'s Nest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12631': {'supplierCode': 'DR13200', 'supplierName': 'Thousand Pillars', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12636': {'supplierCode': 'DR13205', 'supplierName': 'Siamton Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12637': {'supplierCode': 'DR13207', 'supplierName': 'Aditya Hometel', 'parentGroup': 'Sarovar Hotels & Resorts', 'parentGroupId': 'ca84e0be-04a1-465d-95d0-2754a1e65813', 'supplierVersion': 1},
    'HO12639': {'supplierCode': 'DR13214', 'supplierName': 'DVR Mansion', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12640': {'supplierCode': 'DR13215', 'supplierName': 'Garuda India Holidays', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12648': {'supplierCode': 'DR10513', 'supplierName': 'The Orchid', 'parentGroup': 'Kamat Hotels (India) Limited', 'parentGroupId': 'ed653e5e-a1d1-44fe-bf79-0e1a2421a5c8', 'supplierVersion': 4},
    'HO12651': {'supplierCode': 'DR13015', 'supplierName': 'Three Sixty Holidays', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12653': {'supplierCode': 'DR13261', 'supplierName': 'The Rose Goa Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12655': {'supplierCode': 'DR13268', 'supplierName': 'Ayur On The Beach', 'parentGroup': 'Abad Hotels & Resorts', 'parentGroupId': '5fb27ab2-5c97-48be-a6a6-05251369e711', 'supplierVersion': 1},
    'HO12656': {'supplierCode': 'DR13249', 'supplierName': 'Fortune Park Lake City', 'parentGroup': 'Fortune Hotels', 'parentGroupId': 'de3507dd-fa8e-4e22-8057-6d56f54711d5', 'supplierVersion': 2},
    'HO12661': {'supplierCode': 'DR13278', 'supplierName': 'Rai Travels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12663': {'supplierCode': 'DR13294', 'supplierName': 'Sanctuary Amaidiyana', 'parentGroup': 'CGH Earth', 'parentGroupId': '4f3cbf04-baf6-48ae-9110-427e40f39a67', 'supplierVersion': 1},
    'HO12671': {'supplierCode': 'DR12812', 'supplierName': 'Ibis Styles', 'parentGroup': 'Accor Hotels', 'parentGroupId': 'fbc53148-48ab-4faf-8753-d35a3c11455e', 'supplierVersion': 3},
    'HO12672': {'supplierCode': 'DR13327', 'supplierName': 'Fairfield By Marriott Mumbai International Airport', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 1},
    'HO12675': {'supplierCode': 'DR10358', 'supplierName': 'Detours India', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12680': {'supplierCode': 'DR12899', 'supplierName': 'The Fern Residency (Subhash Bridge)', 'parentGroup': 'Fern Hotels & Resorts', 'parentGroupId': '6be198c4-bfec-4170-bd7e-066cc7a5a9ef', 'supplierVersion': 2},
    'HO12681': {'supplierCode': 'DR13035', 'supplierName': 'Single Spot Tourism L.L.C', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12684': {'supplierCode': 'DR13356', 'supplierName': 'Iconiqa', 'parentGroup': 'Royal Orchid Hotels', 'parentGroupId': '025ed948-0a29-4839-9a2a-daaa298dbfad', 'supplierVersion': 1},
    'HO12688': {'supplierCode': 'DR13361', 'supplierName': 'Tornos Destinations India Pvt Ltd', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12691': {'supplierCode': 'DR13364', 'supplierName': 'Eve', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12696': {'supplierCode': 'DR12584', 'supplierName': 'Storii By ITC Hotels Castle Kanota', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 2},
    'HO12698': {'supplierCode': 'DR13385', 'supplierName': 'Landmaster Holidays', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12701': {'supplierCode': 'DR13394', 'supplierName': 'Blue Jelly Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12704': {'supplierCode': 'DR13399', 'supplierName': 'Ramanashree Richmond', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12705': {'supplierCode': 'DR13401', 'supplierName': 'Lemon Tree Premier', 'parentGroup': 'Lemon Tree Hotels', 'parentGroupId': '2d54ccd4-4ee9-4c25-b935-1a2408b3d460', 'supplierVersion': 1},
    'HO12711': {'supplierCode': 'DR13410', 'supplierName': 'Radisson (Sector 29)', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO12714': {'supplierCode': 'DR13398', 'supplierName': 'Radisson Blu', 'parentGroup': 'Radisson Hotels', 'parentGroupId': '0614f697-1cd6-45e0-b22a-e934acb0377d', 'supplierVersion': 1},
    'HO12715': {'supplierCode': 'DR13418', 'supplierName': 'Theory9 - Premium Serviced Apartments', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12716': {'supplierCode': 'DR13420', 'supplierName': 'Hilton Garden Inn Mumbai International Airport', 'parentGroup': 'Hilton Hotels & Resorts', 'parentGroupId': '064b6ff4-e74c-41c8-ab61-660f04ce70d8', 'supplierVersion': 2},
    'HO12717': {'supplierCode': 'DR13424', 'supplierName': 'Anandam Ayurvedic', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12718': {'supplierCode': 'DR13423', 'supplierName': 'Park Regis', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12719': {'supplierCode': 'DR13426', 'supplierName': 'Lagoon Bay At Leonia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12722': {'supplierCode': 'DR13432', 'supplierName': 'Welcomhotel By ITC Hotels', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 1},
    'HO12726': {'supplierCode': 'DR13440', 'supplierName': 'Dia Park Premier', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12729': {'supplierCode': 'DR13445', 'supplierName': 'Flora Vythiri Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12730': {'supplierCode': 'DR13447', 'supplierName': 'Flora Misty Falls', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12734': {'supplierCode': 'DR13463', 'supplierName': 'Yuvarani Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12737': {'supplierCode': 'DR13470', 'supplierName': 'Somerset Greeways', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12743': {'supplierCode': 'DR13477', 'supplierName': 'Royale Lalawi Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12744': {'supplierCode': 'DR13485', 'supplierName': 'Gobindgarh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 2},
    'HO12747': {'supplierCode': 'DR13479', 'supplierName': 'TCI-Tokyo', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12755': {'supplierCode': 'DR13492', 'supplierName': 'Machaan Wilderness Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12761': {'supplierCode': 'DR13502', 'supplierName': 'The Bella Vista', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12765': {'supplierCode': 'DR13515', 'supplierName': 'Ira By Orchid', 'parentGroup': 'Kamat Hotels (India) Limited', 'parentGroupId': 'ed653e5e-a1d1-44fe-bf79-0e1a2421a5c8', 'supplierVersion': 1},
    'HO12777': {'supplierCode': 'DR13528', 'supplierName': 'Welcomhotel by ITC Hotel', 'parentGroup': 'ITC Hotels', 'parentGroupId': 'f45c17de-6691-4c30-8ca3-458a5ba182b7', 'supplierVersion': 2},
    'HO12796': {'supplierCode': 'DR13559', 'supplierName': 'Ritz Carlton', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 1},
    'HO12799': {'supplierCode': 'DR13562', 'supplierName': 'Four Point By Sheraton', 'parentGroup': 'Marriott Hotels & Resorts', 'parentGroupId': '9404aa8a-34ba-4d67-ad0c-6f4a76bc5979', 'supplierVersion': 1},
    'HO12800': {'supplierCode': 'DR13212', 'supplierName': 'Heritage River Journeys Private Limited', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 3},
    'HO12802': {'supplierCode': 'DR13566', 'supplierName': 'Rockholm Ayurveda & Yoga Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12804': {'supplierCode': 'DR13569', 'supplierName': 'Sterling', 'parentGroup': 'Sterling Holiday Resorts', 'parentGroupId': '979205b7-0edb-46f5-8de2-f32a34f7888b', 'supplierVersion': 1},
    'HO12807': {'supplierCode': 'DR13570', 'supplierName': 'Pillow Mint', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12809': {'supplierCode': 'DR13575', 'supplierName': 'Nikko', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12813': {'supplierCode': 'DR13582', 'supplierName': 'Gokulam Park', 'parentGroup': 'Gokulam Hotels & Resorts', 'parentGroupId': '3861346e-bb79-41b6-8f7d-29d058287ce7', 'supplierVersion': 1},
    'HO10015': {'supplierCode': 'DR10176', 'supplierName': 'East View', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10021': {'supplierCode': 'DR10101', 'supplierName': 'Taj Nadesar Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10024': {'supplierCode': 'DR10001', 'supplierName': 'Khandela Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10025': {'supplierCode': 'DR10069', 'supplierName': 'Aurika', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10029': {'supplierCode': 'DR10108', 'supplierName': 'Backwater Ripples', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10031': {'supplierCode': 'DR10190', 'supplierName': 'Lakshmi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10037': {'supplierCode': 'DR10171', 'supplierName': 'Palace on Ganges', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10039': {'supplierCode': 'DR10173', 'supplierName': 'Meraden Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10041': {'supplierCode': 'DR10111', 'supplierName': 'Niraamaya Retreats', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10042': {'supplierCode': 'DR10164', 'supplierName': 'J C Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10044': {'supplierCode': 'DR10183', 'supplierName': 'Le Pondy', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10045': {'supplierCode': 'DR10197', 'supplierName': 'The Coral Tree Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10046': {'supplierCode': 'DR10057', 'supplierName': 'Aloft', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10049': {'supplierCode': 'DR10145', 'supplierName': 'Premkunj', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10050': {'supplierCode': 'DR10090', 'supplierName': 'Caspia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10051': {'supplierCode': 'DR10231', 'supplierName': 'Udai Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10055': {'supplierCode': 'DR10026', 'supplierName': 'Tissas Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10056': {'supplierCode': 'DR10033', 'supplierName': 'Calcutta Bungalow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10057': {'supplierCode': 'DR10233', 'supplierName': 'Vishnupriya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10058': {'supplierCode': 'DR10034', 'supplierName': 'The Peerless Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10060': {'supplierCode': 'DR10146', 'supplierName': 'Rajdarshan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10068': {'supplierCode': 'DR10245', 'supplierName': 'Clarks Amer', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10070': {'supplierCode': 'DR10259', 'supplierName': 'Gandharva By Peppermint', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10071': {'supplierCode': 'DR10212', 'supplierName': 'GRT Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10073': {'supplierCode': 'DR10115', 'supplierName': 'Fateh Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10074': {'supplierCode': 'DR10165', 'supplierName': 'The Madurai Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10075': {'supplierCode': 'DR10295', 'supplierName': 'Treetop', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10080': {'supplierCode': 'DR10127', 'supplierName': 'Howard Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10082': {'supplierCode': 'DR10232', 'supplierName': 'Udaigarh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10083': {'supplierCode': 'DR10080', 'supplierName': 'Bujera Fort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10084': {'supplierCode': 'DR10298', 'supplierName': 'Wildernest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10085': {'supplierCode': 'DR10337', 'supplierName': 'Parkk Boutique', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10087': {'supplierCode': 'DR10294', 'supplierName': 'Dev Villa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10090': {'supplierCode': 'DR10193', 'supplierName': 'Ramada Plaza by Wyndham Agra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10091': {'supplierCode': 'DR10243', 'supplierName': 'Clarion Bella Casa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10093': {'supplierCode': 'DR10273', 'supplierName': 'Indana Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10095': {'supplierCode': 'DR10312', 'supplierName': 'The Desert Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10098': {'supplierCode': 'DR10214', 'supplierName': 'Hyatt Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10102': {'supplierCode': 'DR10391', 'supplierName': 'Surya International', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10112': {'supplierCode': 'DR10058', 'supplierName': 'Aaram Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10116': {'supplierCode': 'DR10225', 'supplierName': 'Accord Metropolitan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10118': {'supplierCode': 'DR10325', 'supplierName': 'Suryaa Villa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10119': {'supplierCode': 'DR10293', 'supplierName': 'Imperial', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10122': {'supplierCode': 'DR10314', 'supplierName': 'Roop Niwas Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10123': {'supplierCode': 'DR10401', 'supplierName': 'Tiger Den Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10128': {'supplierCode': 'DR10207', 'supplierName': 'Clarks', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10129': {'supplierCode': 'DR10489', 'supplierName': 'Samsara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10133': {'supplierCode': 'DR10450', 'supplierName': 'Pilibhit House – IHCL SeleQtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10137': {'supplierCode': 'DR10428', 'supplierName': 'Keys Select', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10141': {'supplierCode': 'DR10126', 'supplierName': 'Radisson Jass', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10146': {'supplierCode': 'DR10433', 'supplierName': 'Ramada By Wyndham Udaipur Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10147': {'supplierCode': 'DR10438', 'supplierName': 'Regency Sameera', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10148': {'supplierCode': 'DR10446', 'supplierName': 'Nidhivan Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10150': {'supplierCode': 'DR10652', 'supplierName': 'Tea Bungalow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10152': {'supplierCode': 'DR10651', 'supplierName': 'Chidambara Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10157': {'supplierCode': 'DR10559', 'supplierName': 'The Elgin Nor-Khill', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10165': {'supplierCode': 'DR10661', 'supplierName': 'Mela Kothi - Chambal Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10170': {'supplierCode': 'DR10560', 'supplierName': 'The Elgin Silver Oaks', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10171': {'supplierCode': 'DR10663', 'supplierName': 'Sparsa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10174': {'supplierCode': 'DR10664', 'supplierName': 'Germanus', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10176': {'supplierCode': 'DR10670', 'supplierName': 'Country Inn & Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10177': {'supplierCode': 'DR10774', 'supplierName': 'Master Paradise', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10180': {'supplierCode': 'DR10180', 'supplierName': 'Welcomhotel By ITC Hotels, Kences Palm Beach', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10183': {'supplierCode': 'DR10673', 'supplierName': 'Sangam', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10187': {'supplierCode': 'DR10675', 'supplierName': 'Greenwoods', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10189': {'supplierCode': 'DR10427', 'supplierName': 'Niraamaya Retreats Surya Samudra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10200': {'supplierCode': 'DR10435', 'supplierName': 'Meluha The Fern', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10201': {'supplierCode': 'DR10091', 'supplierName': 'Fateh Garh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10202': {'supplierCode': 'DR10459', 'supplierName': 'Placid', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10203': {'supplierCode': 'DR10175', 'supplierName': 'Madin', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10205': {'supplierCode': 'DR10596', 'supplierName': 'Udai Bilas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10211': {'supplierCode': 'DR10290', 'supplierName': 'Souvenir Peppermint', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10217': {'supplierCode': 'DR10728', 'supplierName': 'Olive Brook', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10218': {'supplierCode': 'DR10552', 'supplierName': 'Pai Vista', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10219': {'supplierCode': 'DR10583', 'supplierName': 'The Royal Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10220': {'supplierCode': 'DR10614', 'supplierName': 'Naveen', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10224': {'supplierCode': 'DR10410', 'supplierName': 'Xandari Harbour', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10227': {'supplierCode': 'DR10278', 'supplierName': 'Khas Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10230': {'supplierCode': 'DR10576', 'supplierName': 'Castle Bijaipur', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10233': {'supplierCode': 'DR10689', 'supplierName': 'Gateway Chikmagalur - IHCL SeleQtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10241': {'supplierCode': 'DR10746', 'supplierName': 'Clarke\'s Oberoi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10242': {'supplierCode': 'DR10717', 'supplierName': 'WelcomHeritage Traditional Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10245': {'supplierCode': 'DR10726', 'supplierName': 'GRT Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10247': {'supplierCode': 'DR10412', 'supplierName': 'The Elgin Fairlawn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10249': {'supplierCode': 'DR10605', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10250': {'supplierCode': 'DR10549', 'supplierName': 'Castle Bera', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10251': {'supplierCode': 'DR10109', 'supplierName': 'Karni Bhawan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10252': {'supplierCode': 'DR10291', 'supplierName': 'Heritage Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10255': {'supplierCode': 'DR10571', 'supplierName': 'Mayfair Lagoon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10256': {'supplierCode': 'DR10572', 'supplierName': 'Mayfair Spa Resort & Casino', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10260': {'supplierCode': 'DR10707', 'supplierName': 'Le Meridien', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10261': {'supplierCode': 'DR10720', 'supplierName': 'Mayfair Himalayan Spa Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10263': {'supplierCode': 'DR10431', 'supplierName': 'Marine Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10267': {'supplierCode': 'DR10345', 'supplierName': 'Zone By The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10268': {'supplierCode': 'DR10567', 'supplierName': 'Chettinadu Mansion', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10271': {'supplierCode': 'DR10743', 'supplierName': 'Aman-I-Khas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10272': {'supplierCode': 'DR10363', 'supplierName': 'Ramada', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10274': {'supplierCode': 'DR10804', 'supplierName': 'Treehouse Queens Pearl', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10280': {'supplierCode': 'DR10541', 'supplierName': 'The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10281': {'supplierCode': 'DR10751', 'supplierName': 'Asia Health Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10287': {'supplierCode': 'DR10778', 'supplierName': 'Jaypee Siddharth', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10291': {'supplierCode': 'DR10692', 'supplierName': 'The Poovath Beach Front Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10296': {'supplierCode': 'DR10830', 'supplierName': 'Hadoti Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10297': {'supplierCode': 'DR10155', 'supplierName': 'Crystal Sarovar Premiere', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10299': {'supplierCode': 'DR10238', 'supplierName': '28 Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10301': {'supplierCode': 'DR10461', 'supplierName': 'Wayanad Wild', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10307': {'supplierCode': 'DR10373', 'supplierName': 'Park Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10314': {'supplierCode': 'DR10264', 'supplierName': 'El Paradiso', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10317': {'supplierCode': 'DR10328', 'supplierName': 'Trident', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10319': {'supplierCode': 'DR10297', 'supplierName': 'The Postcard Mandalay Hall', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10320': {'supplierCode': 'DR10779', 'supplierName': 'Thikana', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10322': {'supplierCode': 'DR10608', 'supplierName': 'Taj Theog Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10331': {'supplierCode': 'DR10241', 'supplierName': 'Lakes and Lagoon - 2 BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10335': {'supplierCode': 'DR10130', 'supplierName': 'Red Fox', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10337': {'supplierCode': 'DR10716', 'supplierName': 'Red Fox', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10338': {'supplierCode': 'DR10241', 'supplierName': 'Lakes and Lagoon - 3 BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10339': {'supplierCode': 'DR10241', 'supplierName': 'Lakes And Lagoon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10340': {'supplierCode': 'DR10241', 'supplierName': 'Lakes and Lagoon - 2 BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10341': {'supplierCode': 'DR10241', 'supplierName': 'Lakes and Lagoon - 3 BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10342': {'supplierCode': 'DR10813', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10344': {'supplierCode': 'DR10620', 'supplierName': 'Park Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10345': {'supplierCode': 'DR10819', 'supplierName': 'The Promenade', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10348': {'supplierCode': 'DR10244', 'supplierName': 'Emerald Isle', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10349': {'supplierCode': 'DR10742', 'supplierName': 'Abhyaran Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10351': {'supplierCode': 'DR10858', 'supplierName': 'Achrol Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10353': {'supplierCode': 'DR10839', 'supplierName': 'Golden Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10358': {'supplierCode': 'DR10817', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10360': {'supplierCode': 'DR10859', 'supplierName': 'Chancery Pavilion', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10361': {'supplierCode': 'DR10764', 'supplierName': 'Chettinadu Court', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10362': {'supplierCode': 'DR10824', 'supplierName': 'Fortune Park Moksha', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10364': {'supplierCode': 'DR10853', 'supplierName': 'The Fern Ranthambhore Forest Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10367': {'supplierCode': 'DR10266', 'supplierName': 'Golden Tulip Essential', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10368': {'supplierCode': 'DR10874', 'supplierName': 'Gokulam Grand Hotel & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10374': {'supplierCode': 'DR10766', 'supplierName': 'JP Cordial', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10375': {'supplierCode': 'DR10566', 'supplierName': 'Visalam', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10377': {'supplierCode': 'DR10849', 'supplierName': 'Vesta International', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10380': {'supplierCode': 'DR10402', 'supplierName': 'Sawai Madhopur Lodge - IHCL SeleQtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10381': {'supplierCode': 'DR10852', 'supplierName': 'Vesta Maurya Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10384': {'supplierCode': 'DR10483', 'supplierName': 'The Gordon House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10385': {'supplierCode': 'DR10496', 'supplierName': 'Radisson Blu Atria', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10386': {'supplierCode': 'DR10905', 'supplierName': 'Royal Orchid', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10388': {'supplierCode': 'DR10771', 'supplierName': 'The Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10389': {'supplierCode': 'DR10823', 'supplierName': 'Royal Orchid Central', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10392': {'supplierCode': 'DR10889', 'supplierName': 'Forest Canopy', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10394': {'supplierCode': 'DR10897', 'supplierName': 'Vivanta (IT Expressway)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10396': {'supplierCode': 'DR10504', 'supplierName': 'Shangri-La', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10398': {'supplierCode': 'DR10163', 'supplierName': 'Fortune Pandiyan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10408': {'supplierCode': 'DR10124', 'supplierName': 'Karohi Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10410': {'supplierCode': 'DR10320', 'supplierName': 'Golden Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10413': {'supplierCode': 'DR10309', 'supplierName': 'The Gulaal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10415': {'supplierCode': 'DR10174', 'supplierName': 'Mahabs', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10419': {'supplierCode': 'DR10132', 'supplierName': 'Madri Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10420': {'supplierCode': 'DR10863', 'supplierName': 'Bijolai Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10421': {'supplierCode': 'DR10332', 'supplierName': 'Devi Bhawan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10423': {'supplierCode': 'DR10222', 'supplierName': 'The Neem Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10424': {'supplierCode': 'DR10334', 'supplierName': 'Haveli Inn Pal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10428': {'supplierCode': 'DR10335', 'supplierName': 'Indana Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10429': {'supplierCode': 'DR10635', 'supplierName': 'Desert N Dunes', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10433': {'supplierCode': 'DR10901', 'supplierName': 'Trident', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10434': {'supplierCode': 'DR10336', 'supplierName': 'Lords Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10435': {'supplierCode': 'DR10226', 'supplierName': 'Maharaja Ganga Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10438': {'supplierCode': 'DR10343', 'supplierName': 'The Almond Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10439': {'supplierCode': 'DR10513', 'supplierName': 'The Orchid', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10440': {'supplierCode': 'DR10769', 'supplierName': 'Lemon Tree (Guindy)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10443': {'supplierCode': 'DR10881', 'supplierName': 'Le Dupleix', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10444': {'supplierCode': 'DR10584', 'supplierName': 'Chonor House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10451': {'supplierCode': 'DR10117', 'supplierName': 'Garden Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10452': {'supplierCode': 'DR10708', 'supplierName': 'Lemon Tree (Sector 60)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10453': {'supplierCode': 'DR10578', 'supplierName': 'Gajner Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10455': {'supplierCode': 'DR10799', 'supplierName': 'Windflower Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10458': {'supplierCode': 'DR10963', 'supplierName': 'Suroth Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10462': {'supplierCode': 'DR10832', 'supplierName': 'Abad Green Forest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10466': {'supplierCode': 'DR10924', 'supplierName': 'Dune De L\'Orient', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10467': {'supplierCode': 'DR10930', 'supplierName': 'A Beach Symphony', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10470': {'supplierCode': 'DR10857', 'supplierName': 'Guleria Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10471': {'supplierCode': 'DR10025', 'supplierName': 'Baymaas Lakehouse', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10472': {'supplierCode': 'DR10932', 'supplierName': 'SUJAN Sher Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10473': {'supplierCode': 'DR10929', 'supplierName': 'Uday Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10478': {'supplierCode': 'DR10925', 'supplierName': 'Les Hibiscus', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10479': {'supplierCode': 'DR10926', 'supplierName': 'Palais De Mahe', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10482': {'supplierCode': 'DR10975', 'supplierName': 'Dune Eco Village & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10483': {'supplierCode': 'DR10933', 'supplierName': 'Sterling Ooty - Fern Hills', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10484': {'supplierCode': 'DR10073', 'supplierName': 'Bamboo Saa Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10488': {'supplierCode': 'DR10960', 'supplierName': 'Regenta Resort Vanya Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10491': {'supplierCode': 'DR10957', 'supplierName': 'Mihir Garh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10494': {'supplierCode': 'DR10890', 'supplierName': 'Ishwari Niwas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10496': {'supplierCode': 'DR10744', 'supplierName': 'Sawai Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10498': {'supplierCode': 'DR10943', 'supplierName': 'Jagat Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10501': {'supplierCode': 'DR10887', 'supplierName': 'Bundi Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10502': {'supplierCode': 'DR10949', 'supplierName': 'Haveli Braj Bhushanjee', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10503': {'supplierCode': 'DR10931', 'supplierName': 'Shahpura Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10505': {'supplierCode': 'DR10961', 'supplierName': 'Royal Desert Safaries', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10506': {'supplierCode': 'DR10941', 'supplierName': 'Pushkar Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10510': {'supplierCode': 'DR10153', 'supplierName': 'Lemon Tree Premier', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10515': {'supplierCode': 'DR10940', 'supplierName': 'Royal Desert Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10516': {'supplierCode': 'DR10152', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10517': {'supplierCode': 'DR10740', 'supplierName': 'King\'s Abode', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10518': {'supplierCode': 'DR10989', 'supplierName': 'Amrit Bhawan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10519': {'supplierCode': 'DR10586', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10520': {'supplierCode': 'DR10712', 'supplierName': 'Ashhok', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10522': {'supplierCode': 'DR10937', 'supplierName': 'Juna Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10523': {'supplierCode': 'DR10831', 'supplierName': 'Ramathra Fort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10524': {'supplierCode': 'DR10628', 'supplierName': 'Saraca', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10525': {'supplierCode': 'DR10951', 'supplierName': 'Aranya Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10526': {'supplierCode': 'DR10597', 'supplierName': 'Rajakkad Estate', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10527': {'supplierCode': 'DR10321', 'supplierName': 'Jasmin Home', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10530': {'supplierCode': 'DR10698', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10532': {'supplierCode': 'DR10411', 'supplierName': 'Aramness', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10533': {'supplierCode': 'DR10415', 'supplierName': 'The Postcard Gir', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10534': {'supplierCode': 'DR10734', 'supplierName': 'Lemon Tree Premier', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10535': {'supplierCode': 'DR10562', 'supplierName': 'Lemon Tree Amarante Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10536': {'supplierCode': 'DR10735', 'supplierName': 'Coco Lagoon by Great Mount', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10538': {'supplierCode': 'DR10486', 'supplierName': 'The Legend', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10539': {'supplierCode': 'DR10200', 'supplierName': 'Red Fox', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10540': {'supplierCode': 'DR10625', 'supplierName': 'Philipkutty\'s Farm', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10541': {'supplierCode': 'DR10031', 'supplierName': 'Windamere', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10542': {'supplierCode': 'DR10601', 'supplierName': 'Reni Pani Jungle Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10546': {'supplierCode': 'DR10577', 'supplierName': 'Kanha Jungle Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10550': {'supplierCode': 'DR10948', 'supplierName': 'Lemon Tree Vembanad Lake Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10552': {'supplierCode': 'DR11037', 'supplierName': 'Dev Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10553': {'supplierCode': 'DR10736', 'supplierName': 'Isola Di Cocco', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10556': {'supplierCode': 'DR10935', 'supplierName': 'Ranthambhore Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10557': {'supplierCode': 'DR10248', 'supplierName': 'Nelpura Heritage Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10559': {'supplierCode': 'DR10953', 'supplierName': 'Radha Regent', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10562': {'supplierCode': 'DR10976', 'supplierName': 'La Villa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10563': {'supplierCode': 'DR10855', 'supplierName': 'Astoria', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10564': {'supplierCode': 'DR10838', 'supplierName': 'Welcomhotel By ITC Hotels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10566': {'supplierCode': 'DR11006', 'supplierName': 'Zone By The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10567': {'supplierCode': 'DR11017', 'supplierName': 'Royal Orchid Central', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10568': {'supplierCode': 'DR11020', 'supplierName': 'The Fern An Ecotel Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10570': {'supplierCode': 'DR10409', 'supplierName': 'Comfort Inn Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10571': {'supplierCode': 'DR11014', 'supplierName': 'The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10572': {'supplierCode': 'DR10611', 'supplierName': 'Woodville Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10573': {'supplierCode': 'DR11007', 'supplierName': 'Marari Villas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10576': {'supplierCode': 'DR11010', 'supplierName': 'The Pride', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10577': {'supplierCode': 'DR11023', 'supplierName': 'Poetree Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10580': {'supplierCode': 'DR10198', 'supplierName': 'Ramyas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10581': {'supplierCode': 'DR10860', 'supplierName': 'Welcomhotel By ITC Hotels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10582': {'supplierCode': 'DR10985', 'supplierName': 'Kaner Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10584': {'supplierCode': 'DR11009', 'supplierName': 'Karni Bhawan Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10587': {'supplierCode': 'DR11018', 'supplierName': 'Mistletoe', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10588': {'supplierCode': 'DR11011', 'supplierName': 'Zone By The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10589': {'supplierCode': 'DR11027', 'supplierName': 'Hilton Garden Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10590': {'supplierCode': 'DR11033', 'supplierName': 'Taj Chia Kutir Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10592': {'supplierCode': 'DR11030', 'supplierName': 'Hyatt Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10593': {'supplierCode': 'DR10988', 'supplierName': 'Kot Dunara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10594': {'supplierCode': 'DR10954', 'supplierName': 'Talabgaon Castle', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10595': {'supplierCode': 'DR10886', 'supplierName': 'Bundi Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10596': {'supplierCode': 'DR11029', 'supplierName': 'Ikhaya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10598': {'supplierCode': 'DR10756', 'supplierName': 'Blue Jelly', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10599': {'supplierCode': 'DR10756', 'supplierName': 'Blue Jelly-2BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10601': {'supplierCode': 'DR10254', 'supplierName': 'RKV-1BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10602': {'supplierCode': 'DR10756', 'supplierName': 'Blue Jelly-2BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10603': {'supplierCode': 'DR10254', 'supplierName': 'RKV-1BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10604': {'supplierCode': 'DR10756', 'supplierName': 'Blue Jelly-3BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10605': {'supplierCode': 'DR10756', 'supplierName': 'Blue Jelly-3BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10606': {'supplierCode': 'DR10006', 'supplierName': 'White Water Cruise-2BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10607': {'supplierCode': 'DR10254', 'supplierName': 'RKV-2BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10608': {'supplierCode': 'DR10254', 'supplierName': 'RKV-2BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10609': {'supplierCode': 'DR10254', 'supplierName': 'RKV-3BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10610': {'supplierCode': 'DR10254', 'supplierName': 'RKV-3BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10611': {'supplierCode': 'DR10006', 'supplierName': 'White Water Cruise-1BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10612': {'supplierCode': 'DR10649', 'supplierName': 'Xandari Riverscapes', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10613': {'supplierCode': 'DR10649', 'supplierName': 'Xandari Riverscapes', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10614': {'supplierCode': 'DR10006', 'supplierName': 'White Water Cruise-2BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10615': {'supplierCode': 'DR10649', 'supplierName': 'Xandari Riverscapes-2BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10616': {'supplierCode': 'DR10649', 'supplierName': 'Xandari Riverscapes-2BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10617': {'supplierCode': 'DR10006', 'supplierName': 'White Water Cruise-3BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10618': {'supplierCode': 'DR10649', 'supplierName': 'Xandari Riverscapes-3BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10619': {'supplierCode': 'DR10649', 'supplierName': 'Xandari Riverscapes-3BRH', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10620': {'supplierCode': 'DR10432', 'supplierName': 'The Oberoi Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10621': {'supplierCode': 'DR11045', 'supplierName': 'Coral House Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10624': {'supplierCode': 'DR11056', 'supplierName': 'Clarks Safari', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10626': {'supplierCode': 'DR10991', 'supplierName': 'Taj City Centre New Town', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10627': {'supplierCode': 'DR10434', 'supplierName': 'The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10628': {'supplierCode': 'DR11022', 'supplierName': 'Shahpura Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10632': {'supplierCode': 'DR11036', 'supplierName': 'Rann Riders By Kaafila', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10633': {'supplierCode': 'DR10436', 'supplierName': 'The Westin (Rajarhat)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10634': {'supplierCode': 'DR10695', 'supplierName': 'Umaid Lake Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10636': {'supplierCode': 'DR11024', 'supplierName': 'Garh Jaisal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10639': {'supplierCode': 'DR10449', 'supplierName': 'Mantra Koodam', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10641': {'supplierCode': 'DR11040', 'supplierName': 'Justa Rajputana Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10642': {'supplierCode': 'DR11042', 'supplierName': 'Justa Sajjangarh Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10643': {'supplierCode': 'DR10820', 'supplierName': 'Trulyy Rudransh Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10645': {'supplierCode': 'DR10359', 'supplierName': 'Courtyard By Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10647': {'supplierCode': 'DR10550', 'supplierName': 'Gitanjali Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10649': {'supplierCode': 'DR10808', 'supplierName': 'Niraamaya Retreats Cardamom Club', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10651': {'supplierCode': 'DR10405', 'supplierName': 'Hindusthan International', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10652': {'supplierCode': 'DR10318', 'supplierName': 'The Grand Haveli and Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10654': {'supplierCode': 'DR10531', 'supplierName': 'Mahua Kothi, A Taj Safari', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10657': {'supplierCode': 'DR10228', 'supplierName': 'The Udai Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10658': {'supplierCode': 'DR11049', 'supplierName': 'Tree House Hideaway', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10659': {'supplierCode': 'DR10987', 'supplierName': 'Chandela', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10660': {'supplierCode': 'DR11055', 'supplierName': 'The Elgin Mount Pandim', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10661': {'supplierCode': 'DR11047', 'supplierName': 'Bandhav Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10662': {'supplierCode': 'DR11053', 'supplierName': 'Chitvan Jungle Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10664': {'supplierCode': 'DR10530', 'supplierName': 'Royal Safari Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10665': {'supplierCode': 'DR10958', 'supplierName': 'Exotic Luxury Camps', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10666': {'supplierCode': 'DR10397', 'supplierName': 'The Sher Garh Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10669': {'supplierCode': 'DR11028', 'supplierName': 'Chunda Shikar Oudi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10670': {'supplierCode': 'DR11031', 'supplierName': 'Renest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10671': {'supplierCode': 'DR11097', 'supplierName': 'Tuli Tiger Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10672': {'supplierCode': 'DR11091', 'supplierName': 'Sterling Goa Varca Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10673': {'supplierCode': 'DR11067', 'supplierName': 'Comfort Inn Insys', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10674': {'supplierCode': 'DR11098', 'supplierName': 'Tuli Tiger Corridor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10676': {'supplierCode': 'DR11054', 'supplierName': 'Kairali Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10677': {'supplierCode': 'DR10846', 'supplierName': 'Chateau Garli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10678': {'supplierCode': 'DR11079', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10680': {'supplierCode': 'DR11088', 'supplierName': 'Mayfair Waves', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10681': {'supplierCode': 'DR11080', 'supplierName': 'ITC Sonar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10683': {'supplierCode': 'DR10738', 'supplierName': 'Indo Hokke', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10685': {'supplierCode': 'DR11090', 'supplierName': 'The Pride', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10686': {'supplierCode': 'DR11078', 'supplierName': 'Lemon Tree Premier', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10688': {'supplierCode': 'DR10543', 'supplierName': 'Vivanta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10690': {'supplierCode': 'DR11096', 'supplierName': 'Raj Niwas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10691': {'supplierCode': 'DR10624', 'supplierName': 'The Fern Gir Forest Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10692': {'supplierCode': 'DR10554', 'supplierName': 'Red Earth', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10693': {'supplierCode': 'DR11077', 'supplierName': 'Comfort Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10694': {'supplierCode': 'DR11074', 'supplierName': 'Windflower Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10695': {'supplierCode': 'DR10439', 'supplierName': 'Windermere River House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10696': {'supplierCode': 'DR11093', 'supplierName': 'The Serai', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10697': {'supplierCode': 'DR11069', 'supplierName': 'Sarovar Portico Naraina', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10699': {'supplierCode': 'DR11094', 'supplierName': 'Ken River Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10703': {'supplierCode': 'DR10784', 'supplierName': 'Hindusthan International', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10704': {'supplierCode': 'DR11076', 'supplierName': 'Fortune Miramar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10705': {'supplierCode': 'DR10731', 'supplierName': 'Jawai Sagar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10706': {'supplierCode': 'DR10637', 'supplierName': 'Manvar Shergarh, The Desert Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10708': {'supplierCode': 'DR11070', 'supplierName': 'Anopura', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10709': {'supplierCode': 'DR10440', 'supplierName': 'Novotel (Juhu Beach)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10710': {'supplierCode': 'DR10754', 'supplierName': 'Sofitel Mumbai BKC', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10711': {'supplierCode': 'DR10995', 'supplierName': 'Anand Kashi By The Ganges - IHCL SeleQtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10712': {'supplierCode': 'DR11168', 'supplierName': 'GRT Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10714': {'supplierCode': 'DR11104', 'supplierName': 'Neemrana Fort Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10715': {'supplierCode': 'DR11145', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10718': {'supplierCode': 'DR11101', 'supplierName': 'Lemon Tree Premier', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10719': {'supplierCode': 'DR10777', 'supplierName': 'Red Fox (East Delhi)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10721': {'supplierCode': 'DR10971', 'supplierName': 'Lemon Tree Premier (Leisure Valley 2)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10722': {'supplierCode': 'DR10016', 'supplierName': 'Saratha Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10723': {'supplierCode': 'DR10793', 'supplierName': 'Lemon Tree Premiere (City Center)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10724': {'supplierCode': 'DR11138', 'supplierName': 'Lemon Tree (Electronics City)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10725': {'supplierCode': 'DR10368', 'supplierName': 'Taurus Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10726': {'supplierCode': 'DR11148', 'supplierName': 'Lemon Tree (Whitefield)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10728': {'supplierCode': 'DR10565', 'supplierName': 'Ibis', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10729': {'supplierCode': 'DR10871', 'supplierName': 'Bloom Suites (Electronic City)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10730': {'supplierCode': 'DR11154', 'supplierName': 'Harivihar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10732': {'supplierCode': 'DR10822', 'supplierName': 'Red Fox (Sector 60)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10735': {'supplierCode': 'DR11099', 'supplierName': 'The Raviz Kadavu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10736': {'supplierCode': 'DR11073', 'supplierName': 'Malligi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10737': {'supplierCode': 'DR11162', 'supplierName': 'Vivanta Sikkim', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10738': {'supplierCode': 'DR11100', 'supplierName': 'Regenta Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10739': {'supplierCode': 'DR11120', 'supplierName': 'Shikarbadi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10740': {'supplierCode': 'DR11137', 'supplierName': 'Red Fox', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10741': {'supplierCode': 'DR11124', 'supplierName': 'The Royal Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10742': {'supplierCode': 'DR11140', 'supplierName': 'Poovar Island Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10743': {'supplierCode': 'DR10876', 'supplierName': 'Hilton (Embassy Golflinks)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10744': {'supplierCode': 'DR10607', 'supplierName': 'Brij Gaj Kesri', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10745': {'supplierCode': 'DR11146', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10746': {'supplierCode': 'DR10767', 'supplierName': 'Hilton', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10747': {'supplierCode': 'DR10148', 'supplierName': 'Lemon Tree Premier (Leisure Valley 1)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10748': {'supplierCode': 'DR10485', 'supplierName': 'Hilton Garden Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10749': {'supplierCode': 'DR11123', 'supplierName': 'Hindusthan International', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10750': {'supplierCode': 'DR11086', 'supplierName': 'Mayfair Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10751': {'supplierCode': 'DR10837', 'supplierName': 'Hilton Garden Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10752': {'supplierCode': 'DR10677', 'supplierName': 'Hyatt Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10753': {'supplierCode': 'DR10627', 'supplierName': 'La Place Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10754': {'supplierCode': 'DR10360', 'supplierName': 'Crowne Plaza (Adyar Park)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10756': {'supplierCode': 'DR11144', 'supplierName': 'Akhil Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10758': {'supplierCode': 'DR10750', 'supplierName': 'Alsisar Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10759': {'supplierCode': 'DR11167', 'supplierName': 'Abad Brookside', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10760': {'supplierCode': 'DR10877', 'supplierName': 'Hyatt Centric', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10761': {'supplierCode': 'DR11158', 'supplierName': 'Green Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10764': {'supplierCode': 'DR11152', 'supplierName': 'Hyatt Centric', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10765': {'supplierCode': 'DR11102', 'supplierName': 'Sterling Wayanand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10766': {'supplierCode': 'DR10685', 'supplierName': 'Ramada Plaza (Guindy)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10767': {'supplierCode': 'DR11149', 'supplierName': 'MGM Beach Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10768': {'supplierCode': 'DR10509', 'supplierName': 'Hyatt Centric', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10769': {'supplierCode': 'DR10792', 'supplierName': 'Hyatt Place', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10770': {'supplierCode': 'DR10994', 'supplierName': 'Taj tirupati', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10772': {'supplierCode': 'DR11111', 'supplierName': 'Hill Fort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10774': {'supplierCode': 'DR10351', 'supplierName': 'Desert Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10775': {'supplierCode': 'DR11108', 'supplierName': 'The Bagheera', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10776': {'supplierCode': 'DR10473', 'supplierName': 'Pepper County Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10777': {'supplierCode': 'DR11112', 'supplierName': 'Pench Tree Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10778': {'supplierCode': 'DR10347', 'supplierName': 'Sairafort Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10780': {'supplierCode': 'DR10729', 'supplierName': 'Kanan Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10784': {'supplierCode': 'DR11065', 'supplierName': 'Holy City Paradise Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10785': {'supplierCode': 'DR10525', 'supplierName': '7 Apple', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10790': {'supplierCode': 'DR10809', 'supplierName': 'Harsoli Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10793': {'supplierCode': 'DR11106', 'supplierName': 'Ranvas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10794': {'supplierCode': 'DR11129', 'supplierName': 'The Earth', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10796': {'supplierCode': 'DR11122', 'supplierName': 'Paatlidun Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10797': {'supplierCode': 'DR11115', 'supplierName': 'Godwad Safari Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10798': {'supplierCode': 'DR11118', 'supplierName': 'Bijapur Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10799': {'supplierCode': 'DR10814', 'supplierName': 'Fifu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10800': {'supplierCode': 'DR10602', 'supplierName': 'Fort Khejarla', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10801': {'supplierCode': 'DR11161', 'supplierName': 'Chapslee', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10802': {'supplierCode': 'DR10451', 'supplierName': 'Fateh Safari Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10803': {'supplierCode': 'DR11141', 'supplierName': 'Dera Masuda', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10804': {'supplierCode': 'DR11135', 'supplierName': 'Sarai At Toria', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10805': {'supplierCode': 'DR10593', 'supplierName': 'Surya Mcleod', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10806': {'supplierCode': 'DR11128', 'supplierName': 'Sardargarh Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10808': {'supplierCode': 'DR11126', 'supplierName': 'Fateh Safari Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10809': {'supplierCode': 'DR11125', 'supplierName': 'Punjab Village Farm', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10810': {'supplierCode': 'DR11127', 'supplierName': 'Grand Dragon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10811': {'supplierCode': 'DR11169', 'supplierName': '5 By OYO Metropolitan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10813': {'supplierCode': 'DR10574', 'supplierName': 'Kutch Safari Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10815': {'supplierCode': 'DR11113', 'supplierName': 'Bijay Niwas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10816': {'supplierCode': 'DR10730', 'supplierName': 'Lodge At Wah', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10817': {'supplierCode': 'DR11157', 'supplierName': 'Fateh Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10819': {'supplierCode': 'DR11155', 'supplierName': 'Malji Ka Kamra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10824': {'supplierCode': 'DR11160', 'supplierName': 'Sonaar Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10825': {'supplierCode': 'DR11142', 'supplierName': 'Vasundhara Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10826': {'supplierCode': 'DR11114', 'supplierName': '1st Gate Home-Fusion', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10827': {'supplierCode': 'DR11159', 'supplierName': 'Lchang Nang Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10830': {'supplierCode': 'DR10789', 'supplierName': 'Crowne Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10831': {'supplierCode': 'DR11116', 'supplierName': 'Killa Bhawan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10833': {'supplierCode': 'DR11153', 'supplierName': 'Pench Jungle Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10834': {'supplierCode': 'DR10752', 'supplierName': 'Intercontinental Marine Drive', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10836': {'supplierCode': 'DR10519', 'supplierName': 'Fairfield By Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10837': {'supplierCode': 'DR10598', 'supplierName': 'Cardamom House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10840': {'supplierCode': 'DR10490', 'supplierName': 'Fairfield By Marriott (Rajajinagar)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10847': {'supplierCode': 'DR10802', 'supplierName': 'The Westin', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10849': {'supplierCode': 'DR10810', 'supplierName': 'Le Meridien', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10851': {'supplierCode': 'DR10868', 'supplierName': 'Pagoda Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10855': {'supplierCode': 'DR10970', 'supplierName': 'Courtyard Mumbai International Airport', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10856': {'supplierCode': 'DR10502', 'supplierName': 'Hyatt Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10859': {'supplierCode': 'DR11084', 'supplierName': 'Courtyard By Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10860': {'supplierCode': 'DR10407', 'supplierName': 'JW Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10862': {'supplierCode': 'DR10256', 'supplierName': 'Fairmont', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10863': {'supplierCode': 'DR10681', 'supplierName': 'Purple Beds by Vits', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10864': {'supplierCode': 'DR10526', 'supplierName': 'Ambassador Ajanta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10865': {'supplierCode': 'DR11105', 'supplierName': 'Windflower Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10866': {'supplierCode': 'DR10063', 'supplierName': 'The Ananta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10873': {'supplierCode': 'DR10383', 'supplierName': 'Radisson Blu (Paschim Vihar)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10874': {'supplierCode': 'DR10030', 'supplierName': 'Maha Bodhi Hotel, Resort And Convention Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10877': {'supplierCode': 'DR10508', 'supplierName': 'Mizpah', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10879': {'supplierCode': 'DR10585', 'supplierName': 'JW Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10880': {'supplierCode': 'DR10280', 'supplierName': 'Madhuban', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10882': {'supplierCode': 'DR10103', 'supplierName': 'Umaid Bhawan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10883': {'supplierCode': 'DR10303', 'supplierName': 'Ramada By Wyndham (Raja Park)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10884': {'supplierCode': 'DR10113', 'supplierName': 'Crowne Plaza (Mayur Vihar)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10885': {'supplierCode': 'DR10899', 'supplierName': 'Seven Palms Desert Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10891': {'supplierCode': 'DR10841', 'supplierName': 'Manama', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10892': {'supplierCode': 'DR10803', 'supplierName': 'Krishna Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10893': {'supplierCode': 'DR10170', 'supplierName': 'Chariot Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10894': {'supplierCode': 'DR10956', 'supplierName': 'Ramada By Wyndham (Jaisinghpura)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10897': {'supplierCode': 'DR10885', 'supplierName': 'The Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10903': {'supplierCode': 'DR10128', 'supplierName': 'Leela Palace (UNIRE Testing)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10917': {'supplierCode': 'DR11243', 'supplierName': 'Test NG', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10918': {'supplierCode': 'DR10516', 'supplierName': 'The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10919': {'supplierCode': 'DR10599', 'supplierName': 'Narsingh Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10920': {'supplierCode': 'DR11172', 'supplierName': 'Sariska Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10921': {'supplierCode': 'DR11204', 'supplierName': 'Surasena Regal Vista', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10922': {'supplierCode': 'DR10848', 'supplierName': 'Pushkar Fort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10923': {'supplierCode': 'DR11207', 'supplierName': 'Chandra Mahal Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10924': {'supplierCode': 'DR11197', 'supplierName': 'Chamba Camp Diskit', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10925': {'supplierCode': 'DR10590', 'supplierName': 'Grace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10926': {'supplierCode': 'DR10194', 'supplierName': 'Costa Riviera', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10927': {'supplierCode': 'DR11260', 'supplierName': 'The House Of MG', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10928': {'supplierCode': 'DR11200', 'supplierName': 'Lords Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10929': {'supplierCode': 'DR11198', 'supplierName': 'Chamba Camp Thiksey', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10931': {'supplierCode': 'DR11039', 'supplierName': 'The Johri', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10932': {'supplierCode': 'DR10812', 'supplierName': 'WelcomHeritage Sirsi Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10933': {'supplierCode': 'DR10587', 'supplierName': 'Clouds End Villa Cottages', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10934': {'supplierCode': 'DR10216', 'supplierName': 'Amargarh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10936': {'supplierCode': 'DR10579', 'supplierName': 'Kipling Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10937': {'supplierCode': 'DR11170', 'supplierName': 'Windsor Heights', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10939': {'supplierCode': 'DR11202', 'supplierName': 'Tuli Tiger Corridor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10940': {'supplierCode': 'DR11176', 'supplierName': 'The Astor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10941': {'supplierCode': 'DR11201', 'supplierName': 'Tuli Imperial', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10942': {'supplierCode': 'DR11057', 'supplierName': 'Southern Star', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10943': {'supplierCode': 'DR10242', 'supplierName': 'Bissau Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10945': {'supplierCode': 'DR11181', 'supplierName': 'Mount Valley Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10948': {'supplierCode': 'DR11219', 'supplierName': 'Rawai Luxury Tents', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10949': {'supplierCode': 'DR11196', 'supplierName': 'Jungle House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10951': {'supplierCode': 'DR10845', 'supplierName': 'Orchha Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10952': {'supplierCode': 'DR11085', 'supplierName': 'The Ritz - Carlton', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10953': {'supplierCode': 'DR10631', 'supplierName': 'Ahilya fort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10954': {'supplierCode': 'DR10844', 'supplierName': 'Amaryllis', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10956': {'supplierCode': 'DR10316', 'supplierName': 'Orchha Palace & Convention Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10961': {'supplierCode': 'DR11192', 'supplierName': 'Om Vilas Benares', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10963': {'supplierCode': 'DR11147', 'supplierName': 'Lemon tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10964': {'supplierCode': 'DR11318', 'supplierName': 'Yois', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10965': {'supplierCode': 'DR11328', 'supplierName': 'Woods at Sasan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10966': {'supplierCode': 'DR11554', 'supplierName': 'Ranthambore Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10968': {'supplierCode': 'DR11175', 'supplierName': 'Baragarh Resort & Spa - IHCL SeleQtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10969': {'supplierCode': 'DR11267', 'supplierName': 'The Belgadia Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10970': {'supplierCode': 'DR10613', 'supplierName': 'Haveli Hari Ganga', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10971': {'supplierCode': 'DR11215', 'supplierName': 'Fortune Landmark', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10972': {'supplierCode': 'DR11300', 'supplierName': 'Goldfinch Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10973': {'supplierCode': 'DR11325', 'supplierName': 'WelcomHeritage Jungle Home', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10974': {'supplierCode': 'DR10641', 'supplierName': 'Connaught House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10975': {'supplierCode': 'DR11208', 'supplierName': 'Fortune Select S.G. Highway', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10976': {'supplierCode': 'DR10534', 'supplierName': 'Barefoot Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10979': {'supplierCode': 'DR10741', 'supplierName': 'Maharani Bagh Orchard Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10980': {'supplierCode': 'DR11290', 'supplierName': 'HK Clarks Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10982': {'supplierCode': 'DR11313', 'supplierName': 'Sandesh Pride', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10983': {'supplierCode': 'DR11266', 'supplierName': 'A Home For Nature Lover', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10984': {'supplierCode': 'DR10481', 'supplierName': 'Sinclairs', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10986': {'supplierCode': 'DR10521', 'supplierName': 'Park Inn By Radisson', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10987': {'supplierCode': 'DR11278', 'supplierName': 'The Naini Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10988': {'supplierCode': 'DR11262', 'supplierName': 'Abad Metro', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10989': {'supplierCode': 'DR11214', 'supplierName': 'Svasara Jungle Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10990': {'supplierCode': 'DR11263', 'supplierName': 'Dutch Bungalow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10991': {'supplierCode': 'DR11220', 'supplierName': 'Temi Tea Bungalow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10992': {'supplierCode': 'DR11294', 'supplierName': 'The Bamboo Forest Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10993': {'supplierCode': 'DR11279', 'supplierName': 'The Tower House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10994': {'supplierCode': 'DR11205', 'supplierName': 'Vivanta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10995': {'supplierCode': 'DR10895', 'supplierName': 'Colonel\'s Retreat 2', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10996': {'supplierCode': 'DR10936', 'supplierName': 'Accord', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10998': {'supplierCode': 'DR11296', 'supplierName': 'Vismaya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO10999': {'supplierCode': 'DR11150', 'supplierName': 'Badami Court', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11003': {'supplierCode': 'DR11282', 'supplierName': 'Bandhavgarh Jungle Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11006': {'supplierCode': 'DR11283', 'supplierName': 'Syna Tiger Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11008': {'supplierCode': 'DR11216', 'supplierName': 'Udman (Panchshila Park)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11010': {'supplierCode': 'DR11301', 'supplierName': 'MGM Vailankanni Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11011': {'supplierCode': 'DR11329', 'supplierName': 'Panjim Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11012': {'supplierCode': 'DR11259', 'supplierName': 'Maya Luxury', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11014': {'supplierCode': 'DR11217', 'supplierName': 'Yog Niketan By Sanskriti', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11015': {'supplierCode': 'DR10791', 'supplierName': 'Golden Tulip Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11016': {'supplierCode': 'DR10471', 'supplierName': 'Shahpura Kumbhal Villas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11018': {'supplierCode': 'DR11326', 'supplierName': 'Fernhills Royal Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11021': {'supplierCode': 'DR10092', 'supplierName': 'Shahpura Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11022': {'supplierCode': 'DR10302', 'supplierName': 'Taj Usha Kiran Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11023': {'supplierCode': 'DR11226', 'supplierName': 'Shahpura Dev Panache', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11024': {'supplierCode': 'DR11082', 'supplierName': 'Ganga Lahari', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11025': {'supplierCode': 'DR11221', 'supplierName': 'Pal Garh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11026': {'supplierCode': 'DR11286', 'supplierName': 'Jehan Numa Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11027': {'supplierCode': 'DR10903', 'supplierName': 'Ranbanka Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11028': {'supplierCode': 'DR11231', 'supplierName': 'Tiger\'s Den Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11030': {'supplierCode': 'DR11193', 'supplierName': 'Surya Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11031': {'supplierCode': 'DR11334', 'supplierName': 'Lotus Nikko', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11032': {'supplierCode': 'DR11358', 'supplierName': 'Sterling Kodai Lake', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11033': {'supplierCode': 'DR11434', 'supplierName': 'Pride Ananya Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11034': {'supplierCode': 'DR11428', 'supplierName': 'Mayfair Convention', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11036': {'supplierCode': 'DR11441', 'supplierName': 'Pride', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11039': {'supplierCode': 'DR11382', 'supplierName': 'Sterling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11041': {'supplierCode': 'DR11464', 'supplierName': 'Sinclairs Retreat Dooars', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11042': {'supplierCode': 'DR11381', 'supplierName': 'Sterling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11043': {'supplierCode': 'DR11349', 'supplierName': 'Sterling Ooty Elk Hill', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11044': {'supplierCode': 'DR10923', 'supplierName': 'Pride Kadamb Kunj Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11045': {'supplierCode': 'DR11385', 'supplierName': 'Sinclairs Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11046': {'supplierCode': 'DR10687', 'supplierName': 'The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11047': {'supplierCode': 'DR11372', 'supplierName': 'Sterling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11048': {'supplierCode': 'DR11405', 'supplierName': 'The Pride', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11049': {'supplierCode': 'DR11458', 'supplierName': 'Brij Nest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11050': {'supplierCode': 'DR11388', 'supplierName': 'Infinity Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11051': {'supplierCode': 'DR11406', 'supplierName': 'Pride Sun Village Resort And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11052': {'supplierCode': 'DR11390', 'supplierName': 'Quality Inn Bez Krishnaa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11053': {'supplierCode': 'DR11379', 'supplierName': 'Sterling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11054': {'supplierCode': 'DR11373', 'supplierName': 'Sterling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11055': {'supplierCode': 'DR11408', 'supplierName': 'Comfort Inn President', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11056': {'supplierCode': 'DR11380', 'supplierName': 'Sterling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11057': {'supplierCode': 'DR11424', 'supplierName': 'Pride Amber Vilas Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11058': {'supplierCode': 'DR11368', 'supplierName': 'Clarks Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11059': {'supplierCode': 'DR11375', 'supplierName': 'Sterling - Orange Village', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11060': {'supplierCode': 'DR11336', 'supplierName': 'Comfort Inn Sunset', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11061': {'supplierCode': 'DR11387', 'supplierName': 'Comfort Inn Alstonia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11062': {'supplierCode': 'DR11397', 'supplierName': 'WelcomHeritage Taragarh Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11065': {'supplierCode': 'DR11377', 'supplierName': 'Clarion Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11066': {'supplierCode': 'DR11376', 'supplierName': 'Sterling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11067': {'supplierCode': 'DR11179', 'supplierName': 'Quality Inn VIHA', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11068': {'supplierCode': 'DR10246', 'supplierName': 'Clarks Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11069': {'supplierCode': 'DR11332', 'supplierName': 'The Pride', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11071': {'supplierCode': 'DR11341', 'supplierName': 'The Aodhi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11072': {'supplierCode': 'DR11370', 'supplierName': 'Sterling (Kodai Valley)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11073': {'supplierCode': 'DR11366', 'supplierName': 'Comfort Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11076': {'supplierCode': 'DR11348', 'supplierName': 'The Riverview Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11077': {'supplierCode': 'DR11351', 'supplierName': 'Quality Hotel D V Manor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11078': {'supplierCode': 'DR11384', 'supplierName': 'Intercontinental (Mahabalipuram)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11080': {'supplierCode': 'DR11378', 'supplierName': 'Fortune Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11081': {'supplierCode': 'DR11374', 'supplierName': 'Fortune Inn Promenade', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11085': {'supplierCode': 'DR11422', 'supplierName': 'Fortune Inn Sree Kanya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11086': {'supplierCode': 'DR11447', 'supplierName': 'Vivanta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11087': {'supplierCode': 'DR11361', 'supplierName': 'Fortune Select Trinity', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11088': {'supplierCode': 'DR11457', 'supplierName': 'ITC Royal Bengal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11090': {'supplierCode': 'DR11429', 'supplierName': 'Fortune Inn Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11092': {'supplierCode': 'DR11451', 'supplierName': 'Taj Exotica Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11098': {'supplierCode': 'DR11425', 'supplierName': 'ITC Gardenia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11099': {'supplierCode': 'DR11337', 'supplierName': 'Vivanta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11100': {'supplierCode': 'DR11442', 'supplierName': 'The Leela', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11102': {'supplierCode': 'DR11420', 'supplierName': 'Pashan Garh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11103': {'supplierCode': 'DR11269', 'supplierName': 'The Leela Ashtamudi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11105': {'supplierCode': 'DR10702', 'supplierName': 'Radisson Goa Candolim', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11106': {'supplierCode': 'DR11456', 'supplierName': 'The Postcard Velha', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11107': {'supplierCode': 'DR11367', 'supplierName': 'Fortune Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11108': {'supplierCode': 'DR11445', 'supplierName': 'Storii By ITC Hotels Shanti Morada', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11109': {'supplierCode': 'DR11363', 'supplierName': 'Fortune Park Panchwati', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11112': {'supplierCode': 'DR11364', 'supplierName': 'Fortune Select Exotica', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11113': {'supplierCode': 'DR10616', 'supplierName': 'ITC Kakatiya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11114': {'supplierCode': 'DR11449', 'supplierName': 'Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11115': {'supplierCode': 'DR11403', 'supplierName': 'Blue Diamond - IHCL Seleqtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11116': {'supplierCode': 'DR11401', 'supplierName': 'The Westin Resort & Spa, Himalayas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11117': {'supplierCode': 'DR11335', 'supplierName': 'Gir Serai - IHCL SeleQtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11118': {'supplierCode': 'DR11460', 'supplierName': 'Taj Wayanad Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11119': {'supplierCode': 'DR11327', 'supplierName': 'Ging Tea House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11121': {'supplierCode': 'DR11487', 'supplierName': 'The Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11122': {'supplierCode': 'DR11043', 'supplierName': 'Treehouse Amaara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11123': {'supplierCode': 'DR10479', 'supplierName': 'Grand View', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11124': {'supplierCode': 'DR11513', 'supplierName': 'Pepper Trail', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11125': {'supplierCode': 'DR11352', 'supplierName': 'Indra Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11126': {'supplierCode': 'DR10962', 'supplierName': 'Treehouse Rajbagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11127': {'supplierCode': 'DR11480', 'supplierName': 'Iris Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11128': {'supplierCode': 'DR11552', 'supplierName': 'The Untamed Bandhavgarh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11131': {'supplierCode': 'DR11433', 'supplierName': 'Laxmi Vilas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11132': {'supplierCode': 'DR11503', 'supplierName': 'The Mountbatten Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11133': {'supplierCode': 'DR11339', 'supplierName': 'Bhanwar Vilas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11134': {'supplierCode': 'DR11436', 'supplierName': 'The Ummaid Bagh Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11136': {'supplierCode': 'DR11493', 'supplierName': 'Eagles Nest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11138': {'supplierCode': 'DR11556', 'supplierName': 'Castle Khandela', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11140': {'supplierCode': 'DR10569', 'supplierName': 'Ravla Bhenswara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11141': {'supplierCode': 'DR11276', 'supplierName': 'RAAS', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11142': {'supplierCode': 'DR11342', 'supplierName': 'Classic Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11143': {'supplierCode': 'DR11530', 'supplierName': 'The Lavitra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11145': {'supplierCode': 'DR11505', 'supplierName': 'Chobdar Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11146': {'supplierCode': 'DR11273', 'supplierName': 'Ekaanta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11147': {'supplierCode': 'DR11534', 'supplierName': 'Fateh’s Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11148': {'supplierCode': 'DR11549', 'supplierName': 'Vasant Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11150': {'supplierCode': 'DR11550', 'supplierName': 'Aangan Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11151': {'supplierCode': 'DR11547', 'supplierName': 'Renest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11152': {'supplierCode': 'DR11191', 'supplierName': 'Vatsalya Vihar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11153': {'supplierCode': 'DR11321', 'supplierName': 'Cardinal Express Oxmo', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11154': {'supplierCode': 'DR11483', 'supplierName': 'Pleasant Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11155': {'supplierCode': 'DR11343', 'supplierName': 'Bassi Fort Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11156': {'supplierCode': 'DR11548', 'supplierName': 'Blackbuck Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11157': {'supplierCode': 'DR11268', 'supplierName': 'Treehouse The Pugmark', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11158': {'supplierCode': 'DR11553', 'supplierName': 'Chhotaram Prajapat\'s Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11160': {'supplierCode': 'DR11501', 'supplierName': 'Dewa Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11161': {'supplierCode': 'DR11188', 'supplierName': 'Devra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11164': {'supplierCode': 'DR11386', 'supplierName': 'Fortune District Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11166': {'supplierCode': 'DR11357', 'supplierName': 'Ram Pratap Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11167': {'supplierCode': 'DR11471', 'supplierName': 'The Pushkar Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11168': {'supplierCode': 'DR11340', 'supplierName': 'Palanpur Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11169': {'supplierCode': 'DR11544', 'supplierName': 'Abbotsford', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11170': {'supplierCode': 'DR11491', 'supplierName': 'Laxman Jhula Divine Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11171': {'supplierCode': 'DR10394', 'supplierName': 'Comfort Inn Benares', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11172': {'supplierCode': 'DR11498', 'supplierName': 'Anuraga Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11173': {'supplierCode': 'DR11526', 'supplierName': 'Grand Uniara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11174': {'supplierCode': 'DR11359', 'supplierName': 'Mirvana Nature Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11175': {'supplierCode': 'DR11546', 'supplierName': 'Villa Palladio', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11177': {'supplierCode': 'DR11557', 'supplierName': 'Crimson Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11178': {'supplierCode': 'DR10573', 'supplierName': 'Rawla Jojawar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11179': {'supplierCode': 'DR10773', 'supplierName': 'Ranthambhore National Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11180': {'supplierCode': 'DR11541', 'supplierName': 'Arista', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11181': {'supplierCode': 'DR10592', 'supplierName': 'Serkong House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11184': {'supplierCode': 'DR11845', 'supplierName': 'Diphlu River Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11185': {'supplierCode': 'DR11277', 'supplierName': 'Holiday Inn Express & Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11186': {'supplierCode': 'DR12050', 'supplierName': 'Boheda Manor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11187': {'supplierCode': 'DR11800', 'supplierName': 'Dekeling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11188': {'supplierCode': 'DR11324', 'supplierName': 'WelcomHeritage Gurkha Houseboats', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11191': {'supplierCode': 'DR11195', 'supplierName': 'Shahpura Barliyas House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11192': {'supplierCode': 'DR11391', 'supplierName': 'Shahpura Abhaneri Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11195': {'supplierCode': 'DR11646', 'supplierName': 'JW Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11196': {'supplierCode': 'DR11776', 'supplierName': 'Regenta Inn (Devanahalli)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11198': {'supplierCode': 'DR10982', 'supplierName': 'Regenta Central', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11199': {'supplierCode': 'DR11281', 'supplierName': 'Grand Kakinada By Grt Hotels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11200': {'supplierCode': 'DR11605', 'supplierName': 'Kanha Jungle Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11202': {'supplierCode': 'DR10866', 'supplierName': 'Regenta Central Herald', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11204': {'supplierCode': 'DR11346', 'supplierName': 'Abad Fort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11205': {'supplierCode': 'DR11288', 'supplierName': 'Ramada By Wyndham Jaipur (North)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11206': {'supplierCode': 'DR10959', 'supplierName': 'Dhora Desert Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11207': {'supplierCode': 'DR11446', 'supplierName': 'Bloom Hotel (Indiranagar)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11208': {'supplierCode': 'DR11450', 'supplierName': 'Bloomrooms (Calangute)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11209': {'supplierCode': 'DR11462', 'supplierName': 'Bloomrooms (Janpath)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11210': {'supplierCode': 'DR11435', 'supplierName': 'Bloomrooms (Link Road)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11212': {'supplierCode': 'DR11478', 'supplierName': 'Justa Gurgaon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11213': {'supplierCode': 'DR11476', 'supplierName': 'Justa Lake Nahargarh Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11214': {'supplierCode': 'DR12032', 'supplierName': 'Justa Brij Bhoomi Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11215': {'supplierCode': 'DR11423', 'supplierName': 'Krishna Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11216': {'supplierCode': 'DR11439', 'supplierName': 'Harrisons', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11217': {'supplierCode': 'DR11414', 'supplierName': 'Krishna Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11220': {'supplierCode': 'DR11440', 'supplierName': 'Airlink Castle', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11223': {'supplierCode': 'DR11438', 'supplierName': 'The Raviz', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11227': {'supplierCode': 'DR11356', 'supplierName': 'Pappukutty Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11228': {'supplierCode': 'DR11417', 'supplierName': 'Mannaas Veedu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11229': {'supplierCode': 'DR11427', 'supplierName': 'Grande Bay Resort And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11230': {'supplierCode': 'DR11355', 'supplierName': 'Hampis Boulders', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11231': {'supplierCode': 'DR11312', 'supplierName': 'Sagara Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11233': {'supplierCode': 'DR11353', 'supplierName': 'Trinity', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11234': {'supplierCode': 'DR11395', 'supplierName': 'Mallige Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11236': {'supplierCode': 'DR11303', 'supplierName': 'School Estate', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11238': {'supplierCode': 'DR11430', 'supplierName': 'Le Poshe By Sparsa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11240': {'supplierCode': 'DR11315', 'supplierName': 'Sparsa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11241': {'supplierCode': 'DR11856', 'supplierName': 'The Surya Village', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11243': {'supplierCode': 'DR11347', 'supplierName': 'Siolim House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11245': {'supplierCode': 'DR11437', 'supplierName': 'Ahilya By The Sea', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11249': {'supplierCode': 'DR11271', 'supplierName': 'Fortune Resort Benaullim', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11250': {'supplierCode': 'DR11454', 'supplierName': 'Golden Crown Hotel And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11251': {'supplierCode': 'DR11452', 'supplierName': 'Azaya Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11252': {'supplierCode': 'DR11307', 'supplierName': 'Hard Rock', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11253': {'supplierCode': 'DR11463', 'supplierName': 'Sibaya Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11254': {'supplierCode': 'DR11265', 'supplierName': 'Sol Fiesta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11255': {'supplierCode': 'DR12034', 'supplierName': 'The Grand - Ambala', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11257': {'supplierCode': 'DR11297', 'supplierName': 'Jaypee Vasant Continental', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11258': {'supplierCode': 'DR11531', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11259': {'supplierCode': 'DR11528', 'supplierName': 'Lemon Tree Tarudhan Valley', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11260': {'supplierCode': 'DR11522', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11261': {'supplierCode': 'DR11798', 'supplierName': 'Radisson Blu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11262': {'supplierCode': 'DR11517', 'supplierName': 'Lemon Tree Premier, The Atrium', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11263': {'supplierCode': 'DR11213', 'supplierName': 'Red Fox', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11264': {'supplierCode': 'DR11132', 'supplierName': 'Singinawa Jungle Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11265': {'supplierCode': 'DR10638', 'supplierName': 'El Oceano Beach Villas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11267': {'supplierCode': 'DR11741', 'supplierName': 'Manickam Tourist Home', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11269': {'supplierCode': 'DR11866', 'supplierName': 'Four Points By Sheraton', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11271': {'supplierCode': 'DR11536', 'supplierName': 'Lemon Tree Premier', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11272': {'supplierCode': 'DR11482', 'supplierName': 'Cherai Beach Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11274': {'supplierCode': 'DR11527', 'supplierName': 'The Boat Company', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11275': {'supplierCode': 'DR10816', 'supplierName': 'Godwin', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11277': {'supplierCode': 'DR11393', 'supplierName': 'Fort Manor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11278': {'supplierCode': 'DR11511', 'supplierName': 'The Jungle Book Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11279': {'supplierCode': 'DR11497', 'supplierName': 'Netuk House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11280': {'supplierCode': 'DR10035', 'supplierName': 'Holiday Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11282': {'supplierCode': 'DR11392', 'supplierName': 'Gama Heritage Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11283': {'supplierCode': 'DR10548', 'supplierName': 'Radisson Blu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11284': {'supplierCode': 'DR11759', 'supplierName': 'Sea Lagoon Health Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11285': {'supplierCode': 'DR11532', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11286': {'supplierCode': 'DR11489', 'supplierName': 'The Roseate', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11287': {'supplierCode': 'DR10595', 'supplierName': 'Secret Garden', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11288': {'supplierCode': 'DR11508', 'supplierName': 'WelcomHeritage Azora By Ayatana', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11289': {'supplierCode': 'DR11218', 'supplierName': 'Aura', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11291': {'supplierCode': 'DR11529', 'supplierName': 'Lemon Tree (Candolim)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11292': {'supplierCode': 'DR11537', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11293': {'supplierCode': 'DR11314', 'supplierName': 'Morpho Banashree Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11294': {'supplierCode': 'DR10100', 'supplierName': 'Lemon Tree Wildlife Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11295': {'supplierCode': 'DR11794', 'supplierName': 'Holiday Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11296': {'supplierCode': 'DR11612', 'supplierName': 'Anantha Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11297': {'supplierCode': 'DR10438', 'supplierName': 'Novotel Goa Resorts & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11298': {'supplierCode': 'DR11479', 'supplierName': 'The Lalit Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11299': {'supplierCode': 'DR11402', 'supplierName': 'Green Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11300': {'supplierCode': 'DR10406', 'supplierName': 'Hyatt Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11301': {'supplierCode': 'DR10438', 'supplierName': 'Conrad', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11303': {'supplierCode': 'DR10645', 'supplierName': 'Radisson Blu Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11305': {'supplierCode': 'DR11852', 'supplierName': 'Ramada By Wyndham Powai Hotel & Convention Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11308': {'supplierCode': 'DR11807', 'supplierName': 'Residency Fort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11309': {'supplierCode': 'DR11512', 'supplierName': 'Sanskruti Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11310': {'supplierCode': 'DR11519', 'supplierName': 'Lemon Tree (Hinjawadi)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11311': {'supplierCode': 'DR11796', 'supplierName': 'Signature Club Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11312': {'supplierCode': 'DR11015', 'supplierName': 'The Lalit Ashok', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11313': {'supplierCode': 'DR11063', 'supplierName': 'Somatheeram Ayurvedic Health Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11314': {'supplierCode': 'DR11788', 'supplierName': 'Sheraton Grand Pune Bund Garden', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11315': {'supplierCode': 'DR11173', 'supplierName': 'Bodhgaya Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11316': {'supplierCode': 'DR10884', 'supplierName': 'Renai Green Fields', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11317': {'supplierCode': 'DR11518', 'supplierName': 'Lemon Tree (Udyog Vihar)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11318': {'supplierCode': 'DR11467', 'supplierName': 'Lotus Nikko', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11319': {'supplierCode': 'DR11289', 'supplierName': 'Clarks Avadh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11320': {'supplierCode': 'DR10796', 'supplierName': 'Radisson (Udyog Vihar)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11321': {'supplierCode': 'DR11543', 'supplierName': 'Golden Sands', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11323': {'supplierCode': 'DR10477', 'supplierName': 'Fairfield By Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11324': {'supplierCode': 'DR10797', 'supplierName': 'Ramada By Wyndham Gurgaon Central', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11326': {'supplierCode': 'DR11633', 'supplierName': 'Vinayaga By Poppys', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11327': {'supplierCode': 'DR11803', 'supplierName': 'Adityaz', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11328': {'supplierCode': 'DR11490', 'supplierName': 'Mountain Club Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11329': {'supplierCode': 'DR11525', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11330': {'supplierCode': 'DR11465', 'supplierName': 'Sinclairs', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11331': {'supplierCode': 'DR11836', 'supplierName': 'Four Points By Sheraton', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11333': {'supplierCode': 'DR11514', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11334': {'supplierCode': 'DR11535', 'supplierName': 'Reen Resorts Aanavilasam Plantation', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11335': {'supplierCode': 'DR11504', 'supplierName': 'The Panoramic Getaway', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11336': {'supplierCode': 'DR11210', 'supplierName': 'Fortune Resort Grace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11337': {'supplierCode': 'DR10867', 'supplierName': 'Elephant Route', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11338': {'supplierCode': 'DR11323', 'supplierName': 'Kasmanda Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11339': {'supplierCode': 'DR10795', 'supplierName': 'Keys Select', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11341': {'supplierCode': 'DR11502', 'supplierName': 'Grand Gardenia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11342': {'supplierCode': 'DR11749', 'supplierName': 'Welcomhotel By ITC Hotels, The Savoy', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11343': {'supplierCode': 'DR11304', 'supplierName': 'Marasa Sarovar Premiere', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11344': {'supplierCode': 'DR11509', 'supplierName': 'Howard Johnson By Wyndham', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11345': {'supplierCode': 'DR11540', 'supplierName': 'Mundackal Plantation Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11346': {'supplierCode': 'DR11608', 'supplierName': 'Sujata', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11350': {'supplierCode': 'DR11864', 'supplierName': 'Koolwal Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11352': {'supplierCode': 'DR11533', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11353': {'supplierCode': 'DR11799', 'supplierName': 'Benzz Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11354': {'supplierCode': 'DR10630', 'supplierName': 'Radisson Blu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11356': {'supplierCode': 'DR11645', 'supplierName': 'Carmelia Haven', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11357': {'supplierCode': 'DR11523', 'supplierName': 'Lemon Tree (Shimona)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11359': {'supplierCode': 'DR11801', 'supplierName': 'Royal Court', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11360': {'supplierCode': 'DR10763', 'supplierName': 'Mapple Abhay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11361': {'supplierCode': 'DR11064', 'supplierName': 'Park Plaza Chennai OMR', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11362': {'supplierCode': 'DR11412', 'supplierName': 'Kyriad', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11364': {'supplierCode': 'DR11238', 'supplierName': 'Flame Of The Forest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11365': {'supplierCode': 'DR11409', 'supplierName': 'Renest River Country Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11366': {'supplierCode': 'DR11524', 'supplierName': 'Misty Gate', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11368': {'supplierCode': 'DR11886', 'supplierName': 'MPT Malwa Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11369': {'supplierCode': 'DR11415', 'supplierName': 'Quality Hotel Sabari Classic', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11370': {'supplierCode': 'DR10150', 'supplierName': 'Holiday Inn (Mayur Vihar)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11371': {'supplierCode': 'DR11293', 'supplierName': 'Park Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11372': {'supplierCode': 'DR11804', 'supplierName': 'Grand Imperial', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11373': {'supplierCode': 'DR11888', 'supplierName': 'Raisar Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11374': {'supplierCode': 'DR11887', 'supplierName': 'Khandwa Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11375': {'supplierCode': 'DR11475', 'supplierName': 'Kenilworth Resort And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11376': {'supplierCode': 'DR11520', 'supplierName': 'SinQ Edge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11377': {'supplierCode': 'DR10455', 'supplierName': 'Siris 18', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11378': {'supplierCode': 'DR11310', 'supplierName': 'Andores Resort And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11379': {'supplierCode': 'DR11940', 'supplierName': 'Champaner Heritage Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11380': {'supplierCode': 'DR11889', 'supplierName': 'Glitz', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11381': {'supplierCode': 'DR12038', 'supplierName': 'Wada 1', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11382': {'supplierCode': 'DR11311', 'supplierName': 'De Mandarin', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11383': {'supplierCode': 'DR11477', 'supplierName': 'Goan Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11384': {'supplierCode': 'DR11184', 'supplierName': 'Talaibagh Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11385': {'supplierCode': 'DR12049', 'supplierName': 'The Alampara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11386': {'supplierCode': 'DR11884', 'supplierName': 'Casa De Goa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11387': {'supplierCode': 'DR12044', 'supplierName': 'Le Meridien', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11388': {'supplierCode': 'DR11469', 'supplierName': 'Holiday Inn Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11389': {'supplierCode': 'DR10783', 'supplierName': 'Citrus Classic', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11390': {'supplierCode': 'DR10570', 'supplierName': 'Nazri Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11391': {'supplierCode': 'DR11956', 'supplierName': 'Mundota Fort And Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11392': {'supplierCode': 'DR11292', 'supplierName': 'Red Earth', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11393': {'supplierCode': 'DR10524', 'supplierName': 'Anand At The Satluj', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11394': {'supplierCode': 'DR11806', 'supplierName': 'Royal Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11395': {'supplierCode': 'DR12039', 'supplierName': 'Casa Rio', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11396': {'supplierCode': 'DR10703', 'supplierName': 'Santana Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11397': {'supplierCode': 'DR11394', 'supplierName': 'Orchard Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11398': {'supplierCode': 'DR11948', 'supplierName': 'Hariyali Dhani Camps & Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11400': {'supplierCode': 'DR11875', 'supplierName': 'Crossroads', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11401': {'supplierCode': 'DR11542', 'supplierName': 'O Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11402': {'supplierCode': 'DR12008', 'supplierName': 'Cygnett Lite Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11403': {'supplierCode': 'DR11881', 'supplierName': 'Citrus Cunningham', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11404': {'supplierCode': 'DR11651', 'supplierName': 'Narayan Niwas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11405': {'supplierCode': 'DR10558', 'supplierName': 'Surya Vilas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11406': {'supplierCode': 'DR11306', 'supplierName': 'Casa Anjuna', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11407': {'supplierCode': 'DR10284', 'supplierName': 'Landmark', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11408': {'supplierCode': 'DR11308', 'supplierName': 'Sandalwood Hotel & Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11409': {'supplierCode': 'DR10507', 'supplierName': 'Grand Hyatt', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11410': {'supplierCode': 'DR11507', 'supplierName': 'The Crown', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11411': {'supplierCode': 'DR10761', 'supplierName': 'Jaipur Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11412': {'supplierCode': 'DR11609', 'supplierName': 'Caravela Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11413': {'supplierCode': 'DR11516', 'supplierName': 'Amrapali - House Of Grace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11414': {'supplierCode': 'DR10643', 'supplierName': 'Prazeres Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11415': {'supplierCode': 'DR10704', 'supplierName': 'Sonesta Inns', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11416': {'supplierCode': 'DR10642', 'supplierName': 'The Jaipur House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11417': {'supplierCode': 'DR10453', 'supplierName': 'Kumbhalgarh Valley Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11418': {'supplierCode': 'DR10869', 'supplierName': 'Crest Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11419': {'supplierCode': 'DR11890', 'supplierName': 'Aaram Baagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11420': {'supplierCode': 'DR10319', 'supplierName': 'Desert Tulip', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11421': {'supplierCode': 'DR11177', 'supplierName': 'Bhanwar Singh Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11422': {'supplierCode': 'DR11416', 'supplierName': 'Pokaran Desert Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11423': {'supplierCode': 'DR11949', 'supplierName': 'The Baagh Ananta Elite', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11424': {'supplierCode': 'DR11873', 'supplierName': 'Kyron', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11425': {'supplierCode': 'DR11933', 'supplierName': 'The Golden Palms Hotel & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11427': {'supplierCode': 'DR11805', 'supplierName': 'The Narayana Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11428': {'supplierCode': 'DR11950', 'supplierName': 'Glitz Westend Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11429': {'supplierCode': 'DR11400', 'supplierName': 'Nirali Dhani', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11430': {'supplierCode': 'DR11957', 'supplierName': 'Shiv Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11431': {'supplierCode': 'DR11951', 'supplierName': 'Satvik Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11432': {'supplierCode': 'DR11882', 'supplierName': 'Amargarh Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11433': {'supplierCode': 'DR10840', 'supplierName': 'The Zion', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11434': {'supplierCode': 'DR11999', 'supplierName': 'North Avenue', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11435': {'supplierCode': 'DR10331', 'supplierName': 'Chandelao Garh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11436': {'supplierCode': 'DR12029', 'supplierName': 'Rose Mallow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11437': {'supplierCode': 'DR10131', 'supplierName': 'Rajwara Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11439': {'supplierCode': 'DR11180', 'supplierName': 'Patan Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11441': {'supplierCode': 'DR10606', 'supplierName': 'Rahat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11442': {'supplierCode': 'DR11187', 'supplierName': 'Amar Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11443': {'supplierCode': 'DR11611', 'supplierName': 'Sukhdham Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11446': {'supplierCode': 'DR11607', 'supplierName': 'Nikki’s Nest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11447': {'supplierCode': 'DR10206', 'supplierName': 'City Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11448': {'supplierCode': 'DR11396', 'supplierName': 'Garh Kumbha', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11449': {'supplierCode': 'DR11878', 'supplierName': 'Kumbhalgarh Safari Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11450': {'supplierCode': 'DR10633', 'supplierName': 'Manali Heights', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11451': {'supplierCode': 'DR12005', 'supplierName': 'Nimmu House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11452': {'supplierCode': 'DR11185', 'supplierName': 'Mandawa Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11453': {'supplierCode': 'DR12027', 'supplierName': 'Stok Palace Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11454': {'supplierCode': 'DR11061', 'supplierName': 'Sara Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11455': {'supplierCode': 'DR11610', 'supplierName': 'Palace Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11456': {'supplierCode': 'DR12036', 'supplierName': 'Mandawa Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11457': {'supplierCode': 'DR12048', 'supplierName': 'The Legacy Mandawa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11458': {'supplierCode': 'DR11270', 'supplierName': 'Caspia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11460': {'supplierCode': 'DR11330', 'supplierName': 'Amritara Surya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11461': {'supplierCode': 'DR12038', 'supplierName': 'Wada 1', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11462': {'supplierCode': 'DR12042', 'supplierName': 'Parallel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11465': {'supplierCode': 'DR12243', 'supplierName': 'Brij Anayra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11467': {'supplierCode': 'DR12220', 'supplierName': 'Aurika', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11470': {'supplierCode': 'DR12218', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11471': {'supplierCode': 'DR12045', 'supplierName': 'The Postcard Saligao', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11472': {'supplierCode': 'DR12059', 'supplierName': 'The Postcard Hideaway', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11474': {'supplierCode': 'DR12234', 'supplierName': 'The Postcard On The Arabian Sea', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11475': {'supplierCode': 'DR12054', 'supplierName': 'Country Inn & Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11477': {'supplierCode': 'DR12242', 'supplierName': 'Kadamb Kunj Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11478': {'supplierCode': 'DR12204', 'supplierName': 'Kesar Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11479': {'supplierCode': 'DR12211', 'supplierName': 'Sea Shell', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11480': {'supplierCode': 'DR12237', 'supplierName': 'Gogunda Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11482': {'supplierCode': 'DR12230', 'supplierName': 'Ananta Spa & Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11483': {'supplierCode': 'DR12239', 'supplierName': 'The Experience Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11484': {'supplierCode': 'DR12219', 'supplierName': 'Bungalow 2', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11485': {'supplierCode': 'DR12037', 'supplierName': 'Udman', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11486': {'supplierCode': 'DR12223', 'supplierName': 'Sawai Shivir', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11487': {'supplierCode': 'DR12202', 'supplierName': 'The Gateway Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11488': {'supplierCode': 'DR12203', 'supplierName': 'Novotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11489': {'supplierCode': 'DR12058', 'supplierName': 'Royale Regent', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11490': {'supplierCode': 'DR12236', 'supplierName': 'The Kikar Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11492': {'supplierCode': 'DR12055', 'supplierName': 'Royale Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11497': {'supplierCode': 'DR12249', 'supplierName': 'Justa Sarang', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11498': {'supplierCode': 'DR11285', 'supplierName': 'The Royal Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11499': {'supplierCode': 'DR12251', 'supplierName': 'Starlit Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11500': {'supplierCode': 'DR12248', 'supplierName': 'Atulya Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11501': {'supplierCode': 'DR10847', 'supplierName': 'Marvel cruise', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11502': {'supplierCode': 'DR12254', 'supplierName': 'Discovery', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11503': {'supplierCode': 'DR11068', 'supplierName': 'Spice Coast Cruises', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11505': {'supplierCode': 'DR10847', 'supplierName': 'Marvel Cruise', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11506': {'supplierCode': 'DR11068', 'supplierName': 'Spice Coast Cruises', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11507': {'supplierCode': 'DR10257', 'supplierName': 'Spice Routes', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11509': {'supplierCode': 'DR10110', 'supplierName': 'Kumarakom Lake Resort (Houseboat)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11510': {'supplierCode': 'DR12257', 'supplierName': 'The Gateway Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11511': {'supplierCode': 'DR12261', 'supplierName': 'Elakai', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11512': {'supplierCode': 'DR12061', 'supplierName': 'Mastiff Select', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11514': {'supplierCode': 'DR10443', 'supplierName': 'Manaltheeram Ayurvedic Hospital & Research Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11518': {'supplierCode': 'DR10637', 'supplierName': 'Manvar Kumat, Thar Luxury Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11519': {'supplierCode': 'DR12258', 'supplierName': 'Zone Connect By The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11520': {'supplierCode': 'DR12241', 'supplierName': 'Om Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11521': {'supplierCode': 'DR12316', 'supplierName': 'Zone Palace By The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11522': {'supplierCode': 'DR12315', 'supplierName': 'Zone By The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11523': {'supplierCode': 'DR12317', 'supplierName': 'Zone Connect By The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11524': {'supplierCode': 'DR12284', 'supplierName': 'Keys Prima', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11525': {'supplierCode': 'DR12283', 'supplierName': 'Lemon Tree (Banjara Hills)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11529': {'supplierCode': 'DR12286', 'supplierName': 'Justa Palampur Resort & Convention Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11531': {'supplierCode': 'DR12322', 'supplierName': 'Mulberry Shades', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11532': {'supplierCode': 'DR12321', 'supplierName': 'JW Marriott Bengaluru Prestige Golfshire Res.& Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11533': {'supplierCode': 'DR12323', 'supplierName': 'Sheraton Grand Chennai Resorts and Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11535': {'supplierCode': 'DR12235', 'supplierName': 'Chez Lavania', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11538': {'supplierCode': 'DR12192', 'supplierName': 'Lockhart Bungalow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11540': {'supplierCode': 'DR12319', 'supplierName': 'Araiya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11542': {'supplierCode': 'DR12302', 'supplierName': 'Anchorage 42', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11543': {'supplierCode': 'DR12125', 'supplierName': 'Time Square', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11544': {'supplierCode': 'DR12337', 'supplierName': 'Flamingo Marari', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11545': {'supplierCode': 'DR12279', 'supplierName': 'Fort Pokaran', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11547': {'supplierCode': 'DR12233', 'supplierName': 'The Maharana Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11549': {'supplierCode': 'DR12228', 'supplierName': 'Deeppura Garh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11550': {'supplierCode': 'DR11506', 'supplierName': 'The Blackbuck Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11551': {'supplierCode': 'DR12056', 'supplierName': 'Sankhu Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11552': {'supplierCode': 'DR12296', 'supplierName': 'Norbu House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11553': {'supplierCode': 'DR12297', 'supplierName': 'Junglaat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11555': {'supplierCode': 'DR12347', 'supplierName': 'Tree Of Life Darbargadh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11556': {'supplierCode': 'DR12342', 'supplierName': 'Tree Of Life Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11557': {'supplierCode': 'DR12343', 'supplierName': 'Tree Of Life Vantara Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11560': {'supplierCode': 'DR12329', 'supplierName': 'Tree Of Life Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11561': {'supplierCode': 'DR12350', 'supplierName': 'Oneness', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11562': {'supplierCode': 'DR12355', 'supplierName': 'Irai Safari Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11563': {'supplierCode': 'DR12333', 'supplierName': 'Drenmo Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11564': {'supplierCode': 'DR12356', 'supplierName': 'DoubleTree By Hilton', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11565': {'supplierCode': 'DR12210', 'supplierName': 'Classic Sapphire Ananta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11567': {'supplierCode': 'DR12363', 'supplierName': 'Great Trails', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11568': {'supplierCode': 'DR12362', 'supplierName': 'Great Trails', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11569': {'supplierCode': 'DR12287', 'supplierName': 'Coorg International', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11570': {'supplierCode': 'DR12367', 'supplierName': 'Koder House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11572': {'supplierCode': 'DR12366', 'supplierName': 'The Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11574': {'supplierCode': 'DR12357', 'supplierName': 'Ging Tea House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11576': {'supplierCode': 'DR12374', 'supplierName': 'Taj Wellington Mews', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11579': {'supplierCode': 'DR12294', 'supplierName': 'Taj Guras Kutir Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11580': {'supplierCode': 'DR12372', 'supplierName': 'Taj Taal Kutir', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11581': {'supplierCode': 'DR12368', 'supplierName': 'Rose Amer', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11586': {'supplierCode': 'DR12273', 'supplierName': 'Vivaana Museum', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11587': {'supplierCode': 'DR12339', 'supplierName': '35 Sahakar Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11588': {'supplierCode': 'DR12365', 'supplierName': 'Niraamaya Retreats Samroha', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11589': {'supplierCode': 'DR12212', 'supplierName': 'S Hotels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11590': {'supplierCode': 'DR12231', 'supplierName': 'Khumani', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11591': {'supplierCode': 'DR12309', 'supplierName': 'Sitara Himalaya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11593': {'supplierCode': 'DR12299', 'supplierName': 'Tigergarh Wildlife Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11596': {'supplierCode': 'DR12360', 'supplierName': 'Chomu Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11597': {'supplierCode': 'DR12353', 'supplierName': 'Hotel Rudra Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11598': {'supplierCode': 'DR12377', 'supplierName': 'Brij Bageecha', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11599': {'supplierCode': 'DR12395', 'supplierName': 'Aalia Jungle Retreat & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11600': {'supplierCode': 'DR12388', 'supplierName': 'Aalia Villas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11601': {'supplierCode': 'DR12387', 'supplierName': 'Kinwani House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11602': {'supplierCode': 'DR12298', 'supplierName': 'Arch Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11603': {'supplierCode': 'DR12358', 'supplierName': 'The Hill House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11604': {'supplierCode': 'DR12305', 'supplierName': 'The Sawai Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11605': {'supplierCode': 'DR12393', 'supplierName': 'Shivadya Camps', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11606': {'supplierCode': 'DR12263', 'supplierName': 'Clarks Inn Express', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11607': {'supplierCode': 'DR12349', 'supplierName': 'KD Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11608': {'supplierCode': 'DR12301', 'supplierName': 'Bamboo Saa Mulberry Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11609': {'supplierCode': 'DR12371', 'supplierName': 'Aloof Jungle Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11610': {'supplierCode': 'DR12289', 'supplierName': 'Kohima Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11611': {'supplierCode': 'DR12403', 'supplierName': 'The Clement Retreats - Dera Village', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11612': {'supplierCode': 'DR12402', 'supplierName': 'Bloom Boutique (Connaught Place)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11613': {'supplierCode': 'DR12214', 'supplierName': 'Emarald', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11614': {'supplierCode': 'DR12404', 'supplierName': 'Shivansh Green Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11616': {'supplierCode': 'DR12341', 'supplierName': 'Jaagir Manor – IHCL SeleQtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11618': {'supplierCode': 'DR12394', 'supplierName': 'Chhota Haazri', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11619': {'supplierCode': 'DR12412', 'supplierName': 'Godwin', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11620': {'supplierCode': 'DR12413', 'supplierName': 'Shreevanam By Beejom', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11621': {'supplierCode': 'DR12630', 'supplierName': 'Marine House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11622': {'supplierCode': 'DR12513', 'supplierName': 'East Park Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11623': {'supplierCode': 'DR11538', 'supplierName': 'Courtyard By Marriott (Hebbal)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11624': {'supplierCode': 'DR11538', 'supplierName': 'Novotel Sipcot', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11626': {'supplierCode': 'DR12614', 'supplierName': 'Villa 21', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11627': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Blu Chennai City Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11628': {'supplierCode': 'DR12593', 'supplierName': 'Florence', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11629': {'supplierCode': 'DR11538', 'supplierName': 'DoubleTree Suites By Hilton', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11630': {'supplierCode': 'DR11538', 'supplierName': 'Florence Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11631': {'supplierCode': 'DR11538', 'supplierName': 'Crowne Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11632': {'supplierCode': 'DR11538', 'supplierName': 'Iskcon Guest Ashram', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11633': {'supplierCode': 'DR10589', 'supplierName': 'The River Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11634': {'supplierCode': 'DR12546', 'supplierName': 'Madhuban Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11635': {'supplierCode': 'DR12443', 'supplierName': 'Novotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11636': {'supplierCode': 'DR11538', 'supplierName': 'Grand Mercure', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11637': {'supplierCode': 'DR11538', 'supplierName': 'Nuo By Justa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11639': {'supplierCode': 'DR11538', 'supplierName': 'Lakshmi Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11640': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Blu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11641': {'supplierCode': 'DR11538', 'supplierName': 'Greenpark', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11642': {'supplierCode': 'DR11538', 'supplierName': 'Aadrika', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11644': {'supplierCode': 'DR12451', 'supplierName': 'Park Inn By Radisson', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11645': {'supplierCode': 'DR11538', 'supplierName': 'Hilton Bengaluru Embassy Manyata Business', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11646': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Blue Alibaug', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11647': {'supplierCode': 'DR11538', 'supplierName': 'Honeydewwz Exoticaa Hotel & Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11648': {'supplierCode': 'DR11538', 'supplierName': 'The Ashtan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11649': {'supplierCode': 'DR11538', 'supplierName': 'Holiday Inn Express & Suites Racecourse', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11650': {'supplierCode': 'DR11026', 'supplierName': 'Amal Tamara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11651': {'supplierCode': 'DR11538', 'supplierName': 'The Manor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11652': {'supplierCode': 'DR11538', 'supplierName': 'Justa Indra Nagar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11653': {'supplierCode': 'DR11538', 'supplierName': 'Justa Mg Road', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11654': {'supplierCode': 'DR11538', 'supplierName': 'Exotica Cruises', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11655': {'supplierCode': 'DR11538', 'supplierName': 'The Blossom', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11656': {'supplierCode': 'DR12502', 'supplierName': 'Mastiff Infantry', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11657': {'supplierCode': 'DR12673', 'supplierName': 'Zaza Stay G-54', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11658': {'supplierCode': 'DR11538', 'supplierName': 'Novotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11659': {'supplierCode': 'DR12574', 'supplierName': 'Plumeria Lake Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11660': {'supplierCode': 'DR11538', 'supplierName': 'Jhira Bagh Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11661': {'supplierCode': 'DR10249', 'supplierName': 'Punnamada Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11662': {'supplierCode': 'DR11538', 'supplierName': 'The River Front Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11664': {'supplierCode': 'DR11538', 'supplierName': 'Fort Begu (1430 AD)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11665': {'supplierCode': 'DR12795', 'supplierName': 'Sterling Lake Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11666': {'supplierCode': 'DR12643', 'supplierName': 'D\'S Casa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11667': {'supplierCode': 'DR11538', 'supplierName': 'The Padmini Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11668': {'supplierCode': 'DR11538', 'supplierName': 'Justa Birding Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11669': {'supplierCode': 'DR11538', 'supplierName': 'Venezia Kerala Houseboat Cruise', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11670': {'supplierCode': 'DR11538', 'supplierName': 'The Divine Hima', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11671': {'supplierCode': 'DR12448', 'supplierName': 'Beach Gate Bungalows', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11672': {'supplierCode': 'DR11538', 'supplierName': 'Itmenaan Estate', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11674': {'supplierCode': 'DR11538', 'supplierName': 'Club Mahindra Cherai Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11675': {'supplierCode': 'DR11538', 'supplierName': 'Kumaon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11677': {'supplierCode': 'DR12205', 'supplierName': 'Ten Square', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11678': {'supplierCode': 'DR11538', 'supplierName': 'Lakeside Guesthouse', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11679': {'supplierCode': 'DR12482', 'supplierName': 'Radhika Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11680': {'supplierCode': 'DR11538', 'supplierName': 'Dadhikar Fort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11681': {'supplierCode': 'DR11538', 'supplierName': 'Marine Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11682': {'supplierCode': 'DR11538', 'supplierName': 'MPT Holiday Homes Amarkantak', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11684': {'supplierCode': 'DR11538', 'supplierName': 'Fortune Inn Heritage Walk', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11685': {'supplierCode': 'DR11538', 'supplierName': 'Namami Wellness And Health Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11686': {'supplierCode': 'DR11538', 'supplierName': 'Four Points by Sheraton Kochi Infopark', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11687': {'supplierCode': 'DR11538', 'supplierName': 'Regenta Place', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11688': {'supplierCode': 'DR11538', 'supplierName': 'Quality Airport Hotels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11689': {'supplierCode': 'DR11538', 'supplierName': 'Sanjog International', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11690': {'supplierCode': 'DR11538', 'supplierName': 'Old Magazine House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11693': {'supplierCode': 'DR11538', 'supplierName': 'Dukes Forest Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11694': {'supplierCode': 'DR11538', 'supplierName': 'Denzong Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11695': {'supplierCode': 'DR12503', 'supplierName': 'Le Colonial', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11696': {'supplierCode': 'DR11538', 'supplierName': 'Aashyana Lakhanpal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11697': {'supplierCode': 'DR11538', 'supplierName': 'Olive Downtown', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11698': {'supplierCode': 'DR11538', 'supplierName': 'Agonda Serenity Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11700': {'supplierCode': 'DR11538', 'supplierName': 'Alidia Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11701': {'supplierCode': 'DR11538', 'supplierName': 'Anahata Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11702': {'supplierCode': 'DR11538', 'supplierName': 'Athirappilly Green Trees Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11703': {'supplierCode': 'DR11538', 'supplierName': 'Raintree Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11704': {'supplierCode': 'DR11538', 'supplierName': 'Backwoods Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11705': {'supplierCode': 'DR11538', 'supplierName': 'Sherlys Ente Kumbalanghi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11706': {'supplierCode': 'DR12538', 'supplierName': 'Rainforest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11707': {'supplierCode': 'DR11538', 'supplierName': 'Bamboo Yoga Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11708': {'supplierCode': 'DR12519', 'supplierName': 'Bogmallo Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11709': {'supplierCode': 'DR11538', 'supplierName': 'Click Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11711': {'supplierCode': 'DR11538', 'supplierName': 'Cabo Serai Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11712': {'supplierCode': 'DR11538', 'supplierName': 'Casa Jaali', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11713': {'supplierCode': 'DR11538', 'supplierName': 'RBD Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11714': {'supplierCode': 'DR11538', 'supplierName': 'Regenta Place Cunningham', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11715': {'supplierCode': 'DR11538', 'supplierName': 'Renaissance Race Course', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11716': {'supplierCode': 'DR11538', 'supplierName': 'Country Inn & Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11717': {'supplierCode': 'DR11538', 'supplierName': 'Le Meridien', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11718': {'supplierCode': 'DR12515', 'supplierName': 'De Mandarin Beach Resort, Suites And Villas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11719': {'supplierCode': 'DR11538', 'supplierName': 'Sterlings Mac', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11720': {'supplierCode': 'DR11538', 'supplierName': 'Sinnadorai Bungalow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11721': {'supplierCode': 'DR11538', 'supplierName': 'Double Tree By Hilton', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11722': {'supplierCode': 'DR11538', 'supplierName': 'Tea Nest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11723': {'supplierCode': 'DR11538', 'supplierName': 'The Zuri Whitefield', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11724': {'supplierCode': 'DR11538', 'supplierName': 'DoubleTree by Hilton - Panaji', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11725': {'supplierCode': 'DR11538', 'supplierName': 'Wallwood Garden', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11727': {'supplierCode': 'DR11538', 'supplierName': 'Dunhill Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11728': {'supplierCode': 'DR10553', 'supplierName': 'Birders Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11729': {'supplierCode': 'DR11538', 'supplierName': 'Dubare Elephant Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11730': {'supplierCode': 'DR12504', 'supplierName': 'Gokulam Grand Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11732': {'supplierCode': 'DR11538', 'supplierName': 'Evoke Lifestyle', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11733': {'supplierCode': 'DR11538', 'supplierName': 'Fairfield by Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11734': {'supplierCode': 'DR11538', 'supplierName': 'Marriott Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11735': {'supplierCode': 'DR11538', 'supplierName': 'Corbett Ramganga Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11737': {'supplierCode': 'DR11538', 'supplierName': 'Golden Tulip', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11738': {'supplierCode': 'DR11538', 'supplierName': 'Manu Maharani Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11739': {'supplierCode': 'DR12525', 'supplierName': 'Heritage Village Resorts & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11740': {'supplierCode': 'DR11538', 'supplierName': 'Effotel by Sayaji', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11741': {'supplierCode': 'DR11538', 'supplierName': 'Holiday Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11742': {'supplierCode': 'DR11538', 'supplierName': 'JW Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11743': {'supplierCode': 'DR11538', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11744': {'supplierCode': 'DR12454', 'supplierName': 'Taj Corbett Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11745': {'supplierCode': 'DR11538', 'supplierName': 'Fortune Park Sishmo', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11746': {'supplierCode': 'DR12526', 'supplierName': 'Kyriad Prestige', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11747': {'supplierCode': 'DR12633', 'supplierName': 'The Bhuj House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11748': {'supplierCode': 'DR12520', 'supplierName': 'La Grace Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11749': {'supplierCode': 'DR11538', 'supplierName': 'The Hideaway River Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11750': {'supplierCode': 'DR11538', 'supplierName': 'Black Buck Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11751': {'supplierCode': 'DR11538', 'supplierName': 'LaRisa Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11752': {'supplierCode': 'DR11538', 'supplierName': 'Best Western Dalhousie', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11753': {'supplierCode': 'DR12631', 'supplierName': 'Mastiff Select', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11754': {'supplierCode': 'DR11538', 'supplierName': 'LaRiSa Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11755': {'supplierCode': 'DR12354', 'supplierName': '24 MG Road', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11757': {'supplierCode': 'DR12580', 'supplierName': 'Udai Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11758': {'supplierCode': 'DR11538', 'supplierName': 'Maitree Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11759': {'supplierCode': 'DR12746', 'supplierName': 'Mary Budden Estate', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11760': {'supplierCode': 'DR11538', 'supplierName': 'Mandrem Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11761': {'supplierCode': 'DR11538', 'supplierName': 'Annapurna', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11762': {'supplierCode': 'DR11538', 'supplierName': 'Presidium Sarovar Portico Dalhousie', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11763': {'supplierCode': 'DR11538', 'supplierName': 'Nalanda Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11764': {'supplierCode': 'DR11538', 'supplierName': 'Niranjana', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11765': {'supplierCode': 'DR12611', 'supplierName': 'The Gold Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11766': {'supplierCode': 'DR12656', 'supplierName': 'Nanu Beach Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11767': {'supplierCode': 'DR11538', 'supplierName': 'Oaks', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11769': {'supplierCode': 'DR11538', 'supplierName': 'Ramada Darjeeling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11770': {'supplierCode': 'DR12522', 'supplierName': 'Planet Hollywood Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11772': {'supplierCode': 'DR12553', 'supplierName': 'Copper Folia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11773': {'supplierCode': 'DR11538', 'supplierName': 'Clarks Exotica Resort And Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11774': {'supplierCode': 'DR11538', 'supplierName': 'Green Olive', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11775': {'supplierCode': 'DR11538', 'supplierName': 'Resort Rio', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11776': {'supplierCode': 'DR11538', 'supplierName': 'Hyatt Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11777': {'supplierCode': 'DR11538', 'supplierName': 'Renai Kappad Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11778': {'supplierCode': 'DR12516', 'supplierName': 'Sairaj Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11780': {'supplierCode': 'DR11538', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11781': {'supplierCode': 'DR11538', 'supplierName': 'Salcete Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11783': {'supplierCode': 'DR11538', 'supplierName': 'Seashell Suites and Villas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11784': {'supplierCode': 'DR11538', 'supplierName': 'Krinoscco', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11785': {'supplierCode': 'DR11538', 'supplierName': 'Simrose Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11786': {'supplierCode': 'DR11538', 'supplierName': 'Sarovar Portico Dehradun', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11790': {'supplierCode': 'DR12644', 'supplierName': 'Anila', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11791': {'supplierCode': 'DR11538', 'supplierName': 'Sobit Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11792': {'supplierCode': 'DR10070', 'supplierName': 'Bloom Boutique GK1', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11794': {'supplierCode': 'DR12715', 'supplierName': 'Waghoba Eco Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11795': {'supplierCode': 'DR11538', 'supplierName': 'Soul Vacation', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11796': {'supplierCode': 'DR11538', 'supplierName': 'The Westin Goa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11798': {'supplierCode': 'DR12524', 'supplierName': 'Stone Wood Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11799': {'supplierCode': 'DR11538', 'supplierName': 'Bandipur Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11800': {'supplierCode': 'DR11538', 'supplierName': 'The Bay Agonda', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11801': {'supplierCode': 'DR11538', 'supplierName': 'Villa Siolim - Ama Stays & Trails', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11802': {'supplierCode': 'DR11538', 'supplierName': 'The Flora Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11803': {'supplierCode': 'DR12559', 'supplierName': 'Dhole\'s Den', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11804': {'supplierCode': 'DR11538', 'supplierName': 'The HQ', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11806': {'supplierCode': 'DR11538', 'supplierName': 'Fairfield by Marriott Belagavi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11807': {'supplierCode': 'DR11538', 'supplierName': 'Vivenda Dos Palhacos', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11808': {'supplierCode': 'DR11538', 'supplierName': 'W Goa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11809': {'supplierCode': 'DR11538', 'supplierName': 'UK 27 The Fern', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11810': {'supplierCode': 'DR11538', 'supplierName': 'Hampi Heritage & Wilderness Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11813': {'supplierCode': 'DR11538', 'supplierName': 'Zip By Spree Hyde', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11814': {'supplierCode': 'DR12569', 'supplierName': 'Kahani Paradise', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11815': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Hyderabad Hitec City', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11816': {'supplierCode': 'DR11538', 'supplierName': 'Winds Desert Camps', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11817': {'supplierCode': 'DR11538', 'supplierName': 'Isabel Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11818': {'supplierCode': 'DR11538', 'supplierName': 'Fortune Avenue Jalandhar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11819': {'supplierCode': 'DR11538', 'supplierName': 'Palm Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11820': {'supplierCode': 'DR11538', 'supplierName': 'Sitara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11821': {'supplierCode': 'DR12539', 'supplierName': 'Elephant Valley', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11822': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Hotel Jalandhar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11823': {'supplierCode': 'DR11538', 'supplierName': 'The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11824': {'supplierCode': 'DR12507', 'supplierName': 'JC Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11825': {'supplierCode': 'DR12572', 'supplierName': 'The Tamara Kodai', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11826': {'supplierCode': 'DR11538', 'supplierName': 'Ramada', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11827': {'supplierCode': 'DR11538', 'supplierName': 'The Westin Hyderabad Mindspace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11828': {'supplierCode': 'DR11538', 'supplierName': 'Crowne Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11829': {'supplierCode': 'DR12511', 'supplierName': 'Niraamaya Retreats Aradura', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11830': {'supplierCode': 'DR11538', 'supplierName': 'Sayaji', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11831': {'supplierCode': 'DR11538', 'supplierName': 'Jaypee Green Golf & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11834': {'supplierCode': 'DR11538', 'supplierName': 'Fairfield By Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11835': {'supplierCode': 'DR11538', 'supplierName': 'Nataraj Sarovar Portico Jhansi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11836': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Blu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11837': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Blu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11839': {'supplierCode': 'DR11538', 'supplierName': 'Holiday Inn Express', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11840': {'supplierCode': 'DR12485', 'supplierName': 'Welcomhotel By ITC', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11841': {'supplierCode': 'DR12651', 'supplierName': 'Hanuwant Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11842': {'supplierCode': 'DR11538', 'supplierName': 'Novotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11843': {'supplierCode': 'DR11538', 'supplierName': 'Sheraton Grand Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11844': {'supplierCode': 'DR11538', 'supplierName': 'The Corner Courtyard', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11845': {'supplierCode': 'DR11538', 'supplierName': 'DoubleTree by Hilton', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11846': {'supplierCode': 'DR11538', 'supplierName': 'Kana', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11847': {'supplierCode': 'DR11538', 'supplierName': 'Narmada Jacksons', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11848': {'supplierCode': 'DR11538', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11849': {'supplierCode': 'DR11538', 'supplierName': 'Grand Hyatt', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11850': {'supplierCode': 'DR11538', 'supplierName': 'Krishna Prakash Heritage Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11851': {'supplierCode': 'DR12441', 'supplierName': 'Vivanta', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11852': {'supplierCode': 'DR11538', 'supplierName': 'Barwara kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11853': {'supplierCode': 'DR12450', 'supplierName': 'ITC Grand Bharat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11854': {'supplierCode': 'DR11538', 'supplierName': 'Lariya Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11855': {'supplierCode': 'DR12499', 'supplierName': 'Fragrant Nature', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11856': {'supplierCode': 'DR12675', 'supplierName': 'Buena Vista', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11857': {'supplierCode': 'DR10902', 'supplierName': 'The Mandore - A Leafy Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11858': {'supplierCode': 'DR12446', 'supplierName': 'Taj Damdama Lake Resort and Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11859': {'supplierCode': 'DR11538', 'supplierName': 'Khatu Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11860': {'supplierCode': 'DR11538', 'supplierName': 'Blue Wings Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11861': {'supplierCode': 'DR11538', 'supplierName': 'The Westin Sohna Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11862': {'supplierCode': 'DR11538', 'supplierName': 'Rester Select – ITI Circle', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11863': {'supplierCode': 'DR11538', 'supplierName': 'Agastyaa Heritage Ayurvedic Hospital', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11864': {'supplierCode': 'DR11538', 'supplierName': 'The Kothi Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11865': {'supplierCode': 'DR12613', 'supplierName': 'Pearl Palace Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11868': {'supplierCode': 'DR12632', 'supplierName': 'Kyriad Hotel Gulbarga', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11869': {'supplierCode': 'DR12547', 'supplierName': 'Clarks Inn Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11870': {'supplierCode': 'DR11538', 'supplierName': 'Lumbinis Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11871': {'supplierCode': 'DR11538', 'supplierName': 'Pratap Bhawan Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11872': {'supplierCode': 'DR12545', 'supplierName': 'Godwin', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11874': {'supplierCode': 'DR12463', 'supplierName': 'Radisson', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11875': {'supplierCode': 'DR11538', 'supplierName': 'The Ayur Villa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11878': {'supplierCode': 'DR11538', 'supplierName': 'Hyatt Place', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11880': {'supplierCode': 'DR11538', 'supplierName': 'MM Legacy', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11881': {'supplierCode': 'DR11538', 'supplierName': 'Regenta Orkos', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11882': {'supplierCode': 'DR10105', 'supplierName': 'Umaid Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11883': {'supplierCode': 'DR11538', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11884': {'supplierCode': 'DR11538', 'supplierName': 'Aranyak Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11885': {'supplierCode': 'DR11538', 'supplierName': 'KTDC Waterscapes', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11886': {'supplierCode': 'DR11538', 'supplierName': 'Yashail', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11887': {'supplierCode': 'DR12650', 'supplierName': 'Shergarh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11888': {'supplierCode': 'DR11538', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11890': {'supplierCode': 'DR12645', 'supplierName': 'The Bordi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11891': {'supplierCode': 'DR11538', 'supplierName': 'The Kumbha Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11893': {'supplierCode': 'DR11538', 'supplierName': 'Wild Avenue Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11894': {'supplierCode': 'DR11538', 'supplierName': 'Damodra Desert Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11895': {'supplierCode': 'DR12378', 'supplierName': 'Anantya By The Lake', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11897': {'supplierCode': 'DR11538', 'supplierName': 'Courtyard By Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11898': {'supplierCode': 'DR11538', 'supplierName': 'Desert Springs Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11899': {'supplierCode': 'DR11538', 'supplierName': 'Hampshire Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11900': {'supplierCode': 'DR10474', 'supplierName': 'Om Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11901': {'supplierCode': 'DR11538', 'supplierName': 'Avasa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11904': {'supplierCode': 'DR11538', 'supplierName': 'Hyatt Place', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11905': {'supplierCode': 'DR11538', 'supplierName': 'Saboo Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11906': {'supplierCode': 'DR11538', 'supplierName': 'Killa Bhawan Lodge (KB Lodge)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11907': {'supplierCode': 'DR12465', 'supplierName': 'ITC Kohenur', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11908': {'supplierCode': 'DR12542', 'supplierName': 'Sea View', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11909': {'supplierCode': 'DR12456', 'supplierName': 'Mercure Hyderabad KCP', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11910': {'supplierCode': 'DR11538', 'supplierName': 'Fariyas Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11913': {'supplierCode': 'DR11538', 'supplierName': 'Park Hyatt', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11914': {'supplierCode': 'DR12621', 'supplierName': 'Sandhya Kasol', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11915': {'supplierCode': 'DR11538', 'supplierName': 'Priya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11916': {'supplierCode': 'DR11538', 'supplierName': 'Fortune Park Katra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11917': {'supplierCode': 'DR11538', 'supplierName': 'Golden Tulip', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11918': {'supplierCode': 'DR11538', 'supplierName': 'Roopal International', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11919': {'supplierCode': 'DR11538', 'supplierName': 'Hometel Alambagh Lucknow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11920': {'supplierCode': 'DR12647', 'supplierName': 'Grand Dewachen', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11921': {'supplierCode': 'DR11538', 'supplierName': 'Novotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11922': {'supplierCode': 'DR11538', 'supplierName': 'Radisson City Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11923': {'supplierCode': 'DR11538', 'supplierName': 'Ramada', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11925': {'supplierCode': 'DR11538', 'supplierName': 'Hyatt Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11926': {'supplierCode': 'DR11538', 'supplierName': 'Le meridien resort and spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11927': {'supplierCode': 'DR11538', 'supplierName': 'Leela Resort The Exotic Boutique', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11928': {'supplierCode': 'DR11538', 'supplierName': 'Welcom Heritage Shri Mohangarh Fort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11929': {'supplierCode': 'DR11538', 'supplierName': 'Four Points by Sheraton Mahabalipuram', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11930': {'supplierCode': 'DR11538', 'supplierName': 'Kohinoor Continental', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11931': {'supplierCode': 'DR11538', 'supplierName': 'Kaldan Samudhra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11932': {'supplierCode': 'DR11538', 'supplierName': 'Winter Note', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11934': {'supplierCode': 'DR11538', 'supplierName': 'Dekyid', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11935': {'supplierCode': 'DR11538', 'supplierName': 'Krishna Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11936': {'supplierCode': 'DR12417', 'supplierName': 'The Claridges Nabha Residence', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11937': {'supplierCode': 'DR12582', 'supplierName': 'Banons Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11938': {'supplierCode': 'DR12636', 'supplierName': 'Pachewar Garh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11939': {'supplierCode': 'DR11538', 'supplierName': 'Le Sutra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11940': {'supplierCode': 'DR11538', 'supplierName': 'The Johnson\'s', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11941': {'supplierCode': 'DR11538', 'supplierName': 'Kairali -The Ayurvedic Healing Village', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11942': {'supplierCode': 'DR11538', 'supplierName': 'Jw Marriott Mussoorie Walnut Grove Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11943': {'supplierCode': 'DR11538', 'supplierName': 'Comfort inn Snow Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11944': {'supplierCode': 'DR11538', 'supplierName': 'Majestic Court Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11945': {'supplierCode': 'DR11538', 'supplierName': 'Ganpat Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11946': {'supplierCode': 'DR11538', 'supplierName': 'Mastiff Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11947': {'supplierCode': 'DR11538', 'supplierName': 'Rokeby Manor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11948': {'supplierCode': 'DR11538', 'supplierName': 'Nature Vilas Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11949': {'supplierCode': 'DR11538', 'supplierName': 'The Amethyst', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11951': {'supplierCode': 'DR12587', 'supplierName': 'Snow Valley Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11953': {'supplierCode': 'DR11538', 'supplierName': 'Jawai Pugmark Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11955': {'supplierCode': 'DR11538', 'supplierName': 'Udai Vilas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11956': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Blu (Mumbai International Airport)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11958': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Goregaon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11959': {'supplierCode': 'DR11538', 'supplierName': 'Welcomheritage Ramgarh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11963': {'supplierCode': 'DR11538', 'supplierName': 'Centre Point', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11964': {'supplierCode': 'DR11538', 'supplierName': 'Ramada Plaza by Wyndham Palm Grove', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11965': {'supplierCode': 'DR12434', 'supplierName': 'Welcomhotel By ITC Hotels, Bella Vista', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11966': {'supplierCode': 'DR11538', 'supplierName': 'Goldfinch', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11968': {'supplierCode': 'DR12790', 'supplierName': 'Ramee Guestline (Khar)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11970': {'supplierCode': 'DR12629', 'supplierName': 'Tendu Leaf Jungle Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11971': {'supplierCode': 'DR11538', 'supplierName': 'Classic The Mall', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11973': {'supplierCode': 'DR11538', 'supplierName': 'Marari Sea Laps Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11974': {'supplierCode': 'DR11538', 'supplierName': 'Sea Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11975': {'supplierCode': 'DR11538', 'supplierName': 'Marari Village Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11978': {'supplierCode': 'DR11538', 'supplierName': 'Welcomhotel Shimla', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11979': {'supplierCode': 'DR11538', 'supplierName': 'Vijayatej Clarks Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11980': {'supplierCode': 'DR11538', 'supplierName': 'T2 Beacon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11981': {'supplierCode': 'DR11538', 'supplierName': 'Beyond by Sula', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11982': {'supplierCode': 'DR12648', 'supplierName': 'The Lalita Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11983': {'supplierCode': 'DR11538', 'supplierName': 'Banyan Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11984': {'supplierCode': 'DR11538', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11985': {'supplierCode': 'DR11538', 'supplierName': 'Avana Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11986': {'supplierCode': 'DR11538', 'supplierName': 'Courtyard by Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11987': {'supplierCode': 'DR11538', 'supplierName': 'Express Inn Nashik', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11988': {'supplierCode': 'DR12688', 'supplierName': 'By The Riverside', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11990': {'supplierCode': 'DR11538', 'supplierName': 'The Westin Garden City', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11991': {'supplierCode': 'DR11538', 'supplierName': 'Estuary Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11993': {'supplierCode': 'DR11538', 'supplierName': 'The Source', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11995': {'supplierCode': 'DR11538', 'supplierName': 'Bawa Continental', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11997': {'supplierCode': 'DR12509', 'supplierName': 'Blackberry Hills', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11998': {'supplierCode': 'DR11538', 'supplierName': 'Lemon Tree Port Blair', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO11999': {'supplierCode': 'DR11538', 'supplierName': 'Bloom Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12000': {'supplierCode': 'DR11538', 'supplierName': 'Radisson', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12001': {'supplierCode': 'DR11538', 'supplierName': 'Chateau Windsor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12002': {'supplierCode': 'DR11538', 'supplierName': 'Peerless Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12003': {'supplierCode': 'DR12583', 'supplierName': 'Apani Dhani Eco-Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12004': {'supplierCode': 'DR11538', 'supplierName': 'Chandys Drizzle Drops', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12005': {'supplierCode': 'DR11538', 'supplierName': 'Malabar Ocean Front Resort And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12006': {'supplierCode': 'DR11538', 'supplierName': 'Sinclairs Bay View', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12007': {'supplierCode': 'DR11538', 'supplierName': 'Four Points By Sheraton', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12011': {'supplierCode': 'DR11538', 'supplierName': 'Minerva Grand Nellore', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12012': {'supplierCode': 'DR12544', 'supplierName': 'Clouds Valley', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12014': {'supplierCode': 'DR11538', 'supplierName': 'Suba Galaxy', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12015': {'supplierCode': 'DR12666', 'supplierName': 'Kariappa House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12016': {'supplierCode': 'DR11538', 'supplierName': 'Mosaic', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12017': {'supplierCode': 'DR11538', 'supplierName': 'Devonshire Greens', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12018': {'supplierCode': 'DR11538', 'supplierName': 'Hyatt Centric', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12019': {'supplierCode': 'DR11538', 'supplierName': 'La Maison Radha', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12020': {'supplierCode': 'DR11538', 'supplierName': 'Radisson', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12021': {'supplierCode': 'DR11538', 'supplierName': 'Lagoon Sarovar Premiere', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12022': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Blu Mbd', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12024': {'supplierCode': 'DR11538', 'supplierName': 'Rivulet Hospitality', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12026': {'supplierCode': 'DR12510', 'supplierName': 'Rose Gardens Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12027': {'supplierCode': 'DR11538', 'supplierName': 'The Westin Pune Koregaon Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12028': {'supplierCode': 'DR11538', 'supplierName': 'Club Mahindra Derby Green', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12029': {'supplierCode': 'DR11538', 'supplierName': 'Silver Tips', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12030': {'supplierCode': 'DR12457', 'supplierName': 'Vivanta (Hinjawadi)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12031': {'supplierCode': 'DR11538', 'supplierName': 'Star Emirates Luxury Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12032': {'supplierCode': 'DR11538', 'supplierName': 'Gem Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12033': {'supplierCode': 'DR11538', 'supplierName': 'The Hans Coco Palm', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12034': {'supplierCode': 'DR10998', 'supplierName': 'The Richmond', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12035': {'supplierCode': 'DR11538', 'supplierName': 'Highland', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12036': {'supplierCode': 'DR11538', 'supplierName': 'Gulaab Niwaas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12037': {'supplierCode': 'DR11538', 'supplierName': 'Atmantan Wellness Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12039': {'supplierCode': 'DR12318', 'supplierName': 'The Grand Barso', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12040': {'supplierCode': 'DR11538', 'supplierName': 'La Montana by TGI', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12041': {'supplierCode': 'DR11538', 'supplierName': 'Conrad', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12042': {'supplierCode': 'DR12662', 'supplierName': 'Chandera Kothi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12043': {'supplierCode': 'DR11538', 'supplierName': 'Courtyard By Marriott Pune Chakan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12044': {'supplierCode': 'DR11538', 'supplierName': 'Hacra Dhani', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12045': {'supplierCode': 'DR11538', 'supplierName': 'Pushkar Organic', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12046': {'supplierCode': 'DR11538', 'supplierName': 'Sai Towers', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12047': {'supplierCode': 'DR11538', 'supplierName': 'Larisa Hotels & Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12048': {'supplierCode': 'DR11538', 'supplierName': 'Shimla Havens Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12049': {'supplierCode': 'DR11538', 'supplierName': 'Mayfair Lake Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12052': {'supplierCode': 'DR11538', 'supplierName': 'Singhania Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12054': {'supplierCode': 'DR11538', 'supplierName': 'Crowne Plaza Pune City Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12055': {'supplierCode': 'DR11538', 'supplierName': 'Fortune Select Grand Ridge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12056': {'supplierCode': 'DR11538', 'supplierName': 'Manjeera Sarovar Premiere', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12057': {'supplierCode': 'DR11538', 'supplierName': 'E-Square The Fern', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12058': {'supplierCode': 'DR11538', 'supplierName': 'Marasa Sarovar Premiere', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12059': {'supplierCode': 'DR11538', 'supplierName': 'Royal Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12061': {'supplierCode': 'DR11538', 'supplierName': 'Marasa Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12062': {'supplierCode': 'DR11538', 'supplierName': 'Lotus Nikko', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12063': {'supplierCode': 'DR11538', 'supplierName': 'Bungalow On The Beach', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12064': {'supplierCode': 'DR11538', 'supplierName': 'Hyatt', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12065': {'supplierCode': 'DR11538', 'supplierName': 'MGM Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12067': {'supplierCode': 'DR11538', 'supplierName': 'Neelambari Ecotourism', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12068': {'supplierCode': 'DR11538', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12069': {'supplierCode': 'DR12437', 'supplierName': 'Taj Dal View', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12070': {'supplierCode': 'DR11538', 'supplierName': 'Rajah Eco Beach', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12071': {'supplierCode': 'DR11538', 'supplierName': 'Mercure Chennai Sriperumbudur', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12072': {'supplierCode': 'DR11538', 'supplierName': 'Vedaaranya Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12073': {'supplierCode': 'DR11538', 'supplierName': 'Marriott Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12074': {'supplierCode': 'DR11538', 'supplierName': 'Sitaram Beach Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12076': {'supplierCode': 'DR11538', 'supplierName': 'Saptha Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12079': {'supplierCode': 'DR12564', 'supplierName': 'Sunderban Tiger Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12080': {'supplierCode': 'DR11538', 'supplierName': 'Yantra Ayurvedic Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12081': {'supplierCode': 'DR11538', 'supplierName': 'Le Lac Sarovar Portico', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12082': {'supplierCode': 'DR11538', 'supplierName': 'Park Inn by Radisson', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12083': {'supplierCode': 'DR11538', 'supplierName': 'Ramee Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12084': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Blu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12085': {'supplierCode': 'DR12561', 'supplierName': 'KVM Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12086': {'supplierCode': 'DR11538', 'supplierName': 'Surat Marriott Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12087': {'supplierCode': 'DR11538', 'supplierName': 'Royal Orchid Central', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12088': {'supplierCode': 'DR11538', 'supplierName': 'Mayas Kem Pride', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12090': {'supplierCode': 'DR12663', 'supplierName': 'Maitreya Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12091': {'supplierCode': 'DR11538', 'supplierName': 'Srm Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12092': {'supplierCode': 'DR12453', 'supplierName': 'Jungle Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12094': {'supplierCode': 'DR11538', 'supplierName': 'Fortune Park Lakecity', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12095': {'supplierCode': 'DR11538', 'supplierName': 'Ayurbirth Healing Centre', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12096': {'supplierCode': 'DR11538', 'supplierName': 'Dr. Franklin\'S Panchakarma Institute And Research', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12097': {'supplierCode': 'DR11538', 'supplierName': 'Ibis Thane', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12098': {'supplierCode': 'DR11538', 'supplierName': 'Forte Manor By Siel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12099': {'supplierCode': 'DR11538', 'supplierName': 'Hyatt Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12100': {'supplierCode': 'DR11538', 'supplierName': 'Gnanam', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12101': {'supplierCode': 'DR11538', 'supplierName': 'The Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12102': {'supplierCode': 'DR11538', 'supplierName': 'La orchid', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12103': {'supplierCode': 'DR12683', 'supplierName': 'The Earth', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12104': {'supplierCode': 'DR11538', 'supplierName': 'Best Western Vrindavan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12106': {'supplierCode': 'DR11538', 'supplierName': 'Parisutham', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12109': {'supplierCode': 'DR11538', 'supplierName': 'Contour Island Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12110': {'supplierCode': 'DR11538', 'supplierName': 'Hornbill Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12111': {'supplierCode': 'DR11538', 'supplierName': 'Atali Ganga', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12112': {'supplierCode': 'DR11538', 'supplierName': 'DFS Grand Plazas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12113': {'supplierCode': 'DR10996', 'supplierName': 'Edakkal Hermitage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12114': {'supplierCode': 'DR11538', 'supplierName': 'Amaana Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12115': {'supplierCode': 'DR12610', 'supplierName': 'Bhairavgarh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12116': {'supplierCode': 'DR12606', 'supplierName': 'Raga On The Ganges', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12118': {'supplierCode': 'DR11538', 'supplierName': 'Aranya Niwas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12119': {'supplierCode': 'DR11538', 'supplierName': 'The Great Ganga', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12120': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Chandigarh Zirakpur', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12121': {'supplierCode': 'DR11538', 'supplierName': 'Chrissie\'s', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12122': {'supplierCode': 'DR10986', 'supplierName': 'Mewar Villas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12125': {'supplierCode': 'DR11538', 'supplierName': 'Radisson Blu Palace Resort And Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12126': {'supplierCode': 'DR11538', 'supplierName': 'Club Mahindra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12128': {'supplierCode': 'DR11538', 'supplierName': 'Ramada Plaza', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12129': {'supplierCode': 'DR11538', 'supplierName': 'Gateway Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12130': {'supplierCode': 'DR11538', 'supplierName': 'Coffee And Pepper', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12131': {'supplierCode': 'DR11538', 'supplierName': 'Regenta Central Cassia', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12132': {'supplierCode': 'DR11538', 'supplierName': 'The Kehloor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12134': {'supplierCode': 'DR12573', 'supplierName': 'Spicetree Chinnar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12136': {'supplierCode': 'DR12549', 'supplierName': 'Velvet Clarks Exotica', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12137': {'supplierCode': 'DR11538', 'supplierName': 'The Royal Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12138': {'supplierCode': 'DR11538', 'supplierName': 'Shipra Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12142': {'supplierCode': 'DR11538', 'supplierName': 'Athena', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12144': {'supplierCode': 'DR11538', 'supplierName': 'Hyatt Place', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12145': {'supplierCode': 'DR11538', 'supplierName': 'Spree Resort Sariska', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12146': {'supplierCode': 'DR11538', 'supplierName': 'Palette Hill View Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12147': {'supplierCode': 'DR11538', 'supplierName': 'Samudra KTDC', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12148': {'supplierCode': 'DR11538', 'supplierName': 'De Exotica Crest', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12149': {'supplierCode': 'DR12558', 'supplierName': 'Himalayaa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12150': {'supplierCode': 'DR11538', 'supplierName': 'Golden Fern Resorts', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12151': {'supplierCode': 'DR12654', 'supplierName': 'John\'s Farm & Home', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12152': {'supplierCode': 'DR11538', 'supplierName': 'Aiden by Best Western, Hennur', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12153': {'supplierCode': 'DR11538', 'supplierName': 'Baljees Regency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12154': {'supplierCode': 'DR11538', 'supplierName': 'Seashore Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12156': {'supplierCode': 'DR11538', 'supplierName': 'Sangeet', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12157': {'supplierCode': 'DR11538', 'supplierName': 'Church View Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12158': {'supplierCode': 'DR11538', 'supplierName': 'Tijara Fort-Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12159': {'supplierCode': 'DR12600', 'supplierName': 'Vaishali Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12160': {'supplierCode': 'DR11538', 'supplierName': 'Willow Banks', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12161': {'supplierCode': 'DR11538', 'supplierName': 'Maurya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12162': {'supplierCode': 'DR11538', 'supplierName': 'Hindustan Beach Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12164': {'supplierCode': 'DR11538', 'supplierName': 'Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12165': {'supplierCode': 'DR11538', 'supplierName': 'Ganpati Guest House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12166': {'supplierCode': 'DR11538', 'supplierName': 'Sajjoys', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12167': {'supplierCode': 'DR11538', 'supplierName': 'Spoorti Resort & Club House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12168': {'supplierCode': 'DR12872', 'supplierName': 'New Hotel Broadway', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12169': {'supplierCode': 'DR11538', 'supplierName': 'The India Benares', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12170': {'supplierCode': 'DR11538', 'supplierName': 'Pristine', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12171': {'supplierCode': 'DR11538', 'supplierName': 'The Fern Residency', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12172': {'supplierCode': 'DR11538', 'supplierName': 'Sri Omkar Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12173': {'supplierCode': 'DR11538', 'supplierName': 'Rashmi Guest House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12174': {'supplierCode': 'DR11538', 'supplierName': 'Fairfield By Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12175': {'supplierCode': 'DR12303', 'supplierName': 'Dev Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12177': {'supplierCode': 'DR12397', 'supplierName': 'Gateway', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12178': {'supplierCode': 'DR12400', 'supplierName': 'Taj The Trees', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12179': {'supplierCode': 'DR11538', 'supplierName': 'Vembanad House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12180': {'supplierCode': 'DR11538', 'supplierName': 'A.S.Hotels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12181': {'supplierCode': 'DR11538', 'supplierName': 'Nilambagh Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12182': {'supplierCode': 'DR11538', 'supplierName': 'Vijay Vilas Palace', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12183': {'supplierCode': 'DR12334', 'supplierName': 'Holywater By Ganga Kinare', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12184': {'supplierCode': 'DR12487', 'supplierName': 'Renest Shraddha Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12185': {'supplierCode': 'DR12488', 'supplierName': 'The Imperial', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12186': {'supplierCode': 'DR12458', 'supplierName': 'Serene Horizon', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12189': {'supplierCode': 'DR12405', 'supplierName': 'The Ummed', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12191': {'supplierCode': 'DR12340', 'supplierName': 'Vanilla County', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12192': {'supplierCode': 'DR12506', 'supplierName': 'Kurja Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12193': {'supplierCode': 'DR12217', 'supplierName': 'Meri Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12194': {'supplierCode': 'DR10249', 'supplierName': 'Punnamada Resorts (Houseboat)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12195': {'supplierCode': 'DR10974', 'supplierName': 'Royal Lotus View', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12196': {'supplierCode': 'DR12498', 'supplierName': 'Balaram Palace Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12197': {'supplierCode': 'DR12521', 'supplierName': 'Svanir', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12198': {'supplierCode': 'DR12304', 'supplierName': 'Stone Wood Maan Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12199': {'supplierCode': 'DR12505', 'supplierName': 'Neemrana\'s Three Waters', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12203': {'supplierCode': 'DR10968', 'supplierName': 'Bloom Hotel (Worli)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12204': {'supplierCode': 'DR12238', 'supplierName': 'Padmini Bagh Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12206': {'supplierCode': 'DR12615', 'supplierName': 'Wild Mahseer', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12207': {'supplierCode': 'DR12295', 'supplierName': 'Mahal Khandela', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12208': {'supplierCode': 'DR12635', 'supplierName': 'Saura', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12209': {'supplierCode': 'DR12637', 'supplierName': 'Raghu Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12211': {'supplierCode': 'DR12216', 'supplierName': 'The Pearl', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12212': {'supplierCode': 'DR12411', 'supplierName': 'Jagat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12213': {'supplierCode': 'DR12652', 'supplierName': 'Dharana At Shillim', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12214': {'supplierCode': 'DR12518', 'supplierName': 'Heritage Mandawa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12215': {'supplierCode': 'DR12657', 'supplierName': 'Norling House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12216': {'supplierCode': 'DR12661', 'supplierName': 'Pinnacle', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12217': {'supplierCode': 'DR12655', 'supplierName': 'The Nanee', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12218': {'supplierCode': 'DR11538', 'supplierName': 'The Old Inn', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12219': {'supplierCode': 'DR12655', 'supplierName': 'Birethanti Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12220': {'supplierCode': 'DR12655', 'supplierName': 'Heritage Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12221': {'supplierCode': 'DR12665', 'supplierName': 'Costa River', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12223': {'supplierCode': 'DR12655', 'supplierName': 'Atithi Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12224': {'supplierCode': 'DR12655', 'supplierName': 'Soaltee Westend Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12225': {'supplierCode': 'DR12655', 'supplierName': 'Mystic Mountain', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12226': {'supplierCode': 'DR12655', 'supplierName': 'Barahi', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12227': {'supplierCode': 'DR12655', 'supplierName': 'Lake View Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12228': {'supplierCode': 'DR12655', 'supplierName': 'Buddha Maya Garden', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12229': {'supplierCode': 'DR12655', 'supplierName': 'Fish Tail Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12230': {'supplierCode': 'DR12655', 'supplierName': 'Kathmandu Guest House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12231': {'supplierCode': 'DR12655', 'supplierName': 'Le Himalaya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12232': {'supplierCode': 'DR12655', 'supplierName': 'Landmark', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12233': {'supplierCode': 'DR12655', 'supplierName': 'Aabas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12234': {'supplierCode': 'DR12668', 'supplierName': 'Sardargarh Heritage', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12236': {'supplierCode': 'DR12655', 'supplierName': 'Barahi Jungle Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12237': {'supplierCode': 'DR12655', 'supplierName': 'Green Mansions Jungle Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12238': {'supplierCode': 'DR12655', 'supplierName': 'Ghandruk Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12239': {'supplierCode': 'DR12655', 'supplierName': 'Dwarika\'s', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12240': {'supplierCode': 'DR12655', 'supplierName': 'Hyatt Centric', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12241': {'supplierCode': 'DR12655', 'supplierName': 'Taj Safari', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12242': {'supplierCode': 'DR12655', 'supplierName': 'Pokhara Grande', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12243': {'supplierCode': 'DR12655', 'supplierName': 'Neydo', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12244': {'supplierCode': 'DR12655', 'supplierName': 'Monjo Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12245': {'supplierCode': 'DR12655', 'supplierName': 'Basera Boutique', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12246': {'supplierCode': 'DR12655', 'supplierName': 'Namche Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12247': {'supplierCode': 'DR12655', 'supplierName': 'Phakding Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12248': {'supplierCode': 'DR12655', 'supplierName': 'Shambaling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12249': {'supplierCode': 'DR12655', 'supplierName': 'Mount Kailash Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12250': {'supplierCode': 'DR12655', 'supplierName': 'Gorkha Garden', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12251': {'supplierCode': 'DR12655', 'supplierName': 'Tomijong Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12252': {'supplierCode': 'DR12655', 'supplierName': 'Nepali Ghar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12253': {'supplierCode': 'DR12655', 'supplierName': 'Temple Tree Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12254': {'supplierCode': 'DR12655', 'supplierName': 'Jungle Villa Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12255': {'supplierCode': 'DR12655', 'supplierName': 'Maya Manor Boutique', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12256': {'supplierCode': 'DR12655', 'supplierName': 'Chandragiri Hills Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12257': {'supplierCode': 'DR12655', 'supplierName': 'Aagantuk Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12259': {'supplierCode': 'DR12655', 'supplierName': 'Fairfield By Marriott', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12260': {'supplierCode': 'DR12655', 'supplierName': 'Heritage Hotel Suites & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12261': {'supplierCode': 'DR12655', 'supplierName': 'Varnabas Musuem', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12262': {'supplierCode': 'DR12655', 'supplierName': 'Kavya Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12263': {'supplierCode': 'DR12655', 'supplierName': 'The Terraces Resort & Spa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12264': {'supplierCode': 'DR12655', 'supplierName': 'The Inn Patan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12265': {'supplierCode': 'DR12655', 'supplierName': 'The Fort Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12266': {'supplierCode': 'DR12655', 'supplierName': 'Riverside Spring Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12267': {'supplierCode': 'DR12655', 'supplierName': 'Dolmaling', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12268': {'supplierCode': 'DR12655', 'supplierName': 'Hotel Srinagar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12269': {'supplierCode': 'DR12655', 'supplierName': 'Mulberry', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12270': {'supplierCode': 'DR12655', 'supplierName': 'Dhampus Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12271': {'supplierCode': 'DR12655', 'supplierName': 'Gaun Ghar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12272': {'supplierCode': 'DR12701', 'supplierName': 'The Old Court House', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12274': {'supplierCode': 'DR12697', 'supplierName': 'Test Supplier - Audit Walk Through', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12275': {'supplierCode': 'DR12300', 'supplierName': 'Divine Lakshmi Ganga', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12277': {'supplierCode': 'DR11302', 'supplierName': 'Amika', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12278': {'supplierCode': 'DR12713', 'supplierName': 'Roopa Elite', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12280': {'supplierCode': 'DR12716', 'supplierName': 'Denwa Backwater Escape', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12281': {'supplierCode': 'DR12655', 'supplierName': 'Square', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12282': {'supplierCode': 'DR12655', 'supplierName': 'Jagatpur Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12287': {'supplierCode': 'DR12748', 'supplierName': 'Ginger', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12288': {'supplierCode': 'DR12717', 'supplierName': 'Anantara Jewel Bagh', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12289': {'supplierCode': 'DR12747', 'supplierName': 'Da Bungalow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12290': {'supplierCode': 'DR12736', 'supplierName': 'Club Mahindra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12291': {'supplierCode': 'DR12735', 'supplierName': 'Tuli Tiger Corridor', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12293': {'supplierCode': 'DR12730', 'supplierName': 'Park Prime', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12295': {'supplierCode': 'DR12757', 'supplierName': 'Jhadol Vijay Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12296': {'supplierCode': 'DR12758', 'supplierName': 'Lutyens Bungalow', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12297': {'supplierCode': 'DR12755', 'supplierName': 'Amar Bagh Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12300': {'supplierCode': 'DR12756', 'supplierName': 'Club Mahindra', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12301': {'supplierCode': 'DR12734', 'supplierName': 'Kathoni', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12303': {'supplierCode': 'DR12768', 'supplierName': 'Neel Clarks Inn Express', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12304': {'supplierCode': 'DR12275', 'supplierName': 'Planters Homestay', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12305': {'supplierCode': 'DR10696', 'supplierName': 'Thar Oasis Resort & Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12306': {'supplierCode': 'DR12718', 'supplierName': 'Amaya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12307': {'supplierCode': 'DR12763', 'supplierName': 'Ravla Bhenswara', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12308': {'supplierCode': 'DR12655', 'supplierName': 'Baber Mahal Vilas', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12309': {'supplierCode': 'DR12655', 'supplierName': 'Balthali Village Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12310': {'supplierCode': 'DR12655', 'supplierName': 'Heritage Malla', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12311': {'supplierCode': 'DR12655', 'supplierName': 'Safari Narayani', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12312': {'supplierCode': 'DR12655', 'supplierName': 'Majgaon Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12313': {'supplierCode': 'DR12655', 'supplierName': 'Dwarika\'s Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12314': {'supplierCode': 'DR12778', 'supplierName': 'Halez Sparsa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12315': {'supplierCode': 'DR12565', 'supplierName': 'Kaithal Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12316': {'supplierCode': 'DR12293', 'supplierName': 'Country Inn & Suites', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12318': {'supplierCode': 'DR12467', 'supplierName': 'Isla\'s Ridge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12320': {'supplierCode': 'DR12782', 'supplierName': 'Residence De L Eveche', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12321': {'supplierCode': 'DR12808', 'supplierName': 'Welcomhotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12322': {'supplierCode': 'DR12738', 'supplierName': 'Mementos By ITC', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12323': {'supplierCode': 'DR12827', 'supplierName': 'The Crown - IHCL SeleQtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12324': {'supplierCode': 'DR12803', 'supplierName': 'Raajsa - IHCL SeleQtions', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12325': {'supplierCode': 'DR12776', 'supplierName': 'Rawla Bisalpur', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12326': {'supplierCode': 'DR12381', 'supplierName': 'Nine Furlongs', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12327': {'supplierCode': 'DR12773', 'supplierName': 'Bamboo Grove', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12328': {'supplierCode': 'DR12777', 'supplierName': 'Bori Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12330': {'supplierCode': 'DR11538', 'supplierName': 'Java Rain Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12331': {'supplierCode': 'DR12788', 'supplierName': 'Bloom Hotel Juhu', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12332': {'supplierCode': 'DR11538', 'supplierName': 'Sinclairs Retreat', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12333': {'supplierCode': 'DR12791', 'supplierName': 'Anantya In The Village', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12337': {'supplierCode': 'DR10567', 'supplierName': 'Chettinadu Mansion', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12338': {'supplierCode': 'DR12794', 'supplierName': 'Sterling Park', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12339': {'supplierCode': 'DR12749', 'supplierName': 'Sterling V Grand', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12341': {'supplierCode': 'DR12796', 'supplierName': 'Sterling Aravalli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12342': {'supplierCode': 'DR12804', 'supplierName': 'Le Montfort Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12343': {'supplierCode': 'DR12786', 'supplierName': 'Cullinan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12344': {'supplierCode': 'DR12805', 'supplierName': 'Welcomhotel By ITC Hotels (Delhi - Gurugram Highway)', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12345': {'supplierCode': 'DR12801', 'supplierName': 'Storii By ITC Hotels', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12346': {'supplierCode': 'DR12787', 'supplierName': 'The Roseate Ganges', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12347': {'supplierCode': 'DR12792', 'supplierName': 'Fortune Beachfront', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12349': {'supplierCode': 'DR12807', 'supplierName': 'BluSalzz Residence', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12351': {'supplierCode': 'DR12669', 'supplierName': 'Ranthambhore Aangan', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12352': {'supplierCode': 'DR12206', 'supplierName': 'The Ramayana', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12354': {'supplierCode': 'DR12368', 'supplierName': 'Rose Amer - S', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12355': {'supplierCode': 'DR12703', 'supplierName': 'Chunda Haveli', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12356': {'supplierCode': 'DR12752', 'supplierName': 'Atithi Camp & Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12357': {'supplierCode': 'DR12818', 'supplierName': 'Padmaa', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12358': {'supplierCode': 'DR12813', 'supplierName': 'Palighar', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12359': {'supplierCode': 'DR12705', 'supplierName': 'Regenta Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12360': {'supplierCode': 'DR10601', 'supplierName': 'Reni Pani Jungle Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12361': {'supplierCode': 'DR12777', 'supplierName': 'Bori Safari Lodge', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12362': {'supplierCode': 'DR12819', 'supplierName': 'Meru Vann', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12363': {'supplierCode': 'DR12814', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12364': {'supplierCode': 'DR12785', 'supplierName': 'Lemon Tree', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12367': {'supplierCode': 'DR12781', 'supplierName': 'Neemrana\'s Coconut Alley', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12368': {'supplierCode': 'DR12811', 'supplierName': 'Justa Morjim Beach Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12369': {'supplierCode': 'DR12653', 'supplierName': 'Bookmark Resorts, Jogi Mahal', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12370': {'supplierCode': 'DR12281', 'supplierName': 'IRA By Orchid', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12372': {'supplierCode': 'DR12330', 'supplierName': 'Tree Of Life Highlands', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12373': {'supplierCode': 'DR12853', 'supplierName': 'Puratan Qila', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12376': {'supplierCode': 'DR12843', 'supplierName': 'Fortune Beach Resort ECR', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12377': {'supplierCode': 'DR12857', 'supplierName': 'Taj Mussoorie Foothills', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12378': {'supplierCode': 'DR11538', 'supplierName': 'Abhaneri Village Safari Camp', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12380': {'supplierCode': 'DR12040', 'supplierName': 'AMS Test', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12381': {'supplierCode': 'DR11538', 'supplierName': 'TEst VK HOtel API', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12384': {'supplierCode': 'DR12330', 'supplierName': 'Test Aihole Hotel', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12385': {'supplierCode': 'DR12884', 'supplierName': 'AMS hotel test 1', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12388': {'supplierCode': 'DR12901', 'supplierName': 'Bhutan Package', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12389': {'supplierCode': 'DR12902', 'supplierName': 'Sri Lanka Package', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12391': {'supplierCode': 'DR11645', 'supplierName': 'Carmelia Haven Resort', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12393': {'supplierCode': 'DR10497', 'supplierName': 'Surya', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
    'HO12394': {'supplierCode': 'DR10101', 'supplierName': 'Test vk', 'parentGroup': None, 'parentGroupId': None, 'supplierVersion': 1},
}

# ── HOTEL SUPPLIER LOOKUP BY CITY+NAME ───────────────────────────────────────
# Fallback when businessCode is not yet in _HOTEL_SUPPLIER_MAP.
# Key: (city_lower, partial_hotel_name_lower) — checked via 'in' match on hotel name
# Used when portfolio update is pending with full HO code list.
_HOTEL_SUPPLIER_BY_CITY_NAME = {
    # Entries promoted to _HOTEL_SUPPLIER_MAP once HO codes confirmed:
    # HO10355 = Ramada Amritsar (DR10786), HO10336 = The Oberoi Cecil (DR10610)
    # Add future unknowns here until HO code is confirmed from portfolio CSV.
}

def _get_supplier_info(business_code, hotel_name, city_name):
    """Resolve supplier info — businessCode map first, then city+name fallback."""
    sr = _HOTEL_SUPPLIER_MAP.get(business_code)
    if sr:
        return sr
    city_lower = (city_name or '').lower()
    name_lower = (hotel_name or '').lower()
    for (city_key, name_key), sr2 in _HOTEL_SUPPLIER_BY_CITY_NAME.items():
        if city_key in city_lower and name_key in name_lower:
            return sr2
    return {}


# ── FROM: HEADER PATTERN ───────────────────────────────────────────────────────
_FROM_PATTERN = re.compile(
    r'from:\s*(?:"?([^"<\n\r]{2,}?)"?\s+)?<?([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})>?',
    re.IGNORECASE
)


def _extract_start_date(email_raw):
    """Extract travel start date from email body.
    Handles: "Travel dates: 15 May 2026", "15 May 2026", "Travel date: May 15 2026"
    Returns ISO date string or None.
    """
    import re as _re2
    from datetime import datetime as _dt3
    _FR_MONTHS = {
        'janvier':'january','f\xe9vrier':'february','fevrier':'february',
        'mars':'march','avril':'april','mai':'may','juin':'june',
        'juillet':'july','ao\xfbt':'august','aout':'august',
        'septembre':'september','octobre':'october',
        'novembre':'november','d\xe9cembre':'december','decembre':'december',
    }
    patterns = [
        r'travel\s+dates?\s*:?\s*(\d{1,2}\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{4})',
        r'travel\s+dates?\s*:?\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{1,2}\s*,?\s*\d{4})',
        r'depart(?:ure)?\s+dates?\s*:?\s*(\d{1,2}\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{4})',
        r'dates?\s+de\s+voyage\s*:?\s*(\d{1,2}\s+\w+\s+\d{4})',
        r'date\s+de\s+d[e\xe9]part\s*:?\s*(\d{1,2}\s+\w+\s+\d{4})',
        r'reisedatum\s*:?\s*(\d{1,2}\.?\s+\w+\s+\d{4})',
        r'(\d{4}-\d{2}-\d{2})',
        r'(\d{1,2}\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{4})',
    ]
    for pat in patterns:
        m = _re2.search(pat, email_raw, _re2.IGNORECASE)
        if m:
            s = m.group(1).strip()
            sl = s.lower()
            for fr, en in _FR_MONTHS.items():
                sl = _re2.sub(fr, en, sl, flags=_re2.IGNORECASE)
            s = sl.title()
            for fmt in ['%d %B %Y', '%d %b %Y', '%B %d %Y', '%b %d %Y',
                        '%B %d, %Y', '%Y-%m-%d', '%d. %B %Y']:
                try:
                    return _dt3.strptime(s.strip(), fmt).strftime('%Y-%m-%d')
                except ValueError:
                    continue
    return None

def build_api_payload(intent, email_raw):
    """Build the CRM query creation payload from engine intent + raw email.

    Fields derived automatically:
        sourceMarket      — intent['source_market']
        customer          — intent['agency_account_code'] (from CRM portfolio)
        tentativeStartDate — intent['travel_start'] (from PNR)
        currency          — derived from source_market
        emailBody         — raw email text
        senderName/Email  — parsed from From: header
        queryType         — tailormade (≤6 pax) or group (>6 pax); default tailormade
        name              — auto-generated: "Agency – Region – DDMMMYYYY"

    Returns dict ready to POST to /queries endpoint.
    """
    from datetime import datetime as _dt
    now = _dt.utcnow()

    # ── Sender from From: header ───────────────────────────────────────────────
    m = _FROM_PATTERN.search(email_raw)
    sender_name  = m.group(1).strip() if m and m.group(1) else ''
    sender_email = m.group(2).strip() if m else ''

    # ── Individual sender from email signature ─────────────────────────────────
    # Scan email body for a personal name + individual email in signature block
    # Pattern: line with a personal name followed by a title/role line
    # and an individual email (firstname.lastname@domain)
    _sig_email_pattern = re.compile(
        r'([a-zA-Z0-9._%+\-]+\.[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})'
    )
    _sig_name_pattern = re.compile(
        r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s*$'  # 2-4 capitalised words on own line
    )
    # Look for individual email in signature (different from From: email)
    sig_email = ''
    sig_name  = ''
    lines = email_raw.splitlines()
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        # Find an individual email (firstname.lastname format) in signature area
        sig_m = _sig_email_pattern.search(line_stripped)
        if sig_m:
            candidate = sig_m.group(1)
            # Must differ from the From: email and look like firstname.lastname
            local = candidate.split('@')[0]
            if '.' in local and candidate != sender_email:
                sig_email = candidate
                # Look for name in surrounding lines (6 before + 2 after)
                search_range = list(range(max(0, i-6), i)) + list(range(i+1, min(len(lines), i+3)))
                for j in search_range:
                    name_m = _sig_name_pattern.match(lines[j].strip())
                    if name_m:
                        sig_name = name_m.group(1)
                        break
                # Fallback: look for "Firstname Lastname" pattern anywhere nearby
                if not sig_name:
                    for j in range(max(0, i-8), min(len(lines), i+3)):
                        candidate_name = lines[j].strip()
                        parts = candidate_name.split()
                        if (2 <= len(parts) <= 3 and
                            all(p[0].isupper() and p[1:].islower() for p in parts if p and p[0].isalpha())):
                            sig_name = candidate_name
                            break
                if sig_email:
                    break

    # Use individual sender details if found, else fall back to From: header
    actual_sender_name  = sig_name  or sender_name
    actual_sender_email = sender_email or sig_email  # prefer From: header — more reliable

    # ── Guest/client name ──────────────────────────────────────────────────────
    # CSV path: use passenger name directly from parsed intent
    if intent.get('csv_path') and intent.get('guest_name'):
        guest_name = intent['guest_name']
    else:
        guest_name = ''
    _guest_patterns = [
        re.compile(r're\s*:\s*(?:mr|mrs|ms|miss|dr|prof)\.?\s+([A-Za-z\s/\-]+?)(?:\s*[+x]\s*\d|\s*\d\s*pax|\s*$)', re.IGNORECASE),
        re.compile(r'client\s*:\s*([A-Za-z\s/\-]+?)(?:\s*[+x]\s*\d|\s*\d\s*pax|\s*$)', re.IGNORECASE),
        re.compile(r'(?:for|ref|ref\.)\s*:\s*([A-Za-z\s/\-]+?)(?:\s*[+x]\s*\d|\s*\d\s*pax|\s*$)', re.IGNORECASE),
        # Subject: "New Request India - Mr and Mrs. Dittberner"
        re.compile(r'subject:.*?(?:mr|mrs|ms|miss|dr|prof)\.?\s+(?:and\s+(?:mr|mrs|ms)\.?\s+)?([A-Za-z\-]+)', re.IGNORECASE),
        # "* Mr. Jan Dittberner" or "Mr. Jan Dittberner and Mrs."
        re.compile(r'(?:mr|mrs|ms|miss|dr|prof)\.?\s+[A-Za-z]+\s+([A-Za-z\-]+?)(?:\s+and|\s*$|\s*\()', re.IGNORECASE),
        # "client names:" section — "* Mr. Jan Dittberner"
        re.compile(r'client\s+names?\s*:.*?(?:mr|mrs|ms)\.?\s+\w+\s+([A-Za-z\-]+)', re.IGNORECASE | re.DOTALL),
    ]
    for line in email_raw.splitlines()[:30]:  # scan first 30 lines
        for pat in _guest_patterns[:-1]:  # line-by-line patterns (not DOTALL)
            gm = pat.search(line.strip())
            if gm:
                raw = gm.group(1).strip().rstrip('/-')
                guest_name = raw.split('/')[0].strip().title()
                # Reject if too short or looks like a city/keyword
                if len(guest_name) > 2 and not any(
                    kw in guest_name.lower() for kw in ['india', 'delhi', 'tour', 'request', 'feb', 'mar', 'jan']
                ):
                    break
                else:
                    guest_name = ''
        if guest_name:
            break
    # Fallback: scan full email with DOTALL pattern (client names section)
    if not guest_name:
        gm = _guest_patterns[-1].search(email_raw)
        if gm:
            raw = gm.group(1).strip().rstrip('/-').split('/')[0].strip().title()
            if len(raw) > 2:
                guest_name = raw

    # ── queryType rule: ≤6 pax → tailormade, >6 → group, unknown → tailormade ─
    pax = intent.get('num_pax')
    if pax and pax > 6:
        query_type = 'group-range'
    else:
        query_type = 'tailormade'   # default when pax unknown

    # ── Query name ─────────────────────────────────────────────────────────────
    agency   = intent.get('agency_name') or sender_name or 'Unknown Agency'
    # Use travel start date for query name
    _ts = intent.get('travel_start') or _extract_start_date(email_raw)
    # Warn if travel dates are in the past — TCI rejects past tentativeStartDate
    if _ts and str(_ts) < now.strftime('%Y-%m-%d'):
        print(f"  ⚠ Travel start {_ts} is in the past — using today as tentativeStartDate for TCI")
    if _ts:
        from datetime import datetime as _dts
        try:
            date_str = _dts.strptime(str(_ts), '%Y-%m-%d').strftime('%d%b%Y').upper()
        except ValueError:
            date_str = now.strftime('%d%b%Y').upper()
    else:
        date_str = now.strftime('%d%b%Y').upper()

    # Route label: first → last city (skip 0n transits), e.g. "New Delhi – Aurangabad (9 cities)"
    nightly = intent.get('nightly_split') or []
    route_cities = [c for c, n in nightly if n > 0]
    if len(route_cities) >= 2:
        route_label = f"{route_cities[0]} – {route_cities[-1]} ({len(route_cities)} cities)"
    elif len(route_cities) == 1:
        route_label = route_cities[0]
    else:
        regions = intent.get('regions_detected') or []
        REGION_LABELS = {
            'north_india': 'North India', 'south_india': 'South India',
            'central_india': 'Central India', 'east_india': 'East India',
            'west_india': 'West India', 'himalaya': 'Himalaya',
        }
        from collections import Counter as _Counter
        city_regions = [CITY_TO_REGION.get(c, '') for c in intent.get('cities_detected', [])]
        dominant = _Counter(r for r in city_regions if r).most_common(1)
        dom = dominant[0][0] if dominant else (regions[0] if regions else 'india')
        route_label = REGION_LABELS.get(dom, dom.replace('_', ' ').title())

    if guest_name:
        name = f"{agency} – {guest_name} – {route_label} – {date_str}"
    else:
        name = f"{agency} – {route_label} – {date_str}" 

    # ── Ambiguity warning ──────────────────────────────────────────────────────
    if intent.get('agency_ambiguous'):
        all_accts = intent.get('agency_all_accounts', [])
        print(f"  ⚠ API payload: ambiguous domain — using {intent.get('agency_account_code')} "
              f"({intent.get('agency_name')}). Other accounts: "
              f"{', '.join(a[0]+' '+a[1] for a in all_accts[:3])}")

    return {
        'name':                 name,
        'sourceMarket':         intent.get('source_market', 'UNKNOWN'),
        'customerType':         'fto',
        'queryType':            query_type,
        'currency':             'INR',  # Always INR — all queries billed in INR regardless of source market
        'emailDate':            now.strftime('%Y-%m-%d'),
        'emailTime':            now.strftime('%H:%M'),
        'emailBody':            (' '.join(
                email_raw.encode('ascii', 'ignore').decode('ascii').split()
            ) or f"Query created from {intent.get('_input_mode','CSV/PDF')} — Ref: {intent.get('itinerary_ref','N/A')} — Guest: {intent.get('guest_name','N/A')}")[:2000],
        'senderName':           actual_sender_name  or intent.get('agency_name', 'SITA Travel'),
        'senderEmail':          actual_sender_email or 'queries@sita.in',
        'customer':             intent.get('agency_account_code') or _default_account(intent.get('source_market','')),
        'attachments':          [],
        'tentativeStartDate':   (
                lambda d: d if d and d >= now.strftime('%Y-%m-%d')
                else now.strftime('%Y-%m-%d')
            )(
                str(intent.get('travel_start'))
                if intent.get('travel_start')
                else _extract_start_date(email_raw)
            ),
        'inputCreditRequired':  False,
        'customerSeries':       '',
        'seriesName':           '',
        # Guest name stored for reference (not a system field)
        **({'comment': f'Client: {guest_name}'} if guest_name else {}),
        'selfAssign':           False,
    }


def create_query(payload, api_base_url, api_token):
    """POST the query payload to the CRM API (Step 1 only).
    Returns the parsed JSON response (expected to contain data.id).
    Raises requests.HTTPError on non-2xx responses.
    """
    import requests as _req
    r = _req.post(
        f"{api_base_url.rstrip('/')}/queries",
        json=payload,
        headers={
            'Authorization': f'Bearer {api_token}',
            'Content-Type': 'application/json',
        },
        timeout=15,
    )
    r.raise_for_status()
    return r.json()


# ── CITY NAME → SYSTEM CODE MAP ───────────────────────────────────────────────
# Engine city name → legacy system numeric code
# Source: city_codes.csv export from legacy system
CITY_CODE_MAP = {
    'New Delhi':                        '11',
    'Agra':                             '562',
    'Jaipur':                           '141',
    'Udaipur (RJ)':                     '294',
    'Jodhpur':                          '291-1',
    'Jaisalmer':                        '2991',
    'Varanasi':                         '542',
    'Kochi (Cochin)':                   '484',
    'Alappuzha (Alleppey)':             '477',
    'Thekkady (Periyar/Kumily)':        '4869',
    'Madurai':                          '452',
    'Puducherry (Pondicherry)':         '413',
    'Mamallapuram (Mahabalipuram)':     '44-1',
    'Chennai':                          '44',
    'Mumbai':                           '22',
    'Bengaluru':                        '80',
    'Kolkata':                          '33',
    'Hyderabad':                        '40',
    'Goa':                              '832',
    'Mandawa':                          '1592',
    'Bikaner':                          '151',
    'Pushkar':                          '14581',
    'Ranthambore':                      '7462-1',
    'Orchha':                           '7680',
    'Khajuraho':                        '7686',
    'Gwalior':                          '751',
    'Bandhavgarh':                      '7653',
    'Kanha':                            '7642',
    'Pench':                            '7654',
    'Panna':                            '7655',
    'Satpura':                          '7656',
    'Jabalpur':                         '761',
    'Narlai':                           '835',
    'Rohet':                            '2931',
    'Amritsar':                         '183',
    'Shimla':                           '177',
    'Dharamshala':                      '1892',
    'Leh':                              '1982',
    'Srinagar':                         '194',
    'Rishikesh':                        '135',
    'Haridwar':                         '1334',
    'Corbett':                          '5947-1',
    'Darjeeling':                       '354-2',
    'Chandigarh':                       '172',
    'Gurgaon':                          '124',
    'Fatehpur Sikri':                   '0562-1',
    'Mysuru (Mysore)':                  '821',
    'Munnar':                           '4865',
    'Kumarakom':                        '481-1',
    'Kovalam':                          '471-1',
    'Mararikulam':                      '478-1',
    'Thanjavur (Tanjore)':              '4362',
    'Kumbakonam':                       '435',
    'Chettinad':                        '4565-1',
    'Darasuram':                        '435-1',
    'Gangaikondacholapuram':            '4329',
    'Chidambaram':                      '4144',
    'Coorg (Kodagu)':                   '8272-1',
    'Ahmedabad':                        '79',
    'Aurangabad (MH)':                  '2432',
    'Prayagraj':                        '532',
    'Lucknow':                          '522',
    'Bundi':                            '747',
    'Ajmer':                            '145',
    'Kumbhalgarh':                      '2954',
    'Ranakpur':                         '2934',
    'Chittorgarh':                      '1472',
    'Panna':                            '7732',
    'Bhopal':                           '755',
    'Indore':                           '731',
    'Ujjain':                           '734',
    'Guwahati':                         '361',
    'Kaziranga':                        '3774',
    'Pelling':                          '3595',
    'Gangtok':                          '3592',
    'Kathmandu':                        '1',
    'Paro':                             'BT-11',
    'Thimpu':                           'BT-15',
    'Auroville':                        '413-1',
    'Alsisar':                          '1595',
    'Nawalgarh':                        '1594',
    'Deogarh':                          '2904',
    'Shahpura':                         '1484',
    'Samode':                           '1432',
    'Samode':                           '1432',
    'Abhaneri':                         '1420',
    'Amber':                            '1423',
    'Sarnath':                          '542-1',
    'Bodh Gaya':                        '631-1',
    'Patna':                            '612',
    'Hampi (Hosapete)':                 '8394',
    'Badami':                           '8357',
}


# ══════════════════════════════════════════════════════════════════════════
# CSV ITINERARY PARSER
# Parses structured quote request CSVs from agencies like Audley Travel
# Returns intent dict in same schema as parse_email() — zero downstream changes
# ══════════════════════════════════════════════════════════════════════════

def _parse_csv_date(date_str):
    """Parse date strings from Services CSV.
    Handles: "Mon 26 Oct 2026", "Check out: 29 Oct 2026", "2026-10-26"
    """
    if not date_str or not str(date_str).strip():
        return None
    s = str(date_str).strip()
    # Strip prefixes like "Check out: "
    for pfx in ['check out:', 'check in:', 'check-out:', 'check-in:']:
        if s.lower().startswith(pfx):
            s = s[len(pfx):].strip()
    from datetime import datetime as _dt2
    for fmt in ['%a %d %b %Y', '%d %b %Y', '%d-%b-%Y', '%d-%b-%y',
                '%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y']:
        try:
            return _dt2.strptime(s.strip(), fmt).date()
        except ValueError:
            continue
    return None


def _extract_hotel_name(details_str):
    """Extract hotel name from Details of Service column.
    "Ramada by Wyndham - 1x Deluxe room - double - 3 nights" -> ("Ramada by Wyndham", False)
    "(P) The Oberoi Cecil - 1x Deluxe..." -> ("The Oberoi Cecil", True)
    Returns (name, is_primary) -- is_primary=True when (P) prefix present.
    """
    if not details_str:
        return ('', False)
    s = str(details_str).strip()
    # Detect and strip (P) prefix -- agent primary hotel marker
    is_primary = False
    if s.startswith('(P)'):
        is_primary = True
        s = s[3:].strip()
    # Strip <<...>> annotations
    import re as _re2
    s = _re2.sub(r'<<[^>]*>>', '', s).strip()
    # Take text before first " - "
    parts = s.split(' - ')
    name = parts[0].strip()
    # Remove trailing room/pax info if split didn't work cleanly
    name = _re2.sub(r'\s*-\s*\d+x.*$', '', name).strip()
    return (name, is_primary)

def _extract_activity_name(details_str):
    """Extract activity name from Details of Service column.
    "Putting to Bed Ceremony - Private excursion with vehicle and guide. - 2 pax"
    → "Putting to Bed Ceremony"
    """
    if not details_str:
        return ''
    s = str(details_str).strip()
    # Take text before first " - "
    parts = s.split(' - ')
    return parts[0].strip()


def _map_meal_plan(meal_str):
    """Map CSV meal plan description to TCI mealPlan code."""
    if not meal_str:
        return 'cp'
    s = str(meal_str).lower()
    if 'breakfast' in s:       return 'cp'   # Continental Plan
    if 'half board' in s:      return 'map'  # Modified American Plan
    if 'full board' in s:      return 'ap'   # American Plan
    if 'all inclusive' in s:   return 'ai'
    if 'room only' in s:       return 'ep'   # European Plan
    return 'cp'  # default breakfast


# BU code → TCI account code mapping
# Add new agencies here when BU code is known but domain isn't in CRM
_BU_TO_ACCOUNT = {
    'ATUK':  ('ACC0131', 'Audley Travel'),
    'ATUSA': ('ACC0132', 'Audley Travel US'),
}

def _map_business_unit(bu_str, email_domain=''):
    """Map Business Unit / email domain to source market code."""
    if not bu_str:
        bu_str = ''
    bu = str(bu_str).upper()
    if 'UK' in bu or 'ATUK' in bu or 'GBR' in bu: return 'GBR'
    if 'USA' in bu or 'US' in bu:                    return 'USA'
    if 'AUS' in bu or 'AU' in bu:                    return 'AUS'
    if 'DEU' in bu or 'DE' in bu or 'GER' in bu:    return 'DEU'
    if 'FRA' in bu or 'FR' in bu:                    return 'FRA'
    if 'ITA' in bu or 'IT' in bu:                    return 'ITA'
    # Fall back to domain detection
    if '.co.uk' in email_domain or '.uk' in email_domain: return 'GBR'
    if '.de' in email_domain:  return 'DEU'
    if '.fr' in email_domain:  return 'FRA'
    if '.it' in email_domain:  return 'ITA'
    if '.au' in email_domain:  return 'AUS'
    return 'GBR'  # default


def parse_tourlane_email(email_text, agency_path=None):
    """
    Parse a Tourlane structured booking email into a TravartIQ intent dict.
    Format: HTML tables in email body — Accommodations, Activities, Other Transports.
    Maps to ACC1816 Sensation Travel GMBH (tourlane.com domain).
    """
    import re
    from datetime import datetime, date, timedelta

    intent = {
        'market':                   'DEU',
        'source_market':            'DEU',
        'pax':                      1,
        'num_pax':                  1,
        'rooms':                    {'doubles': 0, 'singles': 0, 'twins': 0},
        'tier':                     None,
        'heritage':                 False,
        'duration':                 None,
        'nightly_split':            [],
        'cities_detected':          [],
        'agent_specified_hotels':   [],
        'csv_activities':           [],
        'csv_transfers':            [],
        'csv_rail':                 [],
        'dual_option_hotels':       {},
        'travel_start':             '',
        'travel_end':               '',
        'budget_raw':               '',
        'language_guide':           False,
        'language':                 'DEU',
        'accompanying':             False,
        'agency_account':           'ACC1816',
        'agency_name':              'Sensation Travel GMBH',
        'guest_name':               '',
        'booking_ref':              '',
        '_input_mode':              'TOURLANE',
        '_accommodation_rows':      [],
        'csv_path':                 'tourlane',
    }

    # ── Booking ref from subject or body ─────────────────────────────────────
    ref_m = re.search(r'T-\d{6}-\d{5}', email_text)
    if ref_m:
        intent['booking_ref'] = ref_m.group(0)

    # ── Market from subject ───────────────────────────────────────────────────
    mkt_m = re.search(r'//\s*(\w+)\s+market', email_text, re.IGNORECASE)
    if mkt_m:
        mkt_raw = mkt_m.group(1).upper()
        mkt_map = {'DE': 'DEU', 'GERMANY': 'DEU', 'FR': 'FRA', 'FRANCE': 'FRA',
                   'UK': 'GBR', 'GB': 'GBR', 'US': 'USA', 'IT': 'ITA',
                   'ES': 'ESP', 'CH': 'CHE', 'AT': 'AUT', 'AU': 'AUS'}
        intent['market'] = mkt_map.get(mkt_raw, mkt_raw)
        intent['source_market'] = intent['market']

    # ── Pax from subject ──────────────────────────────────────────────────────
    pax_m = re.search(r'(\d+)\s*[Pp]ax', email_text)
    if pax_m:
        intent['pax']     = int(pax_m.group(1))
        intent['num_pax'] = intent['pax']
        intent['rooms']   = {'doubles': 0, 'singles': intent['pax'], 'twins': 0}

    # ── Guest name from Passenger table ───────────────────────────────────────
    guest_m = re.search(
        r'(?:Full Name|Passenger)[\s\S]{0,30}?\n([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,3})',
        email_text)
    if guest_m:
        intent['guest_name'] = guest_m.group(1).strip()

    # ── Accommodation rows — line-by-line state machine ────────────────────
    # Format: each field on its own line in fixed order:
    # date / nights / area / hotel / rate / room_desc / basis / type / names...
    acc_section = re.search(
        r'Accommodations\s*\n(.*?)(?:\nActivities|\nOther Items|\nOther Transports|$)',
        email_text, re.DOTALL | re.IGNORECASE)

    accommodation_rows = []
    nightly_split = []
    agent_hotels = []
    dmc_cities = set()

    if acc_section:
        _lines = [l.strip() for l in acc_section.group(1).split('\n') if l.strip()]
        # Skip header lines (Check-in, Nights, Area, Hotel Name, Rate, Room, Basis, Type, Names)
        _HEADER_FIELDS = {'check-in','nights','area','hotel name','rate','room',
                          'basis','type','names','full name','birthday'}
        i = 0
        while i < len(_lines) and _lines[i].lower() in _HEADER_FIELDS:
            i += 1
        # Also skip until first date
        while i < len(_lines) and not re.match(r'\d{2}/\d{2}/\d{4}', _lines[i]):
            i += 1

        # Each row: date(0) nights(1) area(2) hotel(3) rate(4) room_desc(5) basis(6) type(7) names(8+)
        # Row ends when next date appears
        while i < len(_lines):
            if not re.match(r'\d{2}/\d{2}/\d{4}', _lines[i]):
                i += 1
                continue
            try:
                ci_str  = _lines[i]
                nights  = int(_lines[i+1]) if i+1 < len(_lines) else 0
                area    = _lines[i+2]      if i+2 < len(_lines) else ''
                hotel   = _lines[i+3]      if i+3 < len(_lines) else ''
                # type is at position 7 (index i+7)
                room_type = _lines[i+7] if i+7 < len(_lines) else 'Double'
                ci = datetime.strptime(ci_str, '%d/%m/%Y').date()
                co = ci + timedelta(days=nights)
                i += 8
                # Skip passenger names until next date or section end
                while i < len(_lines) and not re.match(r'\d{2}/\d{2}/\d{4}', _lines[i]):
                    i += 1

                # Normalise city
                # Normalise city using CITY_CODE_MAP (module-level) + CITY_NORM_CSV
                city = area
                for canon in CITY_CODE_MAP:
                    if area.lower() == canon.lower():
                        city = canon
                        break
                else:
                    for canon, aliases in CITY_NORM_CSV.items():
                        if area.lower() in [a.lower() for a in aliases]:
                            city = canon
                            break
                    else:
                        # Manual common mappings
                        _AREA_MAP = {
                            'ranthambore national park': 'Ranthambore',
                            'new delhi': 'New Delhi', 'delhi': 'New Delhi',
                            'empty area': area,
                        }
                        city = _AREA_MAP.get(area.lower(), area)

                # DMC pre-quoted detection
                if area.lower() in ('empty area', '') or (
                        'corbett' in area.lower() or 'corbett' in hotel.lower()):
                    if 'dmc' in email_text.lower() or area.lower() == 'empty area':
                        dmc_cities.add(city)
                        intent['csv_activities'].append({
                            'city': city, 'date': ci,
                            'name': f'DMC pre-quoted — {hotel}',
                            'description': 'Confirm manually'
                        })

                accommodation_rows.append({
                    'city': city, 'hotel_name': hotel,
                    'check_in': ci, 'check_out': co, 'nights': nights,
                    'room_type': room_type
                })
                nightly_split.append((city, nights))
                if hotel and hotel not in agent_hotels:
                    agent_hotels.append(hotel)

            except Exception:
                i += 1
                continue

    # ── Derive room type from accommodation Type column ─────────────────────
    _room_types = {'doubles': 0, 'singles': 0, 'twins': 0}
    _room_pattern = re.findall(
        r'(?:Double|Single|Twin)',
        acc_section.group(1) if acc_section else '', re.IGNORECASE)
    for _rt in _room_pattern:
        if _rt.lower() == 'double':   _room_types['doubles'] += 1
        elif _rt.lower() == 'single': _room_types['singles'] += 1
        elif _rt.lower() == 'twin':   _room_types['twins']   += 1
    # Deduplicate — count unique room nights not pax lines
    if _room_types['doubles']: _room_types['doubles'] = 1
    if _room_types['singles']: _room_types['singles'] = max(1, intent['pax'] - _room_types['doubles'] * 2)
    if any(_room_types.values()):
        intent['rooms'] = _room_types

    # ── Parse all passenger names ─────────────────────────────────────────────
    pax_section = re.search(
        r'Passenger Information\s*\n(.*?)(?:\nAccommodations|\nBooking Information|$)',
        email_text, re.DOTALL | re.IGNORECASE)
    _all_passengers = []
    _SKIP_PAX_HEADERS = {'full name','birthday','height (cm)','weight (kg)',
                         'shoe size (eu)','clothing size (s/m/l)','height','weight',
                         'shoe size','clothing size'}
    if pax_section:
        _pax_lines = [l.strip() for l in pax_section.group(1).split('\n') if l.strip()]
        _pi = 0
        while _pi < len(_pax_lines):
            _pl = _pax_lines[_pi]
            # Skip header fields
            if _pl.lower() in _SKIP_PAX_HEADERS:
                _pi += 1
                continue
            # A passenger name: text line followed by a date (birthday)
            if (not re.match(r'\d{2}/\d{2}/\d{4}', _pl) and
                    re.match(r'[A-Z][a-zA-Z]+', _pl) and
                    _pi + 1 < len(_pax_lines) and
                    re.match(r'\d{2}/\d{2}/\d{4}', _pax_lines[_pi+1])):
                _all_passengers.append(_pl)
                _pi += 2  # skip name + birthday
                continue
            _pi += 1
    if _all_passengers:
        intent['guest_name'] = _all_passengers[0]
        intent['all_passengers'] = _all_passengers
        if len(_all_passengers) > intent['pax']:
            intent['pax'] = len(_all_passengers)

    # Build csv_hotel_by_city from accommodation_rows for per-city hotel lookup
    # This ensures each city gets its own hotel, not a flat cross-city list
    _tl_hotel_by_city = {}
    for _ar in accommodation_rows:
        _c = _ar['city']
        _h = _ar['hotel_name']
        if _c in _tl_hotel_by_city:
            # Second visit to same city (e.g. Delhi arrival + Delhi return)
            existing = _tl_hotel_by_city[_c]
            if isinstance(existing, list):
                existing.append(_h)
            else:
                _tl_hotel_by_city[_c] = [existing, _h]
        else:
            _tl_hotel_by_city[_c] = _h

    intent['_accommodation_rows']    = accommodation_rows
    intent['nightly_split']          = nightly_split
    intent['agent_specified_hotels'] = agent_hotels
    intent['csv_hotel_by_city']      = _tl_hotel_by_city
    intent['cities_detected']        = [c for c, n in nightly_split]

    # Travel dates
    if accommodation_rows:
        all_ci = [r['check_in'] for r in accommodation_rows if r['check_in']]
        all_co = [r['check_out'] for r in accommodation_rows if r['check_out']]
        if all_ci:
            intent['travel_start'] = str(min(all_ci))
            intent['travel_end']   = str(max(all_co))
            intent['duration']     = (max(all_co) - min(all_ci)).days

    # ── Activities ────────────────────────────────────────────────────────────
    act_section = re.search(
        r'Activities\s*\n(.*?)(?:\nOther Items|\nOther Transports|\nBest regards|$)',
        email_text, re.DOTALL | re.IGNORECASE)

    _ACT_HEADERS = {'date','service','passengers','rate','provider','supplier code',
                    'dimensions','start time','pick-up location','booking notes for dmc'}
    if act_section:
        _act_lines = [l.strip() for l in act_section.group(1).split('\n') if l.strip()]
        _ai = 0
        # Skip headers
        while _ai < len(_act_lines) and _act_lines[_ai].lower() in _ACT_HEADERS:
            _ai += 1
        while _ai < len(_act_lines):
            _al = _act_lines[_ai]
            if not re.match(r'\d{2}/\d{2}/\d{4}', _al):
                _ai += 1
                continue
            try:
                act_date = datetime.strptime(_al, '%d/%m/%Y').date()
            except Exception:
                _ai += 1
                continue
            # Next non-date, non-header line = service name
            _ai += 1
            svc_name = ''
            while _ai < len(_act_lines):
                _candidate = _act_lines[_ai]
                if _candidate.lower() in _ACT_HEADERS or re.match(r'\d{2}/\d{2}/\d{4}', _candidate):
                    break
                if not svc_name and not re.match(r'\d{2}/\d{2}/\d{4}', _candidate):
                    svc_name = _candidate
                    _ai += 1
                    break
                _ai += 1
            # Collect remaining fields until next date
            _act_fields = []
            while _ai < len(_act_lines) and not re.match(r'\d{2}/\d{2}/\d{4}', _act_lines[_ai]):
                _act_fields.append(_act_lines[_ai])
                _ai += 1
            _act_block = ' '.join(_act_fields)

            # Find city from date
            act_city = ''
            for row in accommodation_rows:
                if row['check_in'] and row['check_out']:
                    if row['check_in'] <= act_date < row['check_out']:
                        act_city = row['city']
                        break

            # Language detection
            lang_m = re.search(r'Language:\s*(\w+)', _act_block)
            _act_lang = 'DEU'
            if lang_m:
                _raw_lang = lang_m.group(1).lower()
                _act_lang = {'deutsch':'DEU','german':'DEU','englisch':'ENG',
                             'english':'ENG','francais':'FRA','french':'FRA',
                             'italiano':'ITA','italian':'ITA','espanol':'ESP'}.get(_raw_lang, 'ENG')
                if _act_lang == 'DEU':
                    intent['language_guide'] = 'DEU'
                    intent['language'] = 'DEU'

            if svc_name:
                intent['csv_activities'].append({
                    'city': act_city, 'date': act_date,
                    'name': svc_name, 'description': svc_name,
                    'language': _act_lang
                })

    # ── Transfers + transit city detection ──────────────────────────────────
    _AIRPORT_MAP = {
        'DEL': 'New Delhi',    'BOM': 'Mumbai',         'MAA': 'Chennai',
        'COK': 'Kochi (Cochin)', 'BLR': 'Bengaluru',   'CCU': 'Kolkata',
        'JAI': 'Jaipur',       'VNS': 'Varanasi',       'IDR': 'Indore',
        'JDH': 'Jodhpur',      'UDR': 'Udaipur (RJ)',   'AMD': 'Ahmedabad',
    }

    def _loc_to_city(loc_str):
        iata_m = re.search(r'\[([A-Z]{3})\]', loc_str)
        if iata_m and iata_m.group(1) in _AIRPORT_MAP:
            return _AIRPORT_MAP[iata_m.group(1)]
        for canon in CITY_CODE_MAP:
            if canon.lower() in loc_str.lower():
                return canon
        return None

    xfer_section = re.search(
        r'Other Transports\s*\n(.*?)(?:\nBest regards|$)',
        email_text, re.DOTALL | re.IGNORECASE)

    # Transfer section is line-by-line: date/mode/type/name/from/to each on own line
    _transfer_rows = []
    _XFER_HEADERS = {'date','mode','type','name','start location','end location'}
    if xfer_section:
        _xlines = [l.strip() for l in xfer_section.group(1).split('\n') if l.strip()]
        _xi = 0
        # Skip headers
        while _xi < len(_xlines) and _xlines[_xi].lower() in _XFER_HEADERS:
            _xi += 1
        while _xi < len(_xlines):
            if not re.match(r'\d{2}/\d{2}/\d{4}', _xlines[_xi]):
                _xi += 1
                continue
            try:
                xfer_date = datetime.strptime(_xlines[_xi], '%d/%m/%Y').date()
                # Row: date(0) mode(1) type(2) name(3) from(4) to(5)
                _from = _xlines[_xi+4] if _xi+4 < len(_xlines) else ''
                _to   = _xlines[_xi+5] if _xi+5 < len(_xlines) else ''
                intent['csv_transfers'].append({
                    'description': f"{_xlines[_xi]}  {_from} → {_to}"
                })
                _transfer_rows.append({'date': xfer_date, 'from': _from, 'to': _to})
                _xi += 6
                # Skip any extra lines until next date
                while _xi < len(_xlines) and not re.match(r'\d{2}/\d{2}/\d{4}', _xlines[_xi]):
                    _xi += 1
            except Exception:
                _xi += 1

    # ── Inject 0-night transit cities from transport chain ───────────────────
    _acc_cities = set(c.lower() for c, n in nightly_split)
    if _transfer_rows:
        _transfer_rows.sort(key=lambda x: x['date'])
        # First transfer FROM = arrival city (often airport → first hotel city)
        first_from = _loc_to_city(_transfer_rows[0]['from'])
        _first_is_airport = bool(re.search(r'\[([A-Z]{3})\]', _transfer_rows[0]['from']))
        if first_from and _first_is_airport:
            if not nightly_split or nightly_split[0][0].lower() != first_from.lower():
                nightly_split.insert(0, (first_from, 0))
                _acc_cities.add(first_from.lower())
                print(f"  ✓ Arrival transit: {first_from} (0n) — added to route from transport")
        elif first_from and first_from.lower() not in _acc_cities:
            nightly_split.insert(0, (first_from, 0))
            _acc_cities.add(first_from.lower())
            print(f"  ✓ Arrival transit: {first_from} (0n) — added to route from transport")
        # All remaining transfers — check for any other missing cities
        for _xr in _transfer_rows[1:]:
            for _loc in [_xr['from'], _xr['to']]:
                _tc = _loc_to_city(_loc)
                if _tc and _tc.lower() not in _acc_cities:
                    nightly_split.append((_tc, 0))
                    _acc_cities.add(_tc.lower())
                    print(f"  ✓ Transit city: {_tc} (0n) — added to route from transport")

    intent['nightly_split']   = nightly_split
    intent['cities_detected'] = [c for c, n in nightly_split]

    # ── Agency resolution via domain lookup (same as email/CSV mode) ─────
    intent['agency_account_code'] = ''
    if agency_path and not _AGENCY_BY_DOMAIN:
        load_agency_portfolio(agency_path)
    try:
        _from_domain = 'tourlane.com'
        _from_m = re.search(r'From:[^\n]*@([\w.-]+)', email_text, re.IGNORECASE)
        if _from_m:
            _from_domain = _from_m.group(1).lower().strip('>')
        if _AGENCY_BY_DOMAIN:
            _dom_match = _AGENCY_BY_DOMAIN.get(_from_domain)
            if _dom_match:
                intent['agency_account_code'] = _dom_match.get('account_code', '')
                intent['agency_name']         = _dom_match.get('agency_name', intent['agency_name'])
                print(f"  \u2713 TCI account code: {intent['agency_account_code']}  ({intent['agency_name']})")
            else:
                print(f"  \u26a0 Domain {_from_domain!r} not in agency portfolio")
    except Exception as _ce:
        print(f"  \u26a0 Agency lookup error: {_ce}")


    # ── Ensure source_market and num_pax always set ─────────────────────────
    if not intent.get('source_market'):
        intent['source_market'] = intent.get('market', 'DEU')
    intent['num_pax'] = intent.get('pax', 1)

    # ── Print parsed intent ───────────────────────────────────────────────────
    print(f"\n======================================================================")
    print(f"  TOURLANE BOOKING PARSER")
    print(f"======================================================================")
    print(f"  ✓ Booking ref : {intent['booking_ref']}")
    print(f"  ✓ Guest       : {intent['guest_name']}")
    print(f"  ✓ Agency      : {intent['agency_name']}  [{intent['agency_account']}]")
    print(f"  ✓ Market      : {intent['market']}")
    print(f"  ✓ Pax         : {intent['pax']}  ({intent['rooms']})")
    if intent.get('all_passengers'):
        for _pn in intent['all_passengers']:
            print(f"  ✓ Passenger   : {_pn}")
    print(f"  ✓ Travel      : {intent['travel_start']} → {intent['travel_end']}  ({intent['duration']} nights)")
    print(f"\n  ⚡ Route from Tourlane booking:")
    for city, nights in nightly_split:
        dmc_tag = '  ⚠ DMC pre-quoted' if city in dmc_cities else ''
        print(f"     {city:30} {nights}n{dmc_tag}")
    print(f"\n  Hotels specified:")
    for h in agent_hotels:
        print(f"     → {h}")
    if intent['csv_activities']:
        print(f"\n  Activities ({len([a for a in intent.get('csv_activities',[]) if 'DMC' not in a.get('name','')])}):")
        for a in intent['csv_activities']:
            if 'DMC' not in a['name']:
                print(f"     {str(a.get('city','')):<20} {a['name']}")
    if dmc_cities:
        print(f"\n  ⚠ DMC pre-quoted cities (hotel/tours NOT pushed): {', '.join(dmc_cities)}")
    if intent['csv_transfers']:
        print(f"  Transfers: {len(intent['csv_transfers'])} (flagged — not pushed to TCI)")
    if intent.get('language_guide'):
        print(f"  ✓ German guide requested")

    return intent


def parse_csv_itinerary(services_csv, passenger_csv=None, email_raw='', agency_path=None):
    """
    Parse structured quote request CSVs into intent dict.
    Returns intent dict in same schema as parse_email() — zero downstream changes.

    Args:
        services_csv  : path to *_Services.csv
        passenger_csv : path to *_Passenger_list.csv (optional)
        email_raw     : raw email text (for agency CRM lookup)
        agency_path   : path to agency_crm.xlsx (optional)
    """
    import csv as _csv
    from datetime import datetime as _dt2, timedelta as _td2

    print("\n" + "="*70)
    print("  CSV ITINERARY PARSER")
    print("="*70)

    # Ensure agency portfolio loaded for domain lookup
    if agency_path and not _AGENCY_BY_DOMAIN:
        load_agency_portfolio(agency_path)

    # Ensure agency portfolio is loaded for CRM domain lookup
    if agency_path and not _AGENCY_BY_DOMAIN:
        load_agency_portfolio(agency_path)

    intent = {
        'source_market':         'GBR',
        'num_pax':               2,
        'rooms':                 {'doubles':1, 'singles':0, 'twins':0},
        'tier':                  None,
        'heritage_requested':    False,
        'heritage_keywords':     [],
        'language_guide':        None,
        'budget_raw':            None,
        'cities_detected':       [],
        'regions_detected':      [],
        'nightly_split':         [],
        'agent_specified_hotels':[],
        'activities_requested':  [],
        'transport_notes':       [],
        'named_trains':          [],
        'dates_raw':             [],
        'travel_start':          None,  # set from PNR; body-date fallback applied in recommend()
        'travel_end':            None,
        'duration_nights':       None,
        'csv_path':              True,   # flag for downstream steps
        'agency_account_code':   '',
        'agency_name':           '',
        'guest_name':            '',
        'itinerary_ref':         '',
        'business_unit':         '',
        'passenger_names':       [],
        'csv_activities':        [],   # [{city, date, name, description, pax, guide}]
        'csv_transfers':         [],   # [{city, date, description}]
        'csv_rail':              [],   # [{from, to, train_no, class_, time_}]
        'csv_flights':           [],   # [{flight, from_iata, to_iata, date, time_}]
        'dual_option_hotels':    {},   # {city: [hotel1, hotel2]}
        '_raw_text':             email_raw.lower(),
    }

    # ── PASSENGER CSV ────────────────────────────────────────────────────
    if passenger_csv:
        try:
            with open(passenger_csv, newline='', encoding='utf-8-sig') as f:
                raw = f.read()
            # Parse header metadata lines
            for line in raw.split('\n')[:6]:
                line = line.strip().strip(',')
                if 'Itinerary Owner' in line:
                    intent['agent_name'] = line.split(',')[-1].strip()
                elif 'Business Unit' in line:
                    intent['business_unit'] = line.split(',')[-1].strip()
                elif 'Itinerary Ref' in line:
                    intent['itinerary_ref'] = line.split(',')[-1].strip()

            # Parse passenger rows
            import io as _io
            # Find the actual header row (contains 'Salutation' or 'First Name')
            pax_lines = raw.split('\n')
            header_idx_pax = 0
            for i, line in enumerate(pax_lines):
                if 'Salutation' in line or 'First Name' in line:
                    header_idx_pax = i
                    break
            pax_csv_content = '\n'.join(pax_lines[header_idx_pax:])
            reader = _csv.DictReader(_io.StringIO(pax_csv_content))
            pax_count = 0
            for row in reader:
                sal  = str(row.get('Salutation','')).strip()
                fn   = str(row.get('First Name','')).strip()
                ln   = str(row.get('Last Name','')).strip()
                ptype = str(row.get('Passenger Type','')).strip().lower()
                if fn and ln and ptype in ('adult', ''):
                    full_name = f"{sal} {fn} {ln}".strip()
                    intent['passenger_names'].append(full_name)
                    if pax_count == 0:
                        intent['guest_name'] = f"{fn} {ln}".strip()
                    pax_count += 1
            intent['num_pax'] = pax_count or 2
            # Default room config
            if pax_count == 1:
                intent['rooms'] = {'doubles':0,'singles':1,'twins':0}
            elif pax_count == 2:
                intent['rooms'] = {'doubles':1,'singles':0,'twins':0}
            elif pax_count % 2 == 0:
                intent['rooms'] = {'doubles': pax_count//2, 'singles':0, 'twins':0}
            else:
                intent['rooms'] = {'doubles': pax_count//2, 'singles':1, 'twins':0}

            print(f"  ✓ Passengers: {pax_count} pax — {intent['guest_name']}")
            print(f"  ✓ Agency ref : {intent['itinerary_ref']} | BU: {intent['business_unit']}")
        except Exception as e:
            print(f"  ⚠ Passenger CSV error: {e}")

    # ── AGENCY CRM LOOKUP ────────────────────────────────────────────────
    # Run lookup if email provided — portfolio may already be loaded in memory
    if email_raw and (_AGENCY_BY_DOMAIN or agency_path):
        try:
            # Extract domain directly from From: line — most reliable signal
            import re as _re_ag
            from_m = _re_ag.search(r'From:.*?<([^>]+)>|From:\s*(\S+@\S+)', email_raw, _re_ag.IGNORECASE)
            sender_email = ''
            if from_m:
                sender_email = (from_m.group(1) or from_m.group(2) or '').strip()
            domain = sender_email.split('@')[-1].lower() if '@' in sender_email else ''

            # Look up domain in agency portfolio
            if domain and _AGENCY_BY_DOMAIN:
                match = _AGENCY_BY_DOMAIN.get(domain)
                if match:
                    intent['agency_account_code'] = match.get('account_code','')
                    intent['agency_name']         = match.get('agency_name','')
                    mkt = match.get('market','')
                    if mkt and mkt != 'UNKNOWN':
                        intent['source_market'] = mkt
                    print(f"  ✓ Agency match: {intent['agency_name']} [{intent['agency_account_code']}]  (domain: {domain})")
                else:
                    # Fallback to parse_email for full scoring
                    _tmp = parse_email(email_raw, agency_path=agency_path)
                    if _tmp.get('agency_account_code'):
                        intent['agency_account_code'] = _tmp['agency_account_code']
                        intent['agency_name']         = _tmp.get('agency_name','')
                        intent['source_market']        = _tmp.get('source_market', intent['source_market'])
                        print(f"  ✓ Agency match: {intent['agency_name']} [{intent['agency_account_code']}]")
                    else:
                        print(f"  ⚠ No agency match for domain: {domain}")
        except Exception as e:
            print(f"  ⚠ Agency lookup error: {e}")

    # Market from Business Unit if not already set
    if intent['source_market'] == 'GBR' and intent['business_unit']:
        intent['source_market'] = _map_business_unit(
            intent['business_unit'],
            email_raw.split('\n')[0] if email_raw else ''
        )

    # Account code from BU mapping — fires when CRM domain match didn't find an account
    if not intent.get('agency_account_code') and intent.get('business_unit'):
        _bu_upper = str(intent['business_unit']).upper()
        if _bu_upper in _BU_TO_ACCOUNT:
            _acct, _name = _BU_TO_ACCOUNT[_bu_upper]
            intent['agency_account_code'] = _acct
            intent['agency_name']         = _name
            print(f"  ✓ Agency from BU map: {_bu_upper} → {_acct} ({_name})")

    # ── SERVICES CSV ─────────────────────────────────────────────────────
    accommodation_rows = []   # [(city, check_in, check_out, hotel_name, meal_plan, nights)]
    seen_accom = {}           # city → list of hotels (for dual-option detection)
    import re as _re3

    try:
        with open(services_csv, encoding='utf-8-sig') as f:
            raw_svc = f.read()

        # Skip sep= line and metadata header lines — find the actual CSV header
        lines = raw_svc.split('\n')
        header_idx = 0
        for i, line in enumerate(lines):
            if 'Date,Place,Service' in line or 'Date' in line and 'Service' in line:
                header_idx = i
                break
        csv_content = '\n'.join(lines[header_idx:])

        import io as _io2
        reader = _csv.DictReader(_io2.StringIO(csv_content))

        for row in reader:
            # Strip \r from values — Windows CRLF files leave \r in field values
            def _cell(key): return str(row.get(key,'')).strip().strip('\r')
            date_str    = _cell('Date')
            place       = _cell('Place')
            svc_type    = _cell('Service')
            details     = _cell('Details of Service')
            meal_plan   = _cell('Meal Plan / Add on')
            dropoff     = _cell('Drop off / Time')
            supplier_note = _cell('Supplier Note')
            pickup      = _cell('Pick Up / Time')

            svc_lower = svc_type.lower()
            date_parsed = _parse_csv_date(date_str)

            # ── Accommodation ────────────────────────────────────────────
            if svc_lower == 'accommodation':
                city = None
                for db_city, signals in CITY_NORM_CSV.items():
                    if any(s.lower() in place.lower() for s in signals):
                        city = db_city
                        break
                if not city:
                    # Try CITY_NORM from main engine
                    for norm_city, aliases in _CITY_NORM_MAP.items():
                        if any(a in place.lower() for a in aliases):
                            city = norm_city
                            break
                if not city:
                    city = place  # use as-is if not found

                hotel_name, _hotel_is_primary = _extract_hotel_name(details)
                checkout   = _parse_csv_date(dropoff)
                if not checkout:
                    # Try parsing "X nights" from details
                    m = _re3.search(r'(\d+)\s*nights?', details, _re3.IGNORECASE)
                    if m and date_parsed:
                        from datetime import timedelta as _td3
                        checkout = date_parsed + _td3(days=int(m.group(1)))

                nights = 0
                if date_parsed and checkout:
                    nights = (checkout - date_parsed).days
                # Fallback: extract nights from "X nights" in details string
                if nights == 0:
                    _nm = _re3.search(r'(\d+)\s*nights?', details, _re3.IGNORECASE)
                    if _nm:
                        nights = int(_nm.group(1))
                        if date_parsed and not checkout:
                            from datetime import timedelta as _td4
                            checkout = date_parsed + _td4(days=nights)
                # Derive check_in from checkout - nights when date_parsed is None
                # (CSV Date column uses booking date, not travel date)
                if not date_parsed and checkout and nights > 0:
                    from datetime import timedelta as _td5
                    date_parsed = checkout - _td5(days=nights)

                # Dual-option detection: same city, same dates, different hotel
                city_key = f"{city}_{date_parsed}"
                if city_key not in seen_accom:
                    seen_accom[city_key] = []
                seen_accom[city_key].append(hotel_name)

                accommodation_rows.append({
                    'city':       city,
                    'check_in':   date_parsed,
                    'check_out':  checkout,
                    'hotel_name': hotel_name,
                    'is_primary': _hotel_is_primary,
                    'meal_plan':  _map_meal_plan(meal_plan),
                    'nights':     nights,
                })

            # ── Activity ─────────────────────────────────────────────────
            elif svc_lower == 'activity':
                city = place
                activity_name = _extract_activity_name(details)
                has_guide = 'guide' in details.lower()
                pax_m = _re3.search(r'(\d+)\s*pax', details, _re3.IGNORECASE)
                activity_pax = int(pax_m.group(1)) if pax_m else intent['num_pax']
                intent['csv_activities'].append({
                    'city':        city,
                    'date':        date_parsed,
                    'name':        activity_name,
                    'description': details,
                    'pax':         activity_pax,
                    'guide':       has_guide,
                })

            # ── Transfer ─────────────────────────────────────────────────
            elif svc_lower == 'transfer':
                intent['csv_transfers'].append({
                    'city':        place,
                    'date':        date_parsed,
                    'description': details,
                })

            # ── Rail ─────────────────────────────────────────────────────
            elif svc_lower == 'rail':
                start_stn = str(row.get('Start Station','')).strip()
                end_stn   = str(row.get('End Station','')).strip()
                train_no  = supplier_note  # "train 52452"
                intent['csv_rail'].append({
                    'date':      date_parsed,
                    'from_stn':  start_stn,
                    'to_stn':    end_stn,
                    'train_no':  train_no,
                    'time_':     pickup,
                    'details':   details,
                })
                intent['transport_notes'].append('train mentioned')
                intent['named_trains'].append(details[:60])

            # ── FlightPlaceholder ─────────────────────────────────────────
            elif svc_lower == 'flightplaceholder':
                # Extract IATA codes from flight string
                iata_m = _re3.findall(r'\b([A-Z]{3})\b', details)
                # Filter to likely airport codes (not airline codes like 6E, AI)
                iata_airports = [c for c in iata_m if len(c)==3 and c.isalpha()]
                if len(iata_airports) >= 2:
                    intent['csv_flights'].append({
                        'date':      date_parsed,
                        'from_iata': iata_airports[0],
                        'to_iata':   iata_airports[1],
                        'details':   details,
                    })
                intent['transport_notes'].append('domestic flights mentioned')

    except Exception as e:
        print(f"  ⚠ Services CSV error: {e}")
        import traceback; traceback.print_exc()

    # ── BUILD NIGHTLY SPLIT & CITY SEQUENCE ──────────────────────────────
    # Deduplicate accommodation rows — keep first occurrence per city+checkin
    # but flag dual options
    seen_legs = {}
    nightly_split = []
    agent_hotels  = []
    all_cities    = []
    _csv_hotel_map = {}  # city → exact hotel name from CSV

    # Pass 1: identify (P) primary hotels per leg
    _primary_by_leg = {}
    for row in accommodation_rows:
        leg_key = f"{row['city']}_{row['check_in']}"
        if row.get('is_primary'):
            _primary_by_leg[leg_key] = row['hotel_name']

    for row in accommodation_rows:
        city  = row['city']
        ci    = row['check_in']
        hotel = row['hotel_name']
        is_primary = row.get('is_primary', False)
        leg_key = f"{city}_{ci}"

        _leg_has_explicit_primary = leg_key in _primary_by_leg
        _this_is_primary = (not _leg_has_explicit_primary) or is_primary

        # Track dual options
        if leg_key in seen_legs:
            if _this_is_primary and _leg_has_explicit_primary:
                # (P) row -- swap into primary position
                _prev_hotel = seen_legs[leg_key]['hotel']
                if city not in intent['dual_option_hotels']:
                    intent['dual_option_hotels'][city] = [_prev_hotel]
                elif _prev_hotel not in intent['dual_option_hotels'][city]:
                    intent['dual_option_hotels'][city].insert(0, _prev_hotel)
                seen_legs[leg_key] = {'hotel': hotel, 'row': row}
                # Also update csv_hotel_map so Step 10 uses the (P) primary
                _csv_hotel_map[city] = hotel
                # Promote in agent_hotels list
                if hotel in agent_hotels:
                    agent_hotels.remove(hotel)
                if hotel not in agent_hotels:
                    agent_hotels.insert(0, hotel)
                for i, (nc, nn) in enumerate(nightly_split):
                    if nc == city:
                        nightly_split[i] = (city, row['nights']); break
                print(f"  ✓ PRIMARY (P) {city}: {hotel}  (alt: {_prev_hotel})")
            else:
                # Alternate option for same city+date
                if city not in intent['dual_option_hotels']:
                    intent['dual_option_hotels'][city] = [seen_legs[leg_key]['hotel']]
                if hotel not in intent['dual_option_hotels'][city]:
                    intent['dual_option_hotels'][city].append(hotel)
                print(f"  ⚠ DUAL OPTION {city}: {seen_legs[leg_key]['hotel']}  vs  {hotel}")
            if hotel and hotel not in agent_hotels:
                agent_hotels.append(hotel)
            continue

        seen_legs[leg_key] = {'hotel': hotel, 'row': row}
        nightly_split.append((city, row['nights']))
        if hotel and hotel not in agent_hotels:
            agent_hotels.append(hotel)
        # Allow consecutive same city (e.g. Alleppey 1n houseboat + 3n hotel)
        if not all_cities or all_cities[-1] != city:
            all_cities.append(city)
        elif all_cities[-1] == city:
            # Consecutive same city — add again to create separate route
            all_cities.append(city)
        # Store city→hotel mapping — use list for cities with multiple stays
        if city not in _csv_hotel_map:
            _csv_hotel_map[city] = hotel
        else:
            # Multiple stays in same city — store as list
            existing = _csv_hotel_map[city]
            if isinstance(existing, list):
                existing.append(hotel)
            else:
                _csv_hotel_map[city] = [existing, hotel]

    # ── Add activity-only cities as 0-night transit stops ────────────────────
    _nightly_city_set = set(c for c, n in nightly_split)
    for _act in intent.get('csv_activities', []):
        _act_city_raw = _act.get('city', '').strip()
        if not _act_city_raw:
            continue
        # Check direct match in CITY_CODE_MAP keys
        _act_city = None
        for _canon in CITY_CODE_MAP:
            if _act_city_raw.lower() == _canon.lower():
                _act_city = _canon
                break
        if not _act_city:
            _act_city = _act_city_raw
        if _act_city not in _nightly_city_set and _act_city in CITY_CODE_MAP:
            nightly_split.append((_act_city, 0))
            _nightly_city_set.add(_act_city)

    intent['nightly_split']         = nightly_split
    intent['agent_specified_hotels'] = agent_hotels
    intent['cities_detected']        = all_cities
    intent['csv_hotel_by_city']      = _csv_hotel_map  # city → exact hotel name
    intent['_input_mode']            = 'CSV'
    intent['_accommodation_rows']    = accommodation_rows  # for date lookup in Step 11

    # Travel dates
    if accommodation_rows:
        check_ins  = [r['check_in']  for r in accommodation_rows if r['check_in']]
        check_outs = [r['check_out'] for r in accommodation_rows if r['check_out']]
        if check_ins:
            intent['travel_start'] = str(min(check_ins))
        if check_outs:
            intent['travel_end']   = str(max(check_outs))
        if intent['travel_start'] and intent['travel_end']:
            from datetime import datetime as _dt4
            ts = _dt4.fromisoformat(intent['travel_start']).date()
            te = _dt4.fromisoformat(intent['travel_end']).date()
            intent['duration_nights'] = ((te - ts).days, (te - ts).days)  # tuple (min, max) as int

    # Regions from cities
    for city in all_cities:
        region = CITY_TO_REGION.get(city)
        if region and region not in intent['regions_detected']:
            intent['regions_detected'].append(region)

    # ── PRINT PARSED INTENT ───────────────────────────────────────────────
    print(f"\n── PARSED INTENT (CSV) ─────────────────────────────────────────")
    print(f"  Market       : {intent['source_market']}  [{intent['agency_name'] or 'keyword scoring'}]  [{intent['agency_account_code'] or 'no CRM match'}]")
    print(f"  Pax          : {intent['num_pax']}")
    print(f"  Rooms        : {intent['rooms']}")
    print(f"  Guest        : {intent['guest_name']}")
    print(f"  Ref          : {intent['itinerary_ref']}")
    if intent['travel_start']:
        from datetime import datetime as _dt5
        ts = _dt5.fromisoformat(intent['travel_start']).date()
        te = _dt5.fromisoformat(intent['travel_end']).date()
        nights = (te - ts).days
        print(f"  Travel       : {intent['travel_start']} → {intent['travel_end']}  ({nights} nights)")

    print(f"\n  ⚡ Route from CSV ({sum(n for _,n in nightly_split)} overnight nights):")
    for city, nights in nightly_split:
        print(f"     {city:35} {nights}n")

    print(f"\n  Hotels specified:")
    for h in agent_hotels:
        print(f"     → {h}")

    if intent['dual_option_hotels']:
        print(f"\n  ⚠ DUAL OPTIONS (quote both):")
        for city, opts in intent['dual_option_hotels'].items():
            print(f"     {city}: {' vs '.join(opts)}")

    if intent['csv_activities']:
        print(f"\n  Activities ({len(intent['csv_activities'])}):")
        for a in intent['csv_activities']:
            print(f"     {a['city']:20} {a['name']}")

    if intent['csv_rail']:
        print(f"\n  Rail segments:")
        for r in intent['csv_rail']:
            print(f"     {r['details'][:70]}")

    if intent['csv_flights']:
        print(f"\n  Domestic flights:")
        for f in intent['csv_flights']:
            print(f"     {f['from_iata']} → {f['to_iata']}  {f['details'][:60]}")

    if intent['csv_transfers']:
        print(f"  Transfers: {len(intent['csv_transfers'])} (flagged — not pushed to TCI)")

    print("\n" + "="*70)
    return intent


# City normalisation for CSV place names (supplements main CITY_NORM)
CITY_NORM_CSV = {
    'New Delhi':                        ['delhi', 'new delhi'],
    'Amritsar':                         ['amritsar'],
    'Shimla':                           ['shimla'],
    'Chandigarh':                       ['chandigarh'],
    'Jaipur':                           ['jaipur'],
    'Agra':                             ['agra'],
    'Varanasi':                         ['varanasi'],
    'Goa':                              ['goa', 'south beaches, goa', 'south beaches'],
    'Mumbai':                           ['mumbai', 'bombay'],
    'Dharamshala':                      ['dharamshala', 'dharamsala'],
    'Haridwar':                         ['haridwar'],
    'Rishikesh':                        ['rishikesh'],
    'Udaipur (RJ)':                     ['udaipur'],
    'Jodhpur':                          ['jodhpur'],
    'Corbett':                          ['corbett'],
    'Darjeeling':                       ['darjeeling'],
    'Aurangabad (MH)':                  ['aurangabad'],
    'Kochi (Cochin)':                   ['kochi', 'cochin'],
    'Alappuzha (Alleppey)':             ['alleppey', 'alappuzha', 'alleppey (alappuzha)', 'backwaters', 'houseboat'],
    'Kumarakom':                        ['kumarakom'],
    'Kolkata':                          ['kolkata', 'calcutta', 'kolkata (calcutta)'],
    'Bengaluru':                        ['bengaluru', 'bangalore'],
    # VDM / French-operator additions
    'Chennai':                          ['madras', 'chennai'],
    'Mamallapuram (Mahabalipuram)':     ['mahabalipuram', 'mamallapuram'],
    'Puducherry (Pondicherry)':         ['pondicherry', 'puducherry'],
    'Thanjavur (Tanjore)':              ['tanjore', 'thanjavur'],
    'Kumbakonam':                       ['kumbakonam'],
    'Chettinad':                        ['chettinad', 'karaikudi'],
    'Madurai':                          ['madurai'],
    'Munnar':                           ['munnar'],
    'Mysuru (Mysore)':                  ['mysore', 'mysuru'],
    'Thekkady (Periyar/Kumily)':        ['thekkady', 'periyar', 'kumily'],
    'Leh':                              ['leh', 'ladakh'],
    'Srinagar':                         ['srinagar'],
    'Ranthambore':                      ['ranthambore', 'ranthambhore',
                                         'ranthambhore national park',
                                         'ranthambore national park'],
    'Mararikulam':                      ['marari', 'mararikulam', 'marari beach',
                                         'carnoustie marari'],
    'Khajuraho':                        ['khajuraho'],
    'Mandawa':                          ['mandawa'],
    'Bikaner':                          ['bikaner'],
    'Pushkar':                          ['pushkar'],
    'Bandhavgarh':                      ['bandhavgarh'],
    'Kanha':                            ['kanha'],
}

# Alias map for CSV CITY_NORM lookup
_CITY_NORM_MAP = {
    city: [s.lower() for s in signals]
    for city, signals in CITY_NORM_CSV.items()
}


# ══════════════════════════════════════════════════════════════════════════
# PDF ITINERARY PARSER
# Parses VDM-style (Voyageurs du Monde) itinerary PDFs into intent dict.
# Returns intent dict in same schema as parse_csv_itinerary() — zero
# downstream changes required.  Requires: pip install pdfplumber
# ══════════════════════════════════════════════════════════════════════════

# ── Agency name lookup (for PDF parser — no email domain available) ────────────
def _lookup_agency_by_name(name_str):
    """Fuzzy CRM lookup by agency name — for PDF flows where no email domain exists.
    Tries: exact normalised match → word-subset match → first-word match.
    Returns CRM record dict or None.
    """
    import re as _re_ag
    def _norm(s):
        return _re_ag.sub(r'[\s\-_&,\.\']+', ' ', s.lower()).strip()

    if not name_str or not _AGENCY_BY_NAME:
        return None

    q = _norm(name_str)

    # 1. Exact normalised match
    if q in _AGENCY_BY_NAME:
        return _AGENCY_BY_NAME[q]

    # 2. Query words are all contained in a candidate name (handles word order)
    q_words = set(q.split())
    for key, rec in _AGENCY_BY_NAME.items():
        key_words = set(key.split())
        if q_words and q_words.issubset(key_words):
            return rec

    # 3. Any candidate whose normalised name contains the full query as substring
    for key, rec in _AGENCY_BY_NAME.items():
        if q in key or key in q:
            return rec

    return None


# Known VDM/operator supplier names — stripped from end of service lines
_VDM_SUPPLIERS = {
    'distant frontiers', 'world travelers', 'very local trip',
    'quick visa', 'concierge india', 'sas phenix trade international',
    'books and objects of the world', 'books and objects',
}

# IATA codes that are international (skip domestic-only logic)
_VDM_INTL_IATA = {
    'LHR', 'LGW', 'MAN', 'CDG', 'ORY', 'FCO', 'AMS', 'FRA', 'ZRH',
    'JFK', 'EWR', 'LAX', 'SYD', 'MEL', 'NRT', 'HND', 'DXB', 'SIN',
    'BKK', 'KUL', 'HKG', 'DOH', 'AUH',
}

# Indian airport IATA → canonical engine city name (for flight parsing)
_VDM_IATA_CITY = {
    'DEL': 'New Delhi',   'BOM': 'Mumbai',          'MAA': 'Chennai',
    'BLR': 'Bengaluru',   'COK': 'Kochi (Cochin)',   'CCU': 'Kolkata',
    'HYD': 'Hyderabad',   'JAI': 'Jaipur',           'AMD': 'Ahmedabad',
    'GOI': 'Goa',         'VNS': 'Varanasi',         'IXC': 'Chandigarh',
    'ATQ': 'Amritsar',    'UDR': 'Udaipur (RJ)',      'JDH': 'Jodhpur',
    'JSA': 'Jaisalmer',   'JLR': 'Jabalpur',         'TRV': 'Thiruvananthapuram (Trivandrum)',
    'IXZ': 'Andaman',     'PNQ': 'Pune',
}


def _vdm_norm_city(place_str):
    """Normalise a VDM/PDF place name to engine canonical city name.
    Checks CITY_NORM_CSV first, then falls back to main CITY_NORM aliases.
    """
    p = place_str.strip().lower()
    if not p:
        return ''
    # CITY_NORM_CSV (includes VDM additions like Madras, Mahabalipuram, etc.)
    for canon, aliases in CITY_NORM_CSV.items():
        if any(a in p for a in aliases):
            return canon
    # Main engine CITY_NORM_MAP (broader alias set)
    for canon, aliases in _CITY_NORM_MAP.items():
        if any(a in p for a in aliases):
            return canon
    # Return title-cased original as last resort
    return place_str.strip().title()


def _vdm_strip_supplier(name_str):
    """Remove trailing supplier name from a benefit name string."""
    s = name_str.strip()
    for sup in _VDM_SUPPLIERS:
        # Case-insensitive suffix match
        if s.lower().endswith(sup):
            s = s[: -len(sup)].strip()
            break
    return s.strip()


def _vdm_parse_service_line(line):
    """
    Parse one VDM service line into (city_raw, duration, svc_type, name).

    Handles three patterns:
      Transport (no duration):  "London Transport - LHR - BOM - Air India - AI130"
      With city + duration:     "Agra 1 Accommodation - INARA HOUSE Distant Frontiers"
      No city, duration only:   "0 Visas - E-Visa + E-Arrival card Quick Visa"
    """
    SERVICE_KEYWORDS = [
        'Accommodation', 'Excursion', 'Suggestions', 'Transport',
        'Miscellaneous', 'Transfer', 'Car', 'Traveler Services',
        'Home / Support', 'Visas',
    ]
    import re as _r
    for kw in SERVICE_KEYWORDS:
        idx = line.find(kw)
        if idx == -1:
            # case-insensitive fallback
            idx = line.lower().find(kw.lower())
        if idx == -1:
            continue
        prefix = line[:idx].strip().rstrip()
        suffix = line[idx + len(kw):].strip().lstrip(' -').strip()

        # Parse prefix: optional city + optional duration
        # Patterns:  "Agra 1"  |  "Bombay"  |  "0"  |  ""
        city_raw = ''
        duration = 0
        if prefix:
            # Try "CITY DIGIT" first
            m = _r.match(r'^(.+?)\s+(\d+)$', prefix)
            if m:
                city_raw = m.group(1).strip()
                duration = int(m.group(2))
            else:
                # Pure digit → no city
                m2 = _r.match(r'^(\d+)$', prefix)
                if m2:
                    duration = int(m2.group(1))
                else:
                    city_raw = prefix  # just a city name, Transport style

        # Strip leading date/day prefix from city_raw that sometimes bleeds in
        # e.g. "Tue 17/03/2026 Ranakpur" → "Ranakpur"
        import re as _r2
        city_raw = _r2.sub(
            r'^(?:Mon|Tue|Wed|Thu|Fri|Sat|Sun)?\s*\d{1,2}/\d{2}/\d{4}\s*', '',
            city_raw, flags=_r2.IGNORECASE
        ).strip()

        name = _vdm_strip_supplier(suffix)
        return city_raw, duration, kw, name

    return '', 0, '', ''


def parse_pdf_itinerary(pdf_path, email_raw='', agency_path=None):
    """
    Parse a VDM-style itinerary PDF into intent dict.
    Returns same schema as parse_csv_itinerary() — zero downstream changes.

    VDM PDF line format (pdfplumber text extraction):
      Line 1: "Fri Agra 1 Accommodation - INARA Distant Frontiers"
              [Day] [City?] [Duration?] [ServiceType] - [PartialName] [Supplier?]
      Line 2: "27/03/2026 HOUSE"
              [Date] [NameContinuation?]
      Line 3: "In French"  (further name continuation, e.g. language tag)
      ...description block...

    Args:
        pdf_path   : path to PDF file (VDM / Voyageurs du Monde format)
        email_raw  : raw agent email (for agency CRM domain lookup)
        agency_path: path to agency_crm.xlsx (optional)

    Requires: pip install pdfplumber
    """
    try:
        import pdfplumber as _plumber
    except ImportError:
        print("  ⚠ pdfplumber not installed — run: pip install pdfplumber")
        return None

    from datetime import datetime as _dt, timedelta as _td
    import re as _re

    print("\n" + "=" * 70)
    print("  PDF ITINERARY PARSER  (VDM format)")
    print("=" * 70)

    if agency_path and not _AGENCY_BY_DOMAIN:
        load_agency_portfolio(agency_path)

    # ── Extract full text ─────────────────────────────────────────────────
    full_text = ''
    try:
        with _plumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                full_text += (page.extract_text() or '') + '\n'
    except Exception as e:
        print(f"  ⚠ PDF extraction error: {e}")
        return None

    # ── Intent template ───────────────────────────────────────────────────
    intent = {
        'source_market':          'FRA',
        'num_pax':                2,
        'rooms':                  {'doubles': 1, 'singles': 0, 'twins': 0},
        'tier':                   None,
        'heritage_requested':     False,
        'heritage_keywords':      [],
        'language_guide':         None,
        'budget_raw':             None,
        'cities_detected':        [],
        'regions_detected':       [],
        'nightly_split':          [],
        'agent_specified_hotels': [],
        'activities_requested':   [],
        'transport_notes':        [],
        'named_trains':           [],
        'dates_raw':              [],
        'travel_start':           None,
        'travel_end':             None,
        'duration_nights':        None,
        'csv_path':               True,
        'agency_account_code':    '',
        'agency_name':            '',
        'guest_name':             '',
        'itinerary_ref':          '',
        'business_unit':          '',
        'passenger_names':        [],
        'csv_activities':         [],
        'csv_transfers':          [],
        'csv_rail':               [],
        'csv_flights':            [],
        'dual_option_hotels':     {},
        '_raw_text':              email_raw.lower(),
    }

    def _dmY_to_iso(s):
        try:
            return _dt.strptime(s.strip(), '%d/%m/%Y').date().isoformat()
        except Exception:
            return None

    # ── Agency / market from email domain ────────────────────────────────
    if email_raw and _AGENCY_BY_DOMAIN:
        from_m = _re.search(r'From:.*?<([^>]+)>|From:\s*(\S+@\S+)', email_raw, _re.IGNORECASE)
        if from_m:
            sender_email = (from_m.group(1) or from_m.group(2) or '').strip()
            domain = sender_email.split('@')[-1].lower() if '@' in sender_email else ''
            if domain:
                match = _AGENCY_BY_DOMAIN.get(domain)
                if match:
                    intent['agency_account_code'] = match.get('account_code', '')
                    intent['agency_name']          = match.get('agency_name', '')
                    mkt = match.get('market', '')
                    if mkt and mkt != 'UNKNOWN':
                        intent['source_market'] = mkt

    if 'voyageursdumonde' in full_text.lower() or 'voyageurs du monde' in full_text.lower():
        if not intent['agency_name']:
            intent['agency_name'] = 'Voyageurs du Monde'

    # ── Name-based CRM lookup (no email domain in PDF flow) ───────────────
    # Triggered when agency name is known (from PDF content) but account code
    # wasn't resolved via domain lookup above.
    if intent['agency_name'] and not intent['agency_account_code'] and _AGENCY_BY_NAME:
        crm = _lookup_agency_by_name(intent['agency_name'])
        if crm:
            intent['agency_account_code'] = crm.get('account_code', '')
            mkt = crm.get('market', '')
            if mkt and mkt != 'UNKNOWN':
                intent['source_market'] = mkt
            print(f"  ✓ Agency CRM match (by name): {intent['agency_name']} [{intent['agency_account_code']}]  market={intent['source_market']}")

    # ── Parse header ──────────────────────────────────────────────────────
    ref_m = _re.search(
        r'(?:Billed\s+file|Dossier\s+factur[eé]|REFERENCE\s*:?|File\s+N[o°]?\.?)\s*(\d{5,7})',
        full_text, _re.IGNORECASE
    )
    if ref_m:
        intent['itinerary_ref'] = ref_m.group(1)

    title_m = _re.search(
        r'(?:India|Inde)\s*-\s*([A-ZÀ-Ýa-zà-ý\s\-]+?)\s*-\s*'
        r'(?:Mon|Tue|Wed|Thu|Fri|Sat|Sun|Lun|Mar|Mer|Jeu|Ven|Sam|Dim)',
        full_text, _re.IGNORECASE
    )
    if title_m:
        intent['guest_name'] = ' '.join(p.capitalize() for p in title_m.group(1).strip().split())

    pax_m = _re.search(r'Participants\s*:\s*(\d+)', full_text)
    if pax_m:
        n = int(pax_m.group(1))
        intent['num_pax'] = n
        if n == 1:
            intent['rooms'] = {'doubles': 0, 'singles': 1, 'twins': 0}
        elif n == 2:
            intent['rooms'] = {'doubles': 1, 'singles': 0, 'twins': 0}
        elif n % 2 == 0:
            intent['rooms'] = {'doubles': n // 2, 'singles': 0, 'twins': 0}
        else:
            intent['rooms'] = {'doubles': n // 2, 'singles': 1, 'twins': 0}

    dep_m  = _re.search(r'Departure\s*:\s*(\d{2}/\d{2}/\d{4})', full_text, _re.IGNORECASE)
    back_m = _re.search(r'Back\s*:\s*(\d{2}/\d{2}/\d{4})', full_text, _re.IGNORECASE)
    if dep_m:
        intent['travel_start'] = _dmY_to_iso(dep_m.group(1))
    if back_m:
        intent['travel_end'] = _dmY_to_iso(back_m.group(1))
    if intent['travel_start'] and intent['travel_end']:
        ts = _dt.fromisoformat(intent['travel_start']).date()
        te = _dt.fromisoformat(intent['travel_end']).date()
        intent['duration_nights'] = ((te - ts).days, (te - ts).days)

    # ── Line-by-line service parsing ──────────────────────────────────────
    # pdfplumber puts DAY+CITY+SERVICE on ONE line, DATE on the NEXT line,
    # followed by optional name-continuation lines, then a description block.
    #
    # Pattern:
    #   "Fri Agra 1 Accommodation - INARA Distant Frontiers"   ← svc_line
    #   "27/03/2026 HOUSE"                                     ← date + name_cont
    #   "In French"                                            ← more name_cont (optional)
    #   "Reservation Note: ..."                                ← description body

    DAY_ABBREVS = {'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun',
                   'Lun', 'Mar', 'Mer', 'Jeu', 'Ven', 'Sam', 'Dim'}
    SVC_KEYWORDS = ['Accommodation', 'Excursion', 'Suggestions', 'Transport',
                    'Miscellaneous', 'Transfer', 'Car', 'Traveler Services',
                    'Home / Support', 'Visas']

    # Lines that signal end of name-continuation (start of description block)
    BODY_STARTERS = (
        'Reservation Note', 'Service Description', 'Service Details',
        'Room Details', 'Accommodation Details', 'Observation Booking',
        'Observation Reservation', 'Observation of Performance',
        'Observation of Service', 'Description of Service', 'Description of the',
        'Tour Description', 'Booking No', 'Reservation No', 'Note :',
        'Pickup Location', 'Vehicle Return', 'Departure :', 'Arrival :',
        'Flight operated', 'PNR :', 'Class :', 'Stay booked', 'Meeting point',
        'Performance Observation',
        # VDM additional body starters (File 1746499 format)
        'Booking Remark', 'Service Remark', 'Service description',
        'Tour description', 'Description :', 'Remark(s)',
    )

    # VDM page-break noise lines to skip
    NOISE_PATTERNS = [
        _re.compile(r'^\d{1,2}/\d{2}/\d{2,4},\s+\d+:\d+\s+[AP]M\s+File'),
        _re.compile(r'^DEPARTURE DATE\s+CITY\s+DURATION'),
        _re.compile(r'^https?://'),
        _re.compile(r'^\d+/\d+\s*$'),   # "3/9" page number
    ]

    # Inline noise pattern — page timestamp embedded mid-line
    # e.g. "OBEROI MAIDENS , 12:17 File 1746499 - Services"
    _INLINE_NOISE_RE = _re.compile(
        r',?\s*\d{1,2}:\d{2}\s+File\s+\d+\s*-\s*Services.*$',
        _re.IGNORECASE
    )

    def _strip_inline_noise(s):
        return _INLINE_NOISE_RE.sub('', s).strip()

    def _is_noise(line):
        return any(p.match(line) for p in NOISE_PATTERNS)

    def _is_body_start(line):
        return any(line.startswith(s) for s in BODY_STARTERS)

    def _is_svc_line(line):
        parts = line.split()
        return bool(parts and parts[0] in DAY_ABBREVS and
                    any(kw in line for kw in SVC_KEYWORDS))

    raw_lines = [l.strip() for l in full_text.split('\n')]
    accommodation_rows = []
    seen_legs = {}
    n_lines = len(raw_lines)

    i = 0
    while i < n_lines:
        line = raw_lines[i]
        if not line or _is_noise(line):
            i += 1
            continue

        if not _is_svc_line(line):
            i += 1
            continue

        svc_line = line
        city_raw, duration, svc_type, partial_name = _vdm_parse_service_line(svc_line)
        if not svc_type:
            i += 1
            continue
        partial_name = _strip_inline_noise(partial_name)

        # ── Collect date + name continuation from following lines ─────
        date_iso      = None
        name_cont_parts = []
        body_lines    = []
        j = i + 1

        while j < n_lines:
            nxt = raw_lines[j]
            if not nxt or _is_noise(nxt):
                j += 1
                continue

            # Another service line → stop
            if _is_svc_line(nxt):
                break

            # Date line (with optional name continuation after the date)
            if date_iso is None:
                dm = _re.match(r'^(\d{2}/\d{2}/\d{4})\s*(.*)', nxt)
                if dm:
                    date_iso = _dmY_to_iso(dm.group(1))
                    cont = dm.group(2).strip()
                    if cont:
                        name_cont_parts.append(cont)
                    j += 1
                    continue

            # Body description line → stop name collection, start body
            if _is_body_start(nxt):
                # Collect remaining body lines for Rooming/MealPlan parsing
                while j < n_lines:
                    b = raw_lines[j]
                    if _is_svc_line(b) or _is_noise(b):
                        break
                    body_lines.append(b)
                    j += 1
                break

            # Additional name continuation (e.g. "In French", "HOUSE", "HOMESTAY AND ART GALLERY")
            # Stop if we already have 4 continuation pieces (avoids runaway into description prose)
            if len(name_cont_parts) < 4:
                name_cont_parts.append(nxt)
            j += 1

        # Full benefit name = partial from svc_line + continuation from next lines
        # Strip inline page-break noise from each continuation part
        name_cont_parts = [_strip_inline_noise(p) for p in name_cont_parts]
        full_name = (partial_name + ' ' + ' '.join(name_cont_parts)).strip()
        full_name = _re.sub(r'\s+', ' ', full_name)
        block_text = '\n'.join(body_lines)

        svc_lower = svc_type.lower()

        # ── Accommodation ─────────────────────────────────────────────
        if svc_lower == 'accommodation' and duration > 0:
            city  = _vdm_norm_city(city_raw)
            hotel = _re.sub(r'\s+', ' ', full_name).strip()

            meal_m    = _re.search(r'Meal Plan\s*:\s*(.+?)(?:\n|$)', block_text, _re.IGNORECASE)
            meal_plan = _map_meal_plan(meal_m.group(1) if meal_m else '')

            # Passenger names from Rooming block
            if not intent['passenger_names']:
                rooming_idx = next(
                    (k for k, bl in enumerate(body_lines) if bl.strip() == 'Rooming :'), None
                )
                if rooming_idx is not None:
                    for bl in body_lines[rooming_idx + 1:]:
                        bl = bl.strip()
                        if not bl or _is_body_start(bl):
                            break
                        if any(t in bl for t in ('Mr', 'Mrs', 'Ms', 'Miss', 'M.')):
                            intent['passenger_names'].append(bl)
                    if intent['passenger_names'] and not intent['guest_name']:
                        last = intent['passenger_names'][-1]
                        nm = _re.match(r'(?:Mr|Mrs|Ms|Miss|M\.)\s+(\S+)\s+(\S+)', last)
                        if nm:
                            intent['guest_name'] = f"{nm.group(1)} {nm.group(2)}"

            ci_date  = _dt.fromisoformat(date_iso).date() if date_iso else None
            co_date  = ci_date + _td(days=duration) if ci_date else None
            checkout = co_date.isoformat() if co_date else None

            leg_key = f"{city}_{date_iso}"
            if leg_key in seen_legs:
                if city not in intent['dual_option_hotels']:
                    intent['dual_option_hotels'][city] = [seen_legs[leg_key]]
                intent['dual_option_hotels'][city].append(hotel)
                print(f"  ⚠ DUAL OPTION {city}: {seen_legs[leg_key]}  vs  {hotel}")
                if hotel not in intent['agent_specified_hotels']:
                    intent['agent_specified_hotels'].append(hotel)
                i = j
                continue

            seen_legs[leg_key] = hotel
            accommodation_rows.append({
                'city': city, 'check_in': date_iso, 'check_out': checkout,
                'hotel_name': hotel, 'meal_plan': meal_plan, 'nights': duration,
            })

        # ── Miscellaneous — client-booked accommodation ───────────────
        # City added to route (Step 8) but hotel add skipped (Step 10)
        elif svc_lower == 'miscellaneous' and 'accommodation' in full_name.lower() and duration > 0:
            city    = _vdm_norm_city(city_raw)
            ci_date = _dt.fromisoformat(date_iso).date() if date_iso else None
            co_date = ci_date + _td(days=duration) if ci_date else None
            accommodation_rows.append({
                'city':            city,
                'check_in':        date_iso,
                'check_out':       co_date.isoformat() if co_date else None,
                'hotel_name':      '',          # no hotel to add
                'meal_plan':       'cp',
                'nights':          duration,
                'booked_by_client': True,       # Step 10 will skip hotel add
            })
            print(f"  ℹ {city} ({duration}n) — accommodation booked by client: city added to route, hotel skipped")

        # ── Excursion ─────────────────────────────────────────────────
        elif svc_lower == 'excursion':
            city     = _vdm_norm_city(city_raw)
            act_name = full_name
            # Strip language / duration suffixes: "- In French", "- 2 hours"
            act_name = _re.sub(
                r'\s*-\s*(?:In\s+(?:English|French|German|Italian|Spanish)|'
                r'\d+(?:\.\d+)?\s*h(?:ours?)?|\d+(?:\.\d+)?\s*hours?)\s*'
                r'(?:\([^)]+\))?\s*$',
                '', act_name, flags=_re.IGNORECASE
            ).strip()
            act_name = _re.sub(r'\s*-\s*\d+(?:\.\d+)?\s*h(?:ours?)?\s*$', '', act_name).strip()
            intent['csv_activities'].append({
                'city': city, 'date': date_iso, 'name': act_name,
                'description': full_name,
                'pax': intent['num_pax'],
                'guide': 'guide' in block_text.lower(),
            })

        # ── Suggestions (self-guided / day visits) ────────────────────
        elif svc_lower == 'suggestions':
            city = _vdm_norm_city(city_raw)
            intent['csv_activities'].append({
                'city': city, 'date': date_iso, 'name': full_name,
                'description': full_name, 'pax': intent['num_pax'], 'guide': False,
            })

        # ── Transport / Flights ───────────────────────────────────────
        elif svc_lower == 'transport':
            # Full route string: partial_name from svc_line + continuation from date line
            full_route = (partial_name + ' ' + ' '.join(name_cont_parts)).strip()
            iata_codes = _re.findall(r'\b([A-Z]{3})\b', svc_line + ' ' + full_route)
            iata_codes = [c for c in iata_codes if c.isalpha() and len(c) == 3]
            # Remove non-airport codes that appear in airline names
            _skip_codes = {'AIR', 'AIL', 'IND', 'PNR'}
            iata_codes  = [c for c in iata_codes if c not in _skip_codes]
            if len(iata_codes) >= 2:
                from_iata   = iata_codes[0]
                to_iata     = iata_codes[1]
                is_domestic = (from_iata not in _VDM_INTL_IATA and
                               to_iata   not in _VDM_INTL_IATA)
                pnr_m = _re.search(r'PNR\s*:\s*([A-Z0-9]+)', block_text, _re.IGNORECASE)
                intent['csv_flights'].append({
                    'date': date_iso, 'from_iata': from_iata, 'to_iata': to_iata,
                    'details': (svc_line + ' ' + full_route).strip(),
                    'pnr': pnr_m.group(1) if pnr_m else '',
                    'domestic': is_domestic,
                })
                if is_domestic:
                    intent['transport_notes'].append('domestic flights mentioned')

        i = j  # advance past the block we just processed

    # ── Build nightly_split, city sequence, hotel map ─────────────────────
    nightly_split = []
    agent_hotels  = []
    all_cities    = []
    csv_hotel_map = {}

    client_booked_cities = set()
    for row in accommodation_rows:
        city   = row['city']
        hotel  = row['hotel_name']
        nights = row['nights']
        nightly_split.append((city, nights))
        if row.get('booked_by_client'):
            client_booked_cities.add(city)
        if hotel and hotel not in agent_hotels:
            agent_hotels.append(hotel)
        if not all_cities or all_cities[-1] != city:
            all_cities.append(city)
        else:
            all_cities.append(city)  # consecutive same city (e.g. houseboat + hotel)
        if city not in csv_hotel_map:
            csv_hotel_map[city] = hotel
        else:
            existing = csv_hotel_map[city]
            csv_hotel_map[city] = ([existing, hotel] if not isinstance(existing, list)
                                   else existing + [hotel])

    intent['nightly_split']          = nightly_split
    intent['agent_specified_hotels'] = agent_hotels
    intent['cities_detected']        = all_cities
    intent['_input_mode']            = 'PDF'
    intent['client_booked_cities']   = client_booked_cities  # Step 10 skips hotel add for these
    intent['csv_hotel_by_city']      = csv_hotel_map
    intent['_accommodation_rows']    = accommodation_rows

    # Travel dates from accommodation rows if header parsing missed them
    if accommodation_rows:
        check_ins  = [r['check_in']  for r in accommodation_rows if r['check_in']]
        check_outs = [r['check_out'] for r in accommodation_rows if r['check_out']]
        if check_ins and not intent['travel_start']:
            intent['travel_start'] = min(check_ins)
        if check_outs and not intent['travel_end']:
            intent['travel_end']   = max(check_outs)
        if intent['travel_start'] and intent['travel_end'] and not intent['duration_nights']:
            ts = _dt.fromisoformat(intent['travel_start']).date()
            te = _dt.fromisoformat(intent['travel_end']).date()
            intent['duration_nights'] = ((te - ts).days, (te - ts).days)

    # Regions from detected cities
    for city in all_cities:
        region = CITY_TO_REGION.get(city)
        if region and region not in intent['regions_detected']:
            intent['regions_detected'].append(region)

    # Heritage flag from hotel names
    for hotel in agent_hotels:
        if any(s in hotel.lower() for s in HERITAGE_NAME_SIGNALS):
            intent['heritage_requested'] = True
            intent['heritage_keywords']  = ['heritage property detected in PDF']
            break

    # ── Print summary ─────────────────────────────────────────────────────
    print(f"\n── PARSED INTENT (PDF) ─────────────────────────────────────────")
    print(f"  Agency       : {intent['agency_name'] or 'VDM (from PDF)'}  [{intent['agency_account_code'] or 'no CRM match'}]")
    print(f"  Market       : {intent['source_market']}")
    print(f"  Pax          : {intent['num_pax']}")
    print(f"  Rooms        : {intent['rooms']}")
    print(f"  Guest        : {intent['guest_name']}")
    print(f"  Ref          : {intent['itinerary_ref']}")
    if intent['travel_start'] and intent['travel_end']:
        ts  = _dt.fromisoformat(intent['travel_start']).date()
        te  = _dt.fromisoformat(intent['travel_end']).date()
        print(f"  Travel       : {intent['travel_start']} → {intent['travel_end']}  ({(te-ts).days} nights)")

    print(f"\n  ⚡ Route from PDF ({sum(n for _,n in nightly_split)} overnight nights):")
    for city, nights in nightly_split:
        print(f"     {city:35} {nights}n")

    print(f"\n  Hotels specified:")
    for h in agent_hotels:
        print(f"     → {h}")

    if intent['dual_option_hotels']:
        print(f"\n  ⚠ DUAL OPTIONS (quote both):")
        for city, opts in intent['dual_option_hotels'].items():
            print(f"     {city}: {' vs '.join(opts)}")

    if intent['csv_activities']:
        print(f"\n  Activities / Suggestions ({len(intent['csv_activities'])}):")
        for a in intent['csv_activities']:
            print(f"     {a['city']:25} {a['name'][:55]}")

    dom_flights  = [f for f in intent['csv_flights'] if f.get('domestic')]
    intl_flights = [f for f in intent['csv_flights'] if not f.get('domestic')]
    if dom_flights:
        print(f"\n  Domestic flights ({len(dom_flights)}):")
        for f in dom_flights:
            print(f"     {f['from_iata']} → {f['to_iata']}  [{f['date']}]")
    if intl_flights:
        print(f"  International flights ({len(intl_flights)}) — noted, not pushed to TCI")

    print("\n" + "=" * 70)
    return intent


# ── End PDF Itinerary Parser ───────────────────────────────────────────────


def _api_headers(token):
    """Full browser headers required by API Gateway authorizer."""
    from requests.structures import CaseInsensitiveDict
    h = CaseInsensitiveDict()
    h['accept'] = '*/*'
    h['accept-language'] = 'en-US,en;q=0.9'
    h['authorization'] = f'Bearer {token}'
    h['cache-control'] = 'no-cache'
    h['content-type'] = 'application/json'
    h['origin'] = 'https://dtu5juabzuf4r.cloudfront.net'
    h['pragma'] = 'no-cache'
    h['priority'] = 'u=1, i'
    h['referer'] = 'https://dtu5juabzuf4r.cloudfront.net/'
    h['sec-ch-ua'] = '"Chromium";v="146", "Not-A.Brand";v="24", "Microsoft Edge";v="146"'
    h['sec-ch-ua-mobile'] = '?0'
    h['sec-ch-ua-platform'] = '"Windows"'
    h['sec-fetch-dest'] = 'empty'
    h['sec-fetch-mode'] = 'cors'
    h['sec-fetch-site'] = 'cross-site'
    h['user-agent'] = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/146.0.0.0 Safari/537.36 Edg/146.0.0.0'
    return h


# ── PRECONFIGURED TOUR CODE MAP ──────────────────────────────────────────
# Maps (agency_account_code, source_market) → PT business_code
# Loaded dynamically from preconfigured_tours.csv — no code change needed
# to add new agencies, just update the CSV and reload.
PRECONFIGURED_TOUR_MAP = {}   # populated by load_preconfigured_tour_map()

_PT_CSV_PATH = '/content/preconfigured_tours.csv'


# Market-level PT code fallback — used when agency account not in CRM
# Picks the highest-volume agency PT code for the market
_MARKET_PT_DEFAULT = {
    'GBR': 'PT110000',   # Audley Travel — largest GBR agency
    'DEU': 'PT110005',   # Audley Travel DEU
    'FRA': 'PT110003',   # Comptoir Des Voyages
    'AUS': 'PT110013',   # TravelBridge
    'IND': 'PT110007',   # Aa Bee Holidays
}

def load_preconfigured_tour_map(csv_path=_PT_CSV_PATH):
    """Load agency → PT code mapping from CSV file.
    Falls back to hardcoded defaults if file not found.
    """
    global PRECONFIGURED_TOUR_MAP
    if PRECONFIGURED_TOUR_MAP:
        return  # already loaded
    try:
        import csv as _csv_pt
        with open(csv_path, encoding='utf-8-sig') as f:
            for row in _csv_pt.DictReader(f):
                key = (row['customer_code'].strip(), row['source_market_code'].strip())
                PRECONFIGURED_TOUR_MAP[key] = row['business_code'].strip()
        print(f"  ✓ Preconfigured tour map loaded: {len(PRECONFIGURED_TOUR_MAP)} agency+market combos")
    except FileNotFoundError:
        # Hardcoded fallback
        PRECONFIGURED_TOUR_MAP.update({
            ('ACC0131', 'GBR'): 'PT110000',
            ('ACC1606', 'DEU'): 'PT110005',
            ('ACC1816', 'DEU'): 'PT110005',  # Sensation Travel GMBH (Tourlane)
            ('ACC0272', 'FRA'): 'PT110003',
            ('ACC0861', 'FRA'): 'PT110001',
            ('ACC0538', 'GBR'): 'PT110004',
            ('ACC0042', 'GBR'): 'PT110010',
            ('ACC0806', 'AUS'): 'PT110014',
            ('ACC2679', 'AUS'): 'PT110013',
            ('ACC2344', 'MYS'): 'PT110012',
            ('ACC0014', 'IND'): 'PT110007',
            ('ACC2276', 'IND'): 'PT110015',
        })
        print(f"  ⚠ preconfigured_tours.csv not found — using hardcoded map ({len(PRECONFIGURED_TOUR_MAP)} entries)")
    except Exception as e:
        print(f"  ⚠ Preconfigured tour map load error: {e}")


# Market default account codes — used when agency not in CRM portfolio
_MARKET_DEFAULT_ACCOUNTS = {
    'GBR': 'ACC0090',   # default GBR account
    'DEU': 'ACC0118',   # Asien Special Tours
    'FRA': 'ACC0070',   # Altiplano Voyages
    'ITA': 'ACC0090',
    'USA': 'ACC0090',
    'AUS': 'ACC0090',
}

def _default_account(market):
    return _MARKET_DEFAULT_ACCOUNTS.get(market, 'ACC0090')


def select_route(intent, option=1):
    """Select which route option to use before calling create_full_query().

    The recommend() function stores all route options in intent['all_routes'].
    Option 1 is the primary route, Option 2, 3... are alternatives.
    Call this between recommend() and create_full_query() to choose a route.

    Args:
        intent  : dict returned by recommend()
        option  : int, 1-based route number shown in the recommendation output
                  (1 = primary, 2 = Option 2, 3 = Option 3, etc.)

    Returns:
        Updated intent dict with nightly_split set to the chosen route.

    Example:
        intent = engine.recommend(EMAIL, db_path=..., portfolio_path=...)
        intent = engine.select_route(intent, option=3)   # pick Option 3
        result = engine.create_full_query(intent=intent, ...)
    """
    all_routes = intent.get('all_routes') or []

    if not all_routes:
        print('⚠  No route options found in intent — route unchanged.')
        return intent

    if option < 1 or option > len(all_routes):
        print(f'⚠  Option {option} is out of range (1–{len(all_routes)}) — route unchanged.')
        return intent

    chosen = all_routes[option - 1]
    intent['nightly_split'] = chosen
    intent['cities_detected'] = [c for c, n in chosen]
    intent['_selected_route_option'] = option

    total_nights = sum(n for _, n in chosen)
    print(f'\n  ✓ Route Option {option} selected ({total_nights} nights):')
    for city, nights in chosen:
        label = f'{nights}n' if nights > 0 else 'day visit (en route)'
        print(f'     {city:35} {label}')

    return intent


def create_full_query(intent, email_raw, api_base_url, uat_base_url, api_token,
                      start_date=None, db_path=DB_PATH):
    import time as _time
    _t_start = _time.time()
    """Execute all 8 steps to create and populate a query in the legacy system.

    Steps:
        1. POST /queries                    — create query header
        2. POST /queries/{id}/users         — self-assign to logged-in user
        3. GET  /queries/{id}/basic-details — fetch assigned user details
        4. POST /queries/{id}/basic-details — update basic details
        5. POST /queries/{id}/travel-info   — pax, rooms, vehicle, meal plan
        6. POST /queries/{id}/stay-details  — country, duration, start date
        7. POST /queries/{id}/preferences   — hotel tier classification
        8. POST /queries/{id}/itinerary/routes — city route with dates

    Args:
        intent        — dict returned by engine.recommend()
        email_raw     — raw email text string
        api_base_url  — AWS API Gateway base  e.g. https://c0nibkm2k0…/tciapi
        uat_base_url  — UAT product base       e.g. https://uat-product.tcitech.in
        api_token     — Bearer token for both endpoints
        start_date    — travel start date string YYYY-MM-DD (overrides intent if provided)

    Returns:
        dict with keys: query_id, file_code, steps_completed, errors
    """
    import requests as _req
    from datetime import datetime as _dt, timedelta as _td

    # Guard — parse_pdf_itinerary / parse_csv_itinerary returns None on failure
    if intent is None:
        print("  ❌ create_full_query: intent is None — parsing failed (check pdfplumber is installed)")
        return {'query_id': None, 'file_code': None, 'steps_completed': [], 'errors': ['Intent is None — parsing failed']}

    result = {'query_id': None, 'file_code': None, 'steps_completed': [], 'errors': []}
    hdrs   = _api_headers(api_token)
    load_preconfigured_tour_map()   # ensure PT code map is loaded
    apib   = api_base_url.rstrip('/')
    uatb   = uat_base_url.rstrip('/')

    def _post(label, url, payload):
        try:
            r = _req.post(url, json=payload, headers=hdrs, timeout=15)
            r.raise_for_status()
            body = r.json()
            # API returns HTTP 200 even for errors — check inner status
            inner = body.get('status', 200) if isinstance(body, dict) else 200
            if inner >= 400:
                msg = body.get('message', str(body))[:200]
                result['errors'].append(f"{label}: API error {inner} — {msg}")
                print(f"  ✗ {label}: API {inner} — {msg[:120]}")
                return None
            result['steps_completed'].append(label)
            print(f"  ✓ {label}")
            return body
        except Exception as e:
            # Log response body for 400 errors to diagnose payload issues
            err_str = str(e)
            try:
                if hasattr(e, 'response') and e.response is not None:
                    try:
                        rb = e.response.json()
                        err_detail = str(rb)[:300]
                    except Exception:
                        err_detail = e.response.text[:300]
                    err_str = f"{e} | Response: {err_detail}"
            except Exception:
                pass
            result['errors'].append(f"{label}: {err_str}")
            print(f"  ✗ {label}: {err_str[:200]}")
            return None

    def _get(label, url):
        try:
            r = _req.get(url, headers=hdrs, timeout=15)
            r.raise_for_status()
            body = r.json()
            inner = body.get('status', 200) if isinstance(body, dict) else 200
            if inner >= 400:
                msg = body.get('message', str(body))[:200]
                result['errors'].append(f"{label}: API error {inner} — {msg}")
                print(f"  ✗ {label}: API {inner} — {msg[:120]}")
                return None
            result['steps_completed'].append(label)
            print(f"  ✓ {label}")
            return body
        except Exception as e:
            result['errors'].append(f"{label}: {e}")
            print(f"  ✗ {label}: {e}")
            return None

    # ── Step 1: Create query ───────────────────────────────────────────────────
    payload1 = build_api_payload(intent, email_raw)
    resp1    = _post('Step 1 — Create query', f"{apib}/queries", payload1)
    # Handle both direct response and nested {status, data} wrapper
    if resp1 and 'data' not in resp1 and 'id' in resp1:
        resp1 = {'data': resp1}  # unwrap if API returns data directly
    if not resp1 or 'data' not in resp1 or not resp1['data']:
        print("  ✗ Cannot continue — query creation failed")
        return result

    query_id = resp1['data']['id']          # e.g. N009026030002_1
    file_code = resp1['data']['fileCode']   # e.g. N009026030002
    result['query_id']  = query_id
    result['file_code'] = file_code
    print(f"     Query ID: {query_id}")

    # ── Step 2: Self-assign ────────────────────────────────────────────────────
    _post('Step 2 — Assign query',
          f"{apib}/queries/{query_id}/users",
          {'selfAssign': True, 'comments': ''})

    # ── Step 3: Fetch basic details (to get assignedTo/Email) ─────────────────
    resp3 = _get('Step 3 — Fetch basic details',
                 f"{apib}/queries/{query_id}/basic-details")

    # ── Step 4: Update basic details ─────────────────────────────────────────
    if resp3 and 'data' in resp3:
        basic = resp3['data'].get('basicInfo', {}) or {}
        # emailTime format: strip seconds if present
        et = basic.get('emailTime', '')
        if et and len(et) > 5:
            basic['emailTime'] = et[:5]
        # Patch arrivalDetails and departureDetails with first/last city codes
        # TCI requires these to be populated for query to proceed to next steps
        _route_cities = [c for c, n in (intent.get('nightly_split') or []) if n > 0]
        _arr_code = str(CITY_CODE_MAP.get(_route_cities[0], '')) if _route_cities else ''
        _dep_code = str(CITY_CODE_MAP.get(_route_cities[-1], '')) if _route_cities else ''
        if _arr_code:
            basic['arrivalDetails'] = basic.get('arrivalDetails') or {}
            if not basic['arrivalDetails'].get('city'):
                basic['arrivalDetails']['city'] = _arr_code
        if _dep_code:
            basic['departureDetails'] = basic.get('departureDetails') or {}
            if not basic['departureDetails'].get('city'):
                basic['departureDetails']['city'] = _dep_code
        _post('Step 4 — Update basic details',
              f"{apib}/queries/{query_id}/basic-details", basic)

    # ── Step 5: Travel info (pax, rooms, vehicle) ─────────────────────────────
    rooms   = intent.get('rooms', {})
    pax     = intent.get('num_pax') or 1
    singles = rooms.get('singles', 0)
    doubles = rooms.get('doubles', 0)
    twins   = rooms.get('twins', 0)
    # If no rooms detected, derive from pax
    if singles + doubles + twins == 0:
        doubles = pax // 2
        singles = pax % 2

    room_config = {}
    if singles: room_config['single'] = singles
    if doubles: room_config['double'] = doubles
    if twins:   room_config['double'] = room_config.get('double', 0) + twins  # twin → double (TCI only supports single/double/triple)

    _post('Step 5 — Travel info',
          f"{uatb}/queries/{query_id}/travel-info",
          {
              'isFamilyTour':                '',
              'isShoppingTour':              '',
              'isRepAgentCommissionRemoved': '',
              'travellers':                  [],
              'vehicleCount':                1,
              'guideCount':                  1,
              'ftlCount':                    0,
              'mealPlan':                    'cp',
              'roomConfig':                  room_config,
              'rangeMinPax':                 pax,
              'rangeMaxPax':                 pax,
              'paxRangeSlabs':               [{'min': pax, 'max': pax}],
          })

    # ── Step 6: Stay details ──────────────────────────────────────────────────
    dur     = intent.get('duration_nights')
    nights  = dur[1] if isinstance(dur, (list, tuple)) else (dur or 7)
    t_start = start_date or intent.get('travel_start') or _dt.utcnow().strftime('%Y-%m-%d')

    _post('Step 6 — Stay details',
          f"{uatb}/queries/{query_id}/stay-details",
          {
              'scheduleType': 'flexible',
              'countries': [{'country': 'IND', 'lengthOfStay': nights, 'regions': []}],
              'flexibleSchedule': {'startDate': t_start, 'lengthOfStay': nights},
          })

    # ── Step 7: Preferences (hotel tier) ──────────────────────────────────────
    tier = intent.get('tier') or 'first-class'
    _post('Step 7 — Preferences',
          f"{apib}/queries/{query_id}/preferences",
          {
              'difficultyLevel':          None,
              'tourCategory':             '',
              'tourThemes':               None,
              'tciClassification':        tier,
              'officialClassification':   None,
              'tourThemesSpecializations': [],
          })

    # ── Step 8: Routes ────────────────────────────────────────────────────────
    nightly_split = intent.get('nightly_split') or []
    if nightly_split:
        routes     = []
        order      = 1
        cur_date   = _dt.strptime(t_start, '%Y-%m-%d')
        missing    = []

        for city, nights_n in nightly_split:
            city_code = CITY_CODE_MAP.get(city)
            if not city_code:
                missing.append(city)
                continue
            if nights_n == 0:
                # Transit/enroute city — add as 0-night day visit
                routes.append({
                    'city':         city_code,
                    'cityName':     city,
                    'lengthOfStay': 0,
                    'startDate':    cur_date.strftime('%Y-%m-%d'),
                    'endDate':      cur_date.strftime('%Y-%m-%d'),
                    'routeOrder':   order,
                })
                order += 1
                # Don't advance cur_date — transit is same day
                continue
            end_date = cur_date + _td(days=nights_n)
            routes.append({
                'city':         city_code,
                'cityName':     city,
                'lengthOfStay': nights_n,
                'startDate':    cur_date.strftime('%Y-%m-%d'),
                'endDate':      end_date.strftime('%Y-%m-%d'),
                'routeOrder':   order,
            })
            cur_date = end_date
            order   += 1

        if missing:
            print(f"  ⚠ Cities without system codes (skipped from routes): {missing}")
            result['errors'].append(f"Step 8 — missing city codes: {missing}")

        if routes:
            resp8 = _post('Step 8 — Create routes',
                          f"{uatb}/queries/{query_id}/itinerary/routes",
                          {'routes': routes})
            # Extract routeId per city for Step 9 hotel search+add
            # Step 8 response: { data: { routes: [ { id: "uuid", city: "11", startDate, endDate }, ... ] } }
            route_ids = {}  # city_code -> routeId UUID
            # Step 8 response can be:
            # A: { data: [ {id, city, ...}, ... ] }  — data is a list directly
            # B: { data: { routes: [ {id, city, ...}, ... ] } }  — data is a dict
            # C: [ {id, city, ...}, ... ]  — flat list at top level
            raw_routes = []
            if resp8:
                data = resp8.get('data') if isinstance(resp8, dict) else resp8
                if isinstance(data, list):
                    raw_routes = data
                elif isinstance(data, dict):
                    raw_routes = data.get('routes') or data.get('content') or []
            for r in raw_routes:
                if not isinstance(r, dict):
                    continue
                city_val = r.get('city') or r.get('cityCode', '')
                # city_val may be a full dict {code, name, ...} or just a string code
                if isinstance(city_val, dict):
                    ccode = str(city_val.get('code', ''))
                else:
                    ccode = str(city_val)
                rid = r.get('id') or r.get('routeId')
                if ccode and rid:
                    if ccode in route_ids:
                        existing = route_ids[ccode]
                        if isinstance(existing, list):
                            existing.append(rid)
                        else:
                            route_ids[ccode] = [existing, rid]
                    else:
                        route_ids[ccode] = rid
            if route_ids:
                result['route_ids'] = route_ids
                print(f"     Route IDs captured: {len(route_ids)} cities")
            else:
                print("  ⚠ Step 8 — routeIds not found in response (Step 9 will be skipped)")
                result['step8_raw'] = str(resp8)[:500] if resp8 else 'None'
        else:
            print("  ⚠ Step 8 — no routes to create (all cities missing codes or day visits)")
    else:
        print("  ⚠ Step 8 — no nightly split available, routes skipped")

    # ── Step 9: Monument search + add ─────────────────────────────────────────
    # For each city leg: search monuments → match to engine recommendations →
    # get detail for contract → add top 3 monuments per city
    route_ids    = result.get('route_ids', {})
    nightly_split = intent.get('nightly_split') or []

    # CSV path — skip monument add (agent pre-specifies programme)
    _skip_monuments = bool(intent.get('csv_path'))
    if _skip_monuments:
        print(f"\n  ℹ Step 9 skipped — CSV path: monuments pre-specified by agent")
    CURRENCY     = 'INR'  # Always INR
    num_pax      = intent.get('num_pax') or 2
    # Always read account from intent directly — never cache at function entry
    # so that Colab-level overrides (e.g. AGENCY_OVERRIDES) are respected
    account      = intent.get('agency_account_code') or _default_account(intent.get('source_market', ''))
    market       = intent.get('source_market', 'GBR')

    MON_SEARCH_URL = f"{apib}/monuments/action/search?page=1&size=50"
    MON_DETAIL_URL = f"{apib}/monuments/{{bc}}"

    import random as _random
    from datetime import datetime as _datetime

    def _mon_start_time(_c=[0]):
        """Return staggered start time: 09:00 + 15min per call. Reset by passing reset=True."""
        h = 9 + (_c[0] * 15) // 60
        m = (_c[0] * 15) % 60
        _c[0] += 1
        return f"{h:02d}:{m:02d}"
    def _reset_mon_time(_c=None):
        _mon_start_time.__defaults__[0][0] = 0

    if route_ids and nightly_split and not _skip_monuments:
        print(f"\n── STEP 9: MONUMENT SEARCH & ADD ────────────────────────────────")

        # Build city → recommended monument names from engine output
        # Engine stores monument recommendations per city in master.db
        import sqlite3 as _sql
        _mon_db_recs = {}  # city_name -> [monument_name, ...]
        try:
            _con_m = _sql.connect(db_path)
            _cur_m = _con_m.cursor()
            for city_name, nights_n in nightly_split:
                if nights_n == 0:
                    continue
                # Try market-specific first, fallback to all markets
                _cur_m.execute('''
                    SELECT service_name, COUNT(*) as cnt
                    FROM services
                    WHERE city_name = ? AND record_type = 'Monument'
                    AND source_market = ? AND service_name != ''
                    GROUP BY service_name ORDER BY cnt DESC LIMIT 6
                ''', (city_name, market))
                rows = _cur_m.fetchall()
                if not rows:
                    _cur_m.execute('''
                        SELECT service_name, COUNT(*) as cnt
                        FROM services
                        WHERE city_name = ? AND record_type = 'Monument'
                        AND service_name != ''
                        GROUP BY service_name ORDER BY cnt DESC LIMIT 6
                    ''', (city_name,))
                    rows = _cur_m.fetchall()
                rows = rows or []
                _mon_db_recs[city_name] = [r[0].lower() for r in rows]
            _con_m.close()
        except Exception as _e:
            print(f"  ⚠ Monument DB lookup failed: {_e}")

        cur_date_9 = _dt.strptime(t_start, '%Y-%m-%d')

        for city, nights_n in nightly_split:
            if nights_n == 0:
                cur_date_9 += _td(days=0)
                continue

            city_code = CITY_CODE_MAP.get(city)
            end_date_9 = cur_date_9 + _td(days=nights_n)
            route_id  = route_ids.get(str(city_code)) if city_code else None

            if not city_code:
                print(f"  ⚠ {city} — no city code, skipping monuments")
                cur_date_9 = end_date_9
                continue
            if not route_id:
                print(f"  ⚠ {city} — no routeId, skipping monuments")
                cur_date_9 = end_date_9
                continue

            _mon_start_time.__defaults__[0][0] = 0  # reset 15-min counter for each city
            # Step 9a — Search monuments for city + dates
            visit_date = cur_date_9.strftime('%Y-%m-%d')
            mon_search_payload = {
                'city':                          city_code,
                'customer':                      account,
                'sourceMarket':                  market,
                'adultCount':                    num_pax,
                'childCount':                    0,
                'childAge':                      [None],
                'currencyCode':                  CURRENCY,
                'nationality':                   [],
                'from':                          visit_date,
                'to':                            visit_date,
                'priceRangeFrom':                None,
                'priceRangeTo':                  None,
                'isRangeExtended':               False,
                'recommendedDurationFrom':       None,
                'recommendedDurationTo':         None,
                'recommendedDurationRangeExceeded': False,
                'bestTimeToVisit':               [],
                'includeUnavailable':            True,
            }

            mon_search_resp = _post(
                f"Step 9a — Search monuments ({city})",
                MON_SEARCH_URL,
                mon_search_payload
            )

            if not mon_search_resp:
                cur_date_9 = end_date_9
                continue

            # Unwrap response
            raw_mons = (mon_search_resp.get('data') or {}).get('content') or []
            if not raw_mons:
                print(f"  ⚠ {city} — no monuments returned")
                cur_date_9 = end_date_9
                continue

            # Match against engine-recommended monuments (by name similarity)
            rec_names = _mon_db_recs.get(city, [])
            def _fuzzy_mon_match(api_name, db_names):
                """True if api monument name fuzzy-matches a DB-recommended name."""
                a = api_name.lower()
                # Strip noise words
                a_core = a.replace(' museum', '').replace(' observatory', '')                           .replace('(observatory)', '').replace('- drive past', '')                           .replace('(drive past)', '').strip()
                for db in db_names:
                    d = db.replace(' museum', '').replace(' observatory', '')                           .replace('(observatory)', '').replace('- drive past', '')                           .replace('(drive past)', '').replace('palace of wind ', '')                           .strip()
                    if a_core in d or d in a_core:
                        return True
                    a_words = set(a_core.split()) - {'the','a','of','and','at','in','by','&'}
                    d_words = set(d.split()) - {'the','a','of','and','at','in','by','&'}
                    if len(a_words & d_words) >= 2:
                        return True
                    if len(a_words) == 1 and a_words & d_words:
                        return True
                return False

            # Suppress test entries and monuments better suited for separate day trips
            _MON_SUPPRESS = {'test ', 'sample ', 'dummy ', 'xxx '}
            matched   = []
            unmatched = []
            for m in raw_mons:
                mname = m.get('name', '').lower()
                # Skip test/dummy entries
                if any(mname.startswith(s) for s in _MON_SUPPRESS):
                    continue
                is_rec = _fuzzy_mon_match(mname, rec_names) if rec_names else True
                if is_rec:
                    matched.append(m)
                else:
                    unmatched.append(m)

            # ── Monument distribution rules ─────────────────────────────
            # Rule 1: Top 6 matched monuments max per city
            # Rule 2: Max 4 per day
            # Rule 3: 1-night city (2 days) → start from Day 2; more nights → spread
            # Filter out night/evening monuments unless explicitly requested
            _night_requested = any(kw in (intent.get('_raw_text','') or '')
                                   for kw in ['by night', 'at night', 'night visit',
                                              'night ceremony', 'night show', 'sound and light',
                                              'sound & light', 'evening show'])
            _NIGHT_KEYWORDS = ['by night', 'at night', 'night ceremony', 'night show',
                                'evening ceremony', 'light show', 'sound and light',
                                'sound & light', 'illuminated']
            if not _night_requested:
                _matched_day  = [m for m in matched
                                 if not any(kw in m.get('name','').lower() for kw in _NIGHT_KEYWORDS)]
                _unmatched_day = [m for m in unmatched
                                  if not any(kw in m.get('name','').lower() for kw in _NIGHT_KEYWORDS)]
                # Only apply night filter if we still have enough monuments (≥4 total)
                if len(_matched_day) + len(_unmatched_day) >= 4:
                    matched, unmatched = _matched_day, _unmatched_day
                else:
                    # Not enough daytime monuments — keep night ones (better than gap)
                    matched, unmatched = _matched_day + [m for m in matched if m not in _matched_day],                                          _unmatched_day + [m for m in unmatched if m not in _unmatched_day]

            # Deduplicate by name (TCI sometimes has same monument under two codes)
            # Also remove "Excluding Mausoleum" when "Including Mausoleum" is present
            _all_mon = matched + unmatched
            _has_including = any('including mausoleum' in m.get('name','').lower() for m in _all_mon)
            if _has_including:
                _all_mon = [m for m in _all_mon if 'excluding mausoleum' not in m.get('name','').lower()]

            _seen_mon_names = set()
            _deduped = []
            for _m in _all_mon:
                _mname = _m.get('name','').strip().lower()
                if _mname and _mname not in _seen_mon_names:
                    _seen_mon_names.add(_mname)
                    _deduped.append(_m)
            candidates = _deduped[:4]  # top 4 max — cap per proximity logic

            if not candidates:
                print(f"  ⚠ {city} — no matching monuments")
                cur_date_9 = end_date_9
                continue

            # Calculate available days and distribute
            # nights_n nights = nights_n + 1 days but day 1 = arrival/travel
            # So usable days = nights_n (Day 2 onwards)
            available_days = max(1, nights_n)   # day 2 onwards up to check-out
            max_per_day    = 4
            total_slots    = available_days * max_per_day
            candidates     = candidates[:min(6, total_slots)]  # cap at 6

            # Distribute monuments across days (starting from Day 2 of stay)
            # Day 2 = cur_date_9 + 1 day
            mon_day_map = {}  # day_date -> [mon_stubs]
            for idx, mon_stub in enumerate(candidates):
                day_offset = 1 + (idx // max_per_day)  # start from day 2 (offset=1)
                day_offset = min(day_offset, nights_n)  # don't exceed stay
                day_str = (cur_date_9 + _td(days=day_offset)).strftime('%Y-%m-%d')
                mon_day_map.setdefault(day_str, []).append(mon_stub)

            # Step 9b skipped — detail endpoint returns 405, all data from search stub
            mon_objects = []
            for mon_stub in candidates:
                bc_mon = mon_stub.get('businessCode')
                if not bc_mon:
                    continue
                detail_data = {}  # use stub fields only

                # Pick the best contract — prefer base (no _F1/_F2) within validity
                contracts = detail_data.get('contracts') or []
                chosen_contract = None
                for c in contracts:
                    code = c.get('contractCode', '')
                    # Prefer base contract (no _Fn suffix) valid for travel date
                    v_start = c.get('validityStart') or c.get('startDate', '')
                    v_end   = c.get('validityExpiry') or c.get('endDate', '')
                    in_date = (v_start <= visit_date <= v_end) if v_start and v_end else True
                    if in_date and not any(f'_F{i}' in code for i in range(1,10)):
                        chosen_contract = c
                        break
                # Fallback: first valid contract
                if not chosen_contract:
                    for c in contracts:
                        v_start = c.get('validityStart') or c.get('startDate', '')
                        v_end   = c.get('validityExpiry') or c.get('endDate', '')
                        in_date = (v_start <= visit_date <= v_end) if v_start and v_end else True
                        if in_date:
                            chosen_contract = c
                            break
                if not chosen_contract and contracts:
                    chosen_contract = contracts[0]

                if not chosen_contract:
                    # Fall back to search response contract fields
                    chosen_contract = {
                        'id':          mon_stub.get('contractId'),
                        'contractCode': mon_stub.get('contractCode'),
                        'standardRateCard': {
                            'id':           mon_stub.get('contractStandardRateCardId'),
                            'foreignAdult': mon_stub.get('contractStandardRateCardForeignAdult', 0),
                            'foreignChild': mon_stub.get('contractStandardRateCardForeignChild', 0),
                            'foreignInfant': mon_stub.get('contractStandardRateCardForeignInfant', 0),
                            'indianAdult':  mon_stub.get('contractStandardRateCardIndianAdult', 0),
                            'indianChild':  mon_stub.get('contractStandardRateCardIndianChild', 0),
                            'indianInfant': mon_stub.get('contractStandardRateCardIndianInfant', 0),
                        },
                        'supplements': {
                            'id':              mon_stub.get('supplementsId'),
                            'cameraFee':       mon_stub.get('cameraFee', 0),
                            'videoCameraFee':  mon_stub.get('videoCameraFee', 0),
                            'professionalCameraFee': mon_stub.get('professionalCameraFee', 0),
                        },
                        'currency': mon_stub.get('currency', {}),
                    }

                rc   = chosen_contract.get('standardRateCard') or {}
                supp = chosen_contract.get('supplements') or {}
                cur_obj = chosen_contract.get('currency') or mon_stub.get('currency') or {}

                mon_name = (
                    (detail_data.get('descriptions') or [{}])[0].get('name')
                    if detail_data.get('descriptions')
                    else mon_stub.get('name', bc_mon)
                )

                # Per-monument date based on distribution rule
                _mon_idx = [m.get('businessCode') for m in candidates].index(bc_mon) if bc_mon in [m.get('businessCode') for m in candidates] else 0
                _day_off = 1 + (_mon_idx // max_per_day)
                _day_off = min(_day_off, nights_n)
                _mon_visit_date = (cur_date_9 + _td(days=_day_off)).strftime('%Y-%m-%d')

                mon_obj = {
                    'startDate':           _mon_visit_date,
                    'endDate':             _mon_visit_date,
                    'openingTime':         detail_data.get('openingTime') or '',
                    'closingTime':         detail_data.get('closingTime') or '',
                    'startTime':           _mon_start_time(),
                    'isOptional':          False,
                    'guestCategorisation': {
                        'adultAge':        {'max': detail_data.get('adultMaxAge', 99), 'min': detail_data.get('adultMinAge', 12)},
                        'areChildAllowed': detail_data.get('areChildrenAllowed', True),
                        'childAge':        {'max': detail_data.get('childMaxAge', 11),  'min': detail_data.get('childMinAge', 5)},
                        'areInfantsAllowed': detail_data.get('areInfantsAllowed', True),
                        'infantAge':       {'max': detail_data.get('infantMaxAge', 4),  'min': detail_data.get('infantMinAge', 0)},
                    },
                    'businessCode':        bc_mon,
                    'version':             detail_data.get('version') or mon_stub.get('productVersion') or 1,
                    'name':                mon_name,
                    'recommendedDuration': detail_data.get('recommendedDuration') or mon_stub.get('recommendedDuration') or 60,
                    'contractId':          chosen_contract.get('id') or mon_stub.get('contractId'),
                    'supplementsId':       supp.get('id') or mon_stub.get('supplementsId') or '',
                    'currency':            cur_obj.get('code') or 'INR',
                    'currencySymbol':      cur_obj.get('symbol') or '₹',
                    'cameraFee':           supp.get('cameraFee', 0),
                    'videoCameraFee':      supp.get('videoCameraFee', 0),
                    'professionalCameraFee': supp.get('professionalCameraFee', 0),
                    'standardRateId':      rc.get('id') or mon_stub.get('contractStandardRateCardId') or '',
                    'standardRate':        rc.get('foreignAdult') or mon_stub.get('contractStandardRateCardForeignAdult') or 0,
                    'childRate':           rc.get('foreignChild') or mon_stub.get('contractStandardRateCardForeignChild') or 0,
                    'infantRate':          rc.get('foreignInfant') or mon_stub.get('contractStandardRateCardForeignInfant') or 0,
                    'supplementsData':     [],
                    'mandatoryData':       [],
                    'addOnsData':          [],
                    'additions':           [],
                    'extraChargeIds':      [],
                    'availableWeekDays':   detail_data.get('daysAvailable') or mon_stub.get('daysAvailable') or [],
                    'restrictedDates':     [],
                    'actionType':          'add',
                    'businessCodeKey':     f"{bc_mon}{_random.random()}",
                    'serviceCity':         city_code,
                    'bookingWebsite':      detail_data.get('bookingWebsite') or '',
                    'contractCode':        chosen_contract.get('contractCode') or mon_stub.get('contractCode') or '',
                    'managementType':      detail_data.get('managementType') or mon_stub.get('managementType') or 'pvt',
                    'foreignAdultRate':    rc.get('foreignAdult') or 0,
                    'foreignChildRate':    rc.get('foreignChild') or 0,
                    'foreignInfantRate':   rc.get('foreignInfant') or 0,
                    'indianAdultRate':     rc.get('indianAdult') or 0,
                    'indianChildRate':     rc.get('indianChild') or 0,
                    'indianInfantRate':    rc.get('indianInfant') or 0,
                }
                mon_objects.append(mon_obj)

            if mon_objects:
                # Step 9c — Add monuments
                add_resp = _post(
                    f"Step 9c — Add monuments ({city}: {len(mon_objects)} monuments)",
                    f"{apib}/queries/{query_id}/itinerary/routes/{route_id}/monuments",
                    {'monuments': mon_objects}
                )
                for m in mon_objects:
                    print(f"     {m['name']} · {m['contractCode']} · INR {m['standardRate']}")

            cur_date_9 = end_date_9

    else:
        if not route_ids:
            print("\n  ℹ Step 9 skipped — no routeIds (Step 8 response needed)")

    # ── STEP 10: HOTEL SEARCH & ADD ──────────────────────────────────────────
    _hotels_by_city = {}          # city → hotel_name (primary successfully added)
    _hotels_failed_city = {}      # city → hotel_name (primary failed — needs manual add)
    _hotels_missing_dual = {}     # city → [hotel_names] (dual options not found in TCI)
    _hotels_added_dual = {}       # city → [hotel_names] (dual options successfully added)
    _hotels_primary_requested = {}# city → hotel_name ((P) primary requested by agent)
    _email_acts_by_city = {}      # city → [activity_names] (email mode — Step 13 added)
    _email_guides_by_city = {}    # city → guide_name (email mode — Step 12 added)
    _email_guides_failed = {}     # city → reason (email mode — Step 12 failed)
    HOT_SEARCH_URL = f"{apib}/hotels/search?page=1&size=20"

    # Map engine tier to API desiredHotelCategory
    TIER_TO_CAT = {'lux': ['lux'], 'first-class': ['first-class'], 'moderate': ['moderate'], 'budget': ['budget']}
    # Always search all tiers in TCI so DB-recommended hotels are never
    # excluded by a tier filter. Tier preference is applied at display time
    # in the recommendation output, not during TCI hotel search+add.
    hotel_cat = None  # search all tiers — let DB priority chain pick the right hotel

    # Map rooms to roomBreakDown
    rooms = intent.get('rooms', {})
    room_breakdown = {}
    if rooms.get('doubles', 0) > 0: room_breakdown['double'] = rooms['doubles']
    if rooms.get('singles', 0) > 0: room_breakdown['single'] = rooms['singles']
    if rooms.get('twins', 0) > 0:   room_breakdown['double'] = room_breakdown.get('double', 0) + rooms['twins']
    if not room_breakdown: room_breakdown = {'double': 1}

    if route_ids and nightly_split:
        print(f"\n── STEP 10: HOTEL SEARCH & ADD ──────────────────────────────────────")
        cur_date_10 = _dt.strptime(t_start, '%Y-%m-%d')
        _city_hotel_idx = {}  # tracks how many times each city has been processed (for consecutive same-city)
        _route_city_idx = {}   # tracks route_id index for repeated cities

        for city, nights_n in nightly_split:
            if nights_n == 0:
                continue
            city_code  = CITY_CODE_MAP.get(city)
            _route_raw = route_ids.get(str(city_code)) if city_code else None
            _route_visit_idx = _route_city_idx.get(city, 0)
            _route_city_idx[city] = _route_visit_idx + 1
            if isinstance(_route_raw, list):
                route_id = _route_raw[_route_visit_idx] if _route_visit_idx < len(_route_raw) else _route_raw[-1]
            else:
                route_id = _route_raw
            check_in   = cur_date_10.strftime('%Y-%m-%d')
            check_out  = (cur_date_10 + _td(days=nights_n)).strftime('%Y-%m-%d')

            cur_date_10 = cur_date_10 + _td(days=nights_n)

            if not city_code or not route_id:
                print(f"  ⚠ {city} — no city code or routeId, skipping hotels")
                continue

            # Skip hotel add for client-booked cities
            _client_booked = intent.get('client_booked_cities', set())
            if city in _client_booked:
                print(f"  ℹ {city} — accommodation booked by client, hotel add skipped")
                continue

            # Step 10a — Search hotels for city
            _expect_houseboat = False  # overridden below if houseboat detected
            hot_search_payload = {
                'queryId':                query_id,
                'languageCode':           'en',
                'city':                   str(city_code),
                'from':                   check_in,
                'to':                     check_out,
                'customer':               account,
                'mealPlan':               'cp',
                'sourceMarket':           market,
                'specialNeedPreference':  [],
                'themes':                 [],
                'adultCount':             num_pax or 2,
                'childCount':             0,
                'childAge':               [None],
                'noOfGuestPerRoom':       2,
                'currencyCode':           CURRENCY,
                'desiredHotelCategory':   hotel_cat if hotel_cat is not None else ['lux','first-class','moderate','budget'],
                'ratings':                None,
                'roomBreakDown':          room_breakdown,
                'roomType':               None,
                'priceRangeFrom':         None,
                'priceRangeTo':           None,
                'excludeFSRC':            False,
                'inputCreditRequired':    False,
                'customerSeries':         None,
                'paxSlabs':               None,
                'roomRanges':             None,
            }
            # Backwater cities — always include houseboat categories so boats appear in results
            _BACKWATER_CITIES = {'Alappuzha (Alleppey)', 'Kumarakom', 'Alleppey'}
            if city in _BACKWATER_CITIES:
                hot_search_payload['desiredHotelCategory'] = [
                    'houseboat-lux', 'houseboat-premium', 'houseboat-standard',
                    'lux', 'first-class', 'moderate', 'budget'
                ]
            # Override meal plan for detected houseboat stays
            if _expect_houseboat:
                hot_search_payload['mealPlan'] = 'ap'

            search_resp = _post(f"Step 10a — Hotel search ({city})", HOT_SEARCH_URL, hot_search_payload)
            if not search_resp:
                continue

            hotels_found = (search_resp.get('data') or {}).get('content') or []
            if not hotels_found:
                print(f"  ⚠ {city} — no hotels returned from search")
                continue

            # ── Hotel selection priority ──────────────────────────────────────
            # Priority 1: Agent-specified hotel (matched by name)
            # Priority 2: DB-recommended hotel (matched against TCI search results)
            # Priority 3: Top result from TCI search (fallback)

            # For CSV path: use direct city→hotel mapping (exact, no cross-city risk)
            # For email path: use fuzzy match across full agent hotel list
            _csv_hotel_map_intent = intent.get('csv_hotel_by_city') or {}
            if intent.get('csv_path') and city in _csv_hotel_map_intent:
                # CSV path: only the hotel explicitly specified for this city
                _csv_hotel_raw = _csv_hotel_map_intent[city]
                # Pick hotel for THIS specific stay using index (consecutive same-city support)
                _city_idx = _city_hotel_idx.get(city, 0)
                _city_hotel_idx[city] = _city_idx + 1
                if isinstance(_csv_hotel_raw, list):
                    _this_stay_hotel = _csv_hotel_raw[_city_idx].lower() if _city_idx < len(_csv_hotel_raw) else ''
                else:
                    _this_stay_hotel = _csv_hotel_raw.lower() if _csv_hotel_raw else ''
                agent_hotels_lower = [_this_stay_hotel] if _this_stay_hotel else []

                # Detect houseboat — add houseboat keywords to search signals
                # Houseboats are valid TCI products (kettuvallam/houseboat category)
                _HOUSEBOAT_KW = {'houseboat', 'house boat', 'kettuvallam', 'rice boat',
                                 'backwater boat', 'backwaters boat'}
                _is_houseboat = any(kw in _this_stay_hotel for kw in _HOUSEBOAT_KW)
                if _is_houseboat:
                    print(f"  ℹ {city} — houseboat product ('{_this_stay_hotel}') — searching TCI for houseboat")
                    # After search, if top result is not a houseboat product → skip and flag manually
                    # (set flag to check after hotel search completes)
                    _expect_houseboat = True
                else:
                    _expect_houseboat = False
            else:
                # Email path: full fuzzy list with cross-city exclusions
                _raw_agent_hotels = intent.get('agent_specified_hotels') or []
                _city_lower = city.lower().split('(')[0].strip()
                _dual_all   = intent.get('dual_option_hotels') or {}
                _other_city_hotels = set()
                for _dc, _dh_list in _dual_all.items():
                    if _dc != city:
                        for _dh in _dh_list:
                            _other_city_hotels.add(_dh.lower())
                agent_hotels_lower = []
                for h in _raw_agent_hotels:
                    h_l = h.lower().strip()
                    if h_l in _other_city_hotels:
                        continue
                    _ends_with_city = any(
                        h_l.rstrip().endswith(c.lower().split('(')[0].strip())
                        for c in CITY_CODE_MAP.keys()
                        if len(c) > 3
                    )
                    if _ends_with_city:
                        if _city_lower in h_l:
                            agent_hotels_lower.append(h_l.rsplit(_city_lower, 1)[0].strip())
                    else:
                        agent_hotels_lower.append(h_l)

            def _fuzzy_match(api_name, target_name):
                """Return True if api_name fuzzy-matches target_name."""
                import re as _fre
                # Normalise: lowercase, strip apostrophes/punctuation
                def _norm(s):
                    s2 = _fre.sub(r"[^a-z0-9 ]", '', s.lower()).strip()
                    return _fre.sub(r" +", ' ', s2)
                h = _norm(api_name)
                a = _norm(target_name)
                _stop = {'hotel','resort','the','a','&','and','in','by','at','spa',
                         'retreat','lodge','camp','villa','heritage','palace',
                         'himalayas','himalayan','india','indian',
                         'beach','lake','river','garden','fort','haveli',
                         'suites','suite','inn','house','home','bungalow',
                         'ihcl','seleqtions','new','delhi'}
                if a in h or h in a:
                    return True
                a_words = set(a.split()) - _stop
                h_words = set(h.split()) - _stop
                if len(a_words & h_words) >= 2:
                    return True
                _geo_stop = {'himalayas','himalayan','india','indian','palace','resort',
                             'spa','retreat','lodge','camp','villa','heritage','hotel',
                             'haveli','niwas','vilas','mahal','garh','bagh','kothi'}
                _distinct = (a_words & h_words) - _geo_stop
                if _distinct and max((len(w) for w in _distinct), default=0) >= 5:
                    return True
                if len(a_words & h_words) == 1:
                    _w = list(a_words & h_words)[0]
                    if len(_w) >= 8:
                        return True
                return False

            chosen_stub = None

            # Filter out test/audit/dummy hotels from TCI results before matching
            _TEST_KWS_F = ['test ', 'audit', 'dummy', 'sample ', 'walk through', 'walkthrough']
            hotels_found = [h for h in hotels_found
                            if not any(kw in h.get('name','').lower() for kw in _TEST_KWS_F)]

            # Priority 1 — Agent-specified hotel
            # First pass: exact / near-exact match
            def _norm_simple(s):
                import re as _nr
                s2 = _nr.sub(r'[^a-z0-9 ]', '', s.lower()).strip()
                return _nr.sub(r' +', ' ', s2)
            _ex_stop = {'the','a','by','at','and','in','of','hotel','new','delhi',
                        'ihcl','seleqtions','india','indian'}
            for h in hotels_found:
                h_norm = _norm_simple(h.get('name',''))
                for a in agent_hotels_lower:
                    if not a: continue
                    a_norm = _norm_simple(a)
                    if not a_norm: continue
                    # Full substring match
                    if a_norm in h_norm or h_norm in a_norm:
                        chosen_stub = h
                        print(f"  ★ {city}: matched agent hotel (exact) — {h.get('name')}")
                        break
                    # All significant words of shorter name appear in longer
                    a_sig = [w for w in a_norm.split() if w not in _ex_stop and len(w)>=4]
                    h_sig = [w for w in h_norm.split() if w not in _ex_stop and len(w)>=4]
                    if a_sig and h_sig:
                        shorter = a_sig if len(a_sig) <= len(h_sig) else h_sig
                        longer  = h_sig if len(a_sig) <= len(h_sig) else a_sig
                        if shorter and all(w in longer for w in shorter):
                            chosen_stub = h
                            print(f"  ★ {city}: matched agent hotel (exact) — {h.get('name')}")
                            break
                if chosen_stub:
                    break

            # Second pass: fuzzy match if no exact match found
            if not chosen_stub:
                for h in hotels_found:
                    for a in agent_hotels_lower:
                        if not a: continue
                        if _fuzzy_match(h.get('name',''), a):
                            chosen_stub = h
                            print(f"  ★ {city}: matched agent hotel — {h.get('name')}")
                            break
                    if chosen_stub:
                        break

            # Priority 2 — DB-recommended hotel
            _db_names = []
            if not chosen_stub:
                # Get DB recommendations for this city
                try:
                    import sqlite3 as _sq
                    _con2 = _sq.connect(db_path)
                    _db_hotels, _ = get_hotels(
                        _con2, city,
                        market=intent.get('source_market',''),
                        tier=intent.get('tier'),
                        heritage=intent.get('heritage_requested', False),
                        top_n=4,
                        agency_account=intent.get('agency_account_code') or None,
                        agency_name=intent.get('agency_name') or None
                    )
                    _con2.close()
                    _db_names = [r[0] for r in _db_hotels if r[0]]
                except Exception as _p2e:
                    print(f"  ⚠ {city} — DB lookup failed: {_p2e}")
                    _db_names = []

                for db_name in _db_names:
                    for h in hotels_found:
                        if _fuzzy_match(h.get('name',''), db_name):
                            chosen_stub = h
                            print(f"  ✓ {city}: matched DB recommendation — {h.get('name')} (DB: {db_name})")
                            break
                    if chosen_stub:
                        break
                if not chosen_stub and _db_names:
                    print(f"  ⚠ {city} — Priority 2 no match. DB top: {_db_names[:3]}. TCI top 5: {[h.get('name','') for h in hotels_found[:5]]}")

            # Priority 3 — Top TCI search result (skip test/audit/dummy entries)
            if not chosen_stub:
                _TEST_KWS = ['test ', 'audit', 'dummy', 'sample ', 'walk through', 'walkthrough']
                _real_hotels = [h for h in hotels_found
                                if not any(kw in h.get('name','').lower() for kw in _TEST_KWS)]
                chosen_stub = (_real_hotels or hotels_found)[0]
                # If agent wanted a houseboat but TCI returned a regular hotel — skip
                _matched_name = chosen_stub.get('name','')
                if _expect_houseboat:
                    # Prefer houseboat hotels — filter to those with boat keywords
                    _boat_kws = ['boat', 'houseboat', 'kettu', 'backwater', 'cruise', 'float']
                    _boat_hotels = [h for h in (_real_hotels or hotels_found)
                                    if any(kw in h.get('name','').lower() for kw in _boat_kws)]
                    if _boat_hotels:
                        chosen_stub = _boat_hotels[0]
                        _matched_name = chosen_stub.get('name','')
                        print(f"  🚢 {city}: matched houseboat — {_matched_name}")
                    else:
                        print(f"  ⚠ {city} — no houseboat found in TCI for this date — skipping (confirm manually)")
                        continue
                print(f"  ✓ {city}: using top TCI result — {_matched_name}")

            bc = chosen_stub.get('businessCode')
            if not bc:
                print(f"  ⚠ {city} — no businessCode in search result")
                continue

            # Step 10b — Get hotel detail for contract info
            detail_resp = _post(
                f"Step 10b — Hotel detail ({city}: {chosen_stub.get('name', bc)})",
                f"{apib}/hotels/search/{bc}",
                hot_search_payload
            )

            # Extract contract from detail or stub
            detail_data = detail_resp.get('data') if detail_resp else None
            date_contracts = []
            if detail_data:
                # dateWiseContracts is list of per-date contract objects
                date_contracts = detail_data.get('dateWiseContracts') or []
            if not date_contracts:
                # Fall back to stub dateWiseContracts
                date_contracts = chosen_stub.get('dateWiseContracts') or []

            # Build one hotel object per night
            hotel_objects = []
            for night_offset in range(nights_n):
                night_start = (cur_date_10 - _td(days=nights_n) + _td(days=night_offset)).strftime('%Y-%m-%d')
                night_end   = (cur_date_10 - _td(days=nights_n) + _td(days=night_offset + 1)).strftime('%Y-%m-%d')

                # Find matching contract for this date
                dc = next((c for c in date_contracts if c.get('travelDate') == night_start), None)
                if not dc and date_contracts:
                    dc = date_contracts[0]  # fallback to first available

                if not dc:
                    # Use stub fields directly
                    dc = {}

                occ = dc.get('occupancyRates') or chosen_stub.get('occupancyRates') or {}

                hotel_obj = {
                    'businessCode':              bc,
                    'name':                      chosen_stub.get('name', ''),
                    'roomType':                  dc.get('roomType') or chosen_stub.get('roomType') or '',
                    'roomRate':                  dc.get('perRoomPerNightPrice') or dc.get('totalPrice') or 0,
                    'totalRate':                 dc.get('totalPrice') or 0,
                    'currency':                  'INR',
                    'currencySymbol':            '₹',
                    'guestCategorisation': {
                        'adultAge':              {'max': 99, 'min': 13},
                        'areChildAllowed':       True,
                        'childAge':              {'max': 12, 'min': 9},
                        'areYoungChildAllowed':  True,
                        'youngChildAge':         {'max': 8, 'min': 5},
                        'areInfantsAllowed':     True,
                        'infantAge':             {'max': 4, 'min': 0},
                    },
                    'contractId':                dc.get('contractId') or chosen_stub.get('contractId') or '',
                    'rateCardId':                dc.get('rateCardId') or '',
                    'mealPlan':                  dc.get('mealPlan') or 'cp',
                    'occupancyRates':            occ,
                    'version':                   chosen_stub.get('version') or 1,
                    'dateRates':                 [],
                    'supplierName':              _get_supplier_info(bc, chosen_stub.get('name',''), city).get('supplierName') or None,
                    'supplierCode':              _get_supplier_info(bc, chosen_stub.get('name',''), city).get('supplierCode') or dc.get('supplierCode') or '',
                    'parentGroupId':             _get_supplier_info(bc, chosen_stub.get('name',''), city).get('parentGroupId'),
                    'parentGroup':               _get_supplier_info(bc, chosen_stub.get('name',''), city).get('parentGroup'),
                    'contractType':              dc.get('contractType') or 'gen-mkt',
                    'contractCode':              dc.get('contractCode') or '',
                    'rateTypeCode':              dc.get('rateTypeCode') or dc.get('contractCode') or '',
                    'roomCode':                  dc.get('roomCode') or '',
                    'rateCardCode':              dc.get('rateCardCode') or '',
                    'roomRateId':                dc.get('roomRateId') or '',
                    'checkInTime':               dc.get('checkInTime') or '14:00:00',
                    'checkOutTime':              dc.get('checkOutTime') or '12:00:00',
                    'isTripleOccancyAvailable':  dc.get('isTripleOccancyAvailable', False),
                    'isFirmContract':            dc.get('isFirmContract', True),
                    'updateCounter':             0,
                    'spOccupancyRates':          {},
                    'occupancyTypesWithBAR':     {},
                    'startDate':                 night_start,
                    'endDate':                   night_end,
                    'lengthOfStay':              1,
                    'isBaseCategory':            True,   # ← primary hotel
                    'hotelUnavailableDays':      [],
                    'contractStatus':            None,
                    'branchMarginPercent':       0,
                    'commission':                0,
                }
                hotel_objects.append(hotel_obj)

            if hotel_objects:
                add_resp = _post(
                    f"Step 10c — Add hotel ({city}: {chosen_stub.get('name')} · {nights_n} nights)",
                    f"{uatb}/queries/{query_id}/itinerary/routes/{route_id}/hotels",
                    {'actionType': 'add', 'hotels': hotel_objects}
                )
                # Rate not available for specific roomType — retry with roomType: null
                if add_resp is None:
                    _retry = [{**ho, 'roomType': None} for ho in hotel_objects]
                    add_resp = _post(
                        f"Step 10c — Add hotel retry ({city}: {chosen_stub.get('name')})",
                        f"{uatb}/queries/{query_id}/itinerary/routes/{route_id}/hotels",
                        {'actionType': 'add', 'hotels': _retry}
                    )
                # Still failing — retry with both roomType and mealPlan nulled
                if add_resp is None:
                    _retry2 = [{**ho, 'roomType': None, 'mealPlan': None} for ho in hotel_objects]
                    add_resp = _post(
                        f"Step 10c — Add hotel retry2 ({city}: {chosen_stub.get('name')})",
                        f"{uatb}/queries/{query_id}/itinerary/routes/{route_id}/hotels",
                        {'actionType': 'add', 'hotels': _retry2}
                    )

                # If primary still failed, alt 1 will be promoted to primary below
                _primary_failed = add_resp is None
                if _primary_failed:
                    print(f"  ⚠ {city}: primary hotel unavailable — promoting next best to primary")
                    _hotels_failed_city[city] = chosen_stub.get('name', '')
                else:
                    _hotels_by_city[city] = chosen_stub.get('name', '')
                # Track what the agent originally requested (P) primary
                _csv_hmap = intent.get('csv_hotel_by_city') or {}
                _requested = _csv_hmap.get(city, '')
                if isinstance(_requested, list): _requested = _requested[0] if _requested else ''
                if _requested:
                    _hotels_primary_requested[city] = _requested

                # ── Add alternative hotels (isBaseCategory=False) ────────────
                # CSV path: only add alt if a dual-option was specified for this city
                # Email path: use hotels 2 and 3 from TCI search results
                _dual = (intent.get('dual_option_hotels') or {}).get(city, [])
                if intent.get('csv_path'):
                    # CSV: find ALL alt hotels that are NOT the main (handles multiple duals)
                    alt_stubs = []
                    if _dual:
                        _main_name = chosen_stub.get('name','').lower()
                        _alt_names = [_dh.lower() for _dh in _dual if not _fuzzy_match(_dh, _main_name)]
                        for _alt_name in _alt_names:
                            _found = False
                            for h in hotels_found:
                                if h.get('businessCode') != bc:
                                    if _fuzzy_match(h.get('name',''), _alt_name):
                                        alt_stubs.append(h)
                                        print(f"  ✓ {city}: dual option alt — {h.get('name')}")
                                        _hotels_added_dual.setdefault(city, []).append(h.get('name',''))
                                        _found = True
                                        break
                            if not _found:
                                print(f"  ⚠ {city}: dual option '{_alt_name}' not found in TCI")
                                _hotels_missing_dual.setdefault(city, []).append(_alt_name.title())
                else:
                    # Email path: use remaining DB-recommended hotels as alts
                    # _db_names already populated above from Priority 2 query
                    if _db_names:
                        alt_stubs = []
                        for _db_alt in _db_names:
                            for h in hotels_found:
                                if h.get('businessCode') == bc:
                                    continue  # skip primary
                                if _fuzzy_match(h.get('name',''), _db_alt):
                                    if h not in alt_stubs:
                                        alt_stubs.append(h)
                                    break
                            if len(alt_stubs) >= 2:
                                break
                        # Fallback: if DB match found fewer than 2, pad from TCI results
                        if len(alt_stubs) < 2:
                            for h in hotels_found:
                                if h.get('businessCode') != bc and h not in alt_stubs:
                                    alt_stubs.append(h)
                                if len(alt_stubs) >= 2:
                                    break
                    else:
                        # No DB data — fall back to TCI order
                        alt_stubs = [h for h in hotels_found if h.get('businessCode') != bc][:2]
                for alt_idx, alt_stub in enumerate(alt_stubs):
                    alt_bc = alt_stub.get('businessCode')
                    if not alt_bc:
                        continue
                    # If primary failed, promote first alt to primary (isBaseCategory=True)
                    _is_promoted_primary = _primary_failed and alt_idx == 0
                    if _is_promoted_primary:
                        print(f"  ✓ {city}: promoting to primary — {alt_stub.get('name')}")
                    # Get alt hotel detail
                    alt_detail_resp = _post(
                        f"Step 10b — {'Primary' if _is_promoted_primary else 'Alt'} hotel detail ({city}: {alt_stub.get('name', alt_bc)})",
                        f"{apib}/hotels/search/{alt_bc}",
                        hot_search_payload
                    )
                    alt_detail = alt_detail_resp.get('data') if alt_detail_resp else None
                    alt_date_contracts = (alt_detail.get('dateWiseContracts') or []) if alt_detail else []
                    if not alt_date_contracts:
                        alt_date_contracts = alt_stub.get('dateWiseContracts') or []

                    alt_objects = []
                    for night_offset in range(nights_n):
                        night_start = (cur_date_10 - _td(days=nights_n) + _td(days=night_offset)).strftime('%Y-%m-%d')
                        night_end   = (cur_date_10 - _td(days=nights_n) + _td(days=night_offset + 1)).strftime('%Y-%m-%d')
                        dc = next((c for c in alt_date_contracts if c.get('travelDate') == night_start), None)
                        if not dc and alt_date_contracts:
                            dc = alt_date_contracts[0]
                        if not dc:
                            dc = {}
                        occ = dc.get('occupancyRates') or alt_stub.get('occupancyRates') or {}
                        alt_obj = {
                            'businessCode':              alt_bc,
                            'name':                      alt_stub.get('name', ''),
                            'roomType':                  dc.get('roomType') or alt_stub.get('roomType') or '',
                            'roomRate':                  dc.get('perRoomPerNightPrice') or dc.get('totalPrice') or 0,
                            'totalRate':                 dc.get('totalPrice') or 0,
                            'currency':                  'INR',
                            'currencySymbol':            '₹',
                            'guestCategorisation': {
                                'adultAge':              {'max': 99, 'min': 13},
                                'areChildAllowed':       True,
                                'childAge':              {'max': 12, 'min': 9},
                                'areYoungChildAllowed':  True,
                                'youngChildAge':         {'max': 8, 'min': 5},
                                'areInfantsAllowed':     True,
                                'infantAge':             {'max': 4, 'min': 0},
                            },
                            'contractId':                dc.get('contractId') or alt_stub.get('contractId') or '',
                            'rateCardId':                dc.get('rateCardId') or '',
                            'mealPlan':                  dc.get('mealPlan') or 'cp',
                            'occupancyRates':            occ,
                            'version':                   alt_stub.get('version') or 1,
                            'dateRates':                 [],
                            'supplierName':              _get_supplier_info(alt_bc, alt_stub.get('name',''), city).get('supplierName') or None,
                            'supplierCode':              _get_supplier_info(alt_bc, alt_stub.get('name',''), city).get('supplierCode') or dc.get('supplierCode') or '',
                            'parentGroupId':             _get_supplier_info(alt_bc, alt_stub.get('name',''), city).get('parentGroupId'),
                            'parentGroup':               _get_supplier_info(alt_bc, alt_stub.get('name',''), city).get('parentGroup'),
                            'contractType':              dc.get('contractType') or 'gen-mkt',
                            'contractCode':              dc.get('contractCode') or '',
                            'rateTypeCode':              dc.get('rateTypeCode') or dc.get('contractCode') or '',
                            'roomCode':                  dc.get('roomCode') or '',
                            'rateCardCode':              dc.get('rateCardCode') or '',
                            'roomRateId':                dc.get('roomRateId') or '',
                            'checkInTime':               dc.get('checkInTime') or '14:00:00',
                            'checkOutTime':              dc.get('checkOutTime') or '12:00:00',
                            'isTripleOccancyAvailable':  dc.get('isTripleOccancyAvailable', False),
                            'isFirmContract':            dc.get('isFirmContract', True),
                            'updateCounter':             0,
                            'spOccupancyRates':          {},
                            'occupancyTypesWithBAR':     {},
                            'startDate':                 night_start,
                            'endDate':                   night_end,
                            'lengthOfStay':              1,
                            'isBaseCategory':            _is_promoted_primary,   # True if promoted from failed primary
                            'hotelUnavailableDays':      [],
                            'contractStatus':            None,
                            'branchMarginPercent':       0,
                            'commission':                0,
                        }
                        alt_objects.append(alt_obj)

                    if alt_objects:
                        _alt_resp = _post(
                            f"Step 10c — {'Add promoted primary' if _is_promoted_primary else f'Add alt hotel {alt_idx+2}'} ({city}: {alt_stub.get('name')})",
                            f"{uatb}/queries/{query_id}/itinerary/routes/{route_id}/hotels",
                            {'actionType': 'add', 'hotels': alt_objects}
                        )
                        if _alt_resp is None:
                            _alt_retry = [{**ao, 'roomType': None} for ao in alt_objects]
                            _alt_resp = _post(
                                f"Step 10c — Alt retry roomType=null ({city}: {alt_stub.get('name')})",
                                f"{uatb}/queries/{query_id}/itinerary/routes/{route_id}/hotels",
                                {'actionType': 'add', 'hotels': _alt_retry}
                            )
                        if _alt_resp is None:
                            _alt_retry2 = [{**ao, 'roomType': None, 'mealPlan': None} for ao in alt_objects]
                            _alt_resp = _post(
                                f"Step 10c — Alt retry roomType+mealPlan=null ({city}: {alt_stub.get('name')})",
                                f"{uatb}/queries/{query_id}/itinerary/routes/{route_id}/hotels",
                                {'actionType': 'add', 'hotels': _alt_retry2}
                            )
                        # If promoted primary was successfully added — update hotel tracking
                        if _is_promoted_primary and _alt_resp is not None:
                            _hotels_by_city[city]   = alt_stub.get('name', '')
                            _hotels_failed_city.pop(city, None)
                        elif not _is_promoted_primary and _alt_resp is not None:
                            _hotels_added_dual.setdefault(city, []).append(alt_stub.get('name', ''))

    # ── STEP 11: PRECONFIGURED TOURS / ACTIVITIES ────────────────────────────
    TOUR_PACK_SEARCH_URL = f"{apib}/tour-packs/search?page=1"

    csv_activities = intent.get('csv_activities') or []
    # For email path, build activity list from intent activities_requested
    email_activities = []
    if not csv_activities and intent.get('activities_requested'):
        for act_name in intent.get('activities_requested', []):
            for city in (intent.get('nightly_split') or []):
                email_activities.append({
                    'city': city[0], 'name': act_name,
                    'date': intent.get('travel_start'), 'pax': intent.get('num_pax',2)
                })
    all_activities = csv_activities or email_activities

    if route_ids and all_activities:
        print(f"\n── STEP 11: PRECONFIGURED TOURS SEARCH ──────────────────────────────")
        _pt_summary = {}  # city → {added: [], skipped: [], no_packs: bool}

        # Group activities by city
        from collections import defaultdict as _dd
        acts_by_city = _dd(list)
        for act in all_activities:
            # Normalise city name
            act_city = act.get('city','')
            for norm_city in (intent.get('cities_detected') or []):
                if act_city.lower() in norm_city.lower() or norm_city.lower() in act_city.lower():
                    act_city = norm_city
                    break
            acts_by_city[act_city].append(act)

        # Build city→check_in lookup from CSV accommodation rows
        _city_checkin = {}
        _accom_rows = intent.get('_accommodation_rows') or []
        for _row in _accom_rows:
            if _row.get('city') and _row.get('check_in') is not None:
                _city_checkin[_row['city']] = str(_row['check_in'])

        for city, acts in acts_by_city.items():
            city_code = CITY_CODE_MAP.get(city)
            route_id  = route_ids.get(str(city_code)) if city_code else None
            if not city_code or not route_id:
                print(f"  ⚠ {city} — no city code or routeId, skipping tour packs")
                continue

            # Use city check-in date from pre-built lookup
            act_date = _city_checkin.get(city)
            if not act_date:
                for a in acts:
                    if a.get('date'):
                        act_date = str(a['date'])
                        break
            if not act_date:
                act_date = t_start

            # Step 11a — Search tour packs for this city
            tp_search_payload = {
                'city':         str(city_code),
                'customer':     account or _default_account(market),
                'sourceMarket': market,
                'adultCount':   num_pax or 2,
                'childCount':   0,
                'startDate':    act_date,
                'endDate':      act_date,
            }
            tp_resp = _post(f"Step 11a — Tour pack search ({city})", TOUR_PACK_SEARCH_URL, tp_search_payload)
            if not tp_resp:
                continue

            tour_packs = (tp_resp.get('data') or {}).get('content') or []
            if not tour_packs:
                # Date-based search returned 0 — TCI requires a date but may not have
                # future dates loaded. Retry with a known-good proxy date (Sep 2026)
                # to retrieve the pack catalogue, then add packs to actual itinerary.
                _PROXY_DATES = ['2026-09-03', '2026-08-15', '2026-10-15']
                for _proxy in _PROXY_DATES:
                    print(f"  ℹ {city} — no packs for {act_date}, retrying with proxy date {_proxy}...")
                    tp_search_proxy = {
                        'city':         str(city_code),
                        'customer':     account or _default_account(market),
                        'sourceMarket': market,
                        'adultCount':   num_pax or 2,
                        'childCount':   0,
                        'startDate':    _proxy,
                        'endDate':      _proxy,
                    }
                    tp_resp2 = _post(f"Step 11a — Tour pack search proxy ({city})",
                                     TOUR_PACK_SEARCH_URL, tp_search_proxy)
                    if tp_resp2:
                        tour_packs = (tp_resp2.get('data') or {}).get('content') or []
                        if tour_packs:
                            print(f"  ✓ {city}: {len(tour_packs)} tour packs found (proxy date {_proxy})")
                            break
                if not tour_packs:
                    print(f"  ⚠ {city} — no tour packs returned")
                    _pt_summary.setdefault(city, {'added': [], 'skipped': [], 'no_packs': True})
            else:
                print(f"  ✓ {city}: {len(tour_packs)} tour packs found")

            # Step 11b — Fuzzy match each CSV activity to a tour pack
            _tp_stop = {'tour','of','the','a','and','in','city','day','half','full',
                        'private','excursion','with','your','driver','guide','vehicle',
                        'new','old','delhi','india','visit','including','morning',
                        'evening','afternoon','explore','exploring'}

            def _tp_score(csv_name, pack_name):
                """Score similarity between CSV activity name and tour pack name.
                Returns (score, matched) — higher score = better match.
                """
                c = csv_name.lower().strip()
                p = pack_name.lower().strip()
                # Exact match — highest priority
                if c == p:
                    return 100, True
                # One contains the other
                if c in p:
                    return 80, True
                if p in c:
                    return 75, True
                c_words = set(c.split()) - _tp_stop
                p_words = set(p.split()) - _tp_stop
                if not c_words or not p_words:
                    return 0, False
                overlap = c_words & p_words
                # Score = overlap ratio weighted by word length
                score = sum(len(w) for w in overlap)
                max_score = sum(len(w) for w in c_words)
                ratio = score / max_score if max_score else 0
                # Require meaningful overlap ratio (>40%) or 2+ words
                if ratio >= 0.4 or len(overlap) >= 2:
                    return int(ratio * 70), True
                # Single strong keyword (6+ chars)
                strong = {w for w in overlap if len(w) >= 6}
                if strong:
                    return int(ratio * 50), True
                return 0, False

            def _tp_match(csv_name, pack_name):
                _, matched = _tp_score(csv_name, pack_name)
                return matched

            _date_slots  = {}  # route_id+date → next available start time slot
            _used_bcs    = set()  # businessCodes already matched in this city

            for act in acts:
                act_name = act.get('name','')
                # Use per-activity date from CSV if available
                if act.get('date'):
                    act_date = str(act['date'])
                matched_pack = None
                # Score all packs and pick the best match
                _scored = []
                for tp in tour_packs:
                    score, matched = _tp_score(act_name, tp.get('name',''))
                    if matched:
                        _scored.append((score, tp))
                _scored.sort(key=lambda x: x[0], reverse=True)
                # First pass: best match not already used
                for score, tp in _scored:
                    if tp.get('businessCode','') not in _used_bcs:
                        matched_pack = tp
                        break
                # Second pass fallback: allow reuse
                if not matched_pack and _scored:
                    matched_pack = _scored[0][1]

                if matched_pack:
                    bc      = matched_pack.get('businessCode')
                    tp_name = matched_pack.get('name','')
                    # Assign staggered start times — TCI allows multiple tours per day
                    # but start times must not overlap with other packs or their components
                    # Assign start time — sunrise tours always 05:00,
                    # evening ceremonies always 18:00,
                    # otherwise 08:00 for first tour of day, 14:00 for second
                    _EVENING_KEYWORDS = ['putting to bed', 'wagah', 'flag ceremony',
                                         'retreat ceremony', 'flag lowering', 'beating retreat']
                    _slot_key   = f"{route_id}_{act_date}"
                    _time_slots = ['08:00', '14:00']
                    _slot_idx   = _date_slots.get(_slot_key, 0)
                    if 'sunrise' in tp_name.lower() or 'sunrise' in act_name.lower():
                        _start_time = '05:00'
                        # sunrise does not consume a slot — doesn't block 08:00/14:00
                    elif any(k in tp_name.lower() or k in act_name.lower() for k in _EVENING_KEYWORDS):
                        _start_time = '18:00'
                        # evening ceremony does not consume a daytime slot
                    elif _slot_idx >= len(_time_slots):
                        print(f"  ⚠ Step 11b — No more time slots for {act_date} — '{tp_name}' skipped")
                        continue
                    else:
                        _start_time = _time_slots[_slot_idx]
                        _date_slots[_slot_key] = _slot_idx + 1
                    print(f"  ✓ Step 11b — Matched: '{act_name}' → '{tp_name}' [{_start_time}] {act_date}")

                    # Step 11c — Add preconfigured tour
                    # Endpoint: POST /queries/{id}/itinerary/routes/{routeId}/preconfigured-tours
                    # on API Gateway (apib), not UAT
                    def _str_field(v):
                        """Flatten object → string code for fields the add API expects as strings."""
                        if isinstance(v, dict):
                            return str(v.get('code') or v.get('id') or '')
                        return v

                    # Fields that must be strings (not objects) at any nesting level
                    _STRING_FIELDS = {'serviceCity', 'city', 'language', 'currency'}

                    def _prepare_component(item, atype=None, slot_time='08:00'):
                        """Deep-copy a component from the search result, updating dates
                        and flattening any object fields that the add API expects as strings.
                        Pass all other fields through unchanged — the search result is
                        the authoritative source for contractId, rateCardId, version etc.
                        """
                        import copy as _copy
                        obj = _copy.deepcopy(item)
                        obj['startDate'] = act_date
                        obj['endDate']   = act_date
                        if not obj.get('startTime'):
                            obj['startTime'] = slot_time
                        if atype:
                            obj['actionType'] = atype
                        # Flatten object → string for known string fields
                        for field in _STRING_FIELDS:
                            if field in obj and isinstance(obj[field], dict):
                                obj[field] = _str_field(obj[field])
                        # Flatten nested startPoint/endPoint city fields
                        for pt in ['startPoint', 'endPoint']:
                            if pt in obj and isinstance(obj.get(pt), dict):
                                if isinstance(obj[pt].get('city'), dict):
                                    obj[pt]['city'] = _str_field(obj[pt]['city'])
                        return obj

                    # Flatten top-level city/language fields from search result
                    pack_city = matched_pack.get('city', {})
                    pack_city_code = pack_city.get('code', str(city_code)) if isinstance(pack_city, dict) else str(city_code)

                    # ── Build component arrays for PT payload ────────────────────────────
                    # Monuments: search city monuments and match to PT definition BCs
                    def _build_pt_monument(stub, dt, tm, ccode):
                        return {
                            'startDate': dt, 'endDate': dt, 'startTime': tm,
                            'businessCode':      stub.get('businessCode',''),
                            'name':              stub.get('name',''),
                            'serviceCity':       str(ccode),
                            'contractCode':      stub.get('contractCode',''),
                            'contractId':        stub.get('contractId',''),
                            'currency': 'INR', 'currencyName': 'Indian Rupee', 'currencySymbol': '₹',
                            'foreignAdultRate':  stub.get('contractStandardRateCardForeignAdult', 0),
                            'foreignChildRate':  stub.get('contractStandardRateCardForeignChild', 0),
                            'foreignInfantRate': stub.get('contractStandardRateCardForeignInfant', 0),
                            'indianAdultRate':   stub.get('contractStandardRateCardIndianAdult', 0),
                            'indianChildRate':   stub.get('contractStandardRateCardIndianChild', 0),
                            'indianInfantRate':  stub.get('contractStandardRateCardIndianInfant', 0),
                            'standardRate':      stub.get('contractStandardRateCardForeignAdult', 0),
                            'standardRateId':    stub.get('contractStandardRateCardId',''),
                            'version':           stub.get('productVersion', 1),
                            'recommendedDuration': stub.get('recommendedDuration', 60),
                            'managementType':    stub.get('managementType','') or 'asi',
                            'availableWeekDays': stub.get('daysAvailable', []),
                            'restrictedDates':   [{'startDate': d, 'endDate': d}
                                                   for d in (stub.get('unavailableDates') or [])],
                            'guestCategorisation': {
                                'adultAge':         {'min': stub.get('adultMinAge',13), 'max': stub.get('adultMaxAge',99)},
                                'areChildAllowed':  stub.get('areChildrenAllowed', True),
                                'areInfantsAllowed': stub.get('areInfantsAllowed', True),
                                'childAge':         {'min': stub.get('childMinAge',3), 'max': stub.get('childMaxAge',12)},
                                'infantAge':        {'min': stub.get('infantMinAge',0), 'max': stub.get('infantMaxAge',2)},
                            },
                            'supplementsId':     stub.get('supplementsId',''),
                            'cameraFee':         stub.get('cameraFee', 0),
                            'videoCameraFee':     stub.get('videoCameraFee', 0),
                            'professionalCameraFee': stub.get('professionalCameraFee', 0),
                            'childRate': 0, 'infantRate': 0, 'isOptional': False,
                            'openingTime': stub.get('openingTime'), 'closingTime': stub.get('closingTime'),
                            'bookingWebsite': None,
                            'supplementsData': [], 'addOnsData': [], 'additions': [],
                            'mandatoryData': [], 'extraChargeIds': [],
                        }

                    def _flatten_point(pt, ccode):
                        # Flatten startPoint/endPoint — city must be string not object
                        if not pt:
                            return {'city': str(ccode), 'country': 'IND', 'state': '',
                                    'addressLine1': '', 'addressLine2': None,
                                    'addressLine3': None, 'pinCode': '', 'latitude': 0, 'longitude': 0}
                        flat = dict(pt)
                        if isinstance(flat.get('city'), dict):
                            flat['city'] = str(flat['city'].get('code') or ccode)
                        if isinstance(flat.get('country'), dict):
                            flat['country'] = flat['country'].get('code','IND')
                        if isinstance(flat.get('state'), dict):
                            flat['state'] = flat['state'].get('code','')
                        return flat

                    # Known fields TCI expects as strings but API returns as objects
                    _OBJ_TO_STR = {
                        'language':   'code',
                        'currency':   'code',
                        'city':       'code',
                        'country':    'code',
                        'state':      'code',
                        'vehicle':    'name',
                    }

                    def _deep_flatten(obj):
                        # Recursively flatten known object fields to strings
                        if isinstance(obj, dict):
                            result = {}
                            for k, v in obj.items():
                                if isinstance(v, dict) and k in _OBJ_TO_STR:
                                    result[k] = v.get(_OBJ_TO_STR[k], str(v))
                                elif isinstance(v, dict):
                                    result[k] = _deep_flatten(v)
                                elif isinstance(v, list):
                                    result[k] = [_deep_flatten(i) for i in v]
                                else:
                                    result[k] = v
                            return result
                        return obj

                    def _build_pt_activity(det, dt, tm, ccode, npax):
                        slabs = det.get('slabs') or []
                        slab  = next((s for s in slabs
                                      if s.get('paxSlab',{}).get('min',1) <= npax <= s.get('paxSlab',{}).get('max',99)),
                                     slabs[0] if slabs else {})
                        cap   = det.get('guestCapacity') or {}
                        vcount = max(1, -(-npax // (cap.get('max') or 2)))
                        # Flatten currency and language — may be object or string
                        _curr = det.get('currency','INR')
                        if isinstance(_curr, dict):
                            _curr = _curr.get('code','INR')
                        _lang = det.get('language','ENG')
                        if isinstance(_lang, dict):
                            _lang = _lang.get('code','ENG')
                        return {
                            'startDate': dt, 'endDate': dt, 'startTime': tm,
                            'businessCode':     det.get('businessCode',''),
                            'name':             det.get('name',''),
                            'serviceCity':      str(ccode),
                            'actionType':       'add',
                            'contractType':     det.get('contractType','gen-mkt'),
                            'currency':         _curr,
                            'currencyName':     det.get('currencyName','Indian Rupee'),
                            'currencySymbol':   det.get('currencySymbol','₹'),
                            'duration':         det.get('duration', 60),
                            'startPoint':       _flatten_point(det.get('startPoint') or det.get('endPoint'), ccode),
                            'endPoint':         _flatten_point(det.get('endPoint') or det.get('startPoint'), ccode),
                            'guestCapacity':    cap or {'min':1,'max':40},
                            'guestCategorisation': det.get('guestCategorisation', {}),
                            'hasLanguageAddOn': False,
                            'isBaseCategory':   det.get('isBaseCategory', True),
                            'isFirmContract':   det.get('isFirmContract', True),
                            'isOptional':       False,
                            'language':         _lang,
                            'languagePrice':    0, 'languagePricingId': '', 'languagePricingType': None,
                            'pricingType':      det.get('pricingType','unit'),
                            'rateCardId':       det.get('rateCardId',''),
                            'rateCode':         det.get('rateCode',''),
                            'standardPrice':    slab.get('standardPrice', 0),
                            'standardPricingId': slab.get('standardPricingId',''),
                            'slabs':            slabs,
                            'supplierCode':     det.get('supplierCode',''),
                            'vehicle':          det.get('vehicle',''),
                            'vehicleCount':     vcount,
                            'version':          det.get('version', 1),
                            'branchMarginPercent': 0, 'margin': 0,
                            'maxPaxRange': npax, 'minPaxRange': npax,
                            'childPrice': 0, 'infantPrice': 0, 'expertPrice': 0,
                            'expertPricingId': '', 'coverCharge': None,
                            'isExpertInclusive': False, 'isExpertRequired': False,
                        }

                    # Build monuments/activities in same simple format as PT search response
                    # (with dates added) — TCI pricing engine rejects full monument objects
                    # Fields TCI expects as strings but API returns as objects
                    _OBJ_FIELDS = {'language': 'code', 'currency': 'code',
                                   'city': 'code', 'country': 'code', 'state': 'code'}

                    def _pt_simple_component(item, dt):
                        import copy as _cp
                        obj = _cp.deepcopy(item)
                        obj['startDate'] = dt
                        obj['endDate']   = dt
                        if 'currency' not in obj or obj['currency'] is None:
                            obj['currency'] = 'INR'
                        # Flatten any object fields that TCI expects as strings
                        for _f, _key in _OBJ_FIELDS.items():
                            if _f in obj and isinstance(obj[_f], dict):
                                obj[_f] = obj[_f].get(_key, '')
                        # Also flatten inside slabs
                        for _slab in (obj.get('slabs') or []):
                            for _f, _key in _OBJ_FIELDS.items():
                                if _f in _slab and isinstance(_slab[_f], dict):
                                    _slab[_f] = _slab[_f].get(_key, '')
                        return obj

                    # Build monument objects with real rates from monument search
                    _pt_mon_bcs = [m.get('businessCode') for m in (matched_pack.get('monuments') or [])
                                   if m.get('businessCode')]
                    _pt_monuments = []
                    if _pt_mon_bcs:
                        try:
                            _mon_resp_pt = _post(
                                f"Step 11 — Monument search ({city})",
                                MON_SEARCH_URL,
                                {
                                 'businessCodes':  _pt_mon_bcs,
                                 'city':           str(city_code),
                                 'customer':       account or _default_account(market),
                                 'sourceMarket':   market,
                                 'adultCount':     num_pax or 2,
                                 'childCount':     0,
                                 'childAge':       [None],
                                 'currencyCode':   CURRENCY,
                                 'nationality':    [],
                                 'from':           act_date,
                                 'to':             act_date,
                                 'priceRangeFrom': None,
                                 'priceRangeTo':   None,
                                 'isRangeExtended': False,
                                 'recommendedDurationFrom': None,
                                 'recommendedDurationTo':   None,
                                 'recommendedDurationRangeExceeded': False,
                                 'bestTimeToVisit': [],
                                }
                            )
                            _all_mon_stubs = (_mon_resp_pt.get('data') or {}).get('content') or [] if _mon_resp_pt else []
                            for _bc in _pt_mon_bcs:
                                # Find full stub with rates from monument search
                                _full_stub = next((m for m in _all_mon_stubs if m.get('businessCode') == _bc), None)
                                if _full_stub:
                                    # Build monument with real rates using correct builder
                                    _pt_monuments.append(_build_pt_monument(_full_stub, act_date, _start_time, city_code))
                                else:
                                    # Fallback to PT stub (no rates)
                                    _pt_monuments.append(_build_pt_monument(
                                        next((m for m in (matched_pack.get('monuments') or [])
                                              if m.get('businessCode') == _bc), {}), act_date, _start_time, city_code))
                        except Exception as _me:
                            # Fallback to PT stubs on error
                            _pt_monuments = [_pt_simple_component(m, act_date)
                                             for m in (matched_pack.get('monuments') or [])
                                             if m.get('businessCode')]
                    _pt_activities = [
                        _pt_simple_component(a, act_date)
                        for a in (matched_pack.get('activities') or [])
                        if a.get('businessCode') and str(a.get('businessCode','')).strip()
                    ]

                    tp_payload = {
                        'tourName':                  matched_pack.get('name',''),
                        'excursionType':             matched_pack.get('excursionType','FULL_DAY'),
                        'serviceCity':               pack_city_code,
                        'tourDescription':           matched_pack.get('description',''),
                        'businessCode':              bc,
                        'startDate':                 act_date,
                        'endDate':                   act_date,
                        'currency':                  'INR',
                        'spCost':                    matched_pack.get('spCost', 0),
                        'startTime':                 _start_time,
                        'adultCount':                num_pax or 2,
                        'childCount':                0,
                        'monuments':                 _pt_monuments,
                        'activities':                [],
                        'localGuideServices':        [],
                        'accompanyingGuideServices': [],
                        'sundryCosts':               [
                            {
                                'amount':      sc.get('amount', 0),
                                'description': sc.get('description', ''),
                                'currency':    (sc.get('currency') or {}).get('code', 'INR'),
                            }
                            for sc in (matched_pack.get('sundryCosts') or [])
                            if sc.get('amount')
                        ],
                    }
                    # Only add guide if PT definition includes a guide
                    _pt_has_guide = bool(matched_pack.get('localGuideServices'))
                    _pt_has_activity = bool(matched_pack.get('activities'))

                    # Read conveyanceSupplementType from pack definition
                    # e.g. "am" → only amRate applies, zero pmRate (and vice versa)
                    _pt_guide_def = (matched_pack.get('localGuideServices') or [{}])[0]
                    _conv_type = (_pt_guide_def.get('conveyanceSupplementType') or '').lower()

                    # Look up guide for this city and inject into localGuideServices
                    # Uses same logic as Step 12 — PT needs guide embedded in payload
                    if _pt_has_guide:
                      try:
                        from datetime import datetime as _dt11
                        _g_bc   = 'GU30001'  # Half Day for PT (excursions are typically half day)
                        _g_name = 'Half Day Guide Service'
                        _g_type = 'local-half-day'
                        _g_search = _post(
                            f"Step 11 — Guide lookup ({city})",
                            f"{apib}/guide-services/search",
                            {'city': city_code, 'from': act_date, 'to': act_date,
                             'customer': account, 'sourceMarket': market,
                             'noOfGuest': num_pax, 'languagePreferences': [],
                             'durations': [], 'priceFrom': None, 'priceTo': None}
                        )
                        _g_stubs = (_g_search.get('data') or {}).get('content') or [] if _g_search else []
                        _g_stub  = next((g for g in _g_stubs if g.get('businessCode') == _g_bc), None)
                        if _g_stub:
                            _g_detail = _get(
                                f"Step 11 — Guide detail ({_g_bc})",
                                f"{apib}/guide-services/{_g_bc}/basic-details?stage=issued"
                            )
                            _g_det  = (_g_detail.get('data') or {}) if _g_detail else {}
                            _g_rcs  = _g_det.get('rateCards') or []
                            _g_rc   = next((rc for rc in _g_rcs if rc.get('contractCode') == _g_stub.get('contractCode')), _g_rcs[0] if _g_rcs else {})
                            _g_lrid = (_g_rc.get('localRate') or {}).get('id') or ''
                            _g_stdr = (_g_rc.get('localRate') or {}).get('standardRates') or []
                            _g_city_slabs = [s for s in _g_stdr if str((s.get('city') or {}).get('code','')) == str(city_code)]
                            _g_pan_slabs  = [s for s in _g_stdr if str((s.get('city') or {}).get('code','')) in ('PAN','pan','')]
                            _g_slabs = _g_city_slabs or _g_pan_slabs or _g_stdr[:1]
                            # Use search stub rate — TCI already pax-matched it via noOfGuest
                            # (same approach as Step 12, avoids manual slab rate mismatch)
                            _g_rate  = _g_stub.get('rate', 0)
                            _g_srid  = _g_stub.get('rateId', '')
                            tp_payload['localGuideServices'] = [{
                                'id':                       None,
                                'perPaxGuide':              num_pax,
                                'localRateId':              _g_lrid,
                                'serviceCity':              str(city_code),
                                'startTime':                _start_time,
                                'startDate':                act_date,
                                'endDate':                  act_date,
                                'actionType':               'add',
                                'serviceId':                None,
                                'businessCode':             _g_bc,
                                'name':                     _g_name,
                                'type':                     _g_type,
                                'rateCardId':               _g_stub.get('rateCardId') or '',
                                'standardRate':             _g_rate,
                                'standardRateId':           _g_srid,
                                'contractCode':             _g_stub.get('contractCode') or '',
                                'contractType':             'gen-mkt',
                                'rateTypeCode':             _g_stub.get('rateTypeCode') or '',
                                'isFirmContract':           _g_stub.get('firmContract', True),
                                'isBaseCategory':           True,
                                'version':                  _g_det.get('version') or _g_stub.get('productVersion') or 1,
                                'guestCapacity':            {'min': 1, 'max': 40},
                                'conveyanceSupplementId':   _g_stub.get('conveyanceSupplementId'),
                                'outstationSupplementId':   None,
                                'outstationSupplementRate': None,
                                'amRate':                   _g_stub.get('amRate', 0) if _conv_type == 'am' else 0,
                                'pmRate':                   _g_stub.get('pmRate', 0) if _conv_type == 'pm' else 0,
                                'hasLanguageAddOn':         False,
                                'language':                 'ENG',
                                'languageName':             '',
                                'languagePrice':            0,
                                'languagePricingId':        '',
                                'slabs':                    _g_slabs,
                                'currency':                 'INR',
                                'currencyName':             'Indian Rupee',
                                'branchMarginPercent':      0,
                                'monuments':                [],
                                'activities':               [],
                            }]
                      except Exception as _ge:
                        pass  # guide lookup failed — proceed without guide in PT

                    # If PT has activities but no guide, fetch and populate activities
                    if _pt_has_activity and not _pt_has_guide:
                        _pt_act_bcs = [a.get('businessCode') for a in (matched_pack.get('activities') or [])
                                       if a.get('businessCode') and str(a.get('businessCode','')).strip()]
                        _built_acts = []
                        for _a_bc in _pt_act_bcs:
                            try:
                                # Use activity search with correct payload (startPointCity not city)
                                _a_search_resp = _post(
                                    f"Step 11 — Activity search ({_a_bc})",
                                    f"{apib}/activities/action/search?page=1&size=50",
                                    {
                                        'businessCodes':      [_a_bc],
                                        'queryId':            query_id,
                                        'languageCode':       'en',
                                        'startPointCity':     str(city_code),
                                        'endPointCity':       str(city_code),
                                        'customer':           account or _default_account(market),
                                        'sourceMarket':       market,
                                        'specialNeedPreference': [],
                                        'adultCount':         num_pax or 2,
                                        'childCount':         0,
                                        'childAge':           [None],
                                        'currencyCode':       CURRENCY,
                                        'filterByTCIScore':   True,
                                        'filterByThemeScore': True,
                                        'minPaxRange':        num_pax or 2,
                                        'maxPaxRange':        num_pax or 2,
                                        'from':               act_date,
                                        'to':                 act_date,
                                        'difficultyLevel':    [],
                                        'themes':             [],
                                        'specializationsList': [],
                                        'priceRangeFrom':     None,
                                        'priceRangeTo':       None,
                                        'isRangeExtended':    False,
                                        'filterByPrice':      False,
                                    }
                                )
                                _a_results = (_a_search_resp.get('data') or {}).get('content') or [] if _a_search_resp else []
                                _a_stub = next((a for a in _a_results if a.get('businessCode') == _a_bc), None)
                                if _a_stub:
                                    # Map search response to PT activity payload
                                    _npax = num_pax or 2
                                    # Find pax-matching slab from standardPricing
                                    _std_pricing = _a_stub.get('standardPricing') or []
                                    _slab = next(
                                        (s for s in _std_pricing
                                         if s.get('minimumGuests',1) <= _npax <= s.get('maximumGuests',99)),
                                        _std_pricing[0] if _std_pricing else {}
                                    )
                                    _curr = _a_stub.get('currency') or {}
                                    if isinstance(_curr, dict): _curr = _curr.get('code','INR')
                                    _built_acts.append({
                                        'startDate':          act_date,
                                        'endDate':            act_date,
                                        'startTime':          _start_time,
                                        'isBaseCategory':     True,
                                        'businessCode':       _a_bc,
                                        'version':            _a_stub.get('productVersion', 1),
                                        'name':               _a_stub.get('name',''),
                                        'guestCategorisation': {
                                            'adultAge':       {'min': _a_stub.get('adultMinAge',13), 'max': _a_stub.get('adultMaxAge',99)},
                                            'areChildAllowed': _a_stub.get('areChildrenAllowed', True),
                                            'childAge':       {'min': _a_stub.get('childMinAge',3), 'max': _a_stub.get('childMaxAge',12)},
                                            'areInfantsAllowed': _a_stub.get('areInfantsAllowed', True),
                                            'infantAge':      {'min': _a_stub.get('infantMinAge',0), 'max': _a_stub.get('infantMaxAge',2)},
                                        },
                                        'guestCapacity':      {'min': _a_stub.get('minimumGuests',1), 'max': _a_stub.get('maximumGuests',40)},
                                        'rateCardId':         _a_stub.get('activityRateCardId',''),
                                        'currency':           _curr,
                                        'currencySymbol':     '₹',
                                        'currencyName':       'Indian Rupee',
                                        'pricingType':        _a_stub.get('pricingType','pax'),
                                        'standardPricingId':  _slab.get('id',''),
                                        'vehicle':            None,
                                        'standardPrice':      _slab.get('standardPrice', 0),
                                        'vehicleCount':       0,
                                        'childPrice':         _slab.get('childPrice', 0),
                                        'infantPrice':        _slab.get('infantPrice', 0),
                                        'margin':             0,
                                        'duration':           _a_stub.get('durationInHours', 60),
                                        'coverCharge':        None,
                                        'supplierCode':       _a_stub.get('supplierCode',''),
                                        'contractType':       _a_stub.get('contractType','gen-mkt'),
                                        'rateCode':           _a_stub.get('rateCode',''),
                                        'startPoint': {
                                            'addressLine1':   _a_stub.get('startPointAddressLine1',''),
                                            'addressLine2':   None,
                                            'addressLine3':   None,
                                            'city':           str(_a_stub.get('startPointCity', city_code)),
                                            'country':        'IND',
                                            'latitude':       _a_stub.get('startPointLatitude',0),
                                            'longitude':      _a_stub.get('startPointLongitude',0),
                                            'pinCode':        _a_stub.get('startPointPinCode',''),
                                            'state':          'IN-DL',
                                        },
                                        'endPoint': {
                                            'addressLine1':   _a_stub.get('endPointAddressLine1',''),
                                            'addressLine2':   None,
                                            'addressLine3':   None,
                                            'city':           str(_a_stub.get('endPointCity', city_code)),
                                            'country':        'IND',
                                            'latitude':       _a_stub.get('endPointLatitude',0),
                                            'longitude':      _a_stub.get('endPointLongitude',0),
                                            'pinCode':        _a_stub.get('endPointPinCode',''),
                                            'state':          'IN-DL',
                                        },
                                        'serviceCity':        str(city_code),
                                        'branchMarginPercent': 0,
                                        'isOptional':         False,
                                        'maxPaxRange':        _npax,
                                        'minPaxRange':        _npax,
                                        'isExpertRequired':   False,
                                        'isExpertInclusive':  False,
                                        'isFirmContract':     _a_stub.get('firmContract', True),
                                        'expertPricingId':    '',
                                        'expertPrice':        0,
                                        'actionType':         'add',
                                        'languagePricingType': None,
                                        'languagePricingId':  '',
                                        'language':           'ENG',
                                        'languagePrice':      0,
                                        'hasLanguageAddOn':   False,
                                        'slabs': [{
                                            'standardPrice':      _slab.get('standardPrice', 0),
                                            'standardPricingId':  _slab.get('id',''),
                                            'childPrice':         _slab.get('childPrice', 0),
                                            'infantPrice':        _slab.get('infantPrice', 0),
                                            'expertPrice':        0,
                                            'expertPricingId':    '',
                                            'languagePrice':      0,
                                            'languagePricingId':  '',
                                            'paxSlab':            {'min': _slab.get('minimumGuests',1), 'max': _slab.get('maximumGuests',99)},
                                            'availableWeekDays':  _a_stub.get('availableDays', []),
                                            'restrictedDates':    _a_stub.get('unavailableDates', []),
                                        }] if _slab else [],
                                    })
                            except Exception as _ae:
                                pass
                        if _built_acts:
                            tp_payload['activities'] = _built_acts

                    add_url = f"{apib}/queries/{query_id}/itinerary/routes/{route_id}/preconfigured-tours"

                    _post(f"Step 11c — Add tour ({city}: {tp_name})", add_url, tp_payload)
                    _used_bcs.add(bc)
                    _components = []
                    if _pt_monuments: _components.append(f"{len(_pt_monuments)} monument{'s' if len(_pt_monuments)>1 else ''}")
                    if tp_payload.get('activities'): _components.append("activity")
                    if tp_payload.get('localGuideServices'): _components.append("guide")
                    _pt_summary.setdefault(city, {'added': [], 'skipped': [], 'no_packs': False})
                    _pt_summary[city]['added'].append((tp_name, _components, act_date if 'act_date' in dir() else ''))
                else:
                    print(f"  ⚠ Step 11b — No match for: '{act_name}' (confirm manually)")
                    _pt_summary.setdefault(city, {'added': [], 'skipped': [], 'no_packs': False})
                    _pt_summary[city]['skipped'].append(act_name)

    elif all_activities and not route_ids:
        print(f"\n  ℹ Step 11 skipped — no routeIds")
    elif not all_activities:
        pass  # No activities to add — silent



    # ── STEP 12: GUIDE SERVICE SEARCH & ADD ──────────────────────────────────────
    GUIDE_SEARCH_URL = f"{apib}/guide-services/search"

    # Language code map for guide language supplements
    _MARKET_TO_LANG = {
        'FRA': 'FRA', 'DEU': 'DEU', 'ITA': 'ITA', 'ESP': 'ESP',
        'RUS': 'RUS', 'JPN': 'JPN',
    }
    _guide_lang = _MARKET_TO_LANG.get(market, 'ENG') if market else 'ENG'
    _is_accompanying = intent.get('language_guide') and any(
        w in (intent.get('language_guide') or '').lower()
        for w in ['accompanying', 'accompan', 'tour escort', 'tour manager', 'tour leader']
    )

    # Fetch monuments added in Step 9 + route dates for guide scheduling
    _added_monuments_per_city = {}
    _route_dates = {}  # city_code -> (startDate, lengthOfStay)
    try:
        import requests as _rq12
        _it12 = _rq12.get(
            f"{apib}/queries/{query_id}/itinerary",
            headers={'Authorization': f'Bearer {api_token}', 'Content-Type': 'application/json'}
        )
        if _it12.status_code == 200:
            for _rt in (_it12.json().get('data') or {}).get('routes', []):
                _cn  = (_rt.get('city') or {}).get('name', '')
                _cc  = str((_rt.get('city') or {}).get('code', ''))
                _ms  = [(m.get('name') or '').lower() for m in (_rt.get('monuments') or [])]
                _sd  = _rt.get('startDate') or ''
                _los = int(_rt.get('lengthOfStay') or 1)
                if _ms:
                    _added_monuments_per_city[_cn] = _ms
                # Get the date monuments were actually added on
                _mon_services = _rt.get('monuments') or []
                _mon_date = (_mon_services[0].get('startDate') or _sd) if _mon_services else _sd
                if _cc:
                    _route_dates[_cc] = (_mon_date, _los)
                # Also store by city name for fallback lookup
                if _cn:
                    _route_dates[_cn] = (_mon_date, _los)
    except Exception:
        pass

    if route_ids and not intent.get('csv_path'):
        print(f"\n── STEP 12: GUIDE SERVICE SEARCH & ADD ──────────────────────────────────────")

        # For accompanying guide — add once per trip (not per city)
        if _is_accompanying:
            _first_city = next(iter(route_ids), None)
            _first_route_id = route_ids.get(_first_city) if _first_city else None
            if _first_route_id:
                _guide_cities = [(_first_city, _first_route_id)]
                print(f"  ℹ Accompanying guide requested — adding once for trip")
            else:
                _guide_cities = []
        else:
            # Local guide per city based on monument count
            _guide_cities = []
            for city, nights_n in nightly_split:
                city_code = str(CITY_CODE_MAP.get(city, ''))
                route_id  = route_ids.get(city_code) or route_ids.get(city_code + '-1')
                if city_code and route_id and nights_n > 0:
                    _guide_cities.append((city_code, route_id))

        for _city_code, _route_id in _guide_cities:
            # Count monuments added for this city to determine guide type
            _mon_count = sum(
                1 for m in (_added_monuments_per_city.get(
                    next((c for c, n in nightly_split if str(CITY_CODE_MAP.get(c,'')) == _city_code), ''),
                    []
                ))
            )
            # Determine guide type
            if _is_accompanying:
                _guide_bc   = 'GU30002'
                _guide_name = 'Accompanying Guide Service'
                _guide_type = 'accompanying'
            elif _mon_count <= 2:
                _guide_bc   = 'GU30001'
                _guide_name = 'Half Day Guide Service'
                _guide_type = 'local-half-day'
            else:
                _guide_bc   = 'GU30000'
                _guide_name = 'Full Day Guide Service'
                _guide_type = 'local-full-day'

            # Step 12a — Use actual monument visit date for this city
            # Monuments visited Day 2 for 2+ night cities, Day 1 for 1-night
            from datetime import datetime as _dt12, timedelta as _td12
            # Try city code first, then city name, then t_start
            _city_name_for_lookup = next(
                (c for c, n in nightly_split if str(CITY_CODE_MAP.get(c,'')) == str(_city_code)),
                None
            )
            _route_sd, _route_los = (
                _route_dates.get(str(_city_code))
                or (_route_dates.get(_city_name_for_lookup) if _city_name_for_lookup else None)
                or (t_start, 1)
            )
            if _route_sd:
                _base = _dt12.strptime(_route_sd, '%Y-%m-%d')
                # Use monument visit date directly (already set from monument startDate)
                _guide_search_date = _base.strftime('%Y-%m-%d')
            else:
                _guide_search_date = t_start
            _search_resp = _post(
                f"Step 12a — Guide search ({_guide_bc})",
                GUIDE_SEARCH_URL,
                {
                    'city':               _city_code,
                    'from':               _guide_search_date,
                    'to':                 _guide_search_date,
                    'customer':           account,
                    'sourceMarket':       market,
                    'noOfGuest':          num_pax,
                    'languagePreferences': [],
                    'durations':          [],
                    'priceFrom':          None,
                    'priceTo':            None,
                }
            )
            if not _search_resp:
                continue

            _guide_stubs = (_search_resp.get('data') or {}).get('content') or []
            _stub = next(
                (g for g in _guide_stubs if g.get('businessCode') == _guide_bc),
                None
            )
            if not _stub:
                print(f"  ⚠ {_guide_bc} not found in search results")
                continue

            # Step 12b — Get guide detail for localRateId and slabs
            _detail_resp = _get(
                f"Step 12b — Guide detail ({_guide_bc})",
                f"{apib}/guide-services/{_guide_bc}/basic-details?stage=issued"
            )
            if not _detail_resp:
                continue

            _detail = _detail_resp.get('data') or {}
            # Find gen-mkt rate card
            _rate_cards = _detail.get('rateCards') or []
            _gen_mkt = next(
                (rc for rc in _rate_cards if rc.get('contractCode') == _stub.get('contractCode')),
                _rate_cards[0] if _rate_cards else {}
            )
            _local_rate_id = (_gen_mkt.get('localRate') or {}).get('id') or ''
            _all_std_rates = (_gen_mkt.get('localRate') or {}).get('standardRates') or []

            # Filter to city-specific slabs first, fall back to pan-country
            def _slab_city(s):
                c = s.get('city') or {}
                return str(c.get('code', '') or '')
            _city_slabs = [s for s in _all_std_rates if _slab_city(s) == str(_city_code)]
            _pan_slabs  = [s for s in _all_std_rates if _slab_city(s) in ('PAN', 'pan', '')]
            # Also try without suffix (e.g. "291-1" → "291")
            _base_code  = str(_city_code).split('-')[0]
            _base_slabs = [s for s in _all_std_rates if _slab_city(s) == _base_code] if _base_code != str(_city_code) else []
            # Final fallback: use all slabs sorted by pax range match
            _std_rates  = _city_slabs or _base_slabs or _pan_slabs or _all_std_rates[:3]

            # Build slabs — languagePricingId uses standardRateId of each slab (as per working payload)
            _slabs = []
            for _sr in _std_rates:
                _slab = {
                    'standardRatesId':   _sr.get('id') or '',
                    'standardRate':      _sr.get('rate') or 0,
                    'languagePricingId': _sr.get('id') or '',  # each slab's own ID (not lang supplement)
                    'languagePrice':     None,
                    'type':              'standard',
                    'paxSlabs': {
                        'min': (_sr.get('guestSlab') or {}).get('minGuests') or 1,
                        'max': (_sr.get('guestSlab') or {}).get('maxGuests') or 40,
                    }
                }
                _slabs.append(_slab)

            if not _slabs:
                # Accompanying guides don't use slabs — handled separately below
                # For local guides only: flag as failed
                if not _is_accompanying:
                    print(f"  ⚠ No slabs from detail API for {_guide_bc} city {_city_code}")
                    _city_for_failed = next(
                        (c for c, n in nightly_split if str(CITY_CODE_MAP.get(c,'')) == str(_city_code)),
                        _city_code
                    )
                    _email_guides_failed[_city_for_failed] = f"{_guide_name} · no pricing slabs in TCI"
                    continue

            # Language supplement — use market language if available
            _MARKET_LANG = {'FRA':'FRA','DEU':'DEU','ITA':'ITA','ESP':'ESP','RUS':'RUS'}
            _guide_lang_out = _MARKET_LANG.get(market, 'ENG')
            _lang_pricing   = _stub.get('languagePricing') or []
            _lang_match     = next(
                (lp for lp in _lang_pricing if lp.get('code') == _guide_lang_out),
                None
            )
            _has_lang_addon  = bool(_lang_match)
            _lang_price      = _lang_match.get('rate', 0) if _lang_match else 0
            _lang_pricing_id = _lang_match.get('id', '') if _lang_match else (_stub.get('rateId') or '')

            # Step 12c — Add guide
            # Accompanying guides use a different payload structure — no slabs, trip-wide dates
            if _is_accompanying:
                # Get trip end date and length
                from datetime import datetime as _dt_acc, timedelta as _td_acc
                _trip_start_dt = _dt_acc.strptime(t_start, '%Y-%m-%d')
                _trip_nights   = sum(n for _, n in nightly_split if n > 0)
                _trip_end_dt   = _trip_start_dt + _td_acc(days=_trip_nights)
                _trip_end_str  = _trip_end_dt.strftime('%Y-%m-%d')
                # First city route ID
                _first_city_code = str(list(route_ids.keys())[0]) if route_ids else str(_city_code)
                _first_route_id  = route_ids.get(_first_city_code) or _route_id
                _add_url = f"{apib}/queries/{query_id}/itinerary/acc-guide-services"
                _guide_payload = {
                    'guideServices': [{
                        'id':                       None,
                        'serviceCity':              _city_code,
                        'businessCode':             _guide_bc,
                        'startCityRouteId':         _first_route_id,
                        'name':                     _guide_name,
                        'type':                     _guide_type,
                        'rateCardId':               _stub.get('rateCardId') or '',
                        'currency':                 CURRENCY,
                        'standardRate':             _stub.get('rate') or 0,
                        'standardRateId':           _stub.get('rateId') or '',
                        'contractType':             _stub.get('contractType') or 'gen-mkt',
                        'contractCode':             _stub.get('contractCode') or '',
                        'isFirmContract':           _stub.get('firmContract', True),
                        'rateTypeCode':             _stub.get('rateTypeCode') or '',
                        'branchMarginPercent':      0,
                        'guestCapacity':            {
                            'min': _stub.get('guestSlabMinGuests') or 1,
                            'max': _stub.get('guestSlabMaxGuests') or 40,
                        },
                        'actionType':               'add',
                        'version':                  _stub.get('productVersion') or 1,
                        'perPaxGuide':              num_pax,
                        'serviceId':                None,
                        'startDate':                t_start,
                        'endDate':                  _trip_end_str,
                        'lengthOfStay':             _trip_nights,
                        'startCity':                _city_code,
                        'accompanyingRateId':       _stub.get('rateId') or '',
                        'accommodationType':        'with-acc',
                        'isServingExcursionCity':   False,
                        'isBaseCategory':           True,
                        'monuments':                [],
                        'activities':               [],
                        'hasLanguageAddOn':         _has_lang_addon,
                        'languagePricingId':        _lang_pricing_id,
                        'language':                 _guide_lang_out if _has_lang_addon else 'ENG',
                        'languagePrice':            _lang_price,
                    }]
                }
            else:
                _add_url = f"{apib}/queries/{query_id}/itinerary/routes/{_route_id}/loc-guide-services"
                _guide_payload = {
                    'guideServices': [{
                        'id':                       None,
                        'perPaxGuide':              num_pax,
                        'localRateId':              _local_rate_id,
                        'serviceCity':              _city_code,
                        'startTime':                '09:00',
                        'startDate':                _guide_search_date,
                        'endDate':                  _guide_search_date,
                        'businessCode':             _guide_bc,
                        'name':                     _guide_name,
                        'type':                     _guide_type,
                        'rateCardId':               _stub.get('rateCardId') or '',
                        'serviceId':                None,
                        'currency':                 CURRENCY,
                        'standardRate':             _stub.get('rate') or 0,
                        'standardRateId':           _stub.get('rateId') or '',
                        'contractType':             'gen-mkt',
                        'contractCode':             _stub.get('contractCode') or '',
                        'rateTypeCode':             _stub.get('rateTypeCode') or '',
                        'isFirmContract':           _stub.get('firmContract', True),
                        'branchMarginPercent':      0,
                        'guestCapacity':            {'min': 1, 'max': 40},
                        'outstationSupplementId':   None,
                        'outstationSupplementRate': None,
                        'actionType':               'add',
                        'monuments':               [],
                        'activities':              [],
                        'conveyanceSupplementId':  _stub.get('conveyanceSupplementId'),
                        'amRate':                  _stub.get('amRate', 0),
                        'pmRate':                  0,
                        'version':                 _detail.get('version') or _stub.get('productVersion') or 1,
                        'hasLanguageAddOn':        _has_lang_addon,
                        'languagePrice':           _lang_price,
                        'languagePricingId':       _lang_pricing_id,
                        'language':                _guide_lang_out if _has_lang_addon else 'ENG',
                        'isBaseCategory':          True,
                        'slabs':                   _slabs,
                    }]
                }
            _add_resp = _post(
                f"Step 12c — Add guide ({_guide_bc} · {_guide_name})",
                _add_url,
                _guide_payload
            )
            if _add_resp is not None:
                print(f"  ✓ Guide added: {_guide_name} · {_guide_bc} · INR {_stub.get('rate',0):,}")
                _city_for_guide = next(
                    (c for c, n in nightly_split if str(CITY_CODE_MAP.get(c,'')) == str(_city_code)),
                    _city_code
                )
                _email_guides_by_city[_city_for_guide] = _guide_name
    elif intent.get('csv_path'):
        # ── STEP 11 SUMMARY ──────────────────────────────────────────────────────
        print(f"\n── STEP 11 SUMMARY ───────────────────────────────────────────────────────")
        # Collect all cities in route order
        # ANSI colour — Green = done, Red = needs manual action
        _GRN  = '\033[92m'   # green — added successfully
        _RED  = '\033[91m'   # red   — needs manual action (missing, not added, confirm)
        _DIM  = '\033[2m'    # dim   — informational only
        _RST  = '\033[0m'    # reset
        _CYN  = _GRN         # tours added = same green as hotels
        _YLW  = _RED         # warnings = red (needs action)

        # ── Build hotel date lookup from _accommodation_rows ────────────────
        _hotel_dates = {}  # city → (check_in_str, check_out_str)
        for _ar in (intent.get('_accommodation_rows') or []):
            _ar_city = _ar.get('city','')
            if _ar_city and _ar_city not in _hotel_dates:
                _ci = _ar.get('check_in')
                _co = _ar.get('check_out')
                if _ci and _co:
                    def _fmt_date(d):
                        from datetime import date as _dt
                        if isinstance(d, str):
                            try: d = _dt.fromisoformat(d)
                            except: return str(d)
                        return d.strftime('%d %b') if d else ''
                    _hotel_dates[_ar_city] = (_fmt_date(_ci), _fmt_date(_co))

        _DATE_W = 17  # width of date column

        _pt_summary = locals().get('_pt_summary') or {}
        _all_route_cities = [c for c, n in nightly_split if n > 0 or (_pt_summary or {}).get(c)]
        for _city_idx, _s_city in enumerate(_all_route_cities):
            if _city_idx > 0:
                print(f"  {_DIM}{'─'*69}{_RST}")
            _first = True
            def _line(colour, icon, text, tag='', date_str=''):
                nonlocal _first
                _prefix = f"  {_s_city:20}" if _first else f"  {'':20}"
                _first = False
                _tag = f"  {_DIM}({tag}){_RST}" if tag else ''
                _dcol = f"{_DIM}{date_str:<{_DATE_W}}{_RST}" if date_str else f"{'':>{_DATE_W}} "
                print(f"{_prefix} {_dcol}  {colour}{icon} {text}{_RST}{_tag}")

            # Hotels
            _added_name    = _hotels_by_city.get(_s_city, '')
            _failed_name   = _hotels_failed_city.get(_s_city, '')
            _requested_name = _hotels_primary_requested.get(_s_city, '')

            if _added_name:
                # Check if what was added differs from what was (P) requested
                # Use fuzzy check: strip punctuation/brackets, check word overlap
                import re as _re_h
                def _hotel_match(a, b):
                    _clean = lambda s: _re_h.sub(r'[^a-z0-9 ]', '', s.lower()).strip()
                    _words = lambda s: set(_clean(s).split()) - {'the','hotel','a','and','&','by','at'}
                    _a, _b = _words(a), _words(b)
                    if not _a or not _b: return True
                    return len(_a & _b) / max(len(_a), len(_b)) >= 0.5

                _HOUSEBOAT_KW = {'houseboat','house boat','kettuvallam','rice boat','backwater'}
                _is_houseboat_req = any(kw in _requested_name.lower() for kw in _HOUSEBOAT_KW)
                _hd = _hotel_dates.get(_s_city, ('',''))
                _hdate = f"{_hd[0]} → {_hd[1]}" if _hd[0] and _hd[1] else ''
                if (_requested_name and not _hotel_match(_requested_name, _added_name)
                        and not _is_houseboat_req):
                    # (P) primary not in TCI — show it as missing, then show what was added
                    _line(_RED, '✗', _requested_name.title(), 'hotel · not in TCI — (P) primary requested')
                    _line(_GRN, '✓', _added_name, 'hotel · added as best available', date_str=_hdate)
                else:
                    _line(_GRN, '✓', _added_name, 'hotel', date_str=_hdate)
            elif _failed_name:
                _line(_RED, '✗', _failed_name, 'hotel · NOT added — add manually in TCI')
            # Successfully added dual options
            for _alt in _hotels_added_dual.get(_s_city, []):
                _line(_GRN, '✓', _alt, 'hotel · alt option')
            # Missing dual options
            for _miss in _hotels_missing_dual.get(_s_city, []):
                _line(_RED, '✗', _miss, 'hotel · not in TCI catalogue')
            # Tours/activities
            _pt = (_pt_summary or {}).get(_s_city, {})
            if _pt.get('no_packs'):
                _line(_YLW, '⚠', 'No tour packs in TCI — add manually', '')
            else:
                for _tp_entry in _pt.get('added', []):
                    _tp_name  = _tp_entry[0]
                    _comps    = _tp_entry[1]
                    _tp_date  = _tp_entry[2] if len(_tp_entry) > 2 else ''
                    _comp_str = ', '.join(_comps) if _comps else ''
                    # Format single date DD Mon
                    _tdate = ''
                    if _tp_date:
                        try:
                            from datetime import date as _dtc
                            _td = _dtc.fromisoformat(str(_tp_date))
                            _tdate = _td.strftime('%d %b')
                        except: _tdate = str(_tp_date)[:6]
                    _line(_CYN, '✓', _tp_name, _comp_str, date_str=_tdate)
                for _sk in _pt.get('skipped', []):
                    _line(_YLW, '⚠', _sk, 'confirm manually')
            # Cities with hotel only (no activities in CSV)
            if _s_city not in (_pt_summary or {}) and _s_city in _hotels_by_city:
                _line(_DIM, '—', 'No activities specified by agent', '')

        # ── Unmatched cities (e.g. Mumbai — in activities but no TCI city code) ──
        _route_cities_set = set(c for c, n in nightly_split)
        _csv_acts = intent.get('csv_activities') or []
        _csv_xfers = intent.get('csv_transfers') or []
        _csv_rails = intent.get('csv_rail') or []

        # Find cities mentioned in activities but not in route
        # Also exclude city label variants that map to route cities (e.g. Delhi→New Delhi)
        _label_to_route = {}
        for _rc in _route_cities_set:
            _label_to_route[_rc.lower()] = _rc
            # Map common short labels
            if 'new delhi' in _rc.lower(): _label_to_route['delhi'] = _rc
            if 'mumbai' in _rc.lower(): _label_to_route['mumbai'] = _rc
            if 'goa' in _rc.lower(): _label_to_route['goa'] = _rc; _label_to_route['south beaches, goa'] = _rc

        _unrouted_cities = {}
        for _act in _csv_acts:
            _act_city = _act.get('city', '').strip()
            if not _act_city:
                continue
            # Skip if this city label maps to a city already in the route
            if _act_city.lower() in _label_to_route:
                continue
            if _act_city not in _route_cities_set:
                _unrouted_cities.setdefault(_act_city, []).append(_act.get('name', _act_city))

        if _unrouted_cities or _csv_rails or _csv_xfers:
            print(f"\n  ── Services not pushed to TCI (add manually) ────────────────")
            # Unrouted cities
            for _ucity, _uacts in _unrouted_cities.items():
                print(f"  {_RED}⚠ {_ucity:20}{_RST}  No TCI city code — add manually")
                for _ua in _uacts:
                    print(f"  {'':22}  {_DIM}· {_ua}{_RST}")
            # Rail segments
            if _csv_rails:
                print(f"  {_RED}⚠ Rail segments ({len(_csv_rails)}){_RST}  — confirm manually in TCI")
                for _rail in _csv_rails:
                    if isinstance(_rail, str):
                        _rail_desc = _rail
                    elif isinstance(_rail, dict):
                        _rail_desc = _rail.get('details', _rail.get('description', str(_rail)))
                    else:
                        _rail_desc = str(_rail)
                    print(f"  {'':22}  {_DIM}· {_rail_desc}{_RST}")
            # Transfers
            if _csv_xfers:
                print(f"  {_DIM}ℹ {len(_csv_xfers)} transfers parsed — not pushed to TCI (add if required){_RST}")

        print(f"\n  ℹ Step 12 skipped — CSV path: guides not auto-added for agent-specified itineraries")
    else:
        print(f"\n  ℹ Step 12 skipped — no routeIds")


    # ── STEP 13: ACTIVITY SEARCH & ADD ───────────────────────────────────────────

    # Monument → paired activity keywords (from production data analysis)
    MONUMENT_ACTIVITY_PAIRS = {
        # Delhi monuments → Cycle Rickshaw Ride
        'humayuns tomb':          'cycle rickshaw',
        'jama masjid':            'cycle rickshaw',
        'qutub minar':            'cycle rickshaw',
        'india gate':             'cycle rickshaw',
        'red fort':               'cycle rickshaw',
        'gurudwara bangla sahib': 'cycle rickshaw',
        'chandni chowk':          'cycle rickshaw',
        # Jaipur monuments → Jeep Ride at Amber Fort
        'amber fort':             'jeep ride at amber fort',
        'city palace':            'jeep ride at amber fort',
        'jantar mantar':          'jeep ride at amber fort',
        # Kerala
        'periyar':                'boat ride at periyar',
        # Varanasi
        'dashashwamedh':          'motor boat',
        'ghats':                  'motor boat',
    }

    # _added_monuments_per_city already fetched in Step 12 above

    ACT_SEARCH_URL = f"{apib}/activities/action/search?page=1&size=50"

    # Get DB-recommended activities per city (reuse engine output)
    _act_db_recs = {}  # city_name -> [activity_name, ...]
    try:
        import sqlite3 as _sql
        _con_a = _sql.connect(db_path)
        _cur_a = _con_a.cursor()
        for city_name, nights_n in nightly_split:
            if nights_n == 0:
                continue
            _cur_a.execute('''
                SELECT service_name, COUNT(*) as cnt
                FROM services
                WHERE city_name = ? AND record_type = 'Activity'
                AND source_market = ? AND service_name != ''
                GROUP BY service_name ORDER BY cnt DESC LIMIT 5
            ''', (city_name, market))
            rows = _cur_a.fetchall()
            if not rows:
                _cur_a.execute('''
                    SELECT service_name, COUNT(*) as cnt
                    FROM services
                    WHERE city_name = ? AND record_type = 'Activity'
                    AND service_name != ''
                    GROUP BY service_name ORDER BY cnt DESC LIMIT 5
                ''', (city_name,))
                rows = _cur_a.fetchall()
            _act_db_recs[city_name] = [r[0].lower() for r in (rows or [])]
        _con_a.close()
    except Exception as _e:
        print(f"  ⚠ Activity DB lookup failed: {_e}")

    if route_ids and not intent.get('csv_path'):
        print(f"\n── STEP 13: ACTIVITY SEARCH & ADD ───────────────────────────────────────────")

        # Track used activity times per city per date
        _act_times_used = defaultdict(set)  # (city, date) -> {time_str}
        _ACT_START_TIMES = ['09:00', '10:00', '11:00', '12:00', '14:00', '15:00', '16:00']

        for city, nights_n in nightly_split:
            city_code = str(CITY_CODE_MAP.get(city, ''))
            route_id  = route_ids.get(city_code) or route_ids.get(city_code + '-1')
            if not city_code or not route_id or nights_n == 0:
                continue

            rec_names = _act_db_recs.get(city, [])

            # Use first full touring day
            city_idx  = [c for c, n in nightly_split].index(city)
            nights_before = sum(n for _, n in nightly_split[:city_idx])
            _t_start_dt = _dt.strptime(t_start, '%Y-%m-%d')
            visit_date = (_t_start_dt + _td(days=nights_before + 1)).strftime('%Y-%m-%d')

            # Step 13a — Search activities
            act_search_payload = {
                'minPaxRange':            num_pax,
                'maxPaxRange':            num_pax,
                'queryId':                query_id,
                'languageCode':           'en',
                'startPointCity':         city_code,
                'endPointCity':           city_code,
                'customer':               account,
                'sourceMarket':           market,
                'specialNeedPreference':  [],
                'adultCount':             num_pax,
                'childCount':             0,
                'childAge':               [None],
                'currencyCode':           CURRENCY,
                'filterByTCIScore':       True,
                'filterByThemeScore':     True,
                'from':                   visit_date,
                'to':                     visit_date,
                'difficultyLevel':        [],
                'themes':                 [],
                'specializationsList':    [],
                'priceRangeFrom':         None,
                'priceRangeTo':           None,
                'isRangeExtended':        False,
                'filterByPrice':          False,
            }

            act_search_resp = _post(
                f"Step 13a — Search activities ({city})",
                ACT_SEARCH_URL,
                act_search_payload
            )

            if not act_search_resp:
                continue

            raw_acts = (act_search_resp.get('data') or {}).get('content') or []
            # Suppress venue/logistics/billing activities
            _ACT_NAME_SUPPRESS = ['venue', 'wash & change', 'porterage', 'supplement over',
                                   'charges extra', 'tipping', 'toll', 'parking',
                                   'mineral water', 'drinking water', 'luggage',
                                   'conveyance charges', 'transfer', 'transferring',
                                   'autorickshaw transfers', 'tuk tuk conveyance',
                                   'chair for', 'life jacket', 'entrance fee',
                                   'guide charges', 'guide fee', 'outsourced']
            raw_acts = [a for a in raw_acts
                        if not any(kw in a.get('name','').lower() for kw in _ACT_NAME_SUPPRESS)]
            if not raw_acts:
                print(f"  ⚠ {city} — no activities returned")
                continue

            # Fuzzy match against DB recommendations
            def _fuzzy_act_match(api_name, db_names):
                a = api_name.lower()
                a_core = re.sub(r'\s*\(.*?\)', '', a).strip()
                for db in db_names:
                    d = re.sub(r'\s*\(.*?\)', '', db.lower()).strip()
                    if a_core in d or d in a_core:
                        return True
                    a_words = set(a_core.split()) - {'the','a','of','and','at','in','by','&','for'}
                    d_words = set(d.split()) - {'the','a','of','and','at','in','by','&','for'}
                    if len(a_words & d_words) >= 2:
                        return True
                return False

            matched_acts = []
            unmatched_acts = []
            for a in raw_acts:
                aname = a.get('name', '')
                if _fuzzy_act_match(aname, rec_names) if rec_names else True:
                    matched_acts.append(a)
                else:
                    unmatched_acts.append(a)

            # Check if any monument in this city has a paired activity
            paired_keywords = []
            city_monuments = _added_monuments_per_city.get(city, [])
            for mon_name in city_monuments:
                for mon_key, act_kw in MONUMENT_ACTIVITY_PAIRS.items():
                    if mon_key in mon_name:
                        paired_keywords.append(act_kw)
                        break

            # Prioritise: paired (non-DKC) > paired (DKC) > matched non-DKC > matched DKC > unmatched non-DKC
            def _is_dkc(act):
                return '[a dkc experience]' in act.get('name','').lower()

            paired_acts  = [a for a in raw_acts
                            if any(kw in a.get('name','').lower() for kw in paired_keywords)]
            remaining    = [a for a in raw_acts if a not in paired_acts]
            matched      = [a for a in remaining if _fuzzy_act_match(a.get('name',''), rec_names)]
            unmatched    = [a for a in remaining if not _fuzzy_act_match(a.get('name',''), rec_names)]

            # Sort each group: non-DKC first
            def _non_dkc_first(lst): return sorted(lst, key=lambda a: 1 if _is_dkc(a) else 0)

            ordered = (
                _non_dkc_first(paired_acts) +
                _non_dkc_first(matched) +
                _non_dkc_first(unmatched)
            )

            # Only 1 activity per city — monument-complementing, strictly non-DKC preferred
            _non_dkc_candidates = [a for a in ordered if not _is_dkc(a)]
            candidates = _non_dkc_candidates[:1] if _non_dkc_candidates else ordered[:1]
            if paired_acts:
                print(f"  💡 {city}: paired activity from monument — {candidates[0].get('name','')}")
            if not candidates:
                print(f"  ⚠ {city} — no matching activities")
                continue

            added_acts = []
            for act_stub in candidates:
                bc   = act_stub.get('businessCode') or act_stub.get('business_code')
                name = act_stub.get('name', '')
                if not bc:
                    continue

                # Step 13b — Build payload directly from search stub (detail API returns empty slabs)
                print(f"  ✓ Step 13b — Activity detail ({city}: {name})")

                # All pricing data is in the search result
                pricing_type   = act_stub.get('pricingType') or 'pax'
                rate_card_id   = act_stub.get('activityRateCardId') or ''
                rate_code      = act_stub.get('rateCode') or ''
                is_firm        = act_stub.get('firmContract', False)
                supplier_code  = act_stub.get('supplierCode') or ''
                version        = act_stub.get('productVersion') or 1
                duration_hrs   = act_stub.get('durationInHours') or 60
                duration_min   = int(duration_hrs * 60) if duration_hrs < 24 else 60

                # standardPricing array = slabs
                all_slabs_raw = act_stub.get('standardPricing') or []

                # Build slabs in TCI payload format
                slabs = []
                for s in all_slabs_raw:
                    slab = {
                        'standardPrice':     s.get('standardPrice') or 0,
                        'standardPricingId': s.get('id') or '',
                        'childPrice':        s.get('childPrice') or 0,
                        'infantPrice':       s.get('infantPrice') or 0,
                        'expertPrice':       0,
                        'expertPricingId':   '',
                        'languagePrice':     0,
                        'languagePricingId': '',
                        'paxSlab': {
                            'min': s.get('minimumGuests') or 1,
                            'max': s.get('maximumGuests') or 40,
                        },
                        'branchMarginPercent': 0,
                        'availableWeekDays':   act_stub.get('availableDays') or [],
                        'restrictedDates':     [],
                    }
                    # For unit-based pricing, include vehicle field
                    if s.get('vehicle'):
                        slab['vehicle'] = s.get('vehicle')
                    slabs.append(slab)

                if not slabs:
                    print(f"  ⚠ {city} — no slabs for {name}, skipping")
                    continue

                # Find matching slab for pax count
                matching_slab = next(
                    (s for s in all_slabs_raw
                     if (s.get('minimumGuests') or 0) <= num_pax <= (s.get('maximumGuests') or 9999)),
                    all_slabs_raw[0]
                )
                std_pricing_id = act_stub.get('standardPricingId') or matching_slab.get('id') or ''
                std_price      = act_stub.get('standardPrice') or matching_slab.get('standardPrice') or 0
                child_price    = act_stub.get('standardPricingChildPrice') or matching_slab.get('childPrice') or 0
                infant_price   = act_stub.get('standardPricingInfantPrice') or matching_slab.get('infantPrice') or 0

                # Guest categorisation from search stub
                guest_cat = {
                    'adultAge':       {'max': act_stub.get('adultMaxAge') or 99, 'min': act_stub.get('adultMinAge') or 13},
                    'areChildAllowed': act_stub.get('areChildrenAllowed', True),
                    'childsAge':      {'max': act_stub.get('childMaxAge') or 12, 'min': act_stub.get('childMinAge') or 3},
                    'areInfantsAllowed': act_stub.get('areInfantsAllowed', True),
                    'infantAge':      {'max': act_stub.get('infantMaxAge') or 2, 'min': act_stub.get('infantMinAge') or 0},
                }

                # Pick start time — use first available time slot or default
                time_slots = act_stub.get('timeSlots') or []
                used = _act_times_used[(city, visit_date)]
                # Try activity's own time slots first
                act_time = next((t[:5] for t in time_slots if t[:5] not in used), None)
                if not act_time:
                    act_time = next((t for t in _ACT_START_TIMES if t not in used), '12:00')
                _act_times_used[(city, visit_date)].add(act_time)

                act_payload = [{
                    'startDate':           visit_date,
                    'endDate':             visit_date,
                    'startTime':           act_time + ':00' if len(act_time) == 5 else act_time,
                    'isBaseCategory':      True,
                    'businessCode':        bc,
                    'version':             version,
                    'name':                name,
                    'guestCategorisation': guest_cat,
                    'guestCapacity': {
                        'min': act_stub.get('minimumGuests') or 1,
                        'max': act_stub.get('maximumGuests') or 40,
                    },
                    'rateCardId':          rate_card_id,
                    'currency':            CURRENCY,
                    'currencySymbol':      '₹',
                    'pricingType':         pricing_type,
                    'standardPricingId':   std_pricing_id,
                    'vehicle':             matching_slab.get('vehicle') or act_stub.get('standardPricingVehicle') or None,
                    'standardPrice':       std_price,
                    'vehicleCount':        (
                                              # For unit pricing, calculate vehicles needed
                                              max(1, -(-num_pax // max(matching_slab.get('maximumGuests') or 1, 1)))
                                              if pricing_type == 'unit' else None
                                          ),
                    'childPrice':          child_price,
                    'infantPrice':         infant_price,
                    'slabs':               slabs,
                    'supplement':          None,
                    'actionType':          'add',
                    'maxPaxRange':         num_pax,
                    'minPaxRange':         num_pax,
                    'duration':            duration_min,
                    'supplierCode':        supplier_code,
                    'contractType':        'gen-mkt',
                    'rateCode':            rate_code,
                    'serviceCity':         city_code,
                    'isOptional':          False,
                    'expertPricingUnit':   None,
                    'isExpertRequired':    False,
                    'isExpertInclusive':   False,
                    'expertPricingId':     None,
                    'expertPrice':         0,
                    'languagePricingType': 'pax',
                    'languagePricingId':   '',
                    'language':            'ENG',
                    'languagePrice':       0,
                    'hasLanguageAddOn':    False,
                    'isFirmContract':      is_firm,
                    'suggestedTime':       act_time + ':00' if len(act_time) == 5 else act_time,
                }]

                add_url = f"{apib}/queries/{query_id}/itinerary/routes/{route_id}/activities"
                # TCI expects object wrapper, not bare array
                add_resp = _post(
                    f"Step 13c — Add activity ({city}: {name})",
                    add_url,
                    {'activities': act_payload}
                )
                if add_resp is not None:
                    added_acts.append(name)

            if added_acts:
                print(f"  ✓ Step 13c — Added activities ({city}: {len(added_acts)})")
                for a in added_acts:
                    print(f"     {a}")
                _email_acts_by_city[city] = added_acts
    else:
        print(f"\n  ℹ Step 13 skipped — no routeIds")


    # ── STEP 14: VALIDATE ITINERARY ──────────────────────────────────────────
    if query_id:
        try:
            import requests as _rq14
            _val_r = _rq14.get(
                f"{uatb}/queries/{query_id}/itinerary/validate",
                headers={'Authorization': f'Bearer {api_token}'},
                timeout=30
            )
            if _val_r.status_code == 200:
                _val_issues = (_val_r.json().get('data') or [])
                # Filter to real issues — ignore 'ok' and 'version' which are informational
                _real_issues = []
                for _item in (_val_issues or []):
                    for _vi in (_item.get('validationInfo') or []):
                        if _vi.get('key','') not in ('ok', 'version', ''):
                            _real_issues.append((_item, _vi))
                if not _real_issues:
                    print(f"\n  ✓ Step 14 — Itinerary valid (no issues)")
                else:
                    print(f"\n── STEP 14: VALIDATION ISSUES ───────────────────────────────────────────")
                    for _item, _vi in _real_issues:
                        _city = _item.get('city','')
                        _name = _item.get('name','')
                        print(f"  ⚠ {_city} — {_name}: {_vi.get('key','')} — {_vi.get('value','')}")
            else:
                print(f"  ⚠ Step 14 — Validate returned {_val_r.status_code}")
        except Exception as _ve:
            print(f"  ⚠ Step 14 — Validate skipped: {str(_ve)[:80]}")

    # ── EMAIL/PDF MODE STEP 11 SUMMARY ──────────────────────────────────────
    if not intent.get('csv_path') and route_ids:
        # _pt_summary may not be defined if no activities were found
        try:
            _pt_summary
        except NameError:
            _pt_summary = {}
        print(f"\n── STEP 11 SUMMARY ───────────────────────────────────────────────────────")
        _GRN  = '\033[92m'; _RED = '\033[91m'; _DIM = '\033[2m'
        _RST  = '\033[0m';  _CYN = _GRN;       _YLW = _RED
        _DATE_W = 17

        # ── Build derived date lookup from travel_start + cumulative nights ──
        _e_hotel_dates = {}  # city → (check_in_str, check_out_str)
        _e_dates_estimated = False
        try:
            from datetime import date as _edt, timedelta as _etd
            _ts_raw = intent.get('travel_start') or intent.get('travel_start_date') or ''
            # Fallback: derive from dates_raw tuples e.g. [('02','nov'),('20','nov')]
            if not _ts_raw or str(_ts_raw).lower() in ('none', ''):
                _dr = intent.get('dates_raw') or []
                if _dr:
                    _month_map = {'jan':1,'feb':2,'mar':3,'apr':4,'may':5,'jun':6,
                                  'jul':7,'aug':8,'sep':9,'oct':10,'nov':11,'dec':12}
                    _day, _mon = str(_dr[0][0]), str(_dr[0][1]).lower()[:3]
                    _mon_n = _month_map.get(_mon, 0)
                    if _mon_n:
                        import datetime as _dtt
                        _today = _dtt.date.today()
                        _yr = _today.year if _mon_n >= _today.month else _today.year + 1
                        _ts_raw = f"{_yr}-{_mon_n:02d}-{int(_day):02d}"
            if not _ts_raw or str(_ts_raw).lower() in ('none', ''):
                raise ValueError(f"travel_start not set: {_ts_raw!r}")
            _t_start_d = _edt.fromisoformat(str(_ts_raw)[:10])
            _cursor = _t_start_d
            for _ec, _en in nightly_split:
                if _en > 0:
                    _ci_s = _cursor.strftime('%d %b')
                    _co_s = (_cursor + _etd(days=_en)).strftime('%d %b')
                    _e_hotel_dates[_ec] = (_ci_s, _co_s)
                    _cursor += _etd(days=_en)
            _e_dates_estimated = True
        except Exception as _de:
            print(f"  {_DIM}ℹ Dates not available: {_de}{_RST}")

        _all_route_cities = [c for c, n in nightly_split if n > 0]
        for _city_idx, _s_city in enumerate(_all_route_cities):
            if _city_idx > 0:
                print(f"  {_DIM}{'─'*69}{_RST}")
            _first = True
            def _eline(colour, icon, text, tag='', date_str=''):
                nonlocal _first
                _prefix = f"  {_s_city:20}" if _first else f"  {'':20}"
                _first = False
                _tag = f"  {_DIM}({tag}){_RST}" if tag else ''
                _dcol = f"{_DIM}{date_str:<{_DATE_W}}{_RST}" if date_str else f"{'':>{_DATE_W}} "
                print(f"{_prefix} {_dcol}  {colour}{icon} {text}{_RST}{_tag}")

            # Hotels — primary
            _added_name     = _hotels_by_city.get(_s_city, '')
            _failed_name    = _hotels_failed_city.get(_s_city, '')
            _ehd = _e_hotel_dates.get(_s_city, ('',''))
            _ehdate = f"{_ehd[0]} → {_ehd[1]}" if _ehd[0] and _ehd[1] else ''

            if _added_name:
                _eline(_GRN, '✓', _added_name, 'hotel', date_str=_ehdate)
            elif _failed_name:
                _eline(_RED, '✗', _failed_name, 'hotel · NOT added — add manually in TCI')

            # Alt hotels
            for _alt in _hotels_added_dual.get(_s_city, []):
                _eline(_GRN, '✓', _alt, 'hotel · alt option')
            for _miss in _hotels_missing_dual.get(_s_city, []):
                _eline(_RED, '✗', _miss, 'hotel · not in TCI catalogue')

            # Monuments (Step 9)
            _mons = _added_monuments_per_city.get(_s_city, [])
            if _mons:
                _eline(_GRN, '✓', f"{len(_mons)} monument{'s' if len(_mons)!=1 else ''} added", 'monuments')

            # Guide (Step 12)
            _guide = _email_guides_by_city.get(_s_city, '')
            _guide_fail = _email_guides_failed.get(_s_city, '')
            if _guide:
                _eline(_GRN, '✓', _guide, 'guide')
            elif _guide_fail:
                _eline(_RED, '✗', _guide_fail, 'guide · add manually in TCI')
            elif _is_accompanying and _s_city == _all_route_cities[0]:
                _eline(_RED, '✗', 'Accompanying Guide Service', 'guide · not added — add manually in TCI')

            # Activities (Step 13)
            _acts = _email_acts_by_city.get(_s_city, [])
            for _act in _acts:
                _eline(_GRN, '✓', _act, 'activity')

            # City with no activities or monuments
            if not _mons and not _acts and not _guide:
                _eline(_DIM, '—', 'No activities or monuments added', '')

        if _e_dates_estimated:
            print(f"  {_DIM}ℹ Dates estimated from PNR — verify in TCI{_RST}")

    print(f"\n  ── Summary ─────────────────────────────")
    print(f"  Query ID   : {query_id}")
    print(f"  File Code  : {file_code}")
    total_steps = len(result['steps_completed'])
    _elapsed = round(_time.time() - _t_start)
    result['time_taken'] = _elapsed
    print(f"  Completed  : {total_steps} steps")
    print(f"  Time taken : {_elapsed}s")
    if result['errors']:
        print(f"  Errors     : {len(result['errors'])}")
        for e in result['errors']:
            print(f"    • {e}")

    return result


"""
generate_proposal_docx(intent, output_path)
Generates a formatted Word proposal from engine intent output.
Mirrors the Auchan Voyages sample format.
"""

"""
generate_proposal_docx v2 — matches Auchan Voyages sample format
Verdana font, SITA logo, inline icons per activity/monument
"""

# Icon paths — will be set by engine at runtime
PROPOSAL_ICON_DIR = '/content/proposal_icons'

def _get_icon_dir():
    """Return path to icons directory, checking multiple locations."""
    import os
    for d in [PROPOSAL_ICON_DIR, '/home/claude', '/content']:
        if os.path.exists(os.path.join(d, 'icon_sita_logo.png')):
            return d
    return None

"""
generate_proposal_docx v2 — matches Auchan Voyages sample format
Verdana font, SITA logo, inline icons per activity/monument
"""

# Icon paths — will be set by engine at runtime
PROPOSAL_ICON_DIR = '/content/proposal_icons'

def _get_icon_dir():
    """Return path to icons directory, checking multiple locations."""
    import os
    for d in [PROPOSAL_ICON_DIR, '/home/claude', '/content']:
        if os.path.exists(os.path.join(d, 'icon_sita_logo.png')):
            return d
    return None

"""
generate_proposal_docx v2 — matches Auchan Voyages sample format
Verdana font, SITA logo, inline icons per activity/monument
"""

# Icon paths — will be set by engine at runtime
PROPOSAL_ICON_DIR = '/content/proposal_icons'

def _get_icon_dir():
    """Return path to icons directory, checking multiple locations."""
    import os
    for d in [PROPOSAL_ICON_DIR, '/home/claude', '/content']:
        if os.path.exists(os.path.join(d, 'icon_sita_logo.png')):
            return d
    return None

def generate_proposal_docx(intent, output_path=None, email_text=''):
    """Generate formatted itinerary proposal matching SITA sample format."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import os
    except ImportError:
        print("  ⚠ python-docx not installed")
        return None

    icon_dir = _get_icon_dir()

    # ── COLOURS (from sample) ─────────────────────────────────────────────────
    BLUE      = RGBColor(0x2E, 0x75, 0xB6)  # heading blue
    DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)  # dark blue
    GREY      = RGBColor(0x59, 0x59, 0x59)  # body grey
    BLACK     = RGBColor(0x00, 0x00, 0x00)
    WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
    RED       = RGBColor(0xC0, 0x00, 0x00)  # DKC highlight

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_run(para, text, bold=False, italic=False, size=10,
                color=None, font='Verdana'):
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = font
        if color:
            run.font.color.rgb = color
        return run

    def add_inline_icon(para, icon_path, size_px=16):
        """Add a small inline icon image to a paragraph run."""
        if not icon_path or not os.path.exists(icon_path):
            return
        try:
            run = para.add_run()
            run.add_picture(icon_path, width=Inches(size_px/96))
        except Exception:
            pass

    def icon_line(doc, icon_path, text, bold=False, color=None, size=10, indent=0.25):
        """Add a line with an icon followed by text."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        if indent:
            para.paragraph_format.left_indent = Inches(indent)
        if icon_path:
            add_inline_icon(para, icon_path)
            para.add_run('  ')
        add_run(para, text, bold=bold, color=color or GREY, size=size)
        return para

    # ── EXTRACT FROM INTENT ───────────────────────────────────────────────────
    agency        = intent.get('agency_name') or 'Travel Agency'
    market        = intent.get('source_market', '')
    num_pax       = intent.get('num_pax') or 2
    tier          = (intent.get('tier') or 'first-class').replace('-', ' ').title()
    guide_lang    = intent.get('language_guide') or ''
    heritage      = intent.get('heritage_requested', False)
    nightly_split = intent.get('nightly_split') or []
    cities_nights = [(c, n) for c, n in nightly_split if n > 0]
    all_cities    = [c for c, n in nightly_split]
    total_nights  = sum(n for _, n in nightly_split)
    total_days    = total_nights + 1
    top_hotels    = intent.get('_top_hotels') or {}
    activities    = intent.get('_activities') or {}
    monuments     = intent.get('_monuments') or {}

    # Occasion from email
    occasion = ''
    if email_text:
        for kw, label in [('anniversary','Wedding Anniversary'),('honeymoon','Honeymoon'),
                           ('birthday','Birthday Celebration'),('wedding','Wedding')]:
            if kw in email_text.lower():
                occasion = label
                break

    # Region label
    regions = intent.get('regions_detected') or []
    if 'south_india' in regions and 'north_india' in regions:
        region_label = 'INDIA — NORTH & SOUTH'
    elif 'south_india' in regions:
        region_label = 'SOUTH INDIA'
    elif 'north_india' in regions:
        region_label = 'NORTH INDIA'
    else:
        region_label = 'INDIA'

    # Route string
    city_names_short = [c.replace(' (RJ)', '').replace(' (Alleppey)', '').replace(' (Periyar/Kumily)', '') 
                        for c in all_cities if c != 'Fatehpur Sikri']
    route_str = '  ·  '.join(city_names_short)

    # PNR date
    pnr_dates = intent.get('pnr_dates') or []
    period_str = f"{pnr_dates[0]}  |  {total_days} Days / {total_nights} Nights" if pnr_dates else f"{total_days} Days / {total_nights} Nights"

    # Vehicle
    vehicle = ('Toyota Innova Crysta' if num_pax <= 3 else
               'Force Urbania 7' if num_pax <= 6 else
               'A/C Coach')

    # City descriptions
    CITY_DESC = {
        'New Delhi': 'The capital of India — a layered city of Mughal monuments, colonial avenues, vibrant bazaars and world-class museums that spans from the 17th-century walled city of Shahjahanabad to the imperial boulevards of Lutyens\' Delhi.',
        'Agra': 'Once the imperial capital of the Mughal Empire, Agra is home to some of the finest examples of Mughal architecture, crowned by the Taj Mahal — one of the most transcendent buildings ever created.',
        'Jaipur': 'The Pink City of Rajasthan — a city of palace-forts, bustling bazaars and the living heritage of the Rajput maharajas, enclosed by a crenellated city wall with seven original gates.',
        'Jodhpur': 'The Blue City of Rajasthan — dominated by the magnificent Mehrangarh Fort rising 125 metres above a sea of indigo-painted Brahmin houses, with some of the best rooftop views in India.',
        'Udaipur (RJ)': 'The City of Lakes — Udaipur is Rajasthan\'s most romantic city, built around shimmering lakes and white-marble palaces, with the 18th-century Lake Palace floating in the centre of Lake Pichola.',
        'Jaisalmer': 'The Golden City — a medieval fort town rising from the Thar Desert, with intricately carved havelis, living fort quarters, and vast sand dunes beyond the city walls.',
        'Varanasi': 'The oldest living city in the world and the spiritual capital of Hinduism — Varanasi\'s ghats on the sacred Ganges have drawn pilgrims and seekers for over three thousand years.',
        'Kochi (Cochin)': 'Fort Kochi is a palimpsest of colonial history — Portuguese, Dutch and British layers overlaid on an ancient Indian trading port that has welcomed Arab, Chinese and Jewish merchants for over a thousand years.',
        'Alappuzha (Alleppey)': 'The Kerala backwaters are a 900-kilometre network of lagoons, lakes, rivers and canals running parallel to the Arabian Sea. A houseboat cruise offers one of the most distinctive travel experiences in India.',
        'Thekkady (Periyar/Kumily)': 'The drive from the flat backwater plains up into the Cardamom Hills is one of Kerala\'s great scenic journeys — rising from rice paddies through rubber and tea plantations to dense jungle at Thekkady.',
        'Munnar': 'High-altitude tea country in the Western Ghats — Munnar\'s emerald-green plantations are among the most scenic landscapes in Kerala, set around 1,600 metres above sea level.',
        'Madurai': 'One of the oldest continuously inhabited cities in the world and the cultural capital of Tamil Nadu, anchored by the magnificent Meenakshi Amman Temple.',
        'Samode': 'Nestled in the Aravalli Hills near Jaipur, Samode is home to a spectacular 475-year-old heritage palace decorated with original Mughal and Rajasthani frescoes.',
        'Deogarh': 'A charming Rajasthan town known for its magnificent palace, unique school of miniature paintings, and the scenic Raghosagar Lake with its romantic ruined island temple.',
        'Ranthambore': 'One of India\'s finest tiger reserves — set amid dramatic Rajput ruins and dense jungle in eastern Rajasthan, Ranthambore offers some of the best tiger viewing on the subcontinent.',
        'Khajuraho': 'The medieval temples of Khajuraho — a UNESCO World Heritage Site — are celebrated for their extraordinary erotic carvings, a remarkable artistic achievement of the Chandela dynasty.',
        'Bandhavgarh': 'One of India\'s best parks for tiger sightings, set in a dramatic landscape of sal forests and rocky outcrops in Madhya Pradesh.',
    }

    # ── BUILD DOCUMENT ────────────────────────────────────────────────────────
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── SITA LOGO ─────────────────────────────────────────────────────────────
    logo_path = os.path.join(icon_dir, 'icon_sita_logo.png') if icon_dir else None
    if logo_path and os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Inches(1.8))
        logo_para.paragraph_format.space_after = Pt(8)

    # ── ORIGINAL EMAIL ────────────────────────────────────────────────────────
    if email_text and email_text.strip():
        # Section header
        email_hdr = doc.add_paragraph()
        email_hdr.paragraph_format.space_before = Pt(4)
        email_hdr.paragraph_format.space_after  = Pt(4)
        hdr_run = email_hdr.add_run('AGENT ENQUIRY')
        hdr_run.bold = True
        hdr_run.font.size = Pt(9)
        hdr_run.font.name = 'Verdana'
        hdr_run.font.color.rgb = WHITE
        # Blue background for header
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlElement
        pPr = email_hdr._p.get_or_add_pPr()
        shd = _OxmlElement('w:shd')
        shd.set(_qn('w:val'), 'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'), '2E75B6')
        pPr.append(shd)
        email_hdr.paragraph_format.left_indent = Inches(0.1)

        # Email body in box
        email_para = doc.add_paragraph()
        email_para.paragraph_format.space_before = Pt(0)
        email_para.paragraph_format.space_after  = Pt(8)
        email_para.paragraph_format.left_indent  = Inches(0.1)
        email_para.paragraph_format.right_indent = Inches(0.1)
        # Light grey background
        ePr = email_para._p.get_or_add_pPr()
        eshd = _OxmlElement('w:shd')
        eshd.set(_qn('w:val'), 'clear')
        eshd.set(_qn('w:color'), 'auto')
        eshd.set(_qn('w:fill'), 'F2F2F2')
        ePr.append(eshd)
        # Email text — clean and truncate if very long
        email_clean = email_text.strip()
        if len(email_clean) > 1500:
            email_clean = email_clean[:1500] + '...'
        email_run = email_para.add_run(email_clean)
        email_run.font.size = Pt(8)
        email_run.font.name = 'Courier New'
        email_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── TOUR TITLE ────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title_para, f'{region_label} — {route_str.upper()}',
            bold=True, size=16, color=DARK_BLUE)

    dur_para = doc.add_paragraph()
    dur_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(dur_para, f'{total_days} Days / {total_nights} Nights', size=14, color=BLUE)

    # Route cities detail
    cities_detail = '  ·  '.join(c for c, n in nightly_split if n >= 0)
    cities_para = doc.add_paragraph()
    cities_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cities_para, cities_detail, italic=True, size=10, color=GREY)

    # Divider
    div = doc.add_paragraph('─' * 70)
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in div.runs:
        r.font.color.rgb = BLUE
        r.font.size = Pt(8)

    # Prepared for
    prep_para = doc.add_paragraph()
    prep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(prep_para, f'Prepared for {agency}', size=11, color=DARK_BLUE)
    if occasion:
        add_run(prep_para, f'  ★ {occasion}  ★', bold=True, size=11, color=BLUE)

    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(period_para, period_str, italic=True, size=10, color=GREY)

    div2 = doc.add_paragraph('─' * 70)
    div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in div2.runs:
        r.font.color.rgb = BLUE
        r.font.size = Pt(8)

    # ── ITINERARY HEADING ─────────────────────────────────────────────────────
    doc.add_paragraph()
    itin_h = doc.add_paragraph()
    itin_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(itin_h, 'ITINERARY', bold=True, size=13, color=DARK_BLUE)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    short_route = ' – '.join(c.replace(' (RJ)', '') for c in [c for c,n in cities_nights])
    add_run(sub_para, short_route, italic=True, size=9, color=GREY)
    add_run(sub_para, f'\n{total_days} Days / {total_nights} Nights', italic=True, size=9, color=GREY)

    # Icon paths
    ico_plane    = os.path.join(icon_dir, 'icon_airplane.png') if icon_dir else None
    ico_hand     = os.path.join(icon_dir, 'icon_handshake.png') if icon_dir else None
    ico_car      = os.path.join(icon_dir, 'icon_car.png') if icon_dir else None
    ico_bed      = os.path.join(icon_dir, 'icon_bed.png') if icon_dir else None
    ico_temple   = os.path.join(icon_dir, 'icon_temple.png') if icon_dir else None
    ico_activity = os.path.join(icon_dir, 'icon_activity.png') if icon_dir else None

    # ── DAY BY DAY ────────────────────────────────────────────────────────────
    day_num = 1
    prev_city = None

    for city, nights in nightly_split:
        is_transit = (nights == 0)

        for night_idx in range(max(nights, 1) if not is_transit else 1):
            if is_transit and night_idx > 0:
                break

            # Day header
            day_para = doc.add_paragraph()
            day_para.paragraph_format.space_before = Pt(10)
            add_run(day_para, f'Day {day_num}  ', bold=True, size=13, color=BLUE)
            if prev_city and not is_transit and night_idx == 0:
                add_run(day_para, f'{prev_city} – {city}', bold=True, size=13, color=DARK_BLUE)
            else:
                add_run(day_para, city.replace(' (RJ)', '').replace(' (Periyar/Kumily)', ''), bold=True, size=13, color=DARK_BLUE)

            # Day 1 — arrival line
            if day_num == 1:
                dest_code = 'COK' if 'Kochi' in city else 'DEL' if 'Delhi' in city else 'BOM' if 'Mumbai' in city else 'MAA' if 'Chennai' in city else ''
                arr_text = f'International arrival at {city.replace(" (RJ)","").replace(" (Cochin)","")} ({dest_code})' if dest_code else f'International arrival at {city.replace(" (RJ)","")}'
                icon_line(doc, ico_plane, arr_text, color=GREY)
                icon_line(doc, ico_hand, 'Meet & greet on arrival', color=GREY)
                icon_line(doc, ico_car, 'Transfer to hotel', color=GREY)

            # Drive line for subsequent cities
            elif night_idx == 0 and prev_city and not is_transit:
                icon_line(doc, ico_car, f'Drive from {prev_city.replace(" (RJ)","").replace(" (Periyar/Kumily)","")} to {city.replace(" (RJ)","").replace(" (Periyar/Kumily)","")}', color=GREY)

            # City description (first night only)
            if night_idx == 0 and city in CITY_DESC and not is_transit:
                desc_para = doc.add_paragraph()
                desc_para.paragraph_format.left_indent = Inches(0.25)
                desc_para.paragraph_format.space_before = Pt(3)
                add_run(desc_para, CITY_DESC[city], italic=True, size=9.5, color=GREY)

            # Transit visit
            if is_transit:
                icon_line(doc, ico_temple, f'Visit {city}', color=GREY)
                day_num += 1
                break

            # First night only: show monuments + activities
            if night_idx == 0:
                city_mons = monuments.get(city, [])
                for mon_name, _ in city_mons[:3]:
                    icon_line(doc, ico_temple, mon_name, color=DARK_BLUE)
                city_acts = activities.get(city, [])
            else:
                city_acts = []
                icon_line(doc, ico_activity, 'Free time for exploration and leisure', color=GREY)
            for act_name, _ in city_acts[:3]:
                is_dkc = '[A DKC Experience]' in act_name
                clean_name = act_name.replace(' [A DKC Experience]', '')
                dkc_suffix = '  ★ DKC Experience' if is_dkc else ''
                p = icon_line(doc, ico_activity, clean_name, color=DARK_BLUE)
                if is_dkc:
                    add_run(p, dkc_suffix, bold=True, size=9.5, color=RED)

            # Overnight
            hotel = top_hotels.get(city, '')
            ovn_para = doc.add_paragraph()
            ovn_para.paragraph_format.left_indent = Inches(0.25)
            ovn_para.paragraph_format.space_before = Pt(4)
            if ico_bed and os.path.exists(ico_bed):
                add_inline_icon(ovn_para, ico_bed)
                ovn_para.add_run('  ')
            if hotel:
                add_run(ovn_para, f'Overnight at {hotel}', bold=True, size=10, color=DARK_BLUE)
                if night_idx == 0 and nights > 1:
                    add_run(ovn_para, f'  ({nights} nights)', size=9.5, color=GREY)
            else:
                add_run(ovn_para, 'Overnight in city', bold=True, size=10, color=DARK_BLUE)

            prev_city = city
            day_num += 1

    # Last day departure
    last_city = cities_nights[-1][0] if cities_nights else ''
    dep_para = doc.add_paragraph()
    dep_para.paragraph_format.space_before = Pt(10)
    add_run(dep_para, f'Day {day_num}  ', bold=True, size=13, color=BLUE)
    add_run(dep_para, last_city.replace(' (RJ)', ''), bold=True, size=13, color=DARK_BLUE)

    dep_airport = 'MAA' if 'Chennai' in last_city else 'COK' if 'Kochi' in last_city else 'DEL' if 'Delhi' in last_city else 'BOM'
    icon_line(doc, ico_car, f'Transfer to international airport', color=GREY)
    icon_line(doc, ico_plane, f'International departure from {last_city.replace(" (RJ)","").replace(" (Cochin)","")} ({dep_airport})', color=GREY)

    # ── HOTEL SUMMARY ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    div3 = doc.add_paragraph('─' * 70)
    for r in div3.runs: r.font.color.rgb = BLUE; r.font.size = Pt(8)
    add_run(doc.add_paragraph(), 'HOTEL SUMMARY', bold=True, size=12, color=DARK_BLUE)

    htable = doc.add_table(rows=1, cols=3)
    htable.style = 'Table Grid'
    hdr = htable.rows[0].cells
    for i, txt in enumerate(['City', 'Hotel', 'Nights']):
        set_cell_bg(hdr[i], '2E75B6')
        add_run(hdr[i].paragraphs[0], txt, bold=True, size=10, color=WHITE)

    for city, nights in nightly_split:
        if nights == 0:
            continue
        row = htable.add_row().cells
        set_cell_bg(row[0], 'DEEAF1')
        add_run(row[0].paragraphs[0], city.replace(' (RJ)', ''), size=10, color=DARK_BLUE)
        hotel = top_hotels.get(city, 'TBC')
        add_run(row[1].paragraphs[0], f'{hotel}  ({tier})', size=10, color=DARK_BLUE)
        add_run(row[2].paragraphs[0], str(nights), size=10, color=DARK_BLUE)

    # ── F&B RECOMMENDATIONS ──────────────────────────────────────────────────
    fnb_data = intent.get('_fnb') or {}
    if fnb_data:
        doc.add_paragraph()
        div_fnb = doc.add_paragraph('─' * 70)
        for r in div_fnb.runs: r.font.color.rgb = BLUE; r.font.size = Pt(8)
        add_run(doc.add_paragraph(), 'RECOMMENDED RESTAURANTS', bold=True, size=12, color=DARK_BLUE)

        fnb_table = doc.add_table(rows=1, cols=4)
        fnb_table.style = 'Table Grid'
        hdr_f = fnb_table.rows[0].cells
        for i, txt in enumerate(['City', 'Restaurant', 'Type', 'Rating']):
            set_cell_bg(hdr_f[i], '2E75B6')
            add_run(hdr_f[i].paragraphs[0], txt, bold=True, size=10, color=WHITE)

        for city, city_nights in nightly_split:
            if city_nights == 0:
                continue
            rows_fnb = fnb_data.get(city) or []
            for ri, r in enumerate(rows_fnb[:3]):
                row_f = fnb_table.add_row().cells
                if ri == 0:
                    set_cell_bg(row_f[0], 'DEEAF1')
                    add_run(row_f[0].paragraphs[0], city.replace(' (RJ)', ''), size=9.5, color=DARK_BLUE)
                else:
                    set_cell_bg(row_f[0], 'FFFFFF')
                # Hotel attribution
                if r.get('hotel_linked') and r.get('hotel_code'):
                    hotel_nm = _PORTFOLIO_CODE_NAME.get(r['hotel_code'], '') if '_PORTFOLIO_CODE_NAME' in dir() else ''
                    rest_display = f"{r['name']}  @ {hotel_nm}" if hotel_nm else r['name']
                else:
                    rest_display = r['name']
                type_map = {
                    'specialty-restaurant': 'Specialty',
                    'multi-cuisine': 'Multi-cuisine',
                    'coffee-shop': 'All-day dining',
                    'lounge': 'Lounge',
                    'bar': 'Bar',
                }
                type_label = type_map.get(r.get('restaurant_type', ''), r.get('restaurant_type', '').title())
                add_run(row_f[1].paragraphs[0], rest_display, size=9.5, color=DARK_BLUE)
                add_run(row_f[2].paragraphs[0], type_label, size=9.5, color=GREY)
                add_run(row_f[3].paragraphs[0], f"{r.get('tci_score', 0):.1f} ★", size=9.5, color=BLUE)

    # ── INCLUSIONS ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    div4 = doc.add_paragraph('─' * 70)
    for r in div4.runs: r.font.color.rgb = BLUE; r.font.size = Pt(8)
    add_run(doc.add_paragraph(), 'INCLUSIONS', bold=True, size=12, color=DARK_BLUE)

    guide_txt = f'{guide_lang.replace("guide","").strip().title()}-speaking guide' if guide_lang else 'English-speaking guide'
    for item in [
        'Accommodation as mentioned above on twin/double sharing basis',
        'Daily breakfast at all hotels',
        f'All transfers and sightseeing by private air-conditioned vehicle ({vehicle} for {num_pax} pax)',
        f'{guide_txt} throughout the tour',
        'All monument entry fees as per itinerary',
        'All activities as mentioned in the itinerary',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        add_run(p, item, size=10, color=GREY)

    # ── EXCLUSIONS ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_run(doc.add_paragraph(), 'EXCLUSIONS', bold=True, size=12, color=DARK_BLUE)
    for item in [
        'International airfare',
        'Visa fees',
        'Travel insurance (strongly recommended)',
        'Items of personal nature — laundry, tips, telephone, beverages',
        'Any entrance fees not specified above',
        'Anything not mentioned under inclusions',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        add_run(p, item, size=10, color=GREY)

    # ── CANCELLATION POLICY ───────────────────────────────────────────────────
    doc.add_paragraph()
    add_run(doc.add_paragraph(), 'CANCELLATION POLICY', bold=True, size=12, color=DARK_BLUE)
    for item in [
        '25% total cost of land arrangements at the time of confirming the booking is non-refundable.',
        '50% refund for cancellations between 89–60 days prior to travel.',
        'No refund for cancellations less than 60 days prior to travel.',
        'Less than 15 days or no show: 100% of the tour cost.',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        add_run(p, item, size=10, color=GREY)

    # ── GENERAL INFORMATION ───────────────────────────────────────────────────
    doc.add_paragraph()
    add_run(doc.add_paragraph(), 'GENERAL INFORMATION', bold=True, size=12, color=DARK_BLUE)
    for item in [
        'Hotel check-in: 1400 hrs. Check-out: 1200 hrs. Early check-in / late check-out subject to availability.',
        'Timings specified are indicative and subject to change. Passport details required to confirm all bookings.',
        'This proposal is valid for 15 days from submission date, excluding peak periods, festivals, and blackout dates.',
        'Hotel names indicate rates are based on current availability. Alternatives at similar standard will be confirmed if unavailable.',
        'Foreign nationals require a tourist visa to enter India. Travel insurance is strongly advised.',
        'Sita/TCI/Go Vacation and their subsidiaries shall not be liable for an amount exceeding what the claimant has paid them.',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        add_run(p, item, size=10, color=GREY)

    if occasion:
        doc.add_paragraph()
        note = doc.add_paragraph()
        add_run(note, f'★ Special Occasion: {occasion}', bold=True, size=10, color=BLUE)
        add_run(note, ' — Please advise hotel of the occasion at time of booking for complimentary room decoration.', size=10, color=GREY)

    # ── SAVE ──────────────────────────────────────────────────────────────────
    if not output_path:
        safe = (intent.get('agency_name') or 'Proposal').replace(' ', '_')[:20]
        output_path = f'/content/Proposal_{safe}.docx'

    doc.save(output_path)
    print(f"  ✓ Proposal saved: {output_path}")
    return output_path




# ── LLM FALLBACK PARSER ───────────────────────────────────────────────────────

def check_parse_confidence(intent):
    """
    Score the parser output. Returns (confidence, issues, score).
    confidence: 'high' | 'medium' | 'low'
    """
    issues = []
    score  = 100

    if not intent.get('source_market') or intent['source_market'] == 'UNKNOWN':
        issues.append('market not detected')
        score -= 40

    if not intent.get('num_pax'):
        issues.append('pax count missing')
        score -= 20

    if not intent.get('cities_detected') and not intent.get('regions_detected'):
        issues.append('no cities or region found')
        score -= 30

    if not intent.get('duration_nights'):
        issues.append('duration not detected')
        score -= 10

    if score >= 80:
        confidence = 'high'
    elif score >= 50:
        confidence = 'medium'
    else:
        confidence = 'low'

    return confidence, issues, score


def llm_parse_fallback(email_text, missing_fields):
    """
    Ask Gemini Flash (free tier) for fields the regex parser missed.
    Returns a dict of recovered fields, or {} on failure.

    Cost controls:
    - Hard stops at 1,400/day and 12/min (free tier buffer)
    - Caches by email hash — same email never hits LLM twice per session
    - Sends only first 25 lines of email
    - Falls back to Claude Haiku if Gemini unavailable
    """
    import json as _json, hashlib as _hash

    # Cache check — same email never hits LLM twice
    _cache_key = _hash.md5(email_text[:500].encode('utf-8', errors='ignore')).hexdigest()
    if _cache_key in _LLM_CACHE:
        print(f"  ℹ LLM cache hit — skipping API call")
        return _LLM_CACHE[_cache_key]

    # Quota check — stop before free tier limit
    if not _llm_quota_check():
        return {}

    # Truncate email — first 25 lines only to minimise tokens
    _email_short = '\n'.join([l for l in email_text.split('\n') if l.strip()][:25])

    fields_needed = ', '.join(missing_fields)
    prompt = f"""You are a travel industry email parser for an India inbound tour operator.
Extract ONLY these missing fields: {fields_needed}

Return a JSON object with only these keys (null if unknown):
- source_market: 3-letter ISO code (FRA,GBR,DEU,ITA,USA,JPN,AUS,CHE,NLD,ESP,IND,PRT,AUT,POL,RUS,ROU,KEN etc)
- num_pax: integer
- cities_detected: list of Indian city names
- regions_detected: list from [north_india,south_india,central_india,west_india,himalaya,wildlife]
- duration_nights: integer
- tier: lux|first-class|moderate|budget or null

Return ONLY valid JSON. No explanation, no markdown.

EMAIL:
{_email_short}"""

    result = {}

    # Try Gemini Flash (free tier, pre-installed in Colab)
    try:
        import google.generativeai as _genai
        _genai.configure()   # uses GOOGLE_API_KEY env var or Colab ADC
        _model = _genai.GenerativeModel('gemini-2.0-flash')
        _resp  = _model.generate_content(
            prompt,
            generation_config=_genai.GenerationConfig(max_output_tokens=250, temperature=0.0)
        )
        raw = (_resp.text or '').strip().lstrip('`').rstrip('`')
        if raw.startswith('json'):
            raw = raw[4:]
        result = _json.loads(raw.strip())
        print(f"  ✓ Gemini LLM recovered: {[k for k,v in result.items() if v]}  [{_llm_quota_status()}]")
        _LLM_CACHE[_cache_key] = result
        return result
    except ImportError:
        pass
    except Exception as _ge:
        if '429' in str(_ge) or 'quota' in str(_ge).lower() or 'RESOURCE_EXHAUSTED' in str(_ge):
            print(f"  ℹ Gemini quota hit — using regex parser")
        else:
            print(f"  ⚠ Gemini error: {_ge}")

    # Fallback: Claude Haiku (cheapest Anthropic — only if API key set)
    try:
        import anthropic as _anthropic
        _client   = _anthropic.Anthropic()
        _response = _client.messages.create(
            model='claude-haiku-4-5-20251001',
            max_tokens=250,
            messages=[{'role': 'user', 'content': prompt}]
        )
        raw = _response.content[0].text.strip().lstrip('`').rstrip('`')
        if raw.startswith('json'):
            raw = raw[4:]
        result = _json.loads(raw.strip())
        print(f"  ✓ Claude Haiku recovered: {[k for k,v in result.items() if v]}")
        _LLM_CACHE[_cache_key] = result
    except ImportError:
        print(f"  ℹ No LLM available — install google-generativeai for free tier access")
    except Exception as _ae:
        print(f"  ⚠ LLM fallback error: {_ae}")

    return {}


def merge_intent(regex_intent, llm_intent):
    """
    Merge regex and LLM results — regex wins where it found something,
    LLM fills the gaps.
    """
    merged = dict(regex_intent)

    for key, llm_val in llm_intent.items():
        if not llm_val:
            continue
        regex_val = regex_intent.get(key)

        if key == 'cities_detected':
            existing = set(regex_val or [])
            for city in (llm_val or []):
                if city not in existing:
                    merged.setdefault('cities_detected', []).append(city)
                    existing.add(city)
        elif key == 'regions_detected':
            existing = set(regex_val or [])
            for r in (llm_val or []):
                if r not in existing:
                    merged.setdefault('regions_detected', []).append(r)
                    existing.add(r)
        elif key == 'travel_month_approx':
            # Only store if regex didn't already get a travel_start from PNR
            if not regex_intent.get('travel_start'):
                merged['travel_month_approx'] = llm_val
                merged['_travel_month_source'] = 'llm_fallback'
                # Convert YYYY-MM to a travel_start date (1st of that month)
                try:
                    import datetime as _dt2
                    y, m = llm_val.split('-')
                    merged['travel_start'] = _dt2.date(int(y), int(m), 1)
                    merged['travel_end']   = _dt2.date(int(y), int(m), 28)  # approx month end
                except Exception:
                    pass
        elif key == 'previously_visited_india':
            if llm_val is True:
                merged['_previously_visited_india'] = True
        elif not regex_val or regex_val == 'UNKNOWN':
            merged[key] = llm_val
            merged[f'_{key}_source'] = 'llm_fallback'

    return merged


def _log_llm_fallback(email_text, issues, llm_result):
    """Append fallback event to JSONL log for parser improvement tracking."""
    try:
        import json as _json, datetime as _dt
        entry = {
            'ts':       _dt.datetime.utcnow().isoformat(),
            'issues':   issues,
            'llm_keys': list(llm_result.keys()),
            'snippet':  email_text[:200],
        }
        log_path = '/content/fallback_log.jsonl'
        with open(log_path, 'a') as _f:
            _f.write(_json.dumps(entry) + '\n')
    except Exception:
        pass



def _is_tourlane_booking(email_text):
    """
    Detect a structured Tourlane booking confirmation (NOT a freetext consultation).

    Tourlane uses two email types:
      1. Structured booking: Subject 'NEW BOOKING: T-YYYYMM-NNNNN', Accommodations
         table with check-in/hotel/rate columns, "make a new booking" body.
         → must route to parse_tourlane_email().
      2. Freetext consultation: e.g. "Customer request: elephant safari" — these
         look like any other agent email and must use the regular parser.

    Sender domain alone (@tourlane.com) is NOT sufficient — both types come from
    the same domain. We require >=2 structural signals from the booking format.
    """
    t = email_text.lower()
    # Subject line: "NEW BOOKING: T-202604-90482"
    has_new_booking_subject = bool(re.search(
        r'subject:.*new\s+booking:?\s*t-\d{6}-\d+', t))
    # Body has "Booking Number" header followed (within ~200 chars) by the T-ref
    has_booking_ref_in_body = bool(re.search(
        r'booking\s*number[\s\S]{0,200}t-\d{6}-\d+', t))
    # Accommodations table header signature
    has_accommodation_table = (
        'accommodations' in t and 'check-in' in t and 'hotel name' in t
    )
    # Booking phrasing
    has_make_booking_phrase = (
        'make a new booking' in t or 'make a booking with the details' in t
    )
    score = sum([has_new_booking_subject, has_booking_ref_in_body,
                 has_accommodation_table, has_make_booking_phrase])
    return score >= 2


def _normalise_tourlane_intent(intent, email_text):
    """
    Pad a Tourlane intent dict with the keys the rest of recommend() expects.

    Tourlane parser emits a slightly different shape than parse_email():
      duration (int)        → also expose as duration_nights (tuple)
      heritage (bool)       → also expose as heritage_requested + heritage_keywords
      no regions_detected   → empty list
      no transport_notes    → empty list
      no activities_requested → empty list
      no _raw_text          → lowercase email body

    Mutates and returns intent.
    """
    # Duration: Tourlane uses int days; downstream wants (min, max) tuple
    _dur = intent.get('duration')
    if _dur and not intent.get('duration_nights'):
        intent['duration_nights'] = (int(_dur), int(_dur))

    # Heritage: Tourlane uses bool; downstream uses heritage_requested + keywords list
    intent.setdefault('heritage_requested', bool(intent.get('heritage')))
    intent.setdefault('heritage_keywords', [])

    # Region/activity/transport defaults
    intent.setdefault('regions_detected', [])
    intent.setdefault('cities_excluded', [])
    intent.setdefault('transport_notes', [])
    intent.setdefault('activities_requested', [])
    intent.setdefault('domestic_sectors', [])
    intent.setdefault('named_trains', [])
    intent.setdefault('inline_city_order', [])

    # Raw text — used by region detection and various downstream checks
    intent.setdefault('_raw_text', (email_text or '').lower())

    # Market source label — drives the PARSED INTENT print (not used in Tourlane
    # path but harmless to set for downstream consumers)
    intent.setdefault('_market_source', 'tourlane_booking')

    return intent


# ─────────────────────────────────────────────────────────────────────
def _req_dur_scalar(duration_nights):
    """Normalise duration_nights (int, tuple, or None) to a single int or None."""
    if duration_nights is None:
        return None
    if isinstance(duration_nights, (list, tuple)) and len(duration_nights) == 2:
        return (duration_nights[0] + duration_nights[1]) // 2
    if isinstance(duration_nights, int):
        return duration_nights
    return None


# BL2-α (13 May 2026) — cascade envelope → engine all_routes adapter
# Converts historical_query.assemble_options() output into the shape the
# engine's downstream code consumes: list of routes, where each route is
# a list of (city, nights) tuples. Slot 0 = primary route, rest = alts.
# ─────────────────────────────────────────────────────────────────────
def _envelope_to_all_routes(envelope, requested_duration=None):
    """Convert cascade envelope {options, caveat, diagnostic} to engine
    all_routes shape [[(city, nights), ...], ...].

    Each cascade option carries city_sequence (tuple of city names) and
    _nightly_split (list of ints, length-aligned to city_sequence). Drops
    any malformed option silently — engine's REGIONAL_DEFAULTS fallback
    will fire if this returns an empty list.

    v20260514b: applies _apply_city_night_rules() per option so nightly
    splits are clamped to [rule_min, rule_max] and rebalanced to
    requested_duration. Same clamp runs for agent / market / global options
    equally — the rules are tier-agnostic.
    """
    routes = []
    for opt in envelope.get('options', []) or []:
        cities = opt.get('city_sequence') or ()
        nights = opt.get('_nightly_split') or []
        if not cities or len(nights) != len(cities):
            continue
        route = [(c, int(n)) for c, n in zip(cities, nights)]
        if route:
            # Apply city night rules clamp + duration rebalance
            route = _apply_city_night_rules(route, requested_duration)
            routes.append(route)
    return routes


def recommend(email_text, db_path=DB_PATH, portfolio_path=PORTFOLIO_PATH,
              activity_path=ACTIVITY_PATH, fnb_path=FNB_PATH,
              agency_path=AGENCY_PATH, city_rules_path=CITY_RULES_PATH,
              source_market=None, agency_account_code=None, agency_name=None):
    """
    Main entry point. Takes raw email text, returns full recommendation.
    """
    print("\n" + "="*70)
    print("  ITINERARY RECOMMENDATION ENGINE")
    print("="*70)

    # Load agency portfolio first — must be available before parse_email() runs
    # so domain-based market detection works
    if agency_path:
        load_agency_portfolio(agency_path)

    # Parse email
    # Auto-translate non-English emails before parsing
    email_text, was_translated = detect_and_translate(email_text)
    if was_translated:
        print("  [Translation complete — parsing English version]")

    # ── ROUTING: Tourlane structured booking vs freetext email ──────────────
    # Tourlane sends two distinct email types from @tourlane.com:
    #   1. NEW BOOKING confirmations (structured tables) → parse_tourlane_email()
    #   2. Customer consultation requests (freetext)     → parse_email()
    # Detection is structural (Booking Number table, Accommodations table,
    # "make a new booking" phrasing), NOT sender-based.
    _is_tourlane_mode = _is_tourlane_booking(email_text)

    if _is_tourlane_mode:
        intent = parse_tourlane_email(email_text)
        intent = _normalise_tourlane_intent(intent, email_text)

        # Caller overrides still apply
        if source_market:
            intent['source_market'] = source_market
            print(f"  ℹ Market override: {source_market}")
        if agency_account_code:
            intent['agency_account_code'] = agency_account_code
        if agency_name:
            intent['agency_name'] = agency_name
    else:
        intent = parse_email(email_text)

        # ── LLM fallback — fires only when parser confidence is LOW ──────────────
        # Medium confidence means regex got most fields — LLM not worth the quota cost
        _confidence, _issues, _score = check_parse_confidence(intent)
        if _confidence in ('low',):
            print(f"  ⚡ Invoking LLM fallback (confidence: LOW {_score}/100) — issues: {', '.join(_issues)}")
            _llm = llm_parse_fallback(email_text, _issues)
            if _llm:
                intent = merge_intent(intent, _llm)
                intent['_llm_used'] = True
                _log_llm_fallback(email_text, _issues, _llm)
                _recovered = [k for k in _llm if _llm[k] and not k.startswith('_')]
                print(f"  ✓ LLM recovered: {_recovered}")
        # ─────────────────────────────────────────────────────────────────────────

        # ── Caller overrides — applied before routing so DB queries use correct market
        if source_market:
            intent['source_market'] = source_market
            print(f"  ℹ Market override: {source_market}")
        if agency_account_code:
            intent['agency_account_code'] = agency_account_code
        if agency_name:
            intent['agency_name'] = agency_name

        if intent.get('_duration_from_pnr') and intent.get('travel_start'):
            dn = intent['duration_nights']
            print(f"  ℹ  Duration derived from PNR: {dn[0]} nights "
                  f"({intent['travel_start']} → {intent['travel_end']})")

        # Post-parse region cleanup: if only south india requested, drop north_india
        _raw_lower = (intent.get('_raw_text') or '').lower()
        _south_kws = ['south india', 'southern india', 'kerala', 'tamil nadu',
                      'tamilnadu', 'tamil naadu', 'karnataka', 'karnatka']
        _north_kws = ['north india', 'northern india', 'rajasthan', 'rajistan',
                      'rajisthan', 'rajstan', 'golden triangle']
        _south_only = (any(kw in _raw_lower for kw in _south_kws) and
                       not any(kw in _raw_lower for kw in _north_kws))
        if _south_only and 'south_india' in intent.get('regions_detected', []):
            intent['regions_detected'] = [r for r in intent['regions_detected'] if r != 'north_india']

        # Sub-region detection — when email explicitly names a state or circuit,
        # store it so historical_query uses the tighter city set for overlap scoring.
        # This prevents Golden Triangle + Shimla routes from ranking above pure
        # Rajasthan routes when the agent asked for Rajasthan specifically.
        # Gateway cities (Delhi, Mumbai etc.) are always neutral — not penalised.
        _SUB_REGION_KEYWORDS = {
            'rajasthan':      ['rajasthan'],
            'kerala':         ['kerala'],
            'himachal':       ['himachal', 'himachal pradesh'],
            'uttarakhand':    ['uttarakhand', 'uttaranchal'],
            'gujarat':        ['gujarat', 'kutch', 'rann of kutch'],
            'karnataka':      ['karnataka', 'coorg', 'hampi'],
            'tamil_nadu':     ['tamil nadu', 'tamilnadu', 'tamil naadu'],
            'kashmir':        ['kashmir', 'srinagar', 'dal lake'],
            'ladakh':         ['ladakh', 'leh'],
            'punjab':         ['punjab', 'amritsar', 'golden temple'],
            'madhya_pradesh': ['madhya pradesh', 'mp safari', 'central india'],
            'odisha':         ['odisha', 'orissa', 'puri', 'konark'],
            'northeast':      ['northeast', 'north east', 'assam', 'meghalaya',
                               'arunachal', 'nagaland', 'manipur', 'mizoram'],
            'sikkim':         ['sikkim', 'gangtok'],
            'west_bengal':    ['west bengal', 'darjeeling', 'kolkata'],
            'goa':            ['goa'],
            'andaman':        ['andaman', 'nicobar'],
        }
        # v20260522c (N2): sub-region keyword needs DESTINATION/STAY intent
        # context — not just a passing locator mention like "Chand Baori,
        # Abhneri, Rajasthan" (B3 #003). For each candidate keyword we check
        # whether it sits within ~30 chars of an intent vocabulary token
        # (tour, trip, visit, nights, days, in/to/dans/en + duration). The
        # bare-keyword check is retained as a fallback so we don't over-suppress.
        # Sub-region city-set match (≥2 zone cities mentioned in the email) is
        # *not* used here; that's handled later in historical_query scoring.
        _SUB_REGION_INTENT_PRE = (
            r'\b(?:in|to|of|across|throughout|visiting|visit|tour|trip|travel|'
            r'holiday|honeymoon|spend|stay|nights?|days?|weeks?|nuits?|jours?|'
            r'sem|semanas?|d[ií]as?|noches?|n\d|\d+n|dans\s+le|dans\s+la|en|pendant|'
            r'nach|durch|hotels?\s+in|stay\s+in)\b'
        )
        _SUB_REGION_INTENT_POST = (
            r'\b(?:tour|trip|travel|holiday|circuit|loop|journey|nights?|days?|'
            r'weeks?|safari|jaipur|jodhpur|udaipur|jaisalmer|bikaner|n\d|\d+n|'
            r'tourisme|with|pour|et|including)\b'
        )
        def _sub_region_keyword_hits(text_lower, kws):
            """Return True if any kw in kws has destination/stay intent context
            in text_lower; otherwise False. Multi-word keywords (e.g. 'tamil nadu')
            match as a unit."""
            for kw in kws:
                if kw not in text_lower:
                    continue
                _kw_esc = re.escape(kw)
                _pat = re.compile(
                    rf'{_SUB_REGION_INTENT_PRE}[\w\s,:;\-/\.]{{0,30}}?\b{_kw_esc}\b'
                    rf'|\b{_kw_esc}\b[\w\s,:;\-/\.]{{0,30}}?{_SUB_REGION_INTENT_POST}',
                    re.IGNORECASE
                )
                if _pat.search(text_lower):
                    return True
            return False
        _detected_sub = None
        for _sub, _kws in _SUB_REGION_KEYWORDS.items():
            if _sub_region_keyword_hits(_raw_lower, _kws):
                _detected_sub = _sub
                break
        intent['sub_region'] = _detected_sub

    # Skip PARSED INTENT block for Tourlane mode — the Tourlane parser already
    # prints its own banner with booking ref, guest, agency, market, pax, dates,
    # route, hotels, activities. Re-printing would just duplicate.
    if _is_tourlane_mode:
        _skip_parsed_intent = True
    else:
        _skip_parsed_intent = False

    if not _skip_parsed_intent:
        print("\n── PARSED INTENT ─────────────────────────────────────────────────")
        if intent.get('_market_source') == 'agency_portfolio':
            acct = intent.get('agency_account_code', '')
            acct_str = f'  [{acct}]' if acct else ''
            ambig_str = ''
            if intent.get('agency_ambiguous'):
                others = ', '.join(f"{a[0]} {a[1]}" for a in intent.get('agency_all_accounts', [])[:3])
                ambig_str = f'  ⚠ multiple accounts: {others}'
            city_str = ''
            if intent.get('_city_confirmed_market') == intent['source_market']:
                city_str = f'  ✓ city:{intent["_city_found"]}'
            elif intent.get('_city_confirmed_market'):
                city_str = f'  ⚠ city:{intent["_city_found"]}→{intent["_city_confirmed_market"]}'
            print(f"  Market       : {intent['source_market']}  "
                  f"✓ {intent.get('agency_name', '')} ({intent.get('agency_domain', '')})"
                  f"{acct_str}{city_str}{ambig_str}")
        elif intent.get('_market_source') == 'city_lookup':
            print(f"  Market       : {intent['source_market']}  "
                  f"[city lookup — {intent.get('_city_found', '')}  "
                  f"(no portfolio match — consider adding agency)]")
        else:
            conflict = intent.get('_city_market_conflict', '')
            conflict_str = f'  ⚠ {conflict}' if conflict else ''
            print(f"  Market       : {intent['source_market']}  "
                  f"[keyword scoring{' — consider adding to agency portfolio' if intent['source_market'] == 'UNKNOWN' else ''}"
                  f"]{conflict_str}")
        if intent.get('num_children'):
            _adults = intent.get('num_adults', intent['num_pax'])
            _kids = intent['num_children']
            _total = intent.get('num_pax_total', _adults + _kids)
            print(f"  Pax          : {_total} ({_adults} adults + {_kids} children)")
        else:
            print(f"  Pax          : {intent['num_pax']}")
        print(f"  Rooms        : {intent['rooms']}")
        print(f"  Hotel Tier   : {intent['tier'] or 'not specified — will show top booked'}")
        if intent.get('duration_nights'):
            dn = intent['duration_nights']
            print(f"  Duration     : {dn[0]}-{dn[1]} nights" if dn[0] != dn[1] else f"  Duration     : {dn[0]} nights")
            # B15: when body-stated duration differs from PNR calendar gap,
            # surface the difference. Engine prioritises body-stated; consultant
            # decides whether to enrich with the calendar slack or leave as-is.
            if intent.get('_pnr_calendar_gap') and intent['_pnr_calendar_gap'] != dn[1]:
                _gap = intent['_pnr_calendar_gap']
                _diff = _gap - dn[1]
                if _diff > 0:
                    print(f"  ℹ PNR calendar window: {_gap} nights available — agent specified {dn[1]}n.")
                    print(f"    Engine using agent's {dn[1]}n. {_diff} night(s) slack in the PNR — buffer/rest day, or agent may accept enrichment.")
                else:
                    print(f"  ⚠ PNR calendar window: only {_gap} nights available — agent specified {dn[1]}n.")
                    print(f"    Agent's request exceeds the PNR window by {-_diff} night(s) — consider clarifying with agent.")
        print(f"  Heritage     : {'YES — ' + str(intent['heritage_keywords']) if intent['heritage_requested'] else 'No'}")
        # B24 banner: announce heritage tier upgrade for transparency
        if intent.get('_tier_upgraded_for_heritage'):
            _prior = intent.get('_tier_before_upgrade') or 'not specified'
            print(f"  ℹ Tier upgraded to first-class — heritage hotels requested.")
            print(f"    (Originally {_prior}; heritage properties are typically first-class or above.)")
        print(f"  Language Guide: {intent['language_guide'] or 'Not requested'}")
        print(f"  Budget       : {intent['budget_raw'] or 'Not specified'}")
        print(f"  Cities       : {', '.join(intent['cities_detected']) or 'Not detected'}")
        # B16: when PNR cities were skipped due to region mismatch with body,
        # surface this so consultant knows the engine treated them as gateways
        if intent.get('_pnr_gateway_only'):
            _gw_list = ', '.join(f"{c} ({r.replace('_',' ')})" for c, r in intent['_pnr_gateway_only'])
            _body_r = ', '.join(intent.get('regions_detected') or [])
            print(f"  ⚠ PNR/region check: PNR includes {_gw_list} but agent specifies {_body_r}.")
            print(f"    Engine treating PNR city/cities as international gateways only — NOT touring cities.")
            print(f"    Domestic flight will likely be needed (e.g. DEL→MAA). Confirm with agent if Delhi is a touring stop.")
        if intent.get('cities_excluded'):
            print(f"  ⚠ Excluded   : {', '.join(intent['cities_excluded'])}  (previously visited — removed from recommendations)")
        if intent.get('travel_month_approx'):
            print(f"  Travel Period: {intent['travel_month_approx']}  [LLM — from relative date in email]")
        if intent.get('_previously_visited_india'):
            print(f"  ⚠ Revisit    : Client has visited India before — consider alternative routes")
        if intent.get('regions_detected'):
            print(f"  Regions      : {', '.join(intent['regions_detected'])}")
        if intent.get('sub_region'):
            print(f"  Sub-region   : {intent['sub_region']}  ← overlap scoring uses this city set")

        # ── Always show confidence score inside PARSED INTENT ─────────────────
        _conf, _conf_issues, _conf_score = check_parse_confidence(intent)
        intent['parser_score'] = _conf_score  # expose to frontend
        _conf_icon = '✓' if _conf == 'high' else ('⚠' if _conf == 'medium' else '✗')
        _llm_note  = '  [LLM fallback invoked]' if intent.get('_llm_used') else (
                     '  [LLM fallback will fire]' if _conf == 'low' else '')
        if _conf_issues:
            print(f"  Parser Score : {_conf_icon} {_conf.upper()} ({_conf_score}/100){_llm_note}")
            print(f"  Issues found : {', '.join(_conf_issues)}")
        else:
            print(f"  Parser Score : {_conf_icon} HIGH ({_conf_score}/100) — all fields detected")

        # ── Body-date fallback — set travel_start from email body if PNR didn't provide it
        if not intent.get('travel_start'):
            _body_date = _extract_start_date(email_text)
            if _body_date:
                intent['travel_start'] = _body_date
                # Also set travel_end from duration
                _dur = intent.get('duration_nights')
                _nights = _dur[1] if isinstance(_dur, (list,tuple)) else (_dur or 0)
                if _nights:
                    from datetime import datetime as _dtb, timedelta as _tdb
                    _te = _dtb.strptime(_body_date, '%Y-%m-%d') + _tdb(days=_nights)
                    intent['travel_end'] = _te.strftime('%Y-%m-%d')

    # ── Regional defaults — apply when no specific cities detected ────────────
    # Built from top historical routes per market
    REGIONAL_DEFAULTS = {
        'north_india': {
            'GBR': ['New Delhi', 'Agra', 'Jaipur', 'Fatehpur Sikri'],
            'USA': ['New Delhi', 'Agra', 'Jaipur', 'Fatehpur Sikri'],
            'JPN': ['New Delhi', 'Agra', 'Jaipur'],
            'AUS': ['New Delhi', 'Agra', 'Jaipur', 'Fatehpur Sikri'],
            'DEU': ['New Delhi', 'Mandawa', 'Bikaner', 'Jodhpur', 'Udaipur (RJ)', 'Jaipur', 'Agra'],
            'FRA': ['New Delhi', 'Mandawa', 'Bikaner', 'Jodhpur', 'Pushkar', 'Jaipur', 'Agra'],
            'ITA': ['New Delhi', 'Mandawa', 'Bikaner', 'Jaisalmer', 'Jodhpur', 'Udaipur (RJ)', 'Jaipur', 'Agra'],
            'ESP': ['New Delhi', 'Jaipur', 'Agra', 'Khajuraho', 'Varanasi'],
            'CHE': ['New Delhi', 'Mandawa', 'Bikaner', 'Jodhpur', 'Udaipur (RJ)', 'Jaipur', 'Agra'],
            'NLD': ['New Delhi', 'Agra', 'Jaipur', 'Fatehpur Sikri'],
            'DEFAULT': ['New Delhi', 'Agra', 'Jaipur', 'Fatehpur Sikri'],
        },
        'south_india': {
            'GBR': [
                [('Kochi (Cochin)',2),('Alappuzha (Alleppey)',1),('Thekkady (Periyar/Kumily)',2),('Munnar',2)],
            ],
            'DEU': [
                [('Chennai',2),('Mamallapuram (Mahabalipuram)',1),('Thanjavur (Tanjore)',1),('Madurai',2),('Thekkady (Periyar/Kumily)',2),('Kumarakom',1),('Kochi (Cochin)',2)],
            ],
            'FRA': [
                [('Bengaluru',1),('Mysuru (Mysore)',1),('Madurai',2),('Thekkady (Periyar/Kumily)',2),('Munnar',2),('Alappuzha (Alleppey)',1),('Kochi (Cochin)',2)],
                [('Chennai',2),('Mamallapuram (Mahabalipuram)',1),('Thanjavur (Tanjore)',1),('Madurai',2),('Thekkady (Periyar/Kumily)',2),('Kochi (Cochin)',2)],
            ],
            'ITA': [
                [('Chennai',2),('Mamallapuram (Mahabalipuram)',1),('Thanjavur (Tanjore)',1),('Madurai',2),('Thekkady (Periyar/Kumily)',2),('Kumarakom',1),('Kochi (Cochin)',2)],
            ],
            'DEFAULT': [
                [('Kochi (Cochin)',2),('Alappuzha (Alleppey)',1),('Thekkady (Periyar/Kumily)',2),('Munnar',2)],
            ],
        },
        'himalaya': {
            'GBR': ['New Delhi', 'Shimla', 'Chandigarh', 'Amritsar'],
            'DEFAULT': ['New Delhi', 'Shimla', 'Chandigarh', 'Amritsar'],
        },
        'central_india': {
            'DEFAULT': ['New Delhi', 'Khajuraho', 'Varanasi'],
        },
    }

    # ── Wildlife routes — drawn from DB top routes per wildlife zone ────────────
    # Format: list of (city, nights) tuples. Jabalpur = 0n (flight gateway only).
    # Zones: ranthambore | central_india | corbett | south_india | northeast_india
    WILDLIFE_ROUTES = {
        'ranthambore': {
            'label': 'North India — Ranthambore  (530 tours in DB)',
            'DEFAULT': [
                # 8n — Golden Triangle + Ranthambore (most booked: 31x)
                [('New Delhi',3),('Agra',1),('Ranthambore',2),('Fatehpur Sikri',0),('Jaipur',2)],
                # 10n — GT + Ranthambore + Rajasthan extension
                [('New Delhi',3),('Agra',1),('Ranthambore',2),('Fatehpur Sikri',0),('Jaipur',2),('Jodhpur',2)],
                # 12n — full circuit with Udaipur
                [('New Delhi',3),('Agra',1),('Ranthambore',2),('Fatehpur Sikri',0),('Jaipur',2),('Jodhpur',1),('Udaipur (RJ)',2)],
            ],
        },
        'central_india': {
            'label': 'Central India — Dedicated wildlife  (Bandhavgarh 42 · Kanha 35 · Pench 23 tours)',
            'DEFAULT': [
                # 7n — 1 park with Khajuraho culture stop
                [('New Delhi',2),('Khajuraho',1),('Jabalpur',0),('Bandhavgarh',3),('Jabalpur',0)],
                # 10n — 2 parks
                [('New Delhi',2),('Khajuraho',1),('Jabalpur',0),('Bandhavgarh',3),('Kanha',3),('Jabalpur',0)],
                # 13n — 3 parks
                [('New Delhi',2),('Khajuraho',1),('Jabalpur',0),('Bandhavgarh',3),('Kanha',3),('Pench',3),('Jabalpur',0)],
            ],
        },
        'corbett': {
            'label': 'North India — Jim Corbett  (19 tours in DB)',
            'DEFAULT': [
                # 6n — Corbett + Delhi
                [('New Delhi',2),('Corbett',3),('Rishikesh',1)],
                # 8n — with Himalaya extension
                [('New Delhi',2),('Corbett',3),('Rishikesh',2),('Haridwar',1)],
            ],
        },
        'south_india': {
            'label': 'South India — Kabini + Wayanad  (Kabini 44 · Wayanad 35 tours)',
            'DEFAULT': [
                # 7n — Karnataka wildlife circuit
                [('Bengaluru',1),('Mysuru (Mysore)',1),('Kabini',2),('Wayanad',2),('Kochi (Cochin)',1)],
                # 10n — with Kerala culture
                [('Bengaluru',1),('Mysuru (Mysore)',1),('Kabini',2),('Wayanad',2),
                 ('Kochi (Cochin)',1),('Alappuzha (Alleppey)',1),('Thekkady (Periyar/Kumily)',2)],
            ],
        },
        'northeast_india': {
            'label': 'North East — Kaziranga  (6 tours in DB)',
            'DEFAULT': [
                # 7n — Kaziranga + Meghalaya
                [('Guwahati',1),('Kaziranga',3),('Shillong',2),('Cherrapunji',1)],
                # 5n — shorter
                [('Guwahati',1),('Kaziranga',3),('Shillong',1)],
            ],
        },
    }

    # ── Regional defaults — built from top historical routes per market
    # Format: list of (city, nights) tuples ordered by route sequence
    REGIONAL_ROUTES = {
        'north_india': {
            'GBR': [
                # 7n — Golden Triangle classic (most booked GBR)
                [('New Delhi',3),('Agra',2),('Fatehpur Sikri',0),('Jaipur',2)],
                # 8n — Golden Triangle + Ranthambore
                [('New Delhi',3),('Agra',1),('Ranthambore',2),('Fatehpur Sikri',0),('Jaipur',2)],
                # 9n — GT + Udaipur extension
                [('New Delhi',3),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2),('Udaipur (RJ)',3)],
                # 10n — GT + Varanasi spiritual extension
                [('New Delhi',3),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2),('Varanasi',3),('Khajuraho',1)],
                # 10n — GT + Ranthambore + Udaipur
                [('New Delhi',3),('Agra',1),('Ranthambore',2),('Fatehpur Sikri',0),('Jaipur',2),('Udaipur (RJ)',2)],
                # 11n — Grand North India — Delhi + Rajasthan + Varanasi
                [('New Delhi',3),('Jaipur',2),('Jodhpur',2),('Udaipur (RJ)',2),('Agra',1),('Varanasi',1)],
                # 13n — Full circuit — Delhi + Rajasthan + Central + Varanasi
                [('New Delhi',3),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2),('Ranthambore',2),('Udaipur (RJ)',2),('Varanasi',3)],
            ],
            'USA': [
                [('New Delhi',3),('Agra',1),('Ranthambore',2),('Jaipur',2)],
                [('New Delhi',4),('Agra',2),('Fatehpur Sikri',0),('Jaipur',2)],
            ],
            'DEU': [
                [('New Delhi',2),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2),('Ranthambore',2)],
                [('New Delhi',3),('Mandawa',1),('Jaipur',2),('Fatehpur Sikri',0),('Agra',1)],
            ],
            'FRA': [
                # 10n classic Rajasthan with GT
                [('New Delhi',2),('Mandawa',1),('Bikaner',1),
                 ('Jaisalmer',2),('Jodhpur',1),('Pushkar',1),('Jaipur',2),('Agra',1)],
                # 10n Jodhpur-Udaipur with GT
                [('New Delhi',2),('Mandawa',1),('Jodhpur',2),('Udaipur (RJ)',2),('Jaipur',2),('Agra',1)],
                # 12n GT-visited variant — Rajasthan without Agra/Jaipur + Varanasi
                [('New Delhi',2),('Mandawa',1),('Bikaner',1),
                 ('Jaisalmer',2),('Jodhpur',2),('Udaipur (RJ)',2),('Varanasi',2)],
                # 14n full Rajasthan + GT — Jaipur + Agra + Varanasi + Udaipur
                [('New Delhi',2),('Mandawa',1),('Bikaner',1),
                 ('Jaisalmer',2),('Jodhpur',2),('Udaipur (RJ)',2),
                 ('Jaipur',2),('Fatehpur Sikri',0),('Agra',1),('Varanasi',1)],
            ],
            'ITA': [
                [('New Delhi',3),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2),('Udaipur (RJ)',2)],
            ],
            'JPN': [
                [('New Delhi',3),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2)],
            ],
            'IND': [
                # 5n compact Golden Triangle with Fatehpur Sikri day visit
                [('New Delhi',2),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2)],
                # 7n GT + Udaipur extension
                [('New Delhi',2),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2),('Udaipur (RJ)',2)],
                # 9n GT + Udaipur + Jodhpur
                [('New Delhi',2),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2),('Udaipur (RJ)',2),('Jodhpur',2)],
            ],
            'DEFAULT': [
                # 5n compact Golden Triangle with Fatehpur Sikri day visit
                [('New Delhi',2),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2)],
                # 7n GT + Udaipur suggestion
                [('New Delhi',2),('Agra',1),('Fatehpur Sikri',0),('Jaipur',2),('Udaipur (RJ)',2)],
                # 7n with Ranthambore
                [('New Delhi',3),('Agra',1),('Ranthambore',2),('Jaipur',2)],
            ],
        },
        'south_india': {
            'DEU': [
                # 7n Kerala-first — primary when GT visited or shorter duration
                [('Kochi (Cochin)',2),('Alappuzha (Alleppey)',1),
                 ('Thekkady (Periyar/Kumily)',2),('Madurai',1),
                 ('Mamallapuram (Mahabalipuram)',0),('Puducherry (Pondicherry)',1),
                 ('Chennai',0)],
                # 9n full Tamil Nadu + Kerala with Mamallapuram
                [('Chennai',1),('Mamallapuram (Mahabalipuram)',1),
                 ('Puducherry (Pondicherry)',1),('Thanjavur (Tanjore)',1),
                 ('Madurai',1),('Thekkady (Periyar/Kumily)',2),
                 ('Kumarakom',1),('Kochi (Cochin)',1)],
            ],
            'FRA': [
                # 7n Kerala-first with Chola temple trail enroute to Puducherry
                [('Kochi (Cochin)',2),('Alappuzha (Alleppey)',1),
                 ('Thekkady (Periyar/Kumily)',2),('Madurai',1),
                 ('Darasuram',0),('Gangaikondacholapuram',0),('Chidambaram',0),
                 ('Puducherry (Pondicherry)',1),
                 ('Mamallapuram (Mahabalipuram)',0),('Chennai',0)],
                # 9n full Tamil Nadu + Kerala with Chola temple trail
                [('Chennai',1),('Mamallapuram (Mahabalipuram)',1),
                 ('Puducherry (Pondicherry)',1),
                 ('Chidambaram',0),('Gangaikondacholapuram',0),('Darasuram',0),
                 ('Thanjavur (Tanjore)',1),
                 ('Madurai',1),('Thekkady (Periyar/Kumily)',2),('Kochi (Cochin)',2)],
            ],
            'GBR': [
                # 7n Kerala classic — Kochi + backwaters + Thekkady + Munnar
                [('Kochi (Cochin)',2),('Alappuzha (Alleppey)',1),
                 ('Thekkady (Periyar/Kumily)',2),('Munnar',2)],
                # 9n Kerala + Madurai
                [('Kochi (Cochin)',2),('Alappuzha (Alleppey)',1),
                 ('Thekkady (Periyar/Kumily)',2),('Madurai',2),('Kochi (Cochin)',2)],
                # 10n Tamil Nadu + Kerala — Chennai entry
                [('Chennai',1),('Mamallapuram (Mahabalipuram)',1),
                 ('Thanjavur (Tanjore)',1),('Madurai',2),
                 ('Thekkady (Periyar/Kumily)',2),('Munnar',2),('Kochi (Cochin)',1)],
                # 12n full south — Tamil Nadu + Chola trail + Kerala
                [('Chennai',1),('Mamallapuram (Mahabalipuram)',1),
                 ('Puducherry (Pondicherry)',1),
                 ('Thanjavur (Tanjore)',1),('Madurai',2),
                 ('Thekkady (Periyar/Kumily)',2),('Kumarakom',1),
                 ('Alappuzha (Alleppey)',1),('Kochi (Cochin)',2)],
                # 7n Karnataka + Kerala — Bengaluru entry
                [('Bengaluru',1),('Mysuru (Mysore)',1),
                 ('Thekkady (Periyar/Kumily)',2),('Munnar',2),
                 ('Kochi (Cochin)',1)],
            ],
            'DEFAULT': [
                # 7n Kerala classic
                [('Kochi (Cochin)',2),('Alappuzha (Alleppey)',1),
                 ('Thekkady (Periyar/Kumily)',2),('Munnar',2)],
                # 9n Kerala + Madurai
                [('Kochi (Cochin)',2),('Alappuzha (Alleppey)',1),
                 ('Thekkady (Periyar/Kumily)',2),('Madurai',2),('Kochi (Cochin)',2)],
                # 10n Tamil Nadu + Kerala
                [('Chennai',1),('Mamallapuram (Mahabalipuram)',1),
                 ('Thanjavur (Tanjore)',1),('Madurai',2),
                 ('Thekkady (Periyar/Kumily)',2),('Munnar',2),('Kochi (Cochin)',1)],
                # 7n Karnataka + Kerala
                [('Bengaluru',1),('Mysuru (Mysore)',1),
                 ('Thekkady (Periyar/Kumily)',2),('Munnar',2),
                 ('Kochi (Cochin)',1)],
            ],
        },
        'himalaya': {
            'GBR': [
                [('New Delhi',2),('Shimla',2),('Chandigarh',1),('Amritsar',2)],
            ],
            'DEFAULT': [
                [('New Delhi',2),('Shimla',2),('Chandigarh',1),('Amritsar',2)],
            ],
        },
        'central_india': {
            'DEFAULT': [
                [('New Delhi',2),('Khajuraho',2),('Varanasi',3)],
            ],
        },
        'central_india': {
            # Delhi → Jabalpur (flight) → wildlife parks → Mumbai
            # Jabalpur = 0n flight gateway, not a sightseeing stop
            'GBR': [
                # 3 parks via Jabalpur (11n) — Bandhavgarh + Kanha + Pench
                [('New Delhi',2),('Jabalpur',0),('Bandhavgarh',2),('Kanha',2),('Pench',2),('Mumbai',2)],
                # 2 parks via Jabalpur (9n) — Kanha + Pench
                [('New Delhi',2),('Jabalpur',0),('Kanha',2),('Pench',2),('Mumbai',2)],
                # Classic with Khajuraho temples + 2 parks (12n)
                [('New Delhi',2),('Agra',1),('Khajuraho',2),('Jabalpur',0),('Kanha',2),('Pench',2),('Mumbai',2)],
                # Bandhavgarh + Kanha (9n)
                [('New Delhi',2),('Jabalpur',0),('Bandhavgarh',2),('Kanha',2),('Mumbai',2)],
            ],
            'DEFAULT': [
                [('New Delhi',2),('Jabalpur',0),('Bandhavgarh',2),('Kanha',2),('Pench',2),('Mumbai',2)],
                [('New Delhi',2),('Jabalpur',0),('Kanha',2),('Pench',2),('Mumbai',2)],
                [('New Delhi',2),('Agra',1),('Khajuraho',2),('Jabalpur',0),('Kanha',2),('Pench',2),('Mumbai',2)],
            ],
        },
    }

    # ── STRUCTURED SECTION ITINERARY — overrides all route building ────────────
    # When agent listed each city+nights on its own line, use that directly
    if intent.get('nightly_split_agent') and not intent.get('nightly_split'):
        agent_route = intent['nightly_split_agent']
        # Insert Fatehpur Sikri and apply transit rules
        agent_route = insert_fatehpur_sikri(agent_route)
        agent_route = fix_transit_stops(agent_route, market=intent.get('source_market'))
        intent['nightly_split'] = agent_route
        intent['all_routes'] = [agent_route]
        total = sum(n for _, n in agent_route)
        print(f"\n  ⚡ Agent section itinerary applied ({total} nights):")
        for city, nights in agent_route:
            if nights == 0:
                print(f"     {city:35} day visit (en route)")
            else:
                print(f"     {city:35} {nights}n")

    # ── CITY-BASED ROUTE BUILDER ──────────────────────────────────────────────
    # When agent specifies cities explicitly, build route from those cities
    # with proportional night allocation — skip if bookend already built a route

    # ── Wildlife multi-region: fires FIRST before any city injection ───────────
    # When wildlife keyword + no explicit cities → show all zones
    # P2-04 (21 May 2026 v20260521b): tightened trigger to disambiguate from
    # tourist landmarks. Bare 'tiger' false-positived on "Tiger's Nest"
    # (Paro Taktsang, Bhutan monastery) in B3#003 and likely B2#009. Bare
    # 'jungle' could match metaphors ("concrete jungle"). Now requires
    # wildlife-context phrases. Three confirmed reproductions before fix.
    import re as _re_wl
    _email_lc = email_text.lower()
    _wl_strong_kw = any(s in _email_lc for s in [
        'wildlife', 'wild life', 'safari', 'leopard', 'rhino',
        'game drive', 'national park', 'tiger reserve', 'tiger safari',
        'tiger sighting', 'tigress',
    ])
    # Jungle: only fire when paired with a wildlife noun (safari/lodge/camp/etc.)
    _wl_jungle_kw = bool(_re_wl.search(
        r'\bjungle\s+(?:safari|lodge|camp|trek|drive|resort|stay|experience)\b',
        _email_lc
    ))
    _wl_requested_email = _wl_strong_kw or _wl_jungle_kw
    _wl_no_cities = len(intent.get('cities_detected', [])) < 2

    if _wl_requested_email and _wl_no_cities and not intent.get('nightly_split'):
        intent['_wildlife_multi_region'] = True
        _culture_kw = any(s in email_text.lower() for s in [
            'culture', 'cultural', 'heritage', 'monument', 'history', 'historical',
            'temple', 'fort', 'palace'
        ])

        # Build primary route from Ranthambore (most booked)
        _primary_wl = WILDLIFE_ROUTES['ranthambore']['DEFAULT'][0]
        intent['nightly_split'] = list(_primary_wl)
        intent['cities_source'] = 'wildlife_multi_region'

        # Print all wildlife options — compact one line per route
        print(f"\n  🐯 Wildlife interest detected — showing options across all regions")
        if _culture_kw:
            print(f"  ✓ Culture keyword detected — cultural stops included in routes")
        _opt_n = 1
        for zone, zdata in WILDLIFE_ROUTES.items():
            print(f"\n  {zdata['label']}")
            print(f"  {'─' * len(zdata['label'])}")
            for route in zdata['DEFAULT']:
                _rt = sum(n for _, n in route if n > 0)
                # Build compact city sequence — skip transit stops (0n) and duplicates
                _seen = []
                for city, nights in route:
                    short = city.replace(' (Cochin)', '').replace(' (Mysore)', '') \
                                .replace(' (Alleppey)', '').replace(' (RJ)', '') \
                                .replace('(Periyar/Kumily)', '').replace('Mamallapuram (Mahabalipuram)', 'Mamallapuram') \
                                .replace('Puducherry (Pondicherry)', 'Pondicherry') \
                                .replace('Thanjavur (Tanjore)', 'Thanjavur')
                    if short not in _seen:
                        _seen.append(short)
                _route_str = ' → '.join(_seen)
                print(f"  Option {_opt_n:2}  {_rt:2}n   {_route_str}")
                _opt_n += 1

        # Collect all cities across all zones for hotel/activity recommendations
        _all_wl_cities = []
        for zone, zdata in WILDLIFE_ROUTES.items():
            for route in zdata['DEFAULT']:
                for city, nights in route:
                    if city not in _all_wl_cities and city != 'Jabalpur':
                        _all_wl_cities.append(city)
        intent['cities_detected'] = _all_wl_cities
        intent['all_routes'] = [
            list(r) for zdata in WILDLIFE_ROUTES.values() for r in zdata['DEFAULT']
        ]
    _wl_cities_known = {'Bandhavgarh', 'Kanha', 'Pench', 'Panna', 'Corbett',
                        'Jim Corbett', 'Ranthambore', 'Kaziranga', 'Tadoba',
                        'Kabini', 'Nagarhole', 'Wayanad', 'Satpura'}
    _wl_already = bool(_wl_cities_known & set(intent.get('cities_detected', [])))
    _wl_requested = any(s in email_text.lower() for s in [
        'wildlife', 'wild life', 'safari', 'tiger', 'jungle', 'national park'
    ])
    # Only inject wildlife cities if NO explicit cities were detected
    _has_explicit_cities = len(intent.get('cities_detected', [])) >= 3
    # Only fire central_india injection if the new multi-region block hasn't already handled it
    if (_wl_requested and not _wl_already and not _has_explicit_cities
            and 'central_india' in (intent.get('regions_detected') or [])
            and not intent.get('_wildlife_multi_region')):
        _wl_np = re.compile(
            r'(\d+)\s*nights?\s+(?:in\s+)?(?:central\s+india|wildlife|the\s+parks?|jungle)',
            re.IGNORECASE
        )
        _wl_m = _wl_np.search(email_text)
        if _wl_m:
            _wl_nights = int(_wl_m.group(1))
        else:
            _dn = intent.get('duration_nights')
            _total_n = _dn[1] if isinstance(_dn, tuple) else (_dn or 0)
            _gw_n = sum(2 for c in intent.get('cities_detected', [])
                        if c in {'Mumbai', 'New Delhi', 'Chennai', 'Bengaluru', 'Kolkata'}) +                     sum(1 for c in intent.get('cities_detected', []) if c in {'Agra'})
            _wl_nights = max(0, _total_n - _gw_n)

        # Detect "culture" keyword — add Agra + Khajuraho to wildlife routes
        _culture_requested = any(s in email_text.lower() for s in [
            'culture', 'cultural', 'heritage', 'monument', 'history', 'historical'
        ])

        if _wl_nights == 0:
            # No duration specified — show all 3 park options, let consultant pick
            intent['_wildlife_no_duration'] = True
            intent['_wildlife_culture'] = _culture_requested
            _wl_inject = ['Bandhavgarh']  # primary route uses 1 park
            print(f"  🐯 No duration specified — will show 1, 2 and 3-park options")
        elif _wl_nights >= 8:
            _wl_inject = ['Bandhavgarh', 'Kanha', 'Pench']
        elif _wl_nights >= 6:
            _wl_inject = ['Bandhavgarh', 'Kanha']
        else:
            _wl_inject = ['Bandhavgarh']
        for _wc in _wl_inject:
            if _wc not in intent.get('cities_detected', []):
                intent.setdefault('cities_detected', []).append(_wc)
        # Also add Jabalpur as flight gateway
        if 'Jabalpur' not in intent.get('cities_detected', []):
            intent['cities_detected'].append('Jabalpur')
        print(f"  🐯 Wildlife cities injected ({_wl_nights}n): {', '.join(_wl_inject)}")

        # ── If no gateway city at all, default to Delhi entry ─────────────────
        _gateways_known = {'Mumbai', 'New Delhi', 'Chennai', 'Bengaluru', 'Kolkata', 'Kochi (Cochin)'}
        if not bool(_gateways_known & set(intent.get('cities_detected', []))):
            intent.setdefault('cities_detected', []).insert(0, 'New Delhi')
            print("  ⚠ No gateway city in email — defaulting to Delhi entry. Confirm with agent.")

        # ── Force correct travel order for gateway + wildlife routes ───────────
        # Detect which gateway city is entry vs exit from email context
        _detected = intent.get('cities_detected', [])
        _email_lower = email_text.lower()

        # Detect entry gateway (city mentioned first / "arrive in X" / "start in X")
        _gateways = ['Mumbai', 'New Delhi', 'Chennai', 'Bengaluru', 'Kolkata', 'Kochi (Cochin)']
        _entry_gw = None
        _exit_gw  = None

        # Check explicit "arrive in X" / "start in X" language
        for gw in _gateways:
            gw_lower = gw.lower().replace(' (cochin)', '').replace('new ', '')
            if any(p in _email_lower for p in [
                f'arrive in {gw_lower}', f'arriving in {gw_lower}',
                f'arrive at {gw_lower}', f'start in {gw_lower}',
                f'will arrive in {gw_lower}', f'fly into {gw_lower}',
                f'guest will arrive in {gw_lower}',
            ]):
                _entry_gw = gw
            if any(p in _email_lower for p in [
                f'moving to {gw_lower}', f'ending in {gw_lower}',
                f'fly out of {gw_lower}', f'depart from {gw_lower}',
                f'exit via {gw_lower}',
            ]):
                _exit_gw = gw

        # Fallback: first gateway in cities_detected = entry, last = exit
        _entry_explicit = bool(_entry_gw)  # track if entry was explicitly stated
        if not _entry_gw:
            for c in _detected:
                if c in _gateways:
                    _entry_gw = c
                    break
        if not _exit_gw:
            for c in reversed(_detected):
                if c in _gateways and c != _entry_gw:
                    _exit_gw = c
                    break

        # If entry was NOT explicitly stated and two gateways exist, flag both directions
        if not _entry_explicit and _entry_gw and _exit_gw and _entry_gw != _exit_gw:
            _IATA_GW = {'Mumbai': 'BOM', 'New Delhi': 'DEL', 'Chennai': 'MAA',
                        'Bengaluru': 'BLR', 'Kolkata': 'CCU', 'Kochi (Cochin)': 'COK'}
            _gw1_iata = _IATA_GW.get(_entry_gw, 'DEL')
            _gw2_iata = _IATA_GW.get(_exit_gw, 'BOM')
            print("  ⚠ Entry city not specified — two route options available:")
            print("     Option A: " + _entry_gw + " -> Jabalpur -> Wildlife -> " + _exit_gw + "  [" + _gw1_iata + "->JLR]")
            print("     Option B: " + _exit_gw + " -> Jabalpur -> Wildlife -> " + _entry_gw + "  [" + _gw2_iata + "->JLR]")
            print("     Defaulting to Option A — confirm with agent which direction")

        # Build ordered route: entry → Jabalpur → parks → exit
        if _entry_gw and any(c in _detected for c in ['Bandhavgarh','Kanha','Pench','Panna']):
            _WILDLIFE_PARKS = ['Bandhavgarh', 'Kanha', 'Pench', 'Panna']
            _OTHER = [c for c in _detected
                      if c not in _gateways and c not in _WILDLIFE_PARKS and c != 'Jabalpur']
            _parks_in_route = [c for c in _WILDLIFE_PARKS if c in _detected]

            _ordered = [_entry_gw, 'Jabalpur'] + _parks_in_route + _OTHER
            if _exit_gw and _exit_gw != _entry_gw:
                _ordered.append(_exit_gw)
            # Remove dupes preserving order
            _seen_o = set()
            _final_order = []
            for c in _ordered:
                if c not in _seen_o and c in _detected + ['Jabalpur']:
                    _final_order.append(c)
                    _seen_o.add(c)
            intent['cities_detected'] = _final_order

            # Domestic flight: entry gateway → Jabalpur
            _IATA_GW = {'Mumbai': 'BOM', 'New Delhi': 'DEL', 'Chennai': 'MAA',
                        'Bengaluru': 'BLR', 'Kolkata': 'CCU', 'Kochi (Cochin)': 'COK'}
            _entry_iata = _IATA_GW.get(_entry_gw, 'DEL')
            if 'domestic flights mentioned' not in intent.get('transport_notes', []):
                intent.setdefault('transport_notes', []).append('domestic flights mentioned')
            # Detect correct domestic flight destination based on route cities
            _dom_dest = 'JLR'  # default Jabalpur for wildlife
            _dom_dest_name = 'Jabalpur'
            _route_cities_set = set(intent.get('cities_detected', []))
            if 'Darjeeling' in _route_cities_set or 'Kalimpong' in _route_cities_set:
                _dom_dest = 'IXB'  # Bagdogra (nearest to Darjeeling)
                _dom_dest_name = 'Bagdogra (Darjeeling)'
            elif 'Corbett' in _route_cities_set or 'Jim Corbett' in _route_cities_set:
                _dom_dest = 'IXD'  # Pantnagar (nearest to Corbett)
                _dom_dest_name = 'Pantnagar (Corbett)'
            elif 'Aurangabad (MH)' in _route_cities_set:
                _dom_dest = 'IXU'  # Aurangabad airport
                _dom_dest_name = 'Aurangabad'
            if not intent.get('domestic_sectors'):
                intent['domestic_sectors'] = [{'orig': _entry_iata, 'dest': _dom_dest, 'date': ''}]
                print(f"  ✈ Domestic flight: {_entry_iata} → {_dom_dest} ({_dom_dest_name})")

            print(f"  ✈ Entry: {_entry_gw} → Jabalpur → Wildlife → {_exit_gw or 'exit city'}")
            print(f"  ✈ Domestic flight added: {_entry_iata} → JLR")

    # City-based builder: fires when cities detected, even if duration unknown (uses 14n default)
    # B20: Madhya Pradesh parks have NO airport — Jabalpur (JLR) is the only
    # flight gateway. When agent specifies Bandhavgarh / Kanha / Pench / Panna /
    # Satpura but doesn't mention Jabalpur, inject it as a 0n transit hop so
    # downstream domestic-flight booking + day-use room logic fires correctly.
    # Runs AFTER all explicit-itinerary parsing and wildlife auto-injection,
    # catching cases where the agent listed parks without naming Jabalpur as
    # a flight gateway (operationally required — only airport for these parks).
    _MP_PARKS = {'Bandhavgarh', 'Kanha', 'Pench', 'Panna', 'Satpura'}
    _has_mp_park = bool(_MP_PARKS & set(intent.get('cities_detected', [])))
    if _has_mp_park and 'Jabalpur' not in intent.get('cities_detected', []):
        intent.setdefault('cities_detected', []).append('Jabalpur')
        intent['_jabalpur_auto_added'] = True
        print(f"  ✈ Jabalpur added as flight gateway for MP parks (0n transit)")

    _has_duration = bool(intent.get('duration_nights'))
    if not _has_duration and len(intent.get('cities_detected', [])) >= 3:
        # Estimate duration from city count: ~2n per city
        _est = len(intent['cities_detected']) * 2
        intent['duration_nights'] = (_est, _est)
        if not intent.get('_wildlife_multi_region'):
            print(f"  ℹ  Duration estimated from city count: {_est} nights")
    if (not intent.get('nightly_split')) and len(intent['cities_detected']) >= 2 and intent.get('duration_nights'):
        specified_cities = [
            c for c in intent['cities_detected']
            if c not in set(intent.get('cities_excluded', []))
        ]
        # Samode Haveli is near Jaipur — merge into Jaipur if both present
        if 'Samode' in specified_cities and 'Jaipur' in specified_cities:
            specified_cities = [c for c in specified_cities if c != 'Samode']
            print("  ℹ Samode merged into Jaipur (Samode Haveli is near Jaipur)")
        duration = intent['duration_nights']
        min_n, max_n = duration if isinstance(duration, tuple) else (duration, duration)
        target_n = max_n

        # Minimum nights per city
        CITY_MIN = {
            'New Delhi': 2, 'Jaipur': 2, 'Udaipur (RJ)': 2, 'Jodhpur': 2,
            'Varanasi': 2, 'Mumbai': 2, 'Kochi (Cochin)': 2, 'Agra': 1,
            'Jaisalmer': 2, 'Bikaner': 1, 'Pushkar': 1, 'Ranthambore': 2,
            'Gwalior': 1, 'Orchha': 1, 'Khajuraho': 2, 'Lucknow': 2,
            'Rohet': 1, 'Narlai': 1, 'Jawai': 2, 'Panna': 2,
            # Wildlife parks — minimum 2 nights for meaningful safari experience
            'Bandhavgarh': 2, 'Kanha': 2, 'Pench': 2, 'Corbett': 2,
            'Kaziranga': 2, 'Satpura': 2,
            'Jabalpur': 0,  # always transit
        }
        # ── ROUTE ENRICHMENT ────────────────────────────────────────────────────
        # When agent specifies key Rajasthan cities, suggest the full classic circuit
        # by adding logical intermediate stops if duration allows
        rajasthan_cities = {'Jaipur', 'Jodhpur', 'Jaisalmer', 'Udaipur (RJ)'}
        SOUTH_CITY_SET = {
            'Bengaluru', 'Mysuru (Mysore)', 'Ooty (Udhagamandalam)', 'Coimbatore',
            'Chennai', 'Mamallapuram (Mahabalipuram)', 'Thanjavur (Tanjore)', 'Madurai',
            'Thekkady (Periyar/Kumily)', 'Munnar', 'Kumarakom', 'Alappuzha (Alleppey)',
            'Kochi (Cochin)', 'Kovalam', 'Varkala', 'Goa', 'Mangalore',
            'Hampi', 'Badami', 'Hyderabad', 'Puducherry (Pondicherry)',
        }
        regions_check = intent.get('regions_detected') or []
        if 'north_india' in regions_check and 'south_india' in regions_check:
            specified_cities = [c for c in specified_cities if c not in SOUTH_CITY_SET]
        detected_set = set(specified_cities)
        # Always ensure Delhi is entry city for north india routes
        if 'north_india' in regions_check and 'New Delhi' not in detected_set:
            specified_cities.append('New Delhi')
            detected_set.add('New Delhi')
        if len(detected_set & rajasthan_cities) >= 3 and target_n >= 10:
            # Full Rajasthan circuit — add Mandawa, Bikaner, Agra if not already present
            enriched = list(specified_cities)
            if 'Mandawa' not in detected_set:
                enriched.append('Mandawa')
            if 'Bikaner' not in detected_set:
                enriched.append('Bikaner')
            if 'Jodhpur' not in detected_set:
                enriched.append('Jodhpur')
            if 'Agra' not in detected_set and target_n >= 13:
                enriched.append('Agra')
            specified_cities = enriched

        # Preferred night distribution for classic Rajasthan circuit
        RAJASTHAN_NIGHTS = {
            'New Delhi': 2, 'Mandawa': 1, 'Bikaner': 1,
            'Jaisalmer': 2, 'Jodhpur': 2, 'Rohet': 1, 'Narlai': 1, 'Jawai': 2, 'Udaipur (RJ)': 2,
            'Gwalior': 1, 'Orchha': 1, 'Khajuraho': 2, 'Lucknow': 2, 'Panna': 2,
            'Jaipur': 2, 'Agra': 1, 'Pushkar': 1, 'Ranthambore': 2,
        }

        # ── JODHPUR → UDAIPUR STOPOVER SUGGESTION ────────────────────────────────
        # When Jodhpur and Udaipur both in route and duration >= 20n,
        # suggest adding Rohet or Narlai (1n) as a break between the two
        jodhpur_set = {'Jodhpur'}
        udaipur_set = {'Udaipur (RJ)'}
        stopover_cities = {'Rohet', 'Narlai', 'Jawai'}
        cities_set = set(specified_cities)
        if (jodhpur_set & cities_set and udaipur_set & cities_set
                and not (stopover_cities & cities_set)
                and target_n >= 18):
            # Suggest Rohet as default stopover (most booked)
            specified_cities = list(specified_cities)
            jdh_idx = specified_cities.index('Jodhpur')
            specified_cities.insert(jdh_idx + 1, 'Narlai')
            print('  💡 Jodhpur→Udaipur: adding Narlai as suggested stopover (also consider Rohet or Jawai)')

        # Sort cities into geographic/logical travel order based on historical patterns
        CITY_ORDER = [
            # Northern gateway + Uttarakhand/Himachal
            'Amritsar', 'Chandigarh', 'New Delhi', 'Gurgaon',
            'Shimla', 'Manali', 'Dharamshala', 'Dalhousie',
            'Rishikesh', 'Haridwar', 'Mussoorie', 'Nainital',
            'Corbett', 'Jim Corbett', 'Dehradun',
            'Leh', 'Srinagar',
            # Rajasthan — classic circuit west-first then south, exit via Agra
            'Mandawa', 'Nawalgarh', 'Alsisar', 'Bikaner',
            'Jaisalmer', 'Jodhpur', 'Osian (RJ)', 'Rohet', 'Narlai',
            'Rohet', 'Narlai', 'Jawai', 'Ranakpur', 'Udaipur (RJ)', 'Chittorgarh', 'Barli',
            'Pushkar', 'Ajmer', 'Jaipur', 'Abhaneri', 'Ranthambore',
            'Agra', 'Fatehpur Sikri',
            # Central/East — geographic order Delhi→Agra→Gwalior→Orchha→Khajuraho→Varanasi→Lucknow
            # Jabalpur (JLR airport) sits before MP parks as flight gateway
            'Gwalior', 'Orchha', 'Khajuraho', 'Panna', 'Jabalpur', 'Bandhavgarh', 'Kanha',
            'Varanasi', 'Sarnath', 'Bodh Gaya', 'Patna', 'Lucknow', 'Prayagraj',
            # South
            'Hyderabad', 'Chennai', 'Mamallapuram (Mahabalipuram)',
            'Thanjavur (Tanjore)', 'Madurai', 'Thekkady (Periyar/Kumily)',
            'Munnar', 'Kumarakom', 'Alappuzha (Alleppey)', 'Kochi (Cochin)',
            # West exit cities (flight connections)
            'Mumbai', 'Aurangabad (MH)', 'Goa', 'Pune',
            # East India — end of circuit
            'Darjeeling', 'Kalimpong', 'Kolkata',
            'Bhubaneswar', 'Puri', 'Konark',
        ]
        def city_sort_key(city):
            try:
                return CITY_ORDER.index(city)
            except ValueError:
                return 999  # unknown cities go to end

        # Sort all cities geographically — UNLESS agent provided an explicit ordered sequence
        if intent.get('inline_city_order') and len(intent['inline_city_order']) >= 3:
            # Use agent's stated city order; append any detected cities not in the sequence
            agent_ordered = intent['inline_city_order']
            remainder = [c for c in specified_cities if c not in agent_ordered]
            specified_cities = agent_ordered + remainder
            # Even with agent order, enforce Corbett → Delhi transit rule
            if 'Corbett' in specified_cities:
                _corbett_idx = specified_cities.index('Corbett')
                _next_city = specified_cities[_corbett_idx + 1] if _corbett_idx + 1 < len(specified_cities) else None
                if _next_city and _next_city != 'New Delhi':
                    specified_cities.insert(_corbett_idx + 1, 'New Delhi')
                    print("  ℹ Corbett → Delhi transit inserted (only exit route)")
            # B20: Jabalpur transit insertion — must sit BEFORE the first MP park.
            # Agent's bullet list typically omits Jabalpur (it's a flight not a stop).
            # Engine inserts it as 0n transit so domestic-flight + day-use logic fires.
            _MP_PARKS_RT = ['Bandhavgarh', 'Kanha', 'Pench', 'Panna', 'Satpura']
            if 'Jabalpur' in specified_cities:
                # Find first MP park in current order
                _first_park_idx = None
                for _i, _c in enumerate(specified_cities):
                    if _c in _MP_PARKS_RT:
                        _first_park_idx = _i
                        break
                _jabalpur_idx = specified_cities.index('Jabalpur')
                # If Jabalpur is not immediately before the first park, move it
                if _first_park_idx is not None and _jabalpur_idx != _first_park_idx - 1:
                    specified_cities = [c for c in specified_cities if c != 'Jabalpur']
                    # Recompute first-park-idx after Jabalpur removal
                    _first_park_idx = next((i for i, c in enumerate(specified_cities)
                                            if c in _MP_PARKS_RT), None)
                    if _first_park_idx is not None:
                        specified_cities.insert(_first_park_idx, 'Jabalpur')
                        if intent.get('_jabalpur_auto_added'):
                            print("  ℹ Jabalpur positioned before first MP park (transit gateway)")
            # If Aurangabad is return flight city, move it to end (before Darjeeling if present)
            _email_raw = intent.get('_raw_text', '')
            if 'Aurangabad (MH)' in specified_cities and 'aurangabad' in email_text.lower():
                if any(p in email_text.lower() for p in ['flight from aurangabad', '05.03']):
                    specified_cities = [c for c in specified_cities if c != 'Aurangabad (MH)'] + ['Aurangabad (MH)']
        elif len(specified_cities) >= 3:
            # Skip geographic sort if Mumbai+wildlife — order set by wildlife injection
            _is_mumbai_wl = (
                'Mumbai' in specified_cities and
                any(c in specified_cities for c in ['Bandhavgarh', 'Kanha', 'Pench', 'Panna'])
            )
            if not _is_mumbai_wl:
                specified_cities = sorted(specified_cities, key=city_sort_key)
                # Post-sort: move exit-flight city (Aurangabad) to end
                _exit_flight_cities = ['Aurangabad (MH)', 'Mumbai', 'Goa']
                for _efc in _exit_flight_cities:
                    if _efc in specified_cities and specified_cities[-1] != _efc:
                        # Check if this is the return flight city (mentioned near 'return'/'flight from')
                        _efc_lower = _efc.lower().replace(' (mh)', '')
                        if any(p in email_text.lower() for p in [f'flight from {_efc_lower}', f'return.*{_efc_lower}']):
                            specified_cities = [c for c in specified_cities if c != _efc] + [_efc]

        # Allocate nights — use Rajasthan-specific distribution when applicable
        def get_nights(city):
            if len(detected_set & rajasthan_cities) >= 3:
                return RAJASTHAN_NIGHTS.get(city, CITY_MIN.get(city, 1))
            return CITY_MIN.get(city, 1)

        # Allocate minimum nights first
        route = [(c, get_nights(c)) for c in specified_cities]
        total = sum(n for _, n in route)
        # Distribute remaining nights to cities that benefit from more time
        remaining = target_n - total
        CITY_MAX = {
            'New Delhi': 3, 'Jaipur': 3, 'Udaipur (RJ)': 3, 'Jodhpur': 2,
            'Varanasi': 3, 'Jaisalmer': 2, 'Khajuraho': 2, 'Agra': 1,
            'Mandawa': 1, 'Bikaner': 2, 'Pushkar': 2, 'Ranthambore': 3,
            'Gwalior': 1, 'Orchha': 1, 'Rohet': 1, 'Narlai': 1, 'Jawai': 2,
            # Himalayan / Ganga towns — typically 1-2n stops
            'Haridwar': 1, 'Rishikesh': 2, 'Shimla': 2, 'Dharamshala': 2,
            'Chandigarh': 1, 'Amritsar': 2, 'Leh': 3,
        }
        SKIP_EXTRA = {'Fatehpur Sikri', 'Agra', 'Orchha', 'Gurgaon', 'Jabalpur'}
        if remaining > 0:
            # Add extra nights to non-transit cities, prioritising last city
            for i in range(len(route) - 1, -1, -1):
                city, nights = route[i]
                if remaining <= 0:
                    break
                max_n = CITY_MAX.get(city, 3)
                if city not in SKIP_EXTRA and nights < max_n:
                    extra = min(remaining, max_n - nights)
                    route[i] = (city, nights + extra)
                    remaining -= extra

        route = insert_delhi_after_corbett(route)
        route = insert_fatehpur_sikri(route)
        route = fix_transit_stops(route, market=intent.get('source_market'))
        total = sum(n for _, n in route if n > 0)

        # Build arrow format: City (2n) → City (day) → City (3n)
        _arrow_parts = []
        for city, nights in route:
            _arrow_parts.append(f"{city} ({nights}n)" if nights > 0 else f"{city} (day)")
        print(f"\n  ⚡ Building route from specified cities ({total}n): {' → '.join(_arrow_parts)}")

        intent['nightly_split'] = route
        intent['all_routes'] = [route]
        for city, nights in route:
            if city not in intent['cities_detected']:
                intent['cities_detected'].append(city)

        # ── Wildlife no-duration: generate 3 park options ─────────────────────
        # When no duration was specified, build 1-park / 2-park / 3-park alternatives
        if intent.get('_wildlife_no_duration'):
            _with_culture = intent.get('_wildlife_culture', False)
            _gw = 'New Delhi'  # gateway
            _culture_cities = [('Agra', 1), ('Khajuraho', 2)] if _with_culture else []

            def _build_wl_route(parks, gw, culture):
                r = [(_gw, 3)]
                if culture:
                    r += culture
                r += [('Jabalpur', 0)]
                r += [(p, 3) for p in parks]
                return r

            _opt1 = _build_wl_route(['Bandhavgarh'], _gw, _culture_cities)
            _opt2 = _build_wl_route(['Bandhavgarh', 'Kanha'], _gw, _culture_cities)
            _opt3 = _build_wl_route(['Bandhavgarh', 'Kanha', 'Pench'], _gw, _culture_cities)
            intent['all_routes'] = [_opt1, _opt2, _opt3]
            intent['nightly_split'] = _opt1
            # Print alternatives
            print(f"\n  Alternative routes:")
            for _oi, _or in enumerate([_opt2, _opt3], 2):
                _ot = sum(n for _, n in _or)
                _wl_parts = [f"{c} ({n}n)" if n > 0 else f"{c} (day)" for c, n in _or]
                print(f"  Option {_oi:2}  {_ot:2}n   {' → '.join(_wl_parts)}")

    # ── No region detected: show both North + South India ─────────────────────
    # When email gives no regional signal and no cities, default to both circuits
    # so the consultant sees both options rather than always defaulting to North India
    if (not intent.get('nightly_split')) and not intent.get('regions_detected') and not len(intent.get('cities_detected', [])) >= 2:
        intent['regions_detected'] = ['north_india', 'south_india']
        intent['_mixed_region_display'] = True
        print('\n  ⚠ No region or cities specified — showing North + South India options')

    # Skip regional defaults when explicit cities already gave us a route
    if (not intent.get('nightly_split')) and intent.get('regions_detected') and not len(intent.get('cities_detected', [])) >= 3:
        _excluded = set(intent.get('cities_excluded', []))
        _gt_excluded = {'Agra', 'Jaipur'}.issubset(_excluded)
        _revisit = intent.get('_previously_visited_india', False)

        # Only treat as explicit north india request if agent uses forward-looking
        # north india terms — NOT golden triangle (that's a past-visit signal)
        _raw = intent.get('_raw_text', '')
        _north_india_explicit = any(
            kw in _raw
            for kw in ['north india', 'northern india', 'rajasthan', 'rajistan',
                       'rajisthan', 'rajstan', 'delhi agra jaipur']
        ) or (
            # Catches "north & south india", "north and south india", "north+south india"
            re.search(r'north\s*[&+/and\s]+\s*south\s+india', _raw, re.IGNORECASE) is not None
        ) or (
            # Catches standalone "north india" even when split by punctuation
            re.search(r'\bnorth\b.{0,20}\bindia\b', _raw, re.IGNORECASE) is not None
        )

        # When GT cities are excluded (specific past-visit detection):
        # replace north India with south India as the primary route
        if _gt_excluded and not _north_india_explicit:
            if 'south_india' not in intent['regions_detected']:
                intent['regions_detected'].append('south_india')
            # Remove north_india from regions so it won't fire
            intent['regions_detected'] = [
                r for r in intent['regions_detected'] if r != 'north_india'
            ]
            print('  ℹ GT cities previously visited — primary route redirected to South India')

        # When general "visiting again" / "revisit" signal detected (no specific cities):
        # keep North India as primary but inject South India into all_routes as alternatives
        elif _revisit and not _gt_excluded and 'south_india' not in intent['regions_detected']:
            intent['_revisit_south_india_alts'] = True
            print('  ℹ Client has visited India before — South India alternatives will be shown')

        # When ONLY south_india is explicitly requested (keyword match),
        # remove north_india to prevent unwanted split itinerary
        _raw_lower = _raw.lower()
        _south_explicit = any(kw in _raw_lower for kw in [
            'south india', 'southern india', 'kerala', 'tamil nadu',
            'tamilnadu', 'tamil naadu', 'karnataka', 'karnatka'
        ])
        _north_explicit = _north_india_explicit
        if _south_explicit and not _north_explicit:
            intent['regions_detected'] = [
                r for r in intent['regions_detected'] if r != 'north_india'
            ]

        # ─── BL2-α (13 May 2026): try cascade as primary route source ───
        # Before falling back to REGIONAL_ROUTES template, ask historical_query
        # for the agent's actual booking pattern. Cascade returns an envelope
        # of scored options; if any options exist, we use them as the primary
        # route + alternatives. REGIONAL_ROUTES template stays as the deepest
        # fallback for empty / exception cases (BL0 Appendix B preference:
        # cascade-with-caveat-banner over template when cascade has any signal).
        cascade_envelope = None
        cascade_routes   = None
        try:
            import historical_query as _hq
            _acct = intent.get('agency_account_code') or intent.get('account_code')
            intent['_agent_total_bookings'] = (
                _hq.get_agent_total_bookings(db_path, _acct) if _acct else 0
            )
            _candidates = _hq.historical_query(intent, db_path=db_path)
            cascade_envelope = _hq.assemble_options(_candidates, intent)
            if cascade_envelope and cascade_envelope.get('options'):
                cascade_routes = _envelope_to_all_routes(
                    cascade_envelope,
                    requested_duration=_req_dur_scalar(intent.get('duration_nights'))
                )
        except Exception as _bl2a_e:
            # Silent fall-through to REGIONAL_ROUTES template — engine continues
            cascade_envelope = None
            cascade_routes   = None

        if cascade_routes:
            # ── BL2-α path: cascade IS the route source ──
            # Stash envelope for reuse by the SIMILAR HISTORICAL ROUTES display
            # section (avoids recomputing) and snapshot the agent's must-include
            # cities BEFORE we overwrite cities_detected with cascade output.
            intent['_cascade_envelope']   = cascade_envelope
            intent['_agent_must_include'] = list(intent.get('cities_detected') or [])
            _primary_route  = cascade_routes[0]
            _primary_cities = []
            for _c, _n in _primary_route:
                if _c not in _primary_cities:
                    _primary_cities.append(_c)
            intent['cities_detected'] = _primary_cities
            intent['cities_source']   = 'cascade'
            intent['nightly_split']   = _primary_route
            intent['all_routes']      = cascade_routes
            # BL2-α banner truth: stash the primary option's source_pass so the
            # display banner can honestly say WHERE the route came from (agent /
            # market / region / global). Cascade options carry source_pass set
            # by historical_query.query_pass(). Default to 'global' if missing.
            _primary_opt = (cascade_envelope.get('options') or [{}])[0]
            intent['_cascade_primary_pass'] = _primary_opt.get('source_pass', 'global')
            if cascade_envelope.get('caveat'):
                # Honour caveat banner per BL0 Appendix B — show confidence signal
                print(f"\n  {cascade_envelope['caveat']}")

        all_legs = []
        if not cascade_routes:
            for region in intent['regions_detected']:
                if region not in REGIONAL_ROUTES:
                    continue
                market_routes = REGIONAL_ROUTES[region]
                routes = market_routes.get(
                    intent['source_market'],
                    market_routes.get('DEFAULT', [])
                )
                if not routes:
                    continue

                # Select best route by requested duration
                duration = intent.get('duration_nights')
                best_route = None
                if duration and isinstance(duration, (list, tuple)) and len(duration) == 2:
                    min_n, max_n = duration
                    target_n = max_n  # prefer filling the requested duration

                    # Pass 1 — exact fit within range
                    for route in routes:
                        total = sum(n for _, n in route)
                        if min_n <= total <= max_n:
                            best_route = route
                            break

                    # Pass 2 — closest route to target (within ±4 nights)
                    if not best_route:
                        scored = []
                        for route in routes:
                            total = sum(n for _, n in route)
                            diff = abs(total - target_n)
                            if diff <= 4:
                                scored.append((diff, total, route))
                        if scored:
                            scored.sort(key=lambda x: (x[0], -x[1]))  # closest first, longer preferred
                            best_route = scored[0][2]

                if not best_route:
                    best_route = routes[0]

                # Apply already-visited exclusions to the route
                _excluded = set(intent.get('cities_excluded', []))
                if _excluded:
                    best_route = [(c, n) for c, n in best_route if c not in _excluded]

                all_legs.append((region, best_route, routes))

        # Merge all legs into a single route for hotel/activity lookup
        if all_legs:
            # B12: snapshot agent's must-include cities BEFORE regional default
            # overwrites cities_detected with template cities. These are the
            # cities the agent explicitly named (e.g. "Chidambaram") that we
            # need to inject into routes downstream.
            _agent_must_include = list(intent.get('cities_detected') or [])

            merged_cities = []
            merged_route = []
            for region, best_route, routes in all_legs:
                for city, nights in best_route:
                    if city not in merged_cities:
                        merged_cities.append(city)
                        merged_route.append((city, nights))

            intent['cities_detected'] = merged_cities
            intent['_agent_must_include'] = _agent_must_include  # preserved for B12
            intent['cities_source'] = 'regional_default'
            intent['nightly_split'] = merged_route
            # Filter excluded cities from all stored routes too
            _excl = set(intent.get('cities_excluded', []))
            # Build all_routes — deduplicate by route fingerprint
            # Also sort so the route closest to requested duration appears first
            _dur_target = intent.get('duration_nights')
            _target_n   = (_dur_target[1] if isinstance(_dur_target, (list,tuple)) else (_dur_target or 99))
            _raw_routes = [
                tuple((c, n) for c, n in r if c not in _excl)
                for _, _, routes in all_legs for r in routes
            ]
            _seen_fps = set()
            _deduped_routes = []
            for _r in _raw_routes:
                _fp = tuple(c for c, n in _r)  # fingerprint = city sequence only
                if _fp not in _seen_fps:
                    _seen_fps.add(_fp)
                    _deduped_routes.append(list(_r))
            # Sort: closest to requested duration first
            _deduped_routes.sort(key=lambda r: abs(sum(n for _, n in r) - _target_n))
            intent['all_routes'] = _deduped_routes

            # ── B12: inject must-include cities (those agent explicitly named
            # BEFORE regional default merged in template cities) that aren't
            # already in the regional default routes. These are cities the agent
            # asked for ("must include 2 nights in Chidambaram") that the
            # regional template doesn't naturally include.
            #
            # GEO_NEIGHBOURS guides where to insert each city in the existing
            # sequence. After insertion, redistribute night counts so the route
            # still totals to the requested duration.
            _agent_specified = intent.get('_agent_must_include') or []
            _template_cities = set(merged_cities)  # cities pulled in by regional default
            _must_include = [
                c for c in _agent_specified
                if c not in (intent.get('cities_excluded') or [])
                and c not in _template_cities  # only inject if NOT already in template
            ]
            # Only inject if route-builder didn't already build from cities_detected
            # i.e. cities_source = 'regional_default' (single-city / region-only emails)
            if _must_include and intent.get('cities_source') == 'regional_default':
                # Geographic neighbours — when X is requested, insert it after Y
                # (or before Z if Y absent). Tuple = (city, [insert_after_options], default_nights)
                GEO_NEIGHBOURS = {
                    'Chidambaram':            (['Puducherry (Pondicherry)', 'Mamallapuram (Mahabalipuram)', 'Chennai'], 1),
                    'Kumbakonam':             (['Chidambaram', 'Thanjavur (Tanjore)', 'Mamallapuram (Mahabalipuram)'], 1),
                    'Gangaikondacholapuram':  (['Chidambaram', 'Kumbakonam', 'Thanjavur (Tanjore)'], 0),
                    'Chettinadu':             (['Madurai', 'Thanjavur (Tanjore)'], 1),
                    'Tranquebar':             (['Chidambaram', 'Kumbakonam'], 1),
                    'Mysuru (Mysore)':        (['Bengaluru'], 2),
                    'Coorg':                  (['Mysuru (Mysore)', 'Bengaluru'], 2),
                    'Hampi':                  (['Bengaluru', 'Goa'], 2),
                    'Pushkar':                (['Jaipur', 'Ajmer'], 1),
                    'Bundi':                  (['Ranthambore', 'Jaipur'], 1),
                    'Nawalgarh':              (['Mandawa', 'Bikaner'], 1),
                }

                # Parse explicit night counts from email body for each city
                # e.g. "2 nights in chidambaram" -> {'Chidambaram': 2}
                _email_lower = (intent.get('_raw_text') or '').lower()
                _explicit_nights = {}
                for _ci in _must_include:
                    _ci_low = _ci.split(' (')[0].lower()
                    _m = re.search(rf'(\d+)\s*nights?\s*(?:in|at)\s*{re.escape(_ci_low)}', _email_lower)
                    if not _m:
                        _m = re.search(rf'{re.escape(_ci_low)}.*?(\d+)\s*nights?', _email_lower[:_email_lower.find(_ci_low)+len(_ci_low)+50] if _ci_low in _email_lower else '')
                    if _m:
                        _explicit_nights[_ci] = int(_m.group(1))

                _injected_any = False
                _new_all_routes = []
                for _route in intent['all_routes']:
                    _route_cities = [c for c, n in _route]
                    _modified = list(_route)
                    _route_inj = False
                    for _c in _must_include:
                        if _c in _route_cities:
                            # already in route — just check night count if explicit
                            if _c in _explicit_nights:
                                for _i, (_rc, _rn) in enumerate(_modified):
                                    if _rc == _c and _rn != _explicit_nights[_c]:
                                        _modified[_i] = (_c, _explicit_nights[_c])
                                        _route_inj = True
                            continue
                        # Need to inject — find anchor neighbour
                        _neighbours, _default_n = GEO_NEIGHBOURS.get(_c, ([], 1))
                        _nights_for_c = _explicit_nights.get(_c, _default_n)
                        _insert_idx = None
                        for _nb in _neighbours:
                            for _i, (_rc, _rn) in enumerate(_modified):
                                if _rc == _nb:
                                    _insert_idx = _i + 1
                                    break
                            if _insert_idx is not None:
                                break
                        if _insert_idx is None:
                            # No neighbour found — append at end (better than dropping)
                            _insert_idx = len(_modified)
                        _modified.insert(_insert_idx, (_c, _nights_for_c))
                        _route_inj = True
                        _injected_any = True
                    # Re-balance nights so route total ≈ requested duration
                    if _route_inj and _target_n and _target_n > 0:
                        _cur_total = sum(n for _, n in _modified)
                        _delta = _cur_total - _target_n
                        if _delta > 0:
                            # Trim from cities not in must_include, prefer largest first
                            _trim_candidates = sorted(
                                [(i, c, n) for i, (c, n) in enumerate(_modified)
                                 if c not in _must_include and n > 1],
                                key=lambda x: -x[2]
                            )
                            for _i, _c, _n in _trim_candidates:
                                if _delta <= 0:
                                    break
                                _take = min(_delta, _n - 1)
                                _modified[_i] = (_c, _n - _take)
                                _delta -= _take
                    _new_all_routes.append(_modified)
                if _injected_any:
                    intent['all_routes'] = _new_all_routes
                    intent['nightly_split'] = _new_all_routes[0]
                    intent['cities_source'] = 'regional_default_with_injection'
                    # Add injected cities back into cities_detected so hotel/
                    # activity/monument blocks downstream pick them up
                    for _c in _must_include:
                        if _c not in intent['cities_detected']:
                            intent['cities_detected'].append(_c)
                    print(f"  ℹ  Must-include cities injected into routes: {', '.join(_must_include)}")

            # Update primary nightly_split to the best-matching route
            if _deduped_routes and not intent.get('cities_source','').endswith('injection'):
                intent['nightly_split'] = intent['all_routes'][0]
            intent['legs'] = all_legs

            # ── Revisit: inject South India routes into all_routes NOW,
            # before the display block runs, so they appear in route options
            if intent.get('_revisit_south_india_alts'):
                _si_mkt_routes = REGIONAL_ROUTES.get('south_india', {})
                _si_routes = _si_mkt_routes.get(
                    intent.get('source_market', 'DEFAULT'),
                    _si_mkt_routes.get('DEFAULT', [])
                )
                if _si_routes:
                    _excl2 = set(intent.get('cities_excluded', []))
                    _existing2 = intent['all_routes']
                    _existing_fps = {tuple(c for c, n in r) for r in _existing2}
                    _south_added = []
                    for _sr in _si_routes[:3]:
                        _sr_clean = [(c, n) for c, n in _sr if c not in _excl2]
                        _sr_fp = tuple(c for c, n in _sr_clean)
                        if _sr_fp not in _existing_fps:
                            _south_added.append(_sr_clean)
                            _existing_fps.add(_sr_fp)
                    if _south_added:
                        intent['all_routes'] = _existing2 + _south_added

            # Auto-add domestic flight when Jabalpur in merged route
            if any(c == 'Jabalpur' for c, n in merged_route):
                if 'domestic flights mentioned' not in intent.get('transport_notes', []):
                    intent.setdefault('transport_notes', []).append('domestic flights mentioned')
                if not intent.get('domestic_sectors'):
                    _IATA_GW2 = {'Mumbai': 'BOM', 'New Delhi': 'DEL', 'Chennai': 'MAA',
                                  'Bengaluru': 'BLR', 'Kolkata': 'CCU'}
                    _first_c = merged_route[0][0] if merged_route else 'New Delhi'
                    _orig2 = _IATA_GW2.get(_first_c, 'DEL')
                    intent['domestic_sectors'] = [{'orig': _orig2, 'dest': 'JLR', 'date': ''}]
                    print(f"  ✈ Domestic flight added: {_orig2} -> JLR (Jabalpur gateway)")

            # Print GT redirect notice if applicable
            if _gt_excluded and not _north_india_explicit:
                print(f"\n  ℹ  Golden Triangle previously visited — redirecting to South India / Kerala.")
                print(f"     (To suggest North India instead, mention it explicitly in the request.)")
            # Display all routes sorted by proximity to requested duration
        # intent['all_routes'] is already deduped and duration-sorted
        _display_routes = intent.get('all_routes', [])
        if _display_routes and not intent.get('_wildlife_multi_region'):
            _primary = fix_transit_stops(list(_display_routes[0]))
            _prim_total = sum(n for _, n in _primary)
            if intent.get('_mixed_region_display') and all_legs and len(all_legs) >= 2:
                # Compact one-line-per-route display grouped by region
                _opt_counter = 1
                for _region, _best, _region_routes in all_legs:
                    _region_label = _region.replace('_', ' ').title()
                    print(f"\n  {_region_label} Options")
                    print(f"  {'─' * (len(_region_label) + 8)}")
                    for _rr in _region_routes[:4]:
                        _rr = fix_transit_stops(list(_rr))
                        _rr_total = sum(n for _, n in _rr)
                        _seen = []
                        for city, nights in _rr:
                            short = city.replace(' (RJ)', '').replace(' (Cochin)', '') \
                                        .replace(' (Alleppey)', '').replace('(Periyar/Kumily)', '') \
                                        .replace('Mamallapuram (Mahabalipuram)', 'Mamallapuram') \
                                        .replace('Puducherry (Pondicherry)', 'Pondicherry') \
                                        .replace('Thanjavur (Tanjore)', 'Thanjavur')
                            if short not in _seen:
                                _seen.append(short)
                        print(f"  Option {_opt_counter:2}  {_rr_total:2}n   {' → '.join(_seen)}")
                        _opt_counter += 1
            else:
                def _arrow_route(route):
                    _seen = []
                    for _c, _n in route:
                        _s = _c.replace(' (RJ)', '').replace(' (Cochin)', '') \
                                .replace(' (Alleppey)', '').replace('(Periyar/Kumily)', '') \
                                .replace('Mamallapuram (Mahabalipuram)', 'Mamallapuram') \
                                .replace('Puducherry (Pondicherry)', 'Pondicherry') \
                                .replace('Thanjavur (Tanjore)', 'Thanjavur')
                        _l = f"{_s} ({_n}n)" if _n > 0 else f"{_s} (day)"
                        if _l not in _seen: _seen.append(_l)
                    return ' → '.join(_seen)
                _leg_label = (all_legs[0][0] if all_legs else (
                    (intent.get('regions_detected') or ['north_india'])[0]
                )).replace('_', ' ').title()
                if intent.get('cities_source') == 'cascade':
                    # BL2-α: cascade is route source; banner reflects which pass
                    # actually produced the primary route (agent / market /
                    # region / global) — not always 'agent's pattern' just
                    # because the cascade ran.
                    _pass = intent.get('_cascade_primary_pass', 'global')
                    if _pass == 'agent':
                        _phrase = "Using agent's historical pattern"
                    elif _pass == 'market':
                        _phrase = f"Using {intent.get('source_market', 'market')} market booking pattern"
                    elif _pass == 'region':
                        _phrase = "Using regional booking pattern"
                    else:  # global, or anything unrecognised
                        _phrase = "Using TCI-wide booking pattern"
                    print(f"\n  ✦ {_phrase} — {_leg_label} "
                          f"({_prim_total}n): {_arrow_route(_primary)}")
                    # P2-12 (v20260522b): two-tier duration overshoot note.
                    # ℹ soft: primary is >3n from request (mild mismatch).
                    # ⚠ hard: primary is >30% longer than request (significant overshoot).
                    _req_dur_check = _req_dur_scalar(intent.get('duration_nights'))
                    if _req_dur_check is not None and _req_dur_check > 0:
                        _overshoot = _prim_total - _req_dur_check
                        _overshoot_pct = _overshoot / _req_dur_check
                        if _overshoot > 0 and _overshoot_pct > 0.30:
                            print(f"  ⚠  Primary route is {_prim_total}n — "
                                  f"{_overshoot}n longer than requested {_req_dur_check}n "
                                  f"({round(_overshoot_pct*100)}% overshoot) — "
                                  f"shorter alternatives below may be a better fit")
                        elif abs(_overshoot) > 3:
                            print(f"  ℹ  Primary route is {_prim_total}n vs requested "
                                  f"{_req_dur_check}n — check alternatives")
                else:
                    print(f"\n  ⚡ No cities specified — applying {intent['source_market']} "
                          f"{_leg_label} default route:\n     ({_prim_total}n): {_arrow_route(_primary)}")
                if len(_display_routes) > 1:
                    # v20260514b: routes already clamped + rebalanced by
                    # _apply_city_night_rules() in _envelope_to_all_routes().
                    # Light display filter retained as safety net (±5n) for
                    # any routes that slipped through without a rules match.
                    _req_dur = _req_dur_scalar(intent.get('duration_nights'))
                    print(f"\n  Alternative routes:")
                    for i, alt in enumerate(_display_routes[1:], 2):
                        alt = fix_transit_stops(list(alt))
                        alt_total = sum(n for _, n in alt)
                        if _req_dur is not None and abs(alt_total - _req_dur) > 5:
                            continue
                        print(f"  Option {i:2}  {alt_total:2}n   {_arrow_route(alt)}")
            print(f"  Transport    : {', '.join(intent['transport_notes']) or 'None mentioned'}")
    if intent.get('named_trains'):
        print(f"  Named Trains : {', '.join(intent['named_trains'])}")
    if intent.get('agent_specified_hotels'):
        print(f"  Agent Hotels : {', '.join(intent['agent_specified_hotels'])}")
    if intent['activities_requested']:
        print(f"  Activities   : {', '.join(intent['activities_requested'])}")

    # Vehicle
    print("\n── VEHICLE ───────────────────────────────────────────────────────")
    if intent['num_pax']:
        # B25: vehicle capacity sized by total bodies (adults + children),
        # not just adults — kids occupy seats and need transfers too.
        _vehicle_pax = intent.get('num_pax_total') or intent['num_pax']
        vehicle, vehicle_note = recommend_vehicle(_vehicle_pax)
        intent['_vehicle'] = {"name": vehicle, "note": vehicle_note}
        print(f"  ✓ {vehicle}  ({vehicle_note})")
    else:
        print("  ⚠ Pax count not detected — check email manually")

    con = sqlite3.connect(db_path)
    if portfolio_path:
        load_portfolio(portfolio_path)
    if city_rules_path:
        load_city_rules(city_rules_path)
    if activity_path:
        load_activity_portfolio(activity_path)
    if fnb_path:
        load_fnb_portfolio(fnb_path)
        # Build hotel code→name map now that portfolio is loaded
        if _PORTFOLIO_BY_CITY and not _PORTFOLIO_CODE_NAME:
            for city_hotels in _PORTFOLIO_BY_CITY.values():
                for h in city_hotels:
                    if h.get('code'):
                        _PORTFOLIO_CODE_NAME[h['code']] = h['name']

    # ── Wildlife city injection (EARLY — before route building) ─────────────────
    # When wildlife/safari/tiger mentioned + central_india detected but no park named,
    # inject appropriate parks based on explicitly requested wildlife nights
    _wl_cities_known = {'Bandhavgarh', 'Kanha', 'Pench', 'Panna', 'Corbett',
                        'Jim Corbett', 'Ranthambore', 'Kaziranga', 'Tadoba',
                        'Kabini', 'Nagarhole', 'Wayanad', 'Satpura'}
    _wl_already = bool(_wl_cities_known & set(intent.get('cities_detected', [])))
    # P2-04 (v20260522b): tightened trigger — mirrors display site (~12248) and
    # route-validation site (~13315). Bare 'tiger'/'jungle' removed to avoid
    # false-positives on Tiger's Nest, Tiger Hill, 'concrete jungle' etc.
    _wl_text = email_text.lower()
    _wl_strong_kw3 = any(s in _wl_text for s in [
        'wildlife', 'wild life', 'safari', 'leopard', 'rhino', 'elephant herd',
        'game drive', 'national park', 'tiger reserve', 'tiger safari',
        'tiger sighting', 'tigress',
    ])
    _wl_jungle_kw3 = bool(re.search(
        r'\bjungle\s+(?:safari|lodge|camp|trek|drive|resort|stay|experience)\b',
        _wl_text, re.IGNORECASE
    ))
    _wl_requested = _wl_strong_kw3 or _wl_jungle_kw3
    if _wl_requested and not _wl_already and 'central_india' in (intent.get('regions_detected') or []):
        # Detect explicitly stated wildlife/central india nights from email
        import re as _re
        _wl_night_pattern = _re.compile(
            r'(\d+)\s*nights?\s+(?:in\s+)?(?:central\s+india|wildlife|the\s+parks?|jungle)',
            _re.IGNORECASE
        )
        _wl_match = _wl_night_pattern.search(email_text)
        if _wl_match:
            _wl_nights = int(_wl_match.group(1))
        else:
            # Fallback: total nights minus known gateway nights
            _dn = intent.get('duration_nights')
            _total_n = _dn[1] if isinstance(_dn, tuple) else (_dn or 0)
            _gateway_n = sum(2 for c in intent.get('cities_detected', [])
                             if c in {'Mumbai', 'New Delhi', 'Chennai', 'Bengaluru', 'Kolkata'}) +                          sum(1 for c in intent.get('cities_detected', []) if c in {'Agra'})
            _wl_nights = max(0, _total_n - _gateway_n)

        # Scale parks to wildlife nights
        if _wl_nights >= 8:
            _wl_inject = ['Bandhavgarh', 'Kanha', 'Pench']
        elif _wl_nights >= 6:
            _wl_inject = ['Bandhavgarh', 'Kanha']
        else:
            _wl_inject = ['Bandhavgarh']

        for _wc in _wl_inject:
            if _wc not in intent.get('cities_detected', []):
                intent.setdefault('cities_detected', []).append(_wc)
        print(f"  🐯 Wildlife cities injected ({_wl_nights}n): {', '.join(_wl_inject)}")

    # ── Bookend route logic ────────────────────────────────────────────────────
    # When agent specifies entry + exit cities but not intermediate stops,
    # query DB for historical routes between those bookends within duration range
    ENTRY_CITIES = ['New Delhi', 'Mumbai', 'Chennai', 'Kochi (Cochin)',
                    'Bengaluru', 'Kolkata', 'Goa', 'Amritsar']
    EXIT_CITIES  = ENTRY_CITIES

    detected = intent['cities_detected']
    duration = intent.get('duration_nights')

    entry_city = None
    exit_city  = None

    # Detect entry/exit from email context
    text_lower = intent.get('_raw_text', '')
    if len(detected) >= 2 and duration:
        # Check if first and last detected cities are gateway cities
        if detected[0] in ENTRY_CITIES and detected[-1] in EXIT_CITIES:
            entry_city = detected[0]
            exit_city  = detected[-1]
        # Also check for explicit fly in/out language matched to cities
        for city in detected:
            city_lower = city.lower().replace(' (cochin)', '').replace('new ', '')
            if any(phrase in (email_text.lower()) for phrase in [
                f'arrive into {city_lower}', f'fly in to {city_lower}',
                f'fly into {city_lower}', f'arrive {city_lower}',
                f'arriving {city_lower}'
            ]):
                entry_city = city
            if any(phrase in (email_text.lower()) for phrase in [
                f'depart from {city_lower}', f'fly out from {city_lower}',
                f'fly out of {city_lower}', f'depart {city_lower}',
                f'departing {city_lower}'
            ]):
                exit_city = city

    # Beach/leisure cities — only include if explicitly requested
    BEACH_CITIES = {
        'Goa', 'Kovalam', 'Mararikulam', 'Varkala', 'Pondicherry',
        'Puducherry (Pondicherry)', 'Alappuzha (Alleppey)', 'Puri',
        'Mamallapuram (Mahabalipuram)', 'Diu', 'Tarkarli', 'Alibaug',
    }
    beach_requested = any(s in email_text.lower() for s in [
        'beach', 'coastal', 'sea', 'ocean', 'resort', 'goa', 'leisure',
        'relaxation', 'sunbathe', 'swimming pool', 'pool villa'
    ])

    if (entry_city and exit_city and entry_city != exit_city and duration
            and not _is_tourlane_mode):
        min_n, max_n = duration if isinstance(duration, tuple) else (duration, duration)
        bookend_routes, fallback = get_bookend_routes(
            con, entry_city, exit_city,
            intent['source_market'], min_n, max_n, top_n=6
        )

        # Step 1 — Sort by booking frequency
        bookend_routes = sorted(bookend_routes, key=lambda x: x[1], reverse=True)

        # Step 2 — Filter beach cities unless explicitly requested
        if not beach_requested:
            filtered_routes = [r for r in bookend_routes
                               if not any(c in BEACH_CITIES for c, n in r[0])]
            bookend_routes = filtered_routes if filtered_routes else bookend_routes
        bookend_routes = bookend_routes[:6]

        # Step 3 — Validate direction + reject pure wildlife for sightseeing
        WILDLIFE_ONLY_CITIES = {
            'Pench', 'Kanha', 'Bandhavgarh', 'Toria', 'Tadoba',
            'Corbett', 'Kaziranga', 'Satpura', 'Ranthambore',
        }
        # P2-04 (v20260521b): same tightened trigger as the display site above.
        # Disambiguates Tiger's Nest / Tiger Hill / 'concrete jungle' etc.
        import re as _re_wl2
        _email_lc2 = email_text.lower()
        _wl_strong = any(s in _email_lc2 for s in [
            'wildlife', 'safari', 'leopard', 'rhino', 'game drive',
            'national park', 'tiger reserve', 'tiger safari',
            'tiger sighting', 'tigress',
        ])
        _wl_jungle = bool(_re_wl2.search(
            r'\bjungle\s+(?:safari|lodge|camp|trek|drive|resort|stay|experience)\b',
            _email_lc2
        ))
        wildlife_requested = (
            _wl_strong or _wl_jungle
            or 'central_india' in (intent.get('regions_detected') or [])
        )

        region_is_sightseeing = bool(intent.get('regions_detected')) and not wildlife_requested

        valid_routes = []
        for route, count in bookend_routes:
            cities = [c for c, n in route]
            if entry_city not in cities or exit_city not in cities:
                continue
            ei = cities.index(entry_city)
            xi = cities.index(exit_city)
            if ei >= xi:
                continue
            if region_is_sightseeing:
                interior = set(cities) - {entry_city, exit_city}
                if interior and interior.issubset(WILDLIFE_ONLY_CITIES):
                    continue
            valid_routes.append((route, count))
        bookend_routes = valid_routes

        # Step 4 — Construct from regional defaults if not enough historical routes
        if len(bookend_routes) < 2:
            region = intent.get('regions_detected')
            if isinstance(region, list): region = region[0]
            if region and region in REGIONAL_ROUTES:
                market_routes = REGIONAL_ROUTES[region]
                reg_options = market_routes.get(intent['source_market'],
                              market_routes.get('DEFAULT', []))
                for reg_route in reg_options:
                    reg_cities = [c for c, n in reg_route]
                    if entry_city in reg_cities and exit_city not in reg_cities:
                        constructed = list(reg_route) + [(exit_city, 2)]
                        total = sum(n for _, n in constructed)
                        if total < min_n:
                            constructed[-1] = (exit_city, 2 + (min_n - total))
                            total = min_n
                        if total <= max_n:
                            bookend_routes.append((constructed, 0))
                            if len(bookend_routes) >= 3:
                                break

        # FIX — Redistribute nights: if a city has 3+ nights and Mumbai/exit has only 1,
        # move a night from the longest interior city to the exit city
        # Minimum nights rules — cities that need at least 2 nights to do justice
        CITY_MIN_NIGHTS = {
            'New Delhi':      2,
            'Jaipur':         2,
            'Udaipur (RJ)':   2,
            'Jodhpur':        2,
            'Varanasi':       2,
            'Mumbai':         2,
            'Kochi (Cochin)': 2,
        }

        def enforce_minimum_nights(route, total_max):
            """
            Enforce minimum nights per city — only if overall duration permits.
            Will NOT increase total nights beyond total_max.
            Redistributes from cities with surplus nights only.
            """
            route = list(route)
            for _ in range(10):
                current_total = sum(n for _, n in route)
                deficit_idx = None

                for i, (city, nights) in enumerate(route):
                    if nights == 0:  # skip day visits
                        continue
                    min_required = CITY_MIN_NIGHTS.get(city, 1)
                    if nights >= min_required:
                        continue

                    if current_total < total_max:
                        # Duration has room — add a night directly
                        route[i] = (city, nights + 1)
                        deficit_idx = i
                        break
                    else:
                        # Duration is full — can only redistribute
                        # Find donors: cities with nights strictly above their own minimum
                        donors = [
                            (j, c, n) for j, (c, n) in enumerate(route)
                            if j != i and n > CITY_MIN_NIGHTS.get(c, 1) and n > 1
                        ]
                        if donors:
                            donors.sort(key=lambda x: x[2], reverse=True)
                            donor_idx, donor_city, donor_nights = donors[0]
                            route[donor_idx] = (donor_city, donor_nights - 1)
                            route[i] = (city, nights + 1)
                            deficit_idx = i
                            break
                        # No donors available — cannot enforce minimum, leave as is
                        break

                if deficit_idx is None:
                    break
            return route

        def balance_nights(route, exit_c, min_exit=2):
            route = list(route)
            exit_idx = next((i for i, (c, n) in enumerate(route) if c == exit_c), None)
            if exit_idx is None:
                return route
            # Loop until exit city has minimum nights
            for _ in range(5):  # max 5 iterations
                exit_nights = route[exit_idx][1]
                if exit_nights >= min_exit:
                    break
                # Find longest interior city (not entry or exit) with 2+ nights to spare
                interior = [(i, c, n) for i, (c, n) in enumerate(route)
                            if i != 0 and i != exit_idx and n >= 2]
                if not interior:
                    break
                interior.sort(key=lambda x: x[2], reverse=True)
                donor_idx, donor_city, donor_nights = interior[0]
                route[donor_idx] = (donor_city, donor_nights - 1)
                route[exit_idx] = (exit_c, route[exit_idx][1] + 1)
            return route

        if bookend_routes:
            market_label = f"{'⚠ No ' + intent['source_market'] + ' data — all markets' if fallback else intent['source_market'] + ' market'}"
            # Always show entry → exit in correct direction
            display_entry = entry_city
            display_exit  = exit_city
            print(f"\n── SUGGESTED ROUTES  ({display_entry} → ... → {display_exit})  [{market_label}] ──")
            print(f"  Based on {min_n}-{max_n} night historical bookings\n")

            adjusted_routes = []  # must be initialised before the loop
            for i, (route, count) in enumerate(bookend_routes, 1):
                # Ensure correct direction — entry first, exit last
                cities_in_route = [c for c, n in route]
                if cities_in_route and cities_in_route[0] != entry_city and entry_city in cities_in_route:
                    # Reorder to start from entry_city
                    entry_idx = cities_in_route.index(entry_city)
                    route = route[entry_idx:] + route[:entry_idx]
                # Balance nights — give exit city minimum nights based on context
                # If activities requested in exit city, ensure at least 2 nights
                exit_has_activity = any(
                    s in email_text.lower() for s in [
                        'early morning', 'morning tour', 'full day',
                        'sightseeing', 'tour in ' + exit_city.lower(),
                        exit_city.lower() + ' tour', exit_city.lower() + ' sightseeing'
                    ]
                )
                min_exit_nights = 2 if exit_has_activity else 2
                route = balance_nights(route, exit_city, min_exit=min_exit_nights)
                # Enforce city minimum nights where duration permits
                route = enforce_minimum_nights(route, max_n)
                route = insert_fatehpur_sikri(route)
                route = fix_transit_stops(route, market=intent.get('source_market'))
                adjusted_routes.append(route)
                total = sum(n for _, n in route)
                print(f"  Option {i}  ({total} nights)  —  booked {count}x")
                for city, nights in route:
                    if nights == 0:
                        print(f"     {city:35} day visit (en route)")
                    else:
                        print(f"     {city:35} {nights}n")
                print()

            # Use best route to populate cities for recommendations below
            best_route = bookend_routes[0][0]
            all_route_cities_bookend = []
            for route, count in bookend_routes:
                for city, nights in route:
                    if city not in all_route_cities_bookend:
                        all_route_cities_bookend.append(city)

            intent['nightly_split'] = adjusted_routes[0] if adjusted_routes else best_route
            intent['all_routes'] = adjusted_routes if adjusted_routes else [r for r, c in bookend_routes]
            # Merge with already detected cities
            for city, nights in best_route:
                if city not in intent['cities_detected']:
                    intent['cities_detected'].append(city)

    # Similar historical routes — B27-INT (v20260512a)
    # Wires historical_query.py v0.5.4b into recommend(). Three-pass query
    # (agent → market → global), MAX-merge, score, top-N. Returns scored
    # Options with breakdown, confidence labels, nightly_split projections.
    #
    # Design choices:
    #   - Lazy import: engine remains importable if historical_query is absent
    #   - No fallback to legacy get_similar_routes — fail loud, no fake noise
    #   - Stash on intent['historical_options'] for Word doc / Travart / Outlook
    #
    # Input keys read (with legacy fallback in v0.5.4b C11):
    #   cities_detected, regions_detected, agency_account_code|account_code,
    #   source_market|market, duration_nights, agency_name
    if intent.get('cities_detected') or intent.get('regions_detected'):
        print("\n── SIMILAR HISTORICAL ROUTES ─────────────────────────────────────")
        try:
            import historical_query as _hq
            # BL2-α (13 May): if cascade already ran as primary route source, reuse
            # its envelope here to avoid a second SQL pass and prevent drift between
            # the route-source cascade and the display cascade.
            # Wildlife multi-region exception: when _wildlife_multi_region is set,
            # _cascade_envelope was built against ALL wildlife cities (30+ across zones)
            # which is too broad for meaningful ranking. Instead run a fresh query
            # using the original sub-region + parent region context (stripped of the
            # wildlife city injection) so results reflect genuine agent/market history
            # for this geography (e.g. Rajasthan + Ranthambore sequences for ITA).
            options = intent.get('_cascade_envelope')
            if options is None or intent.get('_wildlife_multi_region'):
                _acct = intent.get('agency_account_code') or intent.get('account_code')
                intent['_agent_total_bookings'] = (
                    _hq.get_agent_total_bookings(db_path, _acct) if _acct else 0
                )
                if intent.get('_wildlife_multi_region'):
                    # Wildlife: run per-zone three-tier queries
                    zone_results = _hq.wildlife_historical_query(intent, db_path=db_path)
                    _hq.print_wildlife_options(zone_results)
                    intent['historical_options'] = zone_results
                else:
                    candidates = _hq.historical_query(intent, db_path=db_path)
                    options = _hq.assemble_options(candidates, intent)
                    intent['historical_options'] = options
                    _hq.print_options(options)
            else:
                # Cascade envelope reused — still must surface on historical_options
                intent['historical_options'] = options
                _hq.print_options(options)
        except Exception as _e:
            print(f"  ⚠ B27 historical query failed: {type(_e).__name__}: {_e}")
            intent['historical_options'] = []

    # Build full city list — union of all route options so hotels/activities/monuments
    # are shown for every city the consultant might choose, not just the primary route
    all_route_cities = list(intent['cities_detected'])  # primary route cities
    if intent.get('all_routes'):
        for route in intent['all_routes'][1:]:
            for city, nights in route:
                if city not in all_route_cities:
                    all_route_cities.append(city)

    # ── SPLIT ITINERARY: add south india default cities if south_india detected
    # but no south india cities in route yet
    SOUTH_INDIA_CITIES = [
        'Bengaluru', 'Mysuru (Mysore)', 'Ooty (Udhagamandalam)', 'Coimbatore',
        'Chennai', 'Mamallapuram (Mahabalipuram)', 'Thanjavur (Tanjore)', 'Madurai',
        'Thekkady (Periyar/Kumily)', 'Munnar', 'Kumarakom', 'Alappuzha (Alleppey)',
        'Kochi (Cochin)', 'Kovalam', 'Varkala', 'Goa', 'Mangalore',
        'Hampi', 'Badami', 'Hyderabad', 'Puducherry (Pondicherry)',
    ]
    regions = intent.get('regions_detected') or []
    if 'south_india' in regions and intent.get('cities_source') not in ('regional_default', 'cascade'):
        # Only expand with Kerala default cities when south_india was NOT already
        # built as the primary route (e.g. GT redirect, or BL2-α cascade output).
        # Prevents duplicate display.
        KERALA_DEFAULT = [
            'Bengaluru', 'Mysuru (Mysore)', 'Thekkady (Periyar/Kumily)',
            'Munnar', 'Kumarakom', 'Alappuzha (Alleppey)', 'Kochi (Cochin)'
        ]
        SOUTH_NIGHTS = {
            'Bengaluru':1, 'Mysuru (Mysore)':1, 'Thekkady (Periyar/Kumily)':2,
            'Munnar':2, 'Kumarakom':2, 'Alappuzha (Alleppey)':1, 'Kochi (Cochin)':2
        }
        detected_south = [c for c in intent['cities_detected'] if c in SOUTH_INDIA_CITIES]
        for c in detected_south:
            if c not in all_route_cities:
                all_route_cities.append(c)
        for c in KERALA_DEFAULT:
            if c not in all_route_cities:
                all_route_cities.append(c)
        intent['south_nights'] = SOUTH_NIGHTS
        print('\n  -- SOUTH INDIA LEG (10-11 nights) --')
        for c in KERALA_DEFAULT:
            n = SOUTH_NIGHTS.get(c, 1)
            print(f'     {c:35} {n}n')

    # Hotels per city
    # Print split itinerary notice if both north + south regions
    regions = intent.get('regions_detected') or []
    if 'north_india' in regions and 'south_india' in regions:
        print("\n  ★ SPLIT ITINERARY DETECTED — North India + South India")
        print("    Hotels shown for both legs below.")

    print("\n── HOTEL RECOMMENDATIONS ─────────────────────────────────────────")
    intent['_top_hotels']    = {}  # city -> top hotel name
    intent['_hotel_details'] = {}  # city -> [(name, tier, bookings)]
    agent_hotels = intent.get('agent_specified_hotels', [])
    if agent_hotels:
        print(f"\n  ★ AGENT SPECIFIED:")
        for h in agent_hotels:
            print(f"    → {h}")
        print(f"\n  (Frequency-based suggestions per city below for reference)")

    # Cities that are structural transit-only stops — never have hotels
    TRANSIT_ONLY_CITIES = {'Fatehpur Sikri', 'Orchha', 'Jabalpur'}

    for city in all_route_cities:
        # Determine nights in primary route
        primary_nights = None
        if intent.get('all_routes'):
            for c, n in intent['all_routes'][0]:
                if c == city:
                    primary_nights = n
                    break
        elif intent.get('nightly_split'):
            for c, n in intent['nightly_split']:
                if c == city:
                    primary_nights = n
                    break

        # Structural transit stops — no hotel, no further output
        if city in TRANSIT_ONLY_CITIES:
            print(f"\n  {city}  [Transit stop — no hotel]")
            continue

        # 0-night real destination (day visit / fly-out city) — skip hotel block only
        if primary_nights == 0:
            print(f"\n  {city}  [Day visit — no overnight]")
            continue
        # Get nightly split if available — check all routes
        nights_str = ''
        route_tag = ''
        if intent.get('all_routes'):
            found_in_opt1 = False
            for ri, route in enumerate(intent['all_routes'], 1):
                for c, n in route:
                    if c == city:
                        if ri == 1:
                            nights_str = f" · {n} nights"
                            found_in_opt1 = True
                        elif not found_in_opt1:
                            nights_str = f" · {n} nights"
                            route_tag = f" [Option {ri}]"
                        break
                if found_in_opt1:
                    break
        elif intent.get('nightly_split'):
            for c, n in intent['nightly_split']:
                if c == city:
                    nights_str = f" · {n} nights"
                    break
        hotels, source = get_hotels(
            con, city,
            market         = intent['source_market'],
            tier           = intent['tier'],
            heritage       = intent['heritage_requested'],
            top_n          = 4,
            agency_account = intent.get('agency_account_code') or None,
            agency_name    = intent.get('agency_name') or None
        )
        # Store top hotel for proposal generation
        if hotels:
            intent['_top_hotels'][city]    = hotels[0][0]  # name
            intent['_hotel_details'][city] = hotels
        print(f"\n  {city}{nights_str}{route_tag}  [{source or 'No hotel data — all markets'}]")
        if hotels:
            for name, tier, bookings in hotels:
                # Heritage: name signals OR portfolio official_classification
                _p_info = _PORTFOLIO_NAME_CITY_TIER.get((name.lower(), city.lower())) \
                    or _PORTFOLIO_NAME_TIER.get(name.lower())
                _is_heritage_official = (_p_info and _p_info.get('official') in
                    {'heritage', 'luxury-heritage-boutique-hotel', '5str-lux'})
                heritage_flag = '🏛' if (
                    any(s in name.lower() for s in HERITAGE_NAME_SIGNALS)
                    or _is_heritage_official
                ) else '  '
                # Flag if agent mentioned this hotel
                agent_flag = ' ← AGENT SPECIFIED' if any(
                    a.lower() in name.lower() or name.lower() in a.lower()
                    for a in agent_hotels
                ) else ''
                print(f"    {heritage_flag} {bookings:4}x  {tier:14}  {name}{agent_flag}")
        else:
            print(f"    ⚠ No hotel data found for {city}")

    # Activities per city
    print("\n── ACTIVITIES ────────────────────────────────────────────────────")
    intent['_activities'] = {}  # city -> [(name, count)]
    intent['_monuments']  = {}  # city -> [(name, count)]
    for city in all_route_cities:
        activities, source = get_activities(con, city, intent['source_market'], top_n=5)
        if activities:
            intent['_activities'][city] = activities[:3]
            print(f"\n  {city}  [{source or 'No market data — all markets'}]")
            is_csv = source and 'CSV' in source
            for name, value in activities:
                # DKC flag
                is_dkc = '[A DKC Experience]' in name
                dkc_tag = '  🎯 DKC' if is_dkc else ''
                # Requested flag
                requested = any(
                    name.lower() in req.lower() or req.lower() in name.lower()
                    for req in intent['activities_requested']
                )
                req_flag = '  ★ REQUESTED' if requested else ''
                if is_csv:
                    print(f"    {value:4.1f}★  {name}{dkc_tag}{req_flag}")
                else:
                    print(f"    {int(value):4}x  {name}{dkc_tag}{req_flag}")

    # Explicitly requested activities not yet matched
    unmatched = []
    for req in intent['activities_requested']:
        matched = False
        for city in all_route_cities:
            acts, _ = get_activities(con, city, intent['source_market'], top_n=10)
            if any(req.lower() in a[0].lower() or a[0].lower() in req.lower() for a in acts):
                matched = True
                break
        if not matched:
            unmatched.append(req)
    if unmatched:
        print(f"\n  ★ REQUESTED (confirm manually):")
        for req in unmatched:
            print(f"    - {req}")

    # Monuments per city
    print("\n── MONUMENTS ─────────────────────────────────────────────────────")
    for city in all_route_cities:
        # Get top hotel for this city (for proximity sorting)
        _top_hotel_name = intent.get('_top_hotels', {}).get(city, '')
        monuments, source = get_monuments(con, city, intent['source_market'], top_n=8)
        # Apply proximity sort — nearest 4 monuments to hotel
        if monuments and _top_hotel_name:
            try:
                _sorted = engine_sort_monuments_by_proximity(
                    monuments, hotel_name=_top_hotel_name, max_per_day=4)
                monuments = [(_n, _b) for _n, _b, _d in _sorted]
            except Exception:
                monuments = monuments[:4]
        else:
            monuments = monuments[:4]
        if monuments:
            intent['_monuments'][city] = monuments[:3]
            print(f"\n  {city}  [{source or 'No market data — all markets'}]")
            for name, bookings in monuments:
                print(f"    {bookings:4}x  {name}")

    # F&B per city
    print("\n── F&B RECOMMENDATIONS ───────────────────────────────────────────")
    intent['_fnb'] = {}  # city -> [fnb records]
    fnb_shown = False
    for city in all_route_cities:
        fnb_rows = get_fnb(city, top_n=4)
        if fnb_rows:
            intent['_fnb'][city] = fnb_rows
            fnb_shown = True
            print(f"\n  {city}")
            for r in fnb_rows:
                type_label = {
                    'specialty-restaurant': 'Specialty',
                    'multi-cuisine':        'Multi-cuisine',
                    'coffee-shop':          'All-day dining',
                    'lounge':               'Lounge',
                    'bar':                  'Bar',
                }.get(r['restaurant_type'], r['restaurant_type'].title())
                # Hotel attribution for hotel-linked restaurants
                if r['hotel_linked']:
                    hotel_name = _PORTFOLIO_CODE_NAME.get(r['hotel_code'], '')
                    hotel_tag  = f"  @ {hotel_name}" if hotel_name else '  @ hotel property'
                else:
                    hotel_tag = '  [standalone]'
                dkc_tag = '  🎯 DKC' if r['tour_category'] == 'Experiential' else ''
                print(f"    {r['tci_score']:4.1f}★  {r['name']}{hotel_tag}  [{type_label}]{dkc_tag}")
    if not fnb_shown:
        print("  No F&B data available for route cities")

    # Language guide flag
    if intent['language_guide']:
        print("\n── GUIDE ─────────────────────────────────────────────────────────")
        lang_signal = intent['language_guide'].lower()
        is_accompanying = any(w in lang_signal for w in [
            'accompanying', 'accompan', 'tour escort', 'tour manager',
            'tour leader', 'full time', 'full-time'
        ])
        if is_accompanying:
            print(f"  ⚠ ACCOMPANYING GUIDE REQUIRED: '{intent['language_guide']}'")
            print(f"     Action: Arrange full-trip accompanying guide/tour escort")
        else:
            print(f"  ⚠ LANGUAGE ESCORT REQUIRED: '{intent['language_guide']}'")
            print(f"     Action: Source specialist language guide before confirming tour")

    # ── DOMESTIC FLIGHTS ─────────────────────────────────────────────────────────
    # Last-chance sector detection from route cities (if not set by PNR or injection)
    if 'domestic flights mentioned' in intent.get('transport_notes', []) and not intent.get('domestic_sectors'):
        _dc2 = intent.get('cities_detected', [])
        _IATA_GW4 = {'Mumbai': 'BOM', 'New Delhi': 'DEL', 'Chennai': 'MAA',
                      'Bengaluru': 'BLR', 'Kolkata': 'CCU', 'Kochi (Cochin)': 'COK'}
        _CITY_DEST2 = {
            'Darjeeling': 'IXB', 'Kalimpong': 'IXB',
            'Corbett': 'IXD', 'Jim Corbett': 'IXD',
            'Aurangabad (MH)': 'IXU', 'Srinagar': 'SXR', 'Leh': 'IXL',
            'Amritsar': 'ATQ', 'Varanasi': 'VNS', 'Jodhpur': 'JDH',
            'Jaisalmer': 'JSA', 'Udaipur (RJ)': 'UDR', 'Khajuraho': 'HJR',
        }
        _orig4 = next((code for c, code in _IATA_GW4.items() if c in _dc2), 'DEL')
        _dest4 = next((code for c, code in _CITY_DEST2.items() if c in _dc2), None)
        if _dest4:
            intent['domestic_sectors'] = [{'orig': _orig4, 'dest': _dest4, 'date': ''}]

    domestic_sectors = intent.get('domestic_sectors', [])
    if domestic_sectors or 'domestic flights mentioned' in intent.get('transport_notes', []):
        print("\n── DOMESTIC FLIGHTS ──────────────────────────────────────────────")
        if not domestic_sectors:
            print("  ⚠ Domestic flights mentioned but no sector detected in PNR")
            print("    Verify via GDS or Google Flights")
        else:
            _IATA_NAMES = {
                'DEL': 'New Delhi',  'BOM': 'Mumbai',    'COK': 'Kochi (Cochin)',
                'BLR': 'Bengaluru',  'MAA': 'Chennai',   'HYD': 'Hyderabad',
                'CCU': 'Kolkata',    'AMD': 'Ahmedabad', 'JAI': 'Jaipur',
                'VNS': 'Varanasi',   'LKO': 'Lucknow',   'ATQ': 'Amritsar',
                'JLR': 'Jabalpur',   'GOI': 'Goa',       'UDR': 'Udaipur (RJ)',
                'JDH': 'Jodhpur',    'JSA': 'Jaisalmer', 'KJB': 'Khajuraho',
                'IXC': 'Chandigarh', 'SXR': 'Srinagar',  'IXB': 'Darjeeling',
            }

            def _pnr_to_iso(ds):
                _mm = {'JAN':'01','FEB':'02','MAR':'03','APR':'04','MAY':'05','JUN':'06',
                       'JUL':'07','AUG':'08','SEP':'09','OCT':'10','NOV':'11','DEC':'12'}
                import datetime as _dt2
                m = re.match(r'(\d{1,2})([A-Z]{3})', ds.upper())
                if not m: return None
                day, mon = m.group(1).zfill(2), _mm.get(m.group(2))
                if not mon: return None
                year = _dt2.date.today().year
                cand = f"{year}-{mon}-{day}"
                if _dt2.date.fromisoformat(cand) < _dt2.date.today():
                    cand = f"{year+1}-{mon}-{day}"
                return cand

            for sector in domestic_sectors:
                orig, dest = sector['orig'], sector['dest']
                date_str   = sector.get('date', '')
                orig_city  = _IATA_NAMES.get(orig, orig)
                dest_city  = _IATA_NAMES.get(dest, dest)
                print(f"\n  {orig} → {dest}  ({orig_city} → {dest_city})  [{date_str}]")
                try:
                    from fast_flights import FlightData, Passengers, get_flights
                    iso_date = _pnr_to_iso(date_str) if date_str else None
                    if not iso_date:
                        raise ValueError("No date")
                    fd        = FlightData(date=iso_date, from_airport=orig, to_airport=dest)
                    result_ff = get_flights(
                        flight_data=[fd],
                        trip='one-way',
                        passengers=Passengers(adults=2),
                        seat='economy',
                        fetch_mode='fallback',
                    )
                    flights = result_ff.flights[:5] if result_ff and result_ff.flights else []
                    if flights:
                        print(f"  {'Airline':<20} {'Dep':>8}  {'Arr':>8}  {'Duration':>10}  {'Stops':>6}  Note")
                        print(f"  {'-'*20} {'-'*8}  {'-'*8}  {'-'*10}  {'-'*6}  {'-'*10}")
                        for f in flights:
                            stops_s = 'direct' if f.stops == 0 else f"{f.stops} stop"
                            note    = '★ best value' if f.is_best else ''
                            # Strip verbose date suffix — keep only the time portion
                            dep = re.sub(r'\s+on\s+.*', '', f.departure).strip()
                            arr = re.sub(r'\s+on\s+.*', '', f.arrival).strip()
                            print(f"  {f.name:<20} {dep:>8}  {arr:>8}  "
                                  f"{f.duration:>10}  {stops_s:>6}  {note}")
                        print(f"  ℹ Flight numbers not shown — search {orig}→{dest} {date_str} in GDS to confirm")
                    else:
                        print(f"  No results — verify via GDS or Google Flights")
                except ImportError:
                    print(f"  fast-flights not installed — verify via GDS or Google Flights")
                except Exception as e:
                    print(f"  Live lookup unavailable — verify via GDS or Google Flights")

    # ── Generate proposal Word document ──────────────────────────────────────
    try:
        import os as _os, shutil as _shutil
        # Copy icons to /content/proposal_icons if not already there
        _icon_src = '/content/drive/MyDrive/SITA_TourIntelligence/proposal_icons'
        _icon_dst = '/content/proposal_icons'
        if not _os.path.exists(_os.path.join(_icon_dst, 'icon_sita_logo.png')):
            try:
                if _os.path.exists(_icon_src):
                    _shutil.copytree(_icon_src, _icon_dst, dirs_exist_ok=True)
            except Exception:
                pass
        safe_agency = (intent.get('agency_name') or 'Proposal').replace(' ', '_')[:20]
        proposal_path = f'/content/Proposal_{safe_agency}.docx'
        generate_proposal_docx(intent, output_path=proposal_path, email_text=email_text)
        intent['_proposal_path'] = proposal_path
        # Also save to Drive if mounted
        drive_path = f'/content/drive/MyDrive/SITA_TourIntelligence/proposals/Proposal_{safe_agency}.docx'
        try:
            _os.makedirs(_os.path.dirname(drive_path), exist_ok=True)
            import shutil as _shutil
            _shutil.copy(proposal_path, drive_path)
            print(f"  ✓ Proposal backed up to Drive")
        except:
            pass
    except Exception as _pe:
        print(f"  ⚠ Proposal generation failed: {_pe}")

    print("\n" + "="*70)
    print("  END OF RECOMMENDATION")
    print("="*70 + "\n")

    return {
        # ── What the engine understood ──────────────────────────────────
        "parsed": {
            "market":             intent.get("source_market"),
            "pax":                intent.get("num_pax"),
            "rooms":              intent.get("rooms"),
            "tier":               intent.get("tier"),
            "duration_nights":    intent.get("duration_nights"),
            "heritage_requested": intent.get("heritage_requested", False),
            "language_guide":     intent.get("language_guide"),
            "budget_raw":         intent.get("budget_raw"),
            "cities_detected":    intent.get("cities_detected", []),
            "regions":            intent.get("regions_detected", []),
            "sub_region":         intent.get("sub_region"),
            "travel_start":       intent.get("travel_start"),
            "travel_end":         intent.get("travel_end"),
            "parser_score":       intent.get("parser_score", 0),
            "parser_confidence":  (
                "high"   if intent.get("parser_score", 0) >= 70 else
                "medium" if intent.get("parser_score", 0) >= 40 else
                "low"
            ),
        },

        # ── Ranked itinerary options ────────────────────────────────────
        "itineraries":      intent.get("historical_options", []),

        # ── Per-city content ────────────────────────────────────────────
        "hotels":           intent.get("_top_hotels", {}),
        "activities":       intent.get("_activities", {}),
        "monuments":        intent.get("_monuments", {}),
        "fnb":              intent.get("_fnb", {}),

        # ── Logistics ──────────────────────────────────────────────────
        "vehicle":          intent.get("_vehicle"),
        "domestic_flights": intent.get("domestic_sectors", []),

        # ── Meta ───────────────────────────────────────────────────────
        "proposal_path":    intent.get("_proposal_path"),
        "llm_used":         intent.get("_llm_used", False),
    }


# ── QUICK TEST ────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    # Test with Moretti email
    test_email = """
    Dear all,
    I need for 2 double + 1 single room. Im here for a very important inquiry
    for our VIP clients. MORETTI X 5

    15-17 Delhi, 17-19 Udaipur, 19-20 Jodhpur, 20-22 Jaipur, 22-23 Agra,
    23-24 Khajuraho, 24-26 Varanasi, 26 Oct Delhi wash and change.

    They love heritage accommodation with a good service standard.
    Budget around 2000/2500 per pax.
    Need a very good Italian speaking escort guide.
    Pls include Lake Pichola sunset cruise, Devra cooking class,
    Taj Mahal sunset visit, Aarti ceremony, Ganges sunrise cruise, Sarnath.
    B&B tour but 4 dinners in typical restaurants.
    """
    recommend(test_email)
